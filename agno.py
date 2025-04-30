import re
import os
import json
import time
import string
import chromadb
import asyncio
import textwrap
import random
import csv
import io
import docx
import pypdf
import uuid
import types
import inspect
import requests
import hashlib
import pathlib
import ollama
import collections
import collections.abc
import shutil
import base64
import zlib
import functools
import docstring_parser
import sqlalchemy
from lxml import etree
from rich.box import HEAVY
from rich.panel import Panel
from rich.prompt import Prompt
from rich.console import Group
from rich.json import JSON
from rich.live import Live
from rich.markdown import Markdown
from rich.status import Status
from rich.text import Text
from enum import Enum
from datetime import datetime
from copy import deepcopy
from typing import Any, AsyncIterator, Callable, Dict, Iterator, List, Literal, Optional, Sequence, Set, Type, Union, Tuple, AsyncGenerator, Mapping, get_type_hints, get_args, get_origin


class SafeFormatter(string.Formatter):
    def get_value(self, key, args, kwargs):
        if key not in kwargs:
            return f'{key}'
        return kwargs[key]

    def format_field(self, value, format_spec):
        try:
            return super().format_field(value, format_spec)
        except ValueError:
            return f'{{{value}:{format_spec}}}'


class Dot(dict):
    def __init__(self, seq=None, **kwargs):
        if not isinstance(seq, dict):
            seq = {'value': seq}
        super(Dot, self).__init__(seq, **kwargs)

    def __getattr__(self, attr):
        res = self.get(attr)
        if isinstance(res, dict):
            return Dot(res)
        return res
    __setattr__ = dict.__setitem__
    __delattr__ = dict.__delitem__

    def __call__(self, keypath: str, *args, **kwargs):
        temp = self
        for i in keypath.split('.'):
            temp = temp[int(i) if i.isnumeric() else i]
        return temp


class Base:
    def __init__(self, *args, **kwargs):
        annotations = {}
        for cls in self.__class__.__mro__:
            annotations.update(cls.__dict__.get('__annotations__', {}))
        names = list(annotations.keys())
        for i, v in enumerate(args[:len(names)]):
            setattr(self, names[i], v)
        for k, v in annotations.items():
            if k in kwargs:
                setattr(self, k, kwargs[k])
            elif not hasattr(self, k):
                if hasattr(self.__class__, k):
                    setattr(self, k, getattr(self.__class__, k))
                else:
                    setattr(self, k, None if v == Any else (get_origin(v) or v)())

    def __repr__(self):
        attrs = ', '.join(f"{attr}={getattr(self, attr)!r}" for attr in self.__class__.__annotations__)
        return f"{self.__class__.__name__}({attrs})"

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        return {k: v for k, v in self.__dict__.items() if not k.startswith('_') and (v is not None or not no_none) and (not include or (k in include))}

    @classmethod
    def model_json_schema(cls):
        schema = {
            "type": "object",
            "properties": {},
            "required": []
        }
        def _get_type_name(field_type):
            if field_type is int:
                return "integer"
            elif field_type is float:
                return "number"
            elif field_type is str:
                return "string"
            elif field_type is bool:
                return "boolean"
            elif field_type is None:
                return "null"
            return "object"
        annotations = {}
        for c in cls.__mro__:
            annotations.update(c.__dict__.get('__annotations__', {}))
        for k, t in annotations.items():
            field_schema = {}
            origin = get_origin(t)
            if origin is None:
                field_schema["type"] = _get_type_name(t)
            elif origin is list:
                item_type = get_args(t)[0]
                field_schema["type"] = "array"
                field_schema["items"] = {"type": _get_type_name(item_type)}
            elif origin is dict:
                key_type, value_type = get_args(t)
                field_schema["type"] = "object"
                field_schema["additionalProperties"] = {"type": _get_type_name(value_type)}
            else:
                field_schema["type"] = _get_type_name(t)
            schema["properties"][k] = field_schema
            if hasattr(cls, k):
                default_value = getattr(cls, k)
                if default_value is not None:
                    field_schema["default"] = default_value
            else:
                schema["required"].append(k)
        return schema


class ResponseEvent(Enum):
    start = 'start'
    end = 'end'
    res = 'res'


class RunEvent(Enum):
    run_started = 'RunStarted'
    run_response = 'RunResponse'
    run_completed = 'RunCompleted'
    run_error = 'RunError'
    run_cancelled = 'RunCancelled'
    tool_call_started = 'ToolCallStarted'
    tool_call_completed = 'ToolCallCompleted'
    reasoning_started = 'ReasoningStarted'
    reasoning_step = 'ReasoningStep'
    reasoning_completed = 'ReasoningCompleted'
    updating_memory = 'UpdatingMemory'
    workflow_started = 'WorkflowStarted'
    workflow_completed = 'WorkflowCompleted'


class Timer:
    def __init__(self):
        self.start_time: float = 0
        self.end_time: float = 0.0
        self.elapsed_time: float = 0.0

    @property
    def elapsed(self) -> float:
        return self.elapsed_time or (time.perf_counter() - self.start_time) if self.start_time else 0.0

    def start(self) -> float:
        self.start_time = time.perf_counter()
        return self.start_time

    def stop(self) -> float:
        self.end_time = time.perf_counter()
        if self.start_time:
            self.elapsed_time = self.end_time - self.start_time
        return self.end_time

    def __enter__(self):
        self.start_time = time.perf_counter()
        return self

    def __exit__(self, *args):
        self.end_time = time.perf_counter()
        if self.start_time:
            self.elapsed_time = self.end_time - self.start_time


def get_json_schema_for_arg(t: Any) -> Optional[Dict[str, Any]]:
    type_args = get_args(t)
    type_origin = get_origin(t)
    from types import UnionType
    if type_origin:
        if type_origin in (list, tuple, set, frozenset):
            json_schema_for_items = get_json_schema_for_arg(type_args[0]) if type_args else {'type': 'string'}
            return {'type': 'array', 'items': json_schema_for_items}
        elif type_origin is dict:
            key_schema = get_json_schema_for_arg(type_args[0]) if type_args else {'type': 'string'}
            value_schema = get_json_schema_for_arg(type_args[1]) if len(type_args) > 1 else {'type': 'string'}
            return {'type': 'object', 'propertyNames': key_schema, 'additionalProperties': value_schema}
        elif type_origin in [Union, UnionType]:
            types = []
            for arg in type_args:
                if schema := get_json_schema_for_arg(arg):
                    types.append(schema)
            return {'anyOf': types} if types else None
    types = {
        'number': ['int', 'float', 'complex', 'Decimal'],
        'string': ['str', 'string'],
        'boolean': ['bool', 'boolean'],
        'null': ['NoneType', 'None'],
        'array': ['list', 'tuple', 'set', 'frozenset'],
        'object': ['dict', 'mapping']}
    json_schema = {'type': object}
    for k, v in types.items():
        if t.__name__ in v:
            json_schema['type'] = k
    if json_schema['type'] == 'object':
        json_schema['properties'] = {}
        json_schema['additionalProperties'] = False
    return json_schema


class Function(Base):
    name: str
    entrypoint: Callable = None
    sanitize_arguments: bool = True
    cache_results = False
    cache_dir = ''
    cache_ttl = 3600
    description = ''
    parameters = {'type': 'object', 'properties': {}, 'required': []}
    strict = False
    skip_entrypoint_processing = False
    show_result = False
    stop_after_tool_call = False
    pre_hook: Callable = None
    post_hook: Callable = None
    _agent = None

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        return super().to_dict(include=include or {'name', 'description', 'parameters', 'strict'})

    @classmethod
    def get_entrypoint_docstring(cls, entrypoint: Callable) -> str:
        if isinstance(entrypoint, functools.partial):
            return str(entrypoint)
        doc = inspect.getdoc(entrypoint)
        if not doc:
            return ''
        parsed = docstring_parser.parse(doc)
        lines = []
        if parsed.short_description:
            lines.append(parsed.short_description)
        if parsed.long_description:
            lines.extend(parsed.long_description.split('\n'))
        return '\n'.join(lines)

    @classmethod
    def get_json_schema(cls, type_hints: Dict[str, Any], param_descriptions: Dict[str, str] = None,
                        strict: bool = False) -> Dict[str, Any]:
        json_schema: Dict[str, Any] = {'type': 'object', 'properties': {}}
        if strict:
            json_schema['additionalProperties'] = False
        for k, v in type_hints.items():
            if k == 'return':
                continue
            try:
                type_origin = get_origin(v)
                type_args = get_args(v)
                is_optional = type_origin is Union and len(type_args) == 2 and any(
                    arg is type(None) for arg in type_args)
                if is_optional:
                    v = next(arg for arg in type_args if arg is not type(None))
                if v:
                    arg_json_schema = get_json_schema_for_arg(v)
                else:
                    arg_json_schema = {}
                if arg_json_schema:
                    if is_optional:
                        if isinstance(arg_json_schema['type'], list):
                            arg_json_schema['type'].append('null')
                        else:
                            arg_json_schema['type'] = [arg_json_schema['type'], 'null']
                    if param_descriptions and k in param_descriptions and param_descriptions[k]:
                        arg_json_schema['description'] = param_descriptions[k]
                    json_schema['properties'][k] = arg_json_schema
                else:
                    print(f'Could not parse argument {k} of type {v}')
            except Exception as e:
                print(f'Error processing argument {k}: {str(e)}')
                continue
        return json_schema

    @classmethod
    def from_callable(cls, c: Callable, strict: bool = False) -> 'Function':
        function_name = c.__name__
        parameters = {'type': 'object', 'properties': {}, 'required': []}
        try:
            sig = inspect.signature(c)
            type_hints = get_type_hints(c)
            if 'agent' in sig.parameters:
                del type_hints['agent']
            param_type_hints = {name: type_hints.get(name) for name in sig.parameters if name != 'return' and name != 'agent'}
            param_descriptions = {}
            if docstring := inspect.getdoc(c):
                parsed_doc = docstring_parser.parse(docstring)
                param_docs = parsed_doc.params
                for param in param_docs:
                    param_name = param.arg_name
                    param_type = param.type_name
                    param_descriptions[param_name] = f'({param_type}) {param.description}'
            parameters = cls.get_json_schema(type_hints=param_type_hints, param_descriptions=param_descriptions, strict=strict)
            if strict:
                parameters['required'] = [name for name in parameters['properties'] if name != 'agent']
            else:
                parameters['required'] = [name for name, param in sig.parameters.items() if param.default == param.empty and name != 'self' and name != 'agent']
        except Exception as e:
            print(f'Could not parse args for {function_name}: {e}', exc_info=True)
        return cls(name=function_name, description=cls.get_entrypoint_docstring(entrypoint=c), parameters=parameters, entrypoint=c)

    def process_entrypoint(self, strict: bool = False):
        if self.skip_entrypoint_processing:
            return
        if self.entrypoint is None:
            return
        parameters = {'type': 'object', 'properties': {}, 'required': []}
        params_set_by_user = False
        if self.parameters != parameters:
            params_set_by_user = True
        try:
            sig = inspect.signature(self.entrypoint)
            type_hints = get_type_hints(self.entrypoint)
            if 'agent' in sig.parameters:
                del type_hints['agent']
            param_type_hints = {name: type_hints.get(name) for name in sig.parameters if name != 'return' and name != 'agent'}
            param_descriptions = {}
            if docstring := inspect.getdoc(self.entrypoint):
                parsed_doc = docstring_parser.parse(docstring)
                param_docs = parsed_doc.params
                for param in param_docs:
                    param_name = param.arg_name
                    param_type = param.type_name
                    param_descriptions[param_name] = f'({param_type}) {param.description}'
            parameters = self.get_json_schema(type_hints=param_type_hints, param_descriptions=param_descriptions, strict=strict)
            if strict:
                parameters['required'] = [name for name in parameters['properties'] if name != 'agent']
            else:
                parameters['required'] = [
                    name
                    for name, param in sig.parameters.items()
                    if param.default == param.empty and name != 'self' and name != 'agent'
                ]
            if params_set_by_user:
                self.parameters['additionalProperties'] = False
                if strict:
                    self.parameters['required'] = [name for name in self.parameters['properties'] if name != 'agent']
                else:
                    self.parameters['required'] = [name for name, param in sig.parameters.items() if param.default == param.empty and name != 'self' and name != 'agent']
        except Exception as e:
            print(f'Could not parse args for {self.name}: {e}', exc_info=True)
        self.description = self.description or self.get_entrypoint_docstring(self.entrypoint)
        if not params_set_by_user:
            self.parameters = parameters

    def get_definition_for_prompt_dict(self) -> Optional[Dict[str, Any]]:
        if self.entrypoint is None:
            return None
        type_hints = get_type_hints(self.entrypoint)
        return_type = type_hints.get('return', None)
        returns = None
        if return_type:
            name = str(return_type)
            if 'list' in name or 'dict' in name:
                returns = name
            else:
                returns = return_type.__name__
        function_info = {'name': self.name, 'description': self.description, 'arguments': self.parameters.get('properties', {}), 'returns': returns}
        return function_info

    def get_definition_for_prompt(self) -> Optional[str]:
        function_info = self.get_definition_for_prompt_dict()
        if function_info:
            return json.dumps(function_info, indent=2)
        return None

    def _get_cache_key(self, entrypoint_args: Dict[str, Any], call_args: Dict[str, Any] = None) -> str:
        copy_entrypoint_args = entrypoint_args.copy()
        if 'agent' in copy_entrypoint_args:
            del copy_entrypoint_args['agent']
        args_str = str(copy_entrypoint_args)
        kwargs_str = str(sorted((call_args or {}).items()))
        key_str = f'{self.name}:{args_str}:{kwargs_str}'
        return hashlib.md5(key_str.encode()).hexdigest()

    def _get_cache_file_path(self, cache_key: str) -> str:
        base_cache_dir = self.cache_dir
        func_cache_dir = pathlib.Path(base_cache_dir) / 'functions' / self.name
        func_cache_dir.mkdir(parents=True, exist_ok=True)
        return str(func_cache_dir / f'{cache_key}.json')

    def _get_cached_result(self, cache_file: str) -> Optional[Any]:
        cache_path = pathlib.Path(cache_file)
        if not cache_path.exists():
            return None
        try:
            with cache_path.open('r') as f:
                cache_data = json.load(f)
            timestamp = cache_data.get('timestamp', 0)
            result = cache_data.get('result')
            if time.time() - timestamp <= self.cache_ttl:
                return result
            cache_path.unlink()
        except Exception as e:
            print(f'Error reading cache: {e}')
        return None

    def _save_to_cache(self, cache_file: str, result: Any):
        try:
            with open(cache_file, 'w') as f:
                json.dump({'timestamp': time.time(), 'result': result}, f)
        except Exception as e:
            print(f'Error writing cache: {e}')


class FunctionCall(Base):
    function: Function = None
    arguments: Dict[str, Any]
    result: Any = None
    call_id = ''
    error = ''

    def get_call_str(self) -> str:
        term_width = shutil.get_terminal_size().columns or 80
        max_arg_len = max(20, (term_width - len(self.function.name) - 4) // 2)
        if self.arguments is None:
            return f'{self.function.name}()'
        trimmed_arguments = {}
        for k, v in self.arguments.items():
            if isinstance(v, str) and len(str(v)) > max_arg_len:
                trimmed_arguments[k] = '...'
            else:
                trimmed_arguments[k] = v
        call_str = f'{self.function.name}({", ".join([f"{k}={v}" for k, v in trimmed_arguments.items()])})'
        if len(call_str) > term_width:
            return f'{self.function.name}(...)'
        return call_str

    def execute(self) -> bool:
        if self.function.entrypoint is None:
            return False
        print(f'Running: {self.get_call_str()}')
        function_call_success = False
        if self.function.pre_hook:
            try:
                pre_hook_args = {}
                if 'agent' in inspect.signature(self.function.pre_hook).parameters:
                    pre_hook_args['agent'] = self.function._agent
                if 'fc' in inspect.signature(self.function.pre_hook).parameters:
                    pre_hook_args['fc'] = self
                self.function.pre_hook(**pre_hook_args)
            except Exception as e:
                print(f'Error in pre-hook callback: {e}')
                print(e)
        entrypoint_args = {}
        if 'agent' in inspect.signature(self.function.entrypoint).parameters:
            entrypoint_args['agent'] = self.function._agent
        if 'fc' in inspect.signature(self.function.entrypoint).parameters:
            entrypoint_args['fc'] = self
        if self.function.cache_results and not inspect.isgenerator(self.function.entrypoint):
            cache_key = self.function._get_cache_key(entrypoint_args, self.arguments)
            cache_file = self.function._get_cache_file_path(cache_key)
            cached_result = self.function._get_cached_result(cache_file)
            if cached_result:
                print(f'Cache hit for: {self.get_call_str()}')
                self.result = cached_result
                function_call_success = True
                return function_call_success
        try:
            if self.arguments == {} or self.arguments is None:
                result = self.function.entrypoint(**entrypoint_args)
            else:
                result = self.function.entrypoint(**entrypoint_args, **self.arguments)
            if inspect.isgenerator(result):
                self.result = result
            else:
                self.result = result
                if self.function.cache_results:
                    cache_key = self.function._get_cache_key(entrypoint_args, self.arguments)
                    cache_file = self.function._get_cache_file_path(cache_key)
                    self.function._save_to_cache(cache_file, self.result)
            function_call_success = True
        except Exception as e:
            print(f'Could not run function {self.get_call_str()}')
            print(e)
            self.error = str(e)
            return function_call_success
        if self.function.post_hook:
            try:
                post_hook_args = {}
                if 'agent' in inspect.signature(self.function.post_hook).parameters:
                    post_hook_args['agent'] = self.function._agent
                if 'fc' in inspect.signature(self.function.post_hook).parameters:
                    post_hook_args['fc'] = self
                self.function.post_hook(**post_hook_args)
            except Exception as e:
                print(f'Error in post-hook callback: {e}')
                print(e)
        return function_call_success

    async def aexecute(self) -> bool:
        if self.function.entrypoint is None:
            return False
        print(f'Running: {self.get_call_str()}')
        function_call_success = False
        if inspect.iscoroutinefunction(self.function.pre_hook):
            if self.function.pre_hook:
                try:
                    pre_hook_args = {}
                    if 'agent' in inspect.signature(self.function.pre_hook).parameters:
                        pre_hook_args['agent'] = self.function._agent
                    if 'fc' in inspect.signature(self.function.pre_hook).parameters:
                        pre_hook_args['fc'] = self
                    await self.function.pre_hook(**pre_hook_args)
                except Exception as e:
                    print(f'Error in pre-hook callback: {e}')
                    print(e)
        else:
            if self.function.pre_hook:
                try:
                    pre_hook_args = {}
                    if 'agent' in inspect.signature(self.function.pre_hook).parameters:
                        pre_hook_args['agent'] = self.function._agent
                    if 'fc' in inspect.signature(self.function.pre_hook).parameters:
                        pre_hook_args['fc'] = self
                    self.function.pre_hook(**pre_hook_args)
                except Exception as e:
                    print(f'Error in pre-hook callback: {e}')
                    print(e)
        entrypoint_args = {}
        if 'agent' in inspect.signature(self.function.entrypoint).parameters:
            entrypoint_args['agent'] = self.function._agent
        if 'fc' in inspect.signature(self.function.entrypoint).parameters:
            entrypoint_args['fc'] = self
        if self.function.cache_results and not (inspect.isasyncgen(self.function.entrypoint) or inspect.isgenerator(self.function.entrypoint)):
            cache_key = self.function._get_cache_key(entrypoint_args, self.arguments)
            cache_file = self.function._get_cache_file_path(cache_key)
            cached_result = self.function._get_cached_result(cache_file)
            if cached_result:
                print(f'Cache hit for: {self.get_call_str()}')
                self.result = cached_result
                function_call_success = True
                return function_call_success
        try:
            if self.arguments == {} or self.arguments is None:
                result = self.function.entrypoint(**entrypoint_args)
                if inspect.isasyncgen(self.function.entrypoint) or inspect.isasyncgenfunction(self.function.entrypoint):
                    self.result = result
                else:
                    self.result = await result
            else:
                result = self.function.entrypoint(**entrypoint_args, **self.arguments)
                if inspect.isasyncgen(self.function.entrypoint) or inspect.isasyncgenfunction(self.function.entrypoint):
                    self.result = result
                else:
                    self.result = await result
            if self.function.cache_results and not (inspect.isgenerator(self.result) or inspect.isasyncgen(self.result)):
                cache_key = self.function._get_cache_key(entrypoint_args, self.arguments)
                cache_file = self.function._get_cache_file_path(cache_key)
                self.function._save_to_cache(cache_file, self.result)
            function_call_success = True
        except Exception as e:
            print(f'Could not run function {self.get_call_str()}')
            print(e)
            self.error = str(e)
            return function_call_success
        if inspect.iscoroutinefunction(self.function.post_hook):
            if self.function.post_hook:
                try:
                    post_hook_args = {}
                    if 'agent' in inspect.signature(self.function.post_hook).parameters:
                        post_hook_args['agent'] = self.function._agent
                    if 'fc' in inspect.signature(self.function.post_hook).parameters:
                        post_hook_args['fc'] = self
                    await self.function.post_hook(**post_hook_args)
                except Exception as e:
                    print(f'Error in post-hook callback: {e}')
                    print(e)
        else:
            if self.function.post_hook:
                try:
                    post_hook_args = {}
                    if 'agent' in inspect.signature(self.function.post_hook).parameters:
                        post_hook_args['agent'] = self.function._agent
                    if 'fc' in inspect.signature(self.function.post_hook).parameters:
                        post_hook_args['fc'] = self
                    self.function.post_hook(**post_hook_args)
                except Exception as e:
                    print(f'Error in post-hook callback: {e}')
                    print(e)
        return function_call_success


class Toolkit:
    def __init__(self, name: str = 'toolkit', cache_results: bool = False, cache_ttl: int = 3600, cache_dir: str = ''):
        self.name: str = name
        self.functions: Dict[str, Function] = collections.OrderedDict()
        self.cache_results: bool = cache_results
        self.cache_ttl: int = cache_ttl
        self.cache_dir: str = cache_dir

    def register(self, function: Callable[..., Any], sanitize_arguments: bool = True):
        try:
            f = Function(name=function.__name__, entrypoint=function, sanitize_arguments=sanitize_arguments, cache_results=self.cache_results, cache_dir=self.cache_dir, cache_ttl=self.cache_ttl)
            self.functions[f.name] = f
            print(f'Function: {f.name} registered with {self.name}')
        except Exception as e:
            print(f'Failed to create Function for: {function.__name__}')
            raise e

    def instructions(self) -> str:
        return ''

    def __repr__(self):
        return f'<{self.__class__.__name__} name={self.name} functions={list(self.functions.keys())}>'

    def __str__(self):
        return self.__repr__()


class Media(Base):
    id: str = ''
    original_prompt: str = ''
    revised_prompt: str = ''


class VideoArtifact(Media):
    url: str = ''
    eta: str = ''
    length: str = ''


class ImageArtifact(Media):
    url: str = ''
    content: bytes = None
    mime_type: str = ''
    alt_text: str = ''


class AudioArtifact(Media):
    url: str = ''
    base64_audio: str = ''
    length: str = ''
    mime_type: str = ''


class Video(Base):
    filepath: Union[pathlib.Path, str] = ''
    content: Any = None
    format = 'mp4'

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        response_dict = {'content': base64.b64encode(zlib.compress(self.content) if isinstance(self.content, bytes) else self.content.encode('utf-8')).decode('utf-8') if self.content else None, 'filepath': self.filepath, 'format': self.format}
        return {k: v for k, v in response_dict.items() if v}

    @classmethod
    def from_artifact(cls, artifact: VideoArtifact) -> 'Video':
        return cls(url=artifact.url)


class Audio(Base):
    filepath: Union[pathlib.Path, str] = ''
    content: Any = None
    url = ''
    format = 'mp3'

    @property
    def audio_url_content(self) -> Optional[bytes]:
        if self.url:
            return requests.get(self.url).content
        else:
            return None

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        content = base64.b64encode(zlib.compress(self.content) if isinstance(self.content, bytes) else self.content.encode('utf-8')).decode('utf-8') if self.content else None
        response_dict = {'content': content, 'filepath': self.filepath, 'format': self.format}
        return {k: v for k, v in response_dict.items() if v}

    @classmethod
    def from_artifact(cls, artifact: AudioArtifact) -> 'Audio':
        return cls(url=artifact.url, content=artifact.base64_audio, format=artifact.mime_type)


class AudioResponse(Base):
    id: str = ''
    content: str = ''
    expires_at: int = 0
    transcript: str = ''
    mime_type: str = ''
    sample_rate = 24000
    channels: int = 1

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        content = base64.b64encode(self.content).decode('utf-8') if isinstance(self.content, bytes) else self.content
        response_dict = {'id': self.id, 'content': content, 'expires_at': self.expires_at, 'transcript': self.transcript, 'mime_type': self.mime_type, 'sample_rate': self.sample_rate, 'channels': self.channels}
        return {k: v for k, v in response_dict.items() if v}


class Image(Base):
    url: str = ''
    filepath: Union[pathlib.Path, str] = ''
    content: Any = None
    format: str = 'jpeg'
    detail: str = ''
    id: str = ''

    @property
    def image_url_content(self) -> Optional[bytes]:
        if self.url:
            return requests.get(self.url).content
        else:
            return None

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        response_dict = {'content': base64.b64encode(zlib.compress(self.content) if isinstance(self.content, bytes) else self.content.encode('utf-8')).decode('utf-8')
            if self.content
            else None, 'filepath': self.filepath, 'url': self.url, 'detail': self.detail}
        return {k: v for k, v in response_dict.items() if v}

    @classmethod
    def from_artifact(cls, artifact: ImageArtifact) -> 'Image':
        return cls(url=artifact.url)


class File(Base):
    url: str = ''
    filepath: Union[pathlib.Path, str] = ''
    content: Any = None
    mime_type: str = ''

    @classmethod
    def valid_mime_types(cls) -> List[str]:
        return ['application/pdf', 'application/x-javascript', 'text/javascript', 'application/x-python', 'text/x-python', 'text/plain', 'text/html', 'text/css', 'text/md', 'text/csv', 'text/xml', 'text/rtf']

    @property
    def file_url_content(self) -> Optional[Tuple[bytes, str]]:
        if self.url:
            response = requests.get(self.url)
            content = response.content
            mime_type = response.headers.get('Content-Type', '').split(';')[0]
            return content, mime_type
        else:
            return None


class MessageReferences(Base):
    query = ''
    time = 0.0
    references: List[Dict[str, Any]]


class DocumentCitation(Base):
    document_title: str = ''
    cited_text: str = ''
    file_name: str = ''


class Citations(Base):
    raw: Any = None
    urls: List[Dict]
    documents: List[DocumentCitation]


class MessageMetrics(Base):
    input_tokens: int = 0
    output_tokens: int = 0
    total_tokens: int = 0
    prompt_tokens: int = 0
    completion_tokens: int = 0
    prompt_tokens_details: dict
    completion_tokens_details: dict
    additional_metrics: dict
    time = 0.0
    time_to_first_token = 0.0
    timer: Timer = None

    def start_timer(self):
        if self.timer is None:
            self.timer = Timer()
        self.timer.start()

    def stop_timer(self, set_time: bool = True):
        if self.timer:
            self.timer.stop()
            if set_time:
                self.time = self.timer.elapsed

    def set_time_to_first_token(self):
        if self.timer:
            self.time_to_first_token = self.timer.elapsed

    def __add__(self, other: 'MessageMetrics') -> 'MessageMetrics':
        result = MessageMetrics(input_tokens=self.input_tokens + other.input_tokens, output_tokens=self.output_tokens + other.output_tokens, total_tokens=self.total_tokens + other.total_tokens, prompt_tokens=self.prompt_tokens + other.prompt_tokens, completion_tokens=self.completion_tokens + other.completion_tokens)
        if self.prompt_tokens_details or other.prompt_tokens_details:
            result.prompt_tokens_details = {}
            if self.prompt_tokens_details:
                result.prompt_tokens_details.update(self.prompt_tokens_details)
            if other.prompt_tokens_details:
                for key, value in other.prompt_tokens_details.items():
                    result.prompt_tokens_details[key] = result.prompt_tokens_details.get(key, 0) + value
        if self.completion_tokens_details or other.completion_tokens_details:
            result.completion_tokens_details = {}
            if self.completion_tokens_details:
                result.completion_tokens_details.update(self.completion_tokens_details)
            if other.completion_tokens_details:
                for key, value in other.completion_tokens_details.items():
                    result.completion_tokens_details[key] = result.completion_tokens_details.get(key, 0) + value
        if self.additional_metrics or other.additional_metrics:
            result.additional_metrics = {}
            if self.additional_metrics:
                result.additional_metrics.update(self.additional_metrics)
            if other.additional_metrics:
                result.additional_metrics.update(other.additional_metrics)
        if self.time and other.time:
            result.time = self.time + other.time
        elif self.time:
            result.time = self.time
        elif other.time:
            result.time = other.time
        result.time_to_first_token = self.time_to_first_token or other.time_to_first_token
        return result

    def __radd__(self, other: 'MessageMetrics') -> 'MessageMetrics':
        if other == 0:
            return self
        return self + other


class Message(Base):
    role: Literal['system', 'user', 'assistant', 'tool'] = 'user'
    content: Union[List[Any], str] = None
    name: str = ''
    tool_call_id: str = ''
    tool_calls: List[Dict[str, Any]]
    audio: Sequence[Audio] = None
    images: Sequence[Image] = None
    videos: Sequence[Video] = None
    files: Sequence[File] = None
    audio_output: AudioResponse = None
    image_output: ImageArtifact = None
    thinking: str = ''
    redacted_thinking: str = ''
    provider_data: Dict[str, Any]
    citations: Citations = None
    reasoning_content: str = ''
    tool_name: str = ''
    tool_args: Any = None
    tool_call_error: bool = False
    stop_after_tool_call: bool = False
    add_to_agent_memory: bool = True
    from_history: bool = False
    metrics: MessageMetrics
    references: MessageReferences = None
    created_at: int = int(time.time())

    def get_content_string(self) -> str:
        if isinstance(self.content, str):
            return self.content
        if isinstance(self.content, list):
            if len(self.content) > 0 and isinstance(self.content[0], dict) and 'text' in self.content[0]:
                return self.content[0].get('text', '')
            else:
                return json.dumps(self.content)
        return ''

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        message_dict = {'content': self.content, 'reasoning_content': self.reasoning_content, 'from_history': self.from_history, 'stop_after_tool_call': self.stop_after_tool_call, 'role': self.role, 'name': self.name, 'tool_call_id': self.tool_call_id, 'tool_name': self.tool_name, 'tool_args': self.tool_args, 'tool_call_error': self.tool_call_error, 'tool_calls': self.tool_calls, 'thinking': self.thinking, 'redacted_thinking': self.redacted_thinking}
        message_dict = {k: v for k, v in message_dict.items() if v and not (isinstance(v, (list, dict)) and len(v) == 0)}
        if self.images:
            message_dict['images'] = [img.to_dict() for img in self.images]
        if self.audio:
            message_dict['audio'] = [aud.to_dict() for aud in self.audio]
        if self.videos:
            message_dict['videos'] = [vid.to_dict() for vid in self.videos]
        if self.audio_output:
            message_dict['audio_output'] = self.audio_output.to_dict()
        if self.references:
            message_dict['references'] = self.references.to_dict()
        if self.metrics:
            mdict = self.metrics.__dict__
            mdict.pop('timer')
            message_dict['metrics'] = {k: v for k, v in mdict.items() if v and (not isinstance(v, (int, float)) or v != 0) and (not isinstance(v, dict) or len(v) > 0)}
            if not message_dict['metrics']:
                message_dict.pop('metrics')
        message_dict['created_at'] = self.created_at
        return message_dict

    def to_function_call_dict(self) -> Dict[str, Any]:
        return {'content': self.content, 'tool_call_id': self.tool_call_id, 'tool_name': self.tool_name, 'tool_args': self.tool_args, 'tool_call_error': self.tool_call_error, 'metrics': self.metrics, 'created_at': self.created_at}

    def log(self, metrics: bool = True, level=None):
        try:
            terminal_width = shutil.get_terminal_size().columns
        except Exception:
            terminal_width = 80
        header = f' {self.role} '
        print(f'{header.center(terminal_width - 20, "=")}')
        if self.name:
            print(f'名字: {self.name}')
        if self.tool_call_id:
            print(f'Tool call Id: {self.tool_call_id}')
        if self.thinking:
            print(f'<thinking>\n{self.thinking}\n</thinking>')
        if self.content:
            if isinstance(self.content, str) or isinstance(self.content, list):
                print(self.content)
            elif isinstance(self.content, dict):
                print(json.dumps(self.content, indent=2))
        if self.tool_calls:
            tool_calls_list = ['Tool Calls:']
            for tool_call in self.tool_calls:
                tool_id = tool_call.get('id')
                function_name = tool_call.get('function', {}).get('name')
                tool_calls_list.append(f'  - ID: "{tool_id}"') if tool_id else None
                tool_calls_list.append(f'    Name: "{function_name}"') if function_name else None
                tool_call_arguments = tool_call.get('function', {}).get('arguments')
                if tool_call_arguments:
                    try:
                        arguments = ', '.join(f'{k}: {v}' for k, v in json.loads(tool_call_arguments).items())
                        tool_calls_list.append(f'    Arguments: "{arguments}"')
                    except json.JSONDecodeError:
                        tool_calls_list.append('    Arguments: "Invalid JSON format"')
            tool_calls_str = '\n'.join(tool_calls_list)
            print(tool_calls_str)
        if self.images:
            print(f'图片: {len(self.images)}')
        if self.videos:
            print(f'视频: {len(self.videos)}')
        if self.audio:
            print(f'声音: {len(self.audio)}')
        if self.files:
            print(f'文件: {len(self.files)}')
        metrics_header = ' TOOL METRICS ' if self.role == 'tool' else ' METRICS '
        if metrics and self.metrics and self.metrics != MessageMetrics():
            print(metrics_header)
            token_metrics = []
            if self.metrics.input_tokens:
                token_metrics.append(f'input={self.metrics.input_tokens}')
            if self.metrics.output_tokens:
                token_metrics.append(f'output={self.metrics.output_tokens}')
            if self.metrics.total_tokens:
                token_metrics.append(f'total={self.metrics.total_tokens}')
            if token_metrics:
                print(f'* Tokens:                      {", ".join(token_metrics)}')
            if self.metrics.prompt_tokens_details:
                print(f'* Prompt tokens details:       {self.metrics.prompt_tokens_details}')
            if self.metrics.completion_tokens_details:
                print(f'* Completion tokens details:   {self.metrics.completion_tokens_details}')
            if self.metrics.time:
                print(f'* Time:                        {self.metrics.time:.4f}s')
            if self.metrics.output_tokens and self.metrics.time:
                print(f'* Tokens per second:           {self.metrics.output_tokens / self.metrics.time:.4f} tokens/s')
            if self.metrics.time_to_first_token:
                print(f'* Time to first token:         {self.metrics.time_to_first_token:.4f}s')
            if self.metrics.additional_metrics:
                print(f'* Additional metrics:          {self.metrics.additional_metrics}')
            print(metrics_header)

    def content_is_valid(self) -> bool:
        return self.content and len(self.content) > 0


class ModelResponse(Base):
    role: str = ''
    content: str = ''
    parsed: Any = None
    audio: AudioResponse = None
    image: ImageArtifact = None
    tool_calls: List[Dict[str, Any]]
    event: str = ResponseEvent.res.value
    provider_data: Dict[str, Any]
    thinking: str = ''
    redacted_thinking: str = ''
    reasoning_content: str = ''
    citations: Citations = None
    response_usage: Any = None
    created_at: int = int(time.time())
    extra: Dict[str, Any]


def get_function_call(name: str, arguments: str = None, call_id: str = None, functions: Dict[str, Function] = None) -> Optional[FunctionCall]:
    print(f'Getting function {name}')
    if functions is None:
        return None
    function_to_call: Function = None
    if name in functions:
        function_to_call = functions[name]
    if function_to_call is None:
        print(f'Function {name} not found')
        return None
    function_call = FunctionCall(function=function_to_call)
    if call_id:
        function_call.call_id = call_id
    if arguments and arguments != '':
        try:
            if function_to_call.sanitize_arguments:
                if 'None' in arguments:
                    arguments = arguments.replace('None', 'null')
                if 'True' in arguments:
                    arguments = arguments.replace('True', 'true')
                if 'False' in arguments:
                    arguments = arguments.replace('False', 'false')
            _arguments = json.loads(arguments)
        except Exception as e:
            print(f'无法解码函数参数:\n{arguments}\nError: {e}')
            function_call.error = f'解码函数参数时出错:{e}\n\n请确保我们可以json.loads()参数并重试。'
            return function_call
        if not isinstance(_arguments, dict):
            print(f'函数参数不是有效的JSON对象: {arguments}')
            function_call.error = '函数参数不是有效的JSON对象。\n\n请修复并重试。'
            return function_call
        try:
            clean_arguments: Dict[str, Any] = {}
            for k, v in _arguments.items():
                if isinstance(v, str):
                    _v = v.strip().lower()
                    if _v in ('none', 'null'):
                        clean_arguments[k] = None
                    elif _v == 'true':
                        clean_arguments[k] = True
                    elif _v == 'false':
                        clean_arguments[k] = False
                    else:
                        clean_arguments[k] = v.strip()
                else:
                    clean_arguments[k] = v
            function_call.arguments = clean_arguments
        except Exception as e:
            print(f'无法解析函数参数:\n{arguments}\nError: {e}')
            function_call.error = f'解析函数参数时出错:{e}\n\n请修复并重试。'
            return function_call
    return function_call


def get_function_call_for_tool_call(tool_call: Dict[str, Any], functions: Dict[str, Function] = None) -> Optional[FunctionCall]:
    if tool_call.get('type') == 'function':
        _tool_call_id = tool_call.get('id')
        _tool_call_function = tool_call.get('function')
        if _tool_call_function:
            _tool_call_function_name = _tool_call_function.get('name')
            _tool_call_function_arguments_str = _tool_call_function.get('arguments')
            if _tool_call_function_name:
                return get_function_call(name=_tool_call_function_name, arguments=_tool_call_function_arguments_str, call_id=_tool_call_id, functions=functions)
    return None


class MessageData(Base):
    response_role: Literal['system', 'user', 'assistant', 'tool'] = 'user'
    response_content: Any = ''
    response_thinking: Any = ''
    response_redacted_thinking: Any = ''
    response_citations: Citations = None
    response_tool_calls: List[Dict[str, Any]]
    response_audio: AudioResponse = None
    response_image: ImageArtifact = None
    response_provider_data: Dict[str, Any]
    extra: Dict[str, Any]


class Model(Base):
    id: str = ''
    name: str = ''
    provider: str = ''
    response_format: Any = None  # 不要直接设置，而是在代理上设置response_model属性
    structured_outputs: bool = False
    supports_native_structured_outputs: bool = False
    supports_json_schema_outputs: bool = False
    tool_choice: Union[str, Dict[str, Any]] = None
    show_tool_calls: bool = False
    tool_call_limit: int = 0
    _tools: List[Dict]
    _functions: Dict[str, Function]
    _function_call_stack: List[FunctionCall]
    system_prompt: str = ''
    instructions: List[str]
    tool_message_role: str = 'tool'
    assistant_message_role: str = 'assistant'

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        _dict = super().to_dict(include=include or {'name', 'id', 'provider'})
        if self._functions:
            _dict['functions'] = {k: v.to_dict() for k, v in self._functions.items()}
            _dict['tool_call_limit'] = self.tool_call_limit
        return _dict

    def get_provider(self) -> str:
        return self.provider or self.name or self.__class__.__name__

    def invoke(self, *args, **kwargs) -> Any:
        pass

    async def ainvoke(self, *args, **kwargs) -> Any:
        pass

    def invoke_stream(self, *args, **kwargs) -> Iterator[Any]:
        pass

    async def ainvoke_stream(self, *args, **kwargs) -> AsyncGenerator[Any, None]:
        pass

    def parse_provider_response(self, response: Any) -> ModelResponse:
        pass

    def parse_provider_response_delta(self, response: Any) -> ModelResponse:
        pass

    def set_tools(self, tools: List[Dict]) -> None:
        self._tools = tools

    def set_functions(self, functions: Dict[str, Function]) -> None:
        if len(functions) > 0:
            self._functions = functions

    def response(self, messages: List[Message]) -> ModelResponse:
        print(f'{self.get_provider()} Response Start\nModel: {self.id}')
        self._log_messages(messages)
        model_response = ModelResponse()
        while True:
            assistant_message, has_tool_calls = self._process_model_response(messages=messages, model_response=model_response)
            if has_tool_calls:
                function_calls_to_run = self._prepare_function_calls(assistant_message=assistant_message, messages=messages, model_response=model_response)
                function_call_results: List[Message] = []
                for function_call_response in self.run_function_calls(function_calls=function_calls_to_run, function_call_results=function_call_results):
                    if function_call_response.event == ResponseEvent.end.value and function_call_response.tool_calls:
                        model_response.tool_calls.extend(function_call_response.tool_calls)
                    elif function_call_response.event not in [ResponseEvent.start.value, ResponseEvent.end.value]:
                        if function_call_response.content:
                            model_response.content += function_call_response.content
                self.format_function_call_results(messages=messages, function_call_results=function_call_results, **model_response.extra or {})
                for function_call_result in function_call_results:
                    function_call_result.log(metrics=True)
                if any(m.stop_after_tool_call for m in function_call_results):
                    break
                continue
            break
        print(f'{self.get_provider()} Response End')
        return model_response

    async def aresponse(self, messages: List[Message]) -> ModelResponse:
        print(f'{self.get_provider()} Async Response Start\nModel: {self.id}')
        self._log_messages(messages)
        model_response = ModelResponse()
        while True:
            assistant_message, has_tool_calls = await self._aprocess_model_response(messages=messages, model_response=model_response)
            if has_tool_calls:
                function_calls_to_run = self._prepare_function_calls(assistant_message=assistant_message, messages=messages, model_response=model_response)
                function_call_results: List[Message] = []
                async for function_call_response in self.arun_function_calls(function_calls=function_calls_to_run, function_call_results=function_call_results):
                    if function_call_response.event == ResponseEvent.end.value and function_call_response.tool_calls:
                        model_response.tool_calls.extend(function_call_response.tool_calls)
                    elif function_call_response.event not in [ResponseEvent.start.value, ResponseEvent.end.value]:
                        if function_call_response.content:
                            model_response.content += function_call_response.content
                self.format_function_call_results(messages=messages, function_call_results=function_call_results, **model_response.extra or {})
                for function_call_result in function_call_results:
                    function_call_result.log(metrics=True)
                if any(m.stop_after_tool_call for m in function_call_results):
                    break
                continue
            break
        print(f'{self.get_provider()} Async Response End')
        return model_response

    def _process_model_response(self, messages: List[Message], model_response: ModelResponse) -> Tuple[Message, bool]:
        assistant_message = Message(role=self.assistant_message_role)
        assistant_message.metrics.start_timer()
        response = self.invoke(messages=messages)
        assistant_message.metrics.stop_timer()
        provider_response: ModelResponse = self.parse_provider_response(response)
        if provider_response.parsed:
            model_response.parsed = provider_response.parsed
        self._populate_assistant_message(assistant_message=assistant_message, provider_response=provider_response)
        messages.append(assistant_message)
        assistant_message.log(metrics=True)
        if assistant_message.content:
            if model_response.content is None:
                model_response.content = assistant_message.get_content_string()
            else:
                model_response.content += assistant_message.get_content_string()
        if assistant_message.thinking:
            model_response.thinking = assistant_message.thinking
        if assistant_message.redacted_thinking:
            model_response.redacted_thinking = assistant_message.redacted_thinking
        if assistant_message.citations:
            model_response.citations = assistant_message.citations
        if assistant_message.audio_output:
            model_response.audio = assistant_message.audio_output
        if assistant_message.image_output:
            model_response.image = assistant_message.image_output
        if provider_response.extra:
            if model_response.extra is None:
                model_response.extra = {}
            model_response.extra.update(provider_response.extra)
        return assistant_message, bool(assistant_message.tool_calls)

    async def _aprocess_model_response(self, messages: List[Message], model_response: ModelResponse) -> Tuple[Message, bool]:
        assistant_message = Message(role=self.assistant_message_role)
        assistant_message.metrics.start_timer()
        response = await self.ainvoke(messages=messages)
        assistant_message.metrics.stop_timer()
        provider_response: ModelResponse = self.parse_provider_response(response)
        if provider_response.parsed:
            model_response.parsed = provider_response.parsed
        self._populate_assistant_message(assistant_message=assistant_message, provider_response=provider_response)
        messages.append(assistant_message)
        assistant_message.log(metrics=True)
        if assistant_message.content:
            if model_response.content is None:
                model_response.content = assistant_message.get_content_string()
            else:
                model_response.content += assistant_message.get_content_string()
        if assistant_message.thinking:
            model_response.thinking = assistant_message.thinking
        if assistant_message.redacted_thinking:
            model_response.redacted_thinking = assistant_message.redacted_thinking
        if assistant_message.citations:
            model_response.citations = assistant_message.citations
        if assistant_message.audio_output:
            model_response.audio = assistant_message.audio_output
        if assistant_message.image_output:
            model_response.image = assistant_message.image_output
        if provider_response.extra:
            if model_response.extra is None:
                model_response.extra = {}
            model_response.extra.update(provider_response.extra)
        return assistant_message, bool(assistant_message.tool_calls)

    def _populate_assistant_message(self, assistant_message: Message, provider_response: ModelResponse) -> Message:
        if provider_response.role:
            assistant_message.role = provider_response.role
        if provider_response.content:
            assistant_message.content = provider_response.content
        if provider_response.tool_calls and len(provider_response.tool_calls) > 0:
            assistant_message.tool_calls = provider_response.tool_calls
        if provider_response.audio:
            assistant_message.audio_output = provider_response.audio
        if provider_response.image:
            assistant_message.image_output = provider_response.image
        if provider_response.thinking:
            assistant_message.thinking = provider_response.thinking
        if provider_response.redacted_thinking:
            assistant_message.redacted_thinking = provider_response.redacted_thinking
        if provider_response.reasoning_content:
            assistant_message.reasoning_content = provider_response.reasoning_content
        if provider_response.provider_data:
            assistant_message.provider_data = provider_response.provider_data
        if provider_response.citations:
            assistant_message.citations = provider_response.citations
        if provider_response.response_usage:
            self._add_usage_metrics_to_assistant_message(assistant_message=assistant_message, response_usage=provider_response.response_usage)
        return assistant_message

    def process_response_stream(self, messages: List[Message], assistant_message: Message, stream_data: MessageData) -> Iterator[ModelResponse]:
        for response_delta in self.invoke_stream(messages=messages):
            model_response_delta = self.parse_provider_response_delta(response_delta)
            yield from self._populate_stream_data_and_assistant_message(stream_data=stream_data, assistant_message=assistant_message, model_response=model_response_delta)

    def response_stream(self, messages: List[Message]) -> Iterator[ModelResponse]:
        print(f'{self.get_provider()} Response Stream Start\nModel: {self.id}')
        self._log_messages(messages)
        while True:
            assistant_message = Message(role=self.assistant_message_role)
            stream_data = MessageData()
            assistant_message.metrics.start_timer()
            yield from self.process_response_stream(messages=messages, assistant_message=assistant_message, stream_data=stream_data)
            assistant_message.metrics.stop_timer()
            if stream_data.response_content:
                assistant_message.content = stream_data.response_content
            if stream_data.response_thinking:
                assistant_message.thinking = stream_data.response_thinking
            if stream_data.response_redacted_thinking:
                assistant_message.redacted_thinking = stream_data.response_redacted_thinking
            if stream_data.response_provider_data:
                assistant_message.provider_data = stream_data.response_provider_data
            if stream_data.response_citations:
                assistant_message.citations = stream_data.response_citations
            if stream_data.response_audio:
                assistant_message.audio_output = stream_data.response_audio
            if stream_data.response_tool_calls and len(stream_data.response_tool_calls) > 0:
                assistant_message.tool_calls = stream_data.response_tool_calls
            messages.append(assistant_message)
            assistant_message.log(metrics=True)
            if assistant_message.tool_calls:
                function_calls_to_run: List[FunctionCall] = self.get_function_calls_to_run(assistant_message, messages)
                function_call_results: List[Message] = []
                for function_call_response in self.run_function_calls(function_calls=function_calls_to_run, function_call_results=function_call_results):
                    yield function_call_response
                if stream_data.extra:
                    self.format_function_call_results(messages=messages, function_call_results=function_call_results, **stream_data.extra)
                else:
                    self.format_function_call_results(messages=messages, function_call_results=function_call_results)
                for function_call_result in function_call_results:
                    function_call_result.log(metrics=True)
                if any(m.stop_after_tool_call for m in function_call_results):
                    break
                continue
            break
        print(f'{self.get_provider()} Response Stream End')

    async def aprocess_response_stream(self, messages: List[Message], assistant_message: Message, stream_data: MessageData) -> AsyncIterator[ModelResponse]:
        async for response_delta in self.ainvoke_stream(messages=messages):
            model_response_delta = self.parse_provider_response_delta(response_delta)
            for model_response in self._populate_stream_data_and_assistant_message(stream_data=stream_data, assistant_message=assistant_message, model_response=model_response_delta):
                yield model_response

    async def aresponse_stream(self, messages: List[Message]) -> AsyncIterator[ModelResponse]:
        print(f'{self.get_provider()} Async Response Stream Start\nModel: {self.id}')
        self._log_messages(messages)
        while True:
            assistant_message = Message(role=self.assistant_message_role)
            stream_data = MessageData()
            assistant_message.metrics.start_timer()
            async for response in self.aprocess_response_stream(messages=messages, assistant_message=assistant_message, stream_data=stream_data):
                yield response
            assistant_message.metrics.stop_timer()
            if stream_data.response_content:
                assistant_message.content = stream_data.response_content
            if stream_data.response_thinking:
                assistant_message.thinking = stream_data.response_thinking
            if stream_data.response_redacted_thinking:
                assistant_message.redacted_thinking = stream_data.response_redacted_thinking
            if stream_data.response_provider_data:
                assistant_message.provider_data = stream_data.response_provider_data
            if stream_data.response_audio:
                assistant_message.audio_output = stream_data.response_audio
            if stream_data.response_tool_calls and len(stream_data.response_tool_calls) > 0:
                assistant_message.tool_calls = stream_data.response_tool_calls
            messages.append(assistant_message)
            assistant_message.log(metrics=True)
            if assistant_message.tool_calls:
                function_calls_to_run: List[FunctionCall] = self.get_function_calls_to_run(assistant_message, messages)
                function_call_results: List[Message] = []
                async for function_call_response in self.arun_function_calls(function_calls=function_calls_to_run, function_call_results=function_call_results):
                    yield function_call_response
                if stream_data.extra:
                    self.format_function_call_results(messages=messages, function_call_results=function_call_results, **stream_data.extra)
                else:
                    self.format_function_call_results(messages=messages, function_call_results=function_call_results)
                for function_call_result in function_call_results:
                    function_call_result.log(metrics=True)
                if any(m.stop_after_tool_call for m in function_call_results):
                    break
                continue
            break
        print(f'{self.get_provider()} Async Response Stream End')

    def _populate_stream_data_and_assistant_message(self, stream_data: MessageData, assistant_message: Message, model_response: ModelResponse) -> Iterator[ModelResponse]:
        if not assistant_message.metrics.time_to_first_token:
            assistant_message.metrics.set_time_to_first_token()
        if model_response.role:
            assistant_message.role = model_response.role
        should_yield = False
        if model_response.content:
            stream_data.response_content += model_response.content
            should_yield = True
        if model_response.thinking:
            stream_data.response_thinking += model_response.thinking
            should_yield = True
        if model_response.redacted_thinking:
            stream_data.response_redacted_thinking += model_response.redacted_thinking
            should_yield = True
        if model_response.citations:
            stream_data.response_citations = model_response.citations
            should_yield = True
        if model_response.provider_data:
            if stream_data.response_provider_data is None:
                stream_data.response_provider_data = {}
            stream_data.response_provider_data.update(model_response.provider_data)
        if model_response.tool_calls:
            if stream_data.response_tool_calls is None:
                stream_data.response_tool_calls = []
            stream_data.response_tool_calls.extend(model_response.tool_calls)
            should_yield = True
        if model_response.audio:
            if stream_data.response_audio is None:
                stream_data.response_audio = AudioResponse(id=str(uuid.uuid4()), content='', transcript='')
            if model_response.audio.id:
                stream_data.response_audio.id = model_response.audio.id
            if model_response.audio.content:
                stream_data.response_audio.content += model_response.audio.content
            if model_response.audio.transcript:
                stream_data.response_audio.transcript += model_response.audio.transcript
            if model_response.audio.expires_at:
                stream_data.response_audio.expires_at = model_response.audio.expires_at
            if model_response.audio.mime_type:
                stream_data.response_audio.mime_type = model_response.audio.mime_type
            stream_data.response_audio.sample_rate = model_response.audio.sample_rate
            stream_data.response_audio.channels = model_response.audio.channels
            should_yield = True
        if model_response.image:
            if stream_data.response_image is None:
                stream_data.response_image = model_response.image
        if model_response.extra:
            if stream_data.extra is None:
                stream_data.extra = {}
            stream_data.extra.update(model_response.extra)
        if model_response.response_usage:
            self._add_usage_metrics_to_assistant_message(assistant_message=assistant_message, response_usage=model_response.response_usage)
        if should_yield:
            yield model_response

    def get_function_calls_to_run(self, assistant_message: Message, messages: List[Message], error_response_role: str = 'user') -> List[FunctionCall]:
        function_calls_to_run: List[FunctionCall] = []
        if assistant_message.tool_calls:
            for tool_call in assistant_message.tool_calls:
                _tool_call_id = tool_call.get('id')
                _function_call = get_function_call_for_tool_call(tool_call, self._functions)
                if _function_call is None:
                    messages.append(Message(role=error_response_role, content='Could not find function to call.'))
                    continue
                if _function_call.error:
                    messages.append(Message(role=error_response_role, tool_call_id=_tool_call_id, content=_function_call.error))
                    continue
                function_calls_to_run.append(_function_call)
        return function_calls_to_run

    def _create_function_call_result(self, fc: FunctionCall, success: bool, output: Union[List[Any], str], timer: Timer) -> Message:
        return Message(role=self.tool_message_role, content=output if success else fc.error, tool_call_id=fc.call_id, tool_name=fc.function.name, tool_args=fc.arguments, tool_call_error=not success, stop_after_tool_call=fc.function.stop_after_tool_call, metrics=MessageMetrics(time=timer.elapsed))

    def run_function_calls(self, function_calls: List[FunctionCall], function_call_results: List[Message]) -> Iterator[ModelResponse]:
        if self._function_call_stack is None:
            self._function_call_stack = []
        additional_messages: List[Message] = []
        for fc in function_calls:
            function_call_timer = Timer()
            function_call_timer.start()
            yield ModelResponse(content=fc.get_call_str(), tool_calls=[{'role': self.tool_message_role, 'tool_call_id': fc.call_id, 'tool_name': fc.function.name, 'tool_args': fc.arguments}], event=ResponseEvent.start.value)
            function_call_success = False
            try:
                function_call_success = fc.execute()
            except Exception as e:
                print(f'Error executing function {fc.function.name}: {e}')
                function_call_success = False
                raise e
            function_call_timer.stop()
            function_call_output: Union[List[Any], str] = ''
            if isinstance(fc.result, (types.GeneratorType, collections.abc.Iterator)):
                for item in fc.result:
                    function_call_output += item
                    if fc.function.show_result:
                        yield ModelResponse(content=item)
            else:
                function_call_output = fc.result
                if fc.function.show_result:
                    yield ModelResponse(content=function_call_output)
            function_call_result = self._create_function_call_result(fc, function_call_success, function_call_output, function_call_timer)
            yield ModelResponse(content=f'{fc.get_call_str()} completed in {function_call_timer.elapsed:.4f}s.', tool_calls=[function_call_result.to_function_call_dict()], event=ResponseEvent.end.value)
            function_call_results.append(function_call_result)
            self._function_call_stack.append(fc)
            if self.tool_call_limit and len(self._function_call_stack) >= self.tool_call_limit:
                self.tool_choice = 'none'
                break
        if additional_messages:
            function_call_results.extend(additional_messages)

    async def _arun_function_call(self, function_call: FunctionCall) -> Tuple[bool, Timer, FunctionCall]:
        function_call_timer = Timer()
        function_call_timer.start()
        success = False
        try:
            if inspect.iscoroutinefunction(function_call.function.entrypoint) or inspect.isasyncgenfunction(function_call.function.entrypoint) or inspect.iscoroutine(function_call.function.entrypoint):
                success = await function_call.aexecute()
            else:
                success = await asyncio.to_thread(function_call.execute)
        except Exception as e:
            print(f'Error executing function {function_call.function.name}: {e}')
            success = False
            raise e
        function_call_timer.stop()
        return success, function_call_timer, function_call

    async def arun_function_calls(self, function_calls: List[FunctionCall], function_call_results: List[Message]):
        if self._function_call_stack is None:
            self._function_call_stack = []
        additional_messages: List[Message] = []
        for fc in function_calls:
            yield ModelResponse(content=fc.get_call_str(), tool_calls=[{'role': self.tool_message_role, 'tool_call_id': fc.call_id, 'tool_name': fc.function.name, 'tool_args': fc.arguments}], event=ResponseEvent.start.value)
        results = await asyncio.gather(*(self._arun_function_call(fc) for fc in function_calls), return_exceptions=True)
        for result in results:
            if isinstance(result, BaseException):
                print(f'Error during function call: {result}')
                raise result
            function_call_success, function_call_timer, fc = result
            function_call_output: Union[List[Any], str] = ''
            if isinstance(fc.result, (types.GeneratorType, collections.abc.Iterator)):
                for item in fc.result:
                    function_call_output += item
                    if fc.function.show_result:
                        yield ModelResponse(content=item)
            elif isinstance(fc.result, (types.AsyncGeneratorType, collections.abc.AsyncIterator)):
                async for item in fc.result:
                    function_call_output += item
                    if fc.function.show_result:
                        yield ModelResponse(content=item)
            else:
                function_call_output = fc.result
                if fc.function.show_result:
                    yield ModelResponse(content=function_call_output)
            function_call_result = self._create_function_call_result(fc, function_call_success, function_call_output, function_call_timer)
            yield ModelResponse(content=f'{fc.get_call_str()} completed in {function_call_timer.elapsed:.4f}s.', tool_calls=[function_call_result.to_function_call_dict()], event=ResponseEvent.end.value)
            function_call_results.append(function_call_result)
            self._function_call_stack.append(fc)
            if self.tool_call_limit and len(self._function_call_stack) >= self.tool_call_limit:
                self.tool_choice = 'none'
                break
        if additional_messages:
            function_call_results.extend(additional_messages)

    def _prepare_function_calls(self, assistant_message: Message, messages: List[Message], model_response: ModelResponse) -> List[FunctionCall]:
        if model_response.content is None:
            model_response.content = ''
        if model_response.tool_calls is None:
            model_response.tool_calls = []
        function_calls_to_run: List[FunctionCall] = self.get_function_calls_to_run(assistant_message, messages)
        return function_calls_to_run

    def format_function_call_results(self, messages: List[Message], function_call_results: List[Message], **kwargs) -> None:
        if len(function_call_results) > 0:
            messages.extend(function_call_results)

    def _add_usage_metrics_to_assistant_message(self, assistant_message: Message, response_usage: Any) -> None:
        if isinstance(response_usage, dict):
            if 'input_tokens' in response_usage:
                assistant_message.metrics.input_tokens = response_usage.get('input_tokens', 0)
            if 'output_tokens' in response_usage:
                assistant_message.metrics.output_tokens = response_usage.get('output_tokens', 0)
            if 'prompt_tokens' in response_usage:
                assistant_message.metrics.input_tokens = response_usage.get('prompt_tokens', 0)
            if 'completion_tokens' in response_usage:
                assistant_message.metrics.output_tokens = response_usage.get('completion_tokens', 0)
            if 'total_tokens' in response_usage:
                assistant_message.metrics.total_tokens = response_usage.get('total_tokens', 0)
            else:
                assistant_message.metrics.total_tokens = (assistant_message.metrics.input_tokens + assistant_message.metrics.output_tokens)
        else:
            if hasattr(response_usage, 'input_tokens') and response_usage.input_tokens:
                assistant_message.metrics.input_tokens = response_usage.input_tokens
            if hasattr(response_usage, 'output_tokens') and response_usage.output_tokens:
                assistant_message.metrics.output_tokens = response_usage.output_tokens
            if hasattr(response_usage, 'prompt_tokens') and response_usage.prompt_tokens:
                assistant_message.metrics.input_tokens = response_usage.prompt_tokens
                assistant_message.metrics.prompt_tokens = response_usage.prompt_tokens
            if hasattr(response_usage, 'completion_tokens') and response_usage.completion_tokens:
                assistant_message.metrics.output_tokens = response_usage.completion_tokens
                assistant_message.metrics.completion_tokens = response_usage.completion_tokens
            if hasattr(response_usage, 'total_tokens') and response_usage.total_tokens:
                assistant_message.metrics.total_tokens = response_usage.total_tokens
            else:
                assistant_message.metrics.total_tokens = (assistant_message.metrics.input_tokens + assistant_message.metrics.output_tokens)
        if isinstance(response_usage, dict) and 'additional_metrics' in response_usage:
            assistant_message.metrics.additional_metrics = response_usage['additional_metrics']
        if hasattr(response_usage, 'prompt_tokens_details'):
            if isinstance(response_usage.prompt_tokens_details, dict):
                assistant_message.metrics.prompt_tokens_details = response_usage.prompt_tokens_details
            elif hasattr(response_usage.prompt_tokens_details, 'model_dump'):
                assistant_message.metrics.prompt_tokens_details = response_usage.prompt_tokens_details.to_dict()
        if hasattr(response_usage, 'completion_tokens_details'):
            if isinstance(response_usage.completion_tokens_details, dict):
                assistant_message.metrics.completion_tokens_details = response_usage.completion_tokens_details
            elif hasattr(response_usage.completion_tokens_details, 'model_dump'):
                assistant_message.metrics.completion_tokens_details = (response_usage.completion_tokens_details.to_dict())

    def _log_messages(self, messages: List[Message]) -> None:
        for m in messages:
            m.log(metrics=False)

    def get_system_message_for_model(self) -> Optional[str]:
        return self.system_prompt

    def get_instructions_for_model(self) -> Optional[List[str]]:
        return self.instructions

    def clear(self) -> None:
        self.response_format = None
        self._functions = None
        self._function_call_stack = None

    def __deepcopy__(self, memo: dict):
        cls = self.__class__
        new_model = cls.__new__(cls)
        memo[id(self)] = new_model
        for k, v in self.__dict__.items():
            if k in {'response_format', 'tools', '_functions', '_function_call_stack'}:
                continue
            setattr(new_model, k, deepcopy(v, memo))
        new_model.clear()
        return new_model


class Ollama(Model):
    id: str = 'llama3.1:8b'
    name: str = 'Ollama'
    provider: str = 'Ollama'
    supports_native_structured_outputs: bool = True
    format: Any = None
    options: Any = None
    keep_alive: Union[float, str] = None
    request_params: Dict[str, Any]
    host: str = 'http://localhost:11434'
    timeout: int = 10
    client_params: Dict[str, Any]
    client: ollama.Client = None
    async_client: ollama.AsyncClient = None

    def _get_client_params(self) -> Dict[str, Any]:
        base_params = {'host': self.host, 'timeout': self.timeout}
        client_params = {k: v for k, v in base_params.items() if v}
        if self.client_params:
            client_params.update(self.client_params)
        return client_params

    def get_client(self) -> ollama.Client:
        if self.client:
            return self.client
        self.client = ollama.Client(**self._get_client_params())
        return self.client

    def get_async_client(self) -> ollama.AsyncClient:
        if self.async_client:
            return self.async_client
        return ollama.AsyncClient(**self._get_client_params())

    @property
    def request_kwargs(self) -> Dict[str, Any]:
        base_params = {'format': self.format, 'options': self.options, 'keep_alive': self.keep_alive, 'request_params': self.request_params}
        request_params = {k: v for k, v in base_params.items() if v}
        if self._tools and len(self._tools) > 0:
            request_params['tools'] = self._tools
            for tool in request_params['tools']:
                if 'parameters' in tool['function'] and 'properties' in tool['function']['parameters']:
                    for _, obj in tool['function']['parameters'].get('properties', {}).items():
                        if 'type' in obj and isinstance(obj['type'], list) and len(obj['type']) > 1:
                            obj['type'] = obj['type'][0]
        if self.request_params:
            request_params.update(self.request_params)
        return request_params

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        model_dict = super().to_dict()
        model_dict.update({'format': self.format, 'options': self.options, 'keep_alive': self.keep_alive, 'request_params': self.request_params})
        if self._tools:
            model_dict['tools'] = self._tools
        cleaned_dict = {k: v for k, v in model_dict.items() if v}
        return cleaned_dict

    def _format_message(self, message: Message) -> Dict[str, Any]:
        _message: Dict[str, Any] = {'role': message.role, 'content': message.content}
        if message.role == 'user':
            if message.images:
                message_images = []
                for image in message.images:
                    if image.url:
                        message_images.append(image.image_url_content)
                    if image.filepath:
                        message_images.append(image.filepath)
                    if image.content and isinstance(image.content, bytes):
                        message_images.append(image.content)
                if message_images:
                    _message['images'] = message_images
        return _message

    def _prepare_request_kwargs_for_invoke(self) -> Dict[str, Any]:
        request_kwargs = self.request_kwargs
        if self.response_format and self.structured_outputs:
            if isinstance(self.response_format, type) and issubclass(self.response_format, Base):
                print('Using structured outputs')
                format_schema = self.response_format.model_json_schema()
                if 'format' not in request_kwargs:
                    request_kwargs['format'] = format_schema
        return request_kwargs

    def invoke(self, messages: List[Message]) -> Mapping[str, Any]:
        request_kwargs = self._prepare_request_kwargs_for_invoke()
        return self.get_client().chat(model=self.id.strip(), messages=[self._format_message(m) for m in messages], **{k: v for k, v in request_kwargs.items() if v})

    async def ainvoke(self, messages: List[Message]) -> Mapping[str, Any]:
        request_kwargs = self._prepare_request_kwargs_for_invoke()
        return await self.get_async_client().chat(model=self.id.strip(), messages=[self._format_message(m) for m in messages], **request_kwargs)

    def invoke_stream(self, messages: List[Message]) -> Iterator[Mapping[str, Any]]:
        yield from self.get_client().chat(model=self.id, messages=[self._format_message(m) for m in messages], stream=True, **self.request_kwargs)

    async def ainvoke_stream(self, messages: List[Message]) -> Any:
        async_stream = await self.get_async_client().chat(model=self.id.strip(), messages=[self._format_message(m) for m in messages], stream=True, **self.request_kwargs)
        async for chunk in async_stream:
            yield chunk

    def parse_provider_response(self, response: ollama._types.ChatResponse) -> ModelResponse:
        model_response = ModelResponse()
        response_message: ollama._types.OllamaMessage = response.get('message')
        try:
            if self.response_format and self.structured_outputs and issubclass(self.response_format, Base):
                parsed_object = response_message.content
                if parsed_object:
                    model_response.parsed = parsed_object
        except Exception as e:
            print(f'Error retrieving structured outputs: {e}')
        if response_message.get('role'):
            model_response.role = response_message.get('role')
        if response_message.get('content'):
            model_response.content = response_message.get('content')
        if response_message.get('tool_calls'):
            if model_response.tool_calls is None:
                model_response.tool_calls = []
            for block in response_message.get('tool_calls'):
                tool_call = block.get('function')
                tool_name = tool_call.get('name')
                tool_args = tool_call.get('arguments')
                function_def = {'name': tool_name, 'arguments': (json.dumps(tool_args) if tool_args else None)}
                model_response.tool_calls.append({'type': 'function', 'function': function_def})
        if response.get('done'):
            model_response.response_usage = {'input_tokens': response.get('prompt_eval_count', 0), 'output_tokens': response.get('eval_count', 0), 'total_tokens': response.get('prompt_eval_count', 0) + response.get('eval_count', 0), 'additional_metrics': {'total_duration': response.get('total_duration', 0), 'load_duration': response.get('load_duration', 0), 'prompt_eval_duration': response.get('prompt_eval_duration', 0), 'eval_duration': response.get('eval_duration', 0)}}
        return model_response

    def parse_provider_response_delta(self, response_delta: ollama._types.ChatResponse) -> ModelResponse:
        model_response = ModelResponse()
        response_message = response_delta.get('message')
        if response_message:
            content_delta = response_message.get('content')
            if content_delta and content_delta != '':
                model_response.content = content_delta
            tool_calls = response_message.get('tool_calls')
            if tool_calls:
                for tool_call in tool_calls:
                    tc = tool_call.get('function')
                    tool_name = tc.get('name')
                    tool_args = tc.get('arguments')
                    function_def = {'name': tool_name, 'arguments': json.dumps(tool_args) if tool_args else None}
                    model_response.tool_calls.append({'type': 'function', 'function': function_def})
        if response_delta.get('done'):
            model_response.response_usage = {'input_tokens': response_delta.get('prompt_eval_count', 0), 'output_tokens': response_delta.get('eval_count', 0), 'total_tokens': response_delta.get('prompt_eval_count', 0) + response_delta.get('eval_count', 0), 'additional_metrics': {'total_duration': response_delta.get('total_duration', 0), 'load_duration': response_delta.get('load_duration', 0), 'prompt_eval_duration': response_delta.get('prompt_eval_duration', 0), 'eval_duration': response_delta.get('eval_duration', 0)}}
        return model_response


class ToolCall:
    def __init__(self, tool_calls: List[Dict[str, Any]] = None, response_usage: Mapping[str, Any] = None, response_is_tool_call=False, is_closing_tool_call_tag=False, tool_calls_counter=0, tool_call_content=''):
        self.tool_calls = tool_calls or []
        self.response_usage = response_usage
        self.response_is_tool_call = response_is_tool_call
        self.is_closing_tool_call_tag = is_closing_tool_call_tag
        self.tool_calls_counter = tool_calls_counter
        self.tool_call_content = tool_call_content


class OllamaTools(Ollama):
    @property
    def request_kwargs(self) -> Dict[str, Any]:
        base_params: Dict[str, Any] = {'format': self.format, 'options': self.options, 'keep_alive': self.keep_alive, 'request_params': self.request_params}
        request_params: Dict[str, Any] = {k: v for k, v in base_params.items() if v}
        if self.request_params:
            request_params.update(self.request_params)
        return request_params

    def extract_tool_call_from_string(self, text: str, start_tag: str = '<tool_call>', end_tag: str = '</tool_call>'):
        start_index = text.find(start_tag) + len(start_tag)
        end_index = text.find(end_tag)
        return text[start_index:end_index].strip()

    def parse_provider_response(self, response: ollama._types.ChatResponse) -> ModelResponse:
        model_response = ModelResponse()
        response_message = response.get('message')
        if response_message.get('role'):
            model_response.role = response_message.get('role')
        content = response_message.get('content')
        if content:
            model_response.content = content
            if '<tool_call>' in content and '</tool_call>' in content:
                if model_response.tool_calls is None:
                    model_response.tool_calls = []
                tool_call_responses = content.split('</tool_call>')
                for tool_call_response in tool_call_responses:
                    if tool_call_response != tool_call_responses[-1]:
                        tool_call_response += '</tool_call>'
                    if '<tool_call>' in tool_call_response and '</tool_call>' in tool_call_response:
                        tool_call_content = self.extract_tool_call_from_string(tool_call_response)
                        try:
                            tool_call_dict = json.loads(tool_call_content)
                        except json.JSONDecodeError:
                            raise ValueError(f'Could not parse tool call from: {tool_call_content}')
                        tool_call_name = tool_call_dict.get('name')
                        tool_call_args = tool_call_dict.get('arguments')
                        function_def = {'name': tool_call_name, 'arguments': json.dumps(tool_call_args) if tool_call_args else None}
                        model_response.tool_calls.append({'type': 'function', 'function': function_def})
        if response.get('done'):
            model_response.response_usage = Dot(input_tokens=response.get('prompt_eval_count', 0), output_tokens=response.get('eval_count', 0), total_duration=response.get('total_duration', 0), load_duration=response.get('load_duration', 0), prompt_eval_duration=response.get('prompt_eval_duration', 0), eval_duration=response.get('eval_duration', 0))
            if model_response.response_usage.input_tokens or model_response.response_usage.output_tokens:
                model_response.response_usage.total_tokens = (model_response.response_usage.input_tokens + model_response.response_usage.output_tokens)
        return model_response

    def _create_function_call_result(self, fc: FunctionCall, success: bool, output: Union[List[Any], str], timer: Timer) -> Message:
        content = ('<tool_response>\n'
            + json.dumps({'name': fc.function.name, 'content': output if success else fc.error})
            + '\n</tool_response>')
        return Message(role=self.tool_message_role, content=content, tool_call_id=fc.call_id, tool_name=fc.function.name, tool_args=fc.arguments, tool_call_error=not success, stop_after_tool_call=fc.function.stop_after_tool_call, metrics=MessageMetrics(time=timer.elapsed))

    def format_function_call_results(self, function_call_results: List[Message], messages: List[Message]) -> None:
        if len(function_call_results) > 0:
            for _fc_message in function_call_results:
                _fc_message.content = ('<tool_response>\n'
                    + json.dumps({'name': _fc_message.tool_name, 'content': _fc_message.content})
                    + '\n</tool_response>')
                messages.append(_fc_message)

    def _prepare_function_calls(self, assistant_message: Message, messages: List[Message], model_response: ModelResponse) -> List[FunctionCall]:
        if model_response.content is None:
            model_response.content = ''
        if model_response.tool_calls is None:
            model_response.tool_calls = []
        start_tag: str = '<tool_call>'
        end_tag: str = '</tool_call>'
        text = assistant_message.get_content_string()
        while start_tag in text and end_tag in text:
            start_index = text.find(start_tag)
            end_index = text.find(end_tag) + len(end_tag)
            text = text[:start_index] + text[end_index:]
        model_response.content = str(text)
        model_response.content += '\n\n'
        function_calls_to_run = self.get_function_calls_to_run(assistant_message, messages)
        return function_calls_to_run

    def process_response_stream(self, messages: List[Message], assistant_message: Message, stream_data) -> Iterator[ModelResponse]:
        tool_call_data = ToolCall()
        for response_delta in self.invoke_stream(messages=messages):
            model_response_delta = self.parse_provider_response_delta(response_delta, tool_call_data)
            if model_response_delta:
                yield from self._populate_stream_data_and_assistant_message(stream_data=stream_data, assistant_message=assistant_message, model_response=model_response_delta)

    async def aprocess_response_stream(self, messages: List[Message], assistant_message: Message, stream_data) -> AsyncIterator[ModelResponse]:
        tool_call_data = ToolCall()
        async for response_delta in self.ainvoke_stream(messages=messages):
            model_response_delta = self.parse_provider_response_delta(response_delta, tool_call_data)
            if model_response_delta:
                for model_response in self._populate_stream_data_and_assistant_message(stream_data=stream_data, assistant_message=assistant_message, model_response=model_response_delta):
                    yield model_response

    def parse_provider_response_delta(self, response_delta, tool_call_data: ToolCall) -> ModelResponse:
        model_response = ModelResponse()
        response_message = response_delta.get('message')
        if response_message:
            content_delta = response_message.get('content', '')
            if content_delta and content_delta != '':
                tool_call_data.tool_call_content += content_delta
            if not tool_call_data.response_is_tool_call and '<tool' in content_delta:
                tool_call_data.response_is_tool_call = True
            if tool_call_data.response_is_tool_call:
                if '<tool' in content_delta:
                    tool_call_data.tool_calls_counter += 1
                if tool_call_data.tool_call_content.strip().endswith('</tool_call>'):
                    tool_call_data.tool_calls_counter -= 1
                if tool_call_data.tool_calls_counter == 0 and content_delta.strip().endswith('>'):
                    tool_call_data.response_is_tool_call = False
                    tool_call_data.is_closing_tool_call_tag = True
                    try:

                        tool_calls = []
                        response_content = tool_call_data.tool_call_content
                        if '<tool_call>' in response_content and '</tool_call>' in response_content:
                            tool_call_responses = response_content.split('</tool_call>')
                            for tool_call_response in tool_call_responses:
                                if tool_call_response != tool_call_responses[-1]:
                                    tool_call_response += '</tool_call>'
                                if '<tool_call>' in tool_call_response and '</tool_call>' in tool_call_response:
                                    tool_call_content = self.extract_tool_call_from_string(tool_call_response)
                                    try:
                                        tool_call_dict = json.loads(tool_call_content)
                                    except json.JSONDecodeError:
                                        raise ValueError(f'Could not parse tool call from: {tool_call_content}')
                                    tool_call_name = tool_call_dict.get('name')
                                    tool_call_args = tool_call_dict.get('arguments')
                                    function_def = {'name': tool_call_name}
                                    if tool_call_args:
                                        function_def['arguments'] = json.dumps(tool_call_args)
                                    tool_calls.append({'type': 'function', 'function': function_def})

                        model_response.tool_calls = tool_calls
                        tool_call_data = ToolCall()
                    except Exception as e:
                        print(e)
                        pass
            if not tool_call_data.response_is_tool_call and content_delta:
                if tool_call_data.is_closing_tool_call_tag and content_delta.strip().endswith('>'):
                    tool_call_data.is_closing_tool_call_tag = False
                model_response.content = content_delta
        if response_delta.get('done'):
            model_response.response_usage = {'input_tokens': response_delta.get('prompt_eval_count', 0), 'output_tokens': response_delta.get('eval_count', 0), 'total_tokens': response_delta.get('prompt_eval_count', 0) + response_delta.get('eval_count', 0), 'additional_metrics': {'total_duration': response_delta.get('total_duration', 0), 'load_duration': response_delta.get('load_duration', 0), 'prompt_eval_duration': response_delta.get('prompt_eval_duration', 0), 'eval_duration': response_delta.get('eval_duration', 0)}}
        return model_response

    def get_instructions_to_generate_tool_calls(self) -> List[str]:
        if self._functions:
            return [
                '在第一回合，你没有<tool_results>，所以你不应该编造结果。',
                '要回复用户消息，一次只能使用一个工具。',
                '使用工具时，只能通过工具调用进行响应。没有别的。不要添加任何额外的注释、解释或空格',
                '在任务完成或达到最大迭代次数10之前，不要停止调用函数。']
        return []

    def get_tool_call_prompt(self) -> Optional[str]:
        if self._functions and len(self._functions) > 0:
            tool_call_prompt = textwrap.dedent('''
            您是一个使用语言模型调用的函数。
            您将在<tools></tools>XML标签中获得函数签名。
            您可以使用代理框架进行推理和规划，以帮助用户查询。
            请调用一个函数，并等待在下一次迭代中向您提供函数结果。
            不要对要插入函数的值做出假设。
            调用函数时，不要添加任何额外的注释、解释或空格。
            调用函数后，结果将在<tool_response></tool_response>XML标签中提供给您。
            如果由于函数尚未执行而不存在<tool_response>XML标记，则不要对工具结果做出假设。
            一旦你得到结果，就分析它们，并在需要时调用另一个函数。
            您的最终响应应该直接回答用户查询，并对函数调用的结果进行分析或总结。
            ''')
            tool_call_prompt += '\nHere are the available tools:'
            tool_call_prompt += '\n<tools>\n'
            tool_definitions: List[str] = []
            for _f_name, _function in self._functions.items():
                _function_def = _function.get_definition_for_prompt()
                if _function_def:
                    tool_definitions.append(_function_def)
            tool_call_prompt += '\n'.join(tool_definitions)
            tool_call_prompt += '\n</tools>\n\n'
            tool_call_prompt += textwrap.dedent('''\
            对您将进行的每个工具调用使用以下pydantic模型json模式: {'title': 'FunctionCall', 'type': 'object', 'properties': {'arguments': {'title': 'Arguments', 'type': 'object'}, 'name': {'title': 'Name', 'type': 'string'}}, 'required': ['arguments', 'name']}
            对于每个函数调用，返回一个json对象，其中包含函数名和参数 <tool_call></tool_call> XML标签如下:
            <tool_call>
            {'arguments': <args-dict>, 'name': <function-name>}
            </tool_call>\n
            ''')
            return tool_call_prompt
        return None

    def get_system_message_for_model(self) -> Optional[str]:
        return self.get_tool_call_prompt()

    def get_instructions_for_model(self) -> Optional[List[str]]:
        return self.get_instructions_to_generate_tool_calls()


class RunMessages:
    def __init__(self, messages: List[Message] = None, system_message: Message = None, user_message: Message = None, extra_messages: List[Message] = None):
        self.messages = messages or []
        self.system_message = system_message
        self.user_message = user_message
        self.extra_messages = extra_messages or []

    def get_input_messages(self) -> List[Message]:
        return [i for i in [self.system_message, self.user_message, *self.extra_messages] if i]


class ReasoningStep(Base):
    title: str = ''  # '概括步骤目的的简明标题')
    action: str = ''  # '此步骤衍生出的操作。用第一人称说话')
    result: str = ''  # '执行动作的结果。用第一人称说话')
    reasoning: str = ''  # '这一步骤背后的思考过程和考虑因素')
    over: bool = False  # '指示是继续推理、验证提供的结果，还是确认结果是最终答案')
    confidence: float = 0.0  # '此步骤的置信度得分(0.0至1.0)')


class ReasoningSteps(Base):
    reasoning_steps: List[ReasoningStep]  # '推理步骤列表')


class RunResponseExtraData:
    def __init__(self, references: List[MessageReferences] = None, add_messages: List[Message] = None, reasoning_steps: List['ReasoningStep'] = None, reasoning_messages: List[Message] = None):
        self.references = references or []
        self.add_messages = add_messages or []
        self.reasoning_steps = reasoning_steps or []
        self.reasoning_messages = reasoning_messages or []

    def to_dict(self) -> Dict[str, Any]:
        return {'add_messages': [m.to_dict() for m in self.add_messages], 'reasoning_messages': [m.to_dict() for m in self.reasoning_messages],
                'reasoning_steps': [rs.to_dict() for rs in self.reasoning_steps], 'references': [r.to_dict() for r in self.references]}


class RunResponse:
    def __init__(self, content=None, content_type: str = 'str', thinking: str = None, event: str = RunEvent.run_response.value,
            messages: List[Message] = None, metrics: Dict[str, Any] = None,
            model: str = None, run_id: str = None, agent_id: str = None, session_id: str = None,
            workflow_id: str = None, tools: List[Dict[str, Any]] = None,
            formatted_tool_calls: List[str] = None, images: List[ImageArtifact] = None,
            videos: List[VideoArtifact] = None, audio: List[AudioArtifact] = None,
            response_audio: AudioResponse = None, citations: Citations = None,
            extra_data: RunResponseExtraData = None, created_at=0):
        self.content = content
        self.content_type = content_type
        self.thinking = thinking
        self.event = event
        self.messages = messages or []
        self.metrics = metrics or {}
        self.model = model
        self.run_id = run_id
        self.agent_id = agent_id
        self.session_id = session_id
        self.workflow_id = workflow_id
        self.tools = tools or []
        self.formatted_tool_calls = formatted_tool_calls or []
        self.images = images or []
        self.videos = videos or []
        self.audio = audio or []
        self.response_audio = response_audio
        self.citations = citations
        self.extra_data = extra_data
        self.created_at = created_at or int(time.time())

    def to_dict(self) -> Dict[str, Any]:
        res = self.__dict__.copy()
        res.update({'messages': [m.to_dict() for m in self.messages],
                    'extra_data': (self.extra_data.to_dict() if isinstance(self.extra_data, RunResponseExtraData) else self.extra_data),
                    'images': [img.to_dict() for img in self.images],
                    'videos': [vid.to_dict() for vid in self.videos],
                    'audio': [aud.to_dict() for aud in self.audio],
                    'response_audio': self.response_audio.to_dict() if isinstance(self.response_audio, AudioResponse) else self.response_audio})
        if isinstance(self.content, Base):
            res['content'] = self.content.to_dict()
        return {k: v for k, v in res.items() if v}

    def get_content_as_string(self, **kwargs) -> str:
        if isinstance(self.content, str):
            return self.content
        elif isinstance(self.content, Base):
            return self.content.to_dict()
        else:
            return json.dumps(self.content, **kwargs)


class TeamRunResponse(RunResponse):
    def __init__(self, member_responses: List[Union['TeamRunResponse', RunResponse]] = None, team_id: str = None, **kwargs):
        super().__init__(**kwargs)
        self.member_responses = member_responses or []
        self.team_id = team_id

    def to_dict(self) -> Dict[str, Any]:
        res = {k: v for k, v in self.__dict__.items() if v and k not in ['extra_data', 'response_audio']}
        if self.extra_data:
            res['extra_data'] = self.extra_data.to_dict()
        if self.response_audio:
            res['response_audio'] = self.response_audio.to_dict()
        if isinstance(self.content, Base):
            res['content'] = self.content.to_dict()
        res.update({'messages': [m.to_dict() for m in self.messages],
                    'images': [img.to_dict() for img in self.images],
                    'videos': [vid.to_dict() for vid in self.videos],
                    'audio': [aud.to_dict() for aud in self.audio],
                    'member_responses': [response.to_dict() for response in self.member_responses]})
        return res

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'TeamRunResponse':
        messages = [Message.model_validate(message) for message in data.pop('messages', [])]
        member_responses = data.pop('member_responses', None)
        parsed_member_responses: List[Union['TeamRunResponse', RunResponse]] = []
        if member_responses:
            for response in member_responses:
                if 'agent_id' in response:
                    parsed_member_responses.append(RunResponse(**response))
                else:
                    parsed_member_responses.append(cls.from_dict(response))
        extra_data = data.pop('extra_data', None)
        if extra_data:
            extra_data = RunResponseExtraData(**extra_data)
        images = [ImageArtifact.model_validate(image) for image in data.pop('images', [])]
        videos = [VideoArtifact.model_validate(video) for video in data.pop('videos', [])]
        audio = [AudioArtifact.model_validate(audio) for audio in data.pop('audio', [])]
        response_audio = data.pop('response_audio', None)
        response_audio = AudioResponse.model_validate(response_audio) if response_audio else None
        return cls(messages=messages, member_responses=parsed_member_responses, extra_data=extra_data, images=images, videos=videos, audio=audio, response_audio=response_audio, **data)

    def get_content_as_string(self, **kwargs) -> str:
        if isinstance(self.content, str):
            return self.content
        elif isinstance(self.content, Base):
            return self.content.to_dict()
        else:
            return json.dumps(self.content, **kwargs)

    def add_member_run(self, run_response: Union['TeamRunResponse', RunResponse]) -> None:
        self.member_responses.append(run_response)
        self.images.extend(run_response.images)
        self.videos.extend(run_response.videos)
        self.audio.extend(run_response.audio)


class AgentRun(Base):
    message: Message = None
    messages: List[Message]
    response: RunResponse = None

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        response = {'message': self.message.to_dict() if self.message else None, 'messages': [message.to_dict() for message in self.messages] if self.messages else None, 'response': self.response.to_dict() if self.response else None}
        return {k: v for k, v in response.items() if v}


class MemoryRow(Base):
    memory: Dict[str, Any]
    user_id: str = ''
    created_at: datetime = None
    updated_at: datetime = None
    id: str = ''

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        _dict = self.to_dict()
        _dict['created_at'] = self.created_at.isoformat() if self.created_at else None
        _dict['updated_at'] = self.updated_at.isoformat() if self.updated_at else None
        if self.id is None:
            memory_str = json.dumps(self.memory, sort_keys=True)
            cleaned_memory = memory_str.replace(' ', '').replace('\n', '').replace('\t', '')
            self.id = hashlib.md5(cleaned_memory.encode()).hexdigest()
        return _dict


class MemoryDb:
    def __init__(self, db_url: str, table_name='memory'):
        self.table_name = table_name
        self.db_url = db_url
        self.db_engine: sqlalchemy.Engine = sqlalchemy.create_engine(db_url)
        self.metadata: sqlalchemy.MetaData = sqlalchemy.MetaData()
        self.inspector = sqlalchemy.inspect(self.db_engine)
        self.Session = sqlalchemy.orm.scoped_session(sqlalchemy.orm.sessionmaker(bind=self.db_engine))
        self.table: sqlalchemy.Table = self.get_table()

    def get_table(self) -> sqlalchemy.Table:
        return sqlalchemy.Table(
            self.table_name, self.metadata, sqlalchemy.Column('id', sqlalchemy.String, primary_key=True),
            sqlalchemy.Column('user_id', sqlalchemy.String), sqlalchemy.Column('memory', sqlalchemy.String),
            sqlalchemy.Column('created_at', sqlalchemy.DateTime, server_default=sqlalchemy.text('CURRENT_TIMESTAMP')),
            sqlalchemy.Column('updated_at', sqlalchemy.DateTime, server_default=sqlalchemy.text('CURRENT_TIMESTAMP'), onupdate=sqlalchemy.text('CURRENT_TIMESTAMP')), extend_existing=True)

    def create(self) -> None:
        if not self.table_exists():
            try:
                print(f'Creating table: {self.table_name}')
                self.table.create(self.db_engine, checkfirst=True)
            except Exception as e:
                print(f'Error creating table "{self.table_name}": {e}')
                raise

    def memory_exists(self, memory: MemoryRow) -> bool:
        with self.Session() as session:
            stmt = sqlalchemy.select(self.table.c.id).where(self.table.c.id == memory.id)
            result = session.execute(stmt).first()
            return result

    def read_memories(self, user_id: str = None, limit: int = None, sort: str = None) -> List[MemoryRow]:
        memories: List[MemoryRow] = []
        try:
            with self.Session() as session:
                stmt = sqlalchemy.select(self.table)
                if user_id:
                    stmt = stmt.where(self.table.c.user_id == user_id)
                if sort == 'asc':
                    stmt = stmt.order_by(self.table.c.created_at.asc())
                else:
                    stmt = stmt.order_by(self.table.c.created_at.desc())
                if limit:
                    stmt = stmt.limit(limit)
                result = session.execute(stmt)
                for row in result:
                    memories.append(MemoryRow(id=row.id, user_id=row.user_id, memory=eval(row.memory)))
        except sqlalchemy.exc.SQLAlchemyError as e:
            self.create()
        return memories

    def upsert_memory(self, memory: MemoryRow, create_and_retry: bool = True) -> None:
        try:
            with self.Session() as session:
                existing = session.execute(sqlalchemy.select(self.table).where(self.table.c.id == memory.id)).first()
                if existing:
                    stmt = (self.table.update()
                        .where(self.table.c.id == memory.id)
                        .values(user_id=memory.user_id, memory=str(memory.memory), updated_at=sqlalchemy.text('CURRENT_TIMESTAMP')))
                else:
                    stmt = self.table.insert().values(id=memory.id, user_id=memory.user_id, memory=str(memory.memory))
                session.execute(stmt)
                session.commit()
        except sqlalchemy.exc.SQLAlchemyError as e:
            print(f'Exception upserting into table: {e}')
            if not self.table_exists():
                print(f'Table does not exist: {self.table_name}')
                print('Creating table for future transactions')
                self.create()
                if create_and_retry:
                    return self.upsert_memory(memory, create_and_retry=False)
            else:
                raise

    def delete_memory(self, id: str) -> None:
        with self.Session() as session:
            stmt = sqlalchemy.delete(self.table).where(self.table.c.id == id)
            session.execute(stmt)
            session.commit()

    def drop_table(self) -> None:
        if self.table_exists():
            print(f'Deleting table: {self.table_name}')
            self.table.drop(self.db_engine)

    def table_exists(self) -> bool:
        print(f'Checking if table exists: {self.table.name}')
        try:
            return self.inspector.has_table(self.table.name)
        except Exception as e:
            print(e)
            return False

    def clear(self) -> bool:
        with self.Session() as session:
            stmt = sqlalchemy.delete(self.table)
            session.execute(stmt)
            session.commit()
        return True


class MemoryManager(Base):
    model: Model = None
    user_id: str = ''
    limit: int = 0
    system_prompt: str = ''
    db: MemoryDb = None
    input_message: str = ''
    _tools_for_model: List[Dict]
    _functions_for_model: Dict[str, Function]

    def update_model(self):
        if self.model is None:
            self.model = Ollama()
        self.add_tools_to_model(model=self.model)

    def add_tools_to_model(self, model: Model):
        if self._tools_for_model is None:
            self._tools_for_model = []
        if self._functions_for_model is None:
            self._functions_for_model = {}
        for tool in [
            self.add_memory, self.update_memory, self.delete_memory, self.clear_memory, ]:
            try:
                function_name = tool.__name__
                if function_name not in self._functions_for_model:
                    func = Function.from_callable(tool)
                    self._functions_for_model[func.name] = func
                    self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                    print(f'Included function {func.name}')
            except Exception as e:
                print(f'Could not add function {tool}: {e}')
        model.set_tools(tools=self._tools_for_model)
        model.set_functions(functions=self._functions_for_model)

    def get_existing_memories(self) -> Optional[List[MemoryRow]]:
        if self.db is None:
            return None
        return self.db.read_memories(user_id=self.user_id, limit=self.limit)

    def add_memory(self, memory: str) -> str:
        try:
            if self.db:
                self.db.upsert_memory(MemoryRow(user_id=self.user_id, memory=Memory(memory=memory, input=self.input_message).to_dict()))
            return 'Memory added successfully'
        except Exception as e:
            print(f'Error storing memory in db: {e}')
            return f'Error adding memory: {e}'

    def delete_memory(self, id: str) -> str:
        try:
            if self.db:
                self.db.delete_memory(id=id)
            return 'Memory deleted successfully'
        except Exception as e:
            print(f'Error deleting memory in db: {e}')
            return f'Error deleting memory: {e}'

    def update_memory(self, id: str, memory: str) -> str:
        try:
            if self.db:
                self.db.upsert_memory(MemoryRow(id=id, user_id=self.user_id, memory=Memory(memory=memory, input=self.input_message).to_dict()))
            return 'Memory updated successfully'
        except Exception as e:
            print(f'Error updating memory in db: {e}')
            return f'Error updating memory: {e}'

    def clear_memory(self) -> str:
        try:
            if self.db:
                self.db.clear()
            return 'Memory cleared successfully'
        except Exception as e:
            print(f'Error clearing memory in db: {e}')
            return f'Error clearing memory: {e}'

    def get_system_message(self) -> Message:
        system_prompt_lines = [
            '你的任务是为用户的消息生成一个简洁的记忆。',
            '创建一个能够捕获用户提供的关键信息的内存，就像您正在存储它以供将来参考一样',
            '记忆应该是一个简短的第三人称陈述，概括了用户输入中最重要的方面，没有添加任何无关的细节。',
            '这些记忆将用于增强用户在后续对话中的体验。',
            '您还将收到一份现有记忆列表。您可以：\n1.使用`Add_memory`工具添加新内存。\n2.使用`Update_memory`工具更新内存。\n3.使用`Delete_memory`工具删除内存。\n4.使用“Clear_memory”工具清除所有记忆。使用时要格外小心，因为它会从数据库中删除所有内存。'
        ]
        existing_memories = self.get_existing_memories()
        if existing_memories and len(existing_memories) > 0:
            system_prompt_lines.extend([
                    '\nExisting memories:', '<existing_memories>\n'
                    + '\n'.join([f'  - id: {m.id} | memory: {m.memory}' for m in existing_memories])
                    + '\n</existing_memories>', ])
        return Message(role='system', content='\n'.join(system_prompt_lines))

    def run(self, message: str = None, **kwargs: Any) -> Optional[str]:
        print('*********** MemoryManager Start ***********')
        self.update_model()
        messages_for_model: List[Message] = [self.get_system_message()]
        user_prompt_message = Message(role='user', content=message, **kwargs) if message else None
        if user_prompt_message:
            messages_for_model += [user_prompt_message]
        self.input_message = message
        response = self.model.response(messages=messages_for_model)
        print('*********** MemoryManager End ***********')
        return response.content

    async def arun(self, message: str = None, **kwargs: Any) -> Optional[str]:
        print('*********** Async MemoryManager Start ***********')
        self.update_model()
        messages_for_model: List[Message] = [self.get_system_message()]
        user_prompt_message = Message(role='user', content=message, **kwargs) if message else None
        if user_prompt_message:
            messages_for_model += [user_prompt_message]
        self.input_message = message
        response = await self.model.aresponse(messages=messages_for_model)
        print('*********** Async MemoryManager End ***********')
        return response.content


class Memory(Base):
    memory: str
    id: str
    topic: str
    input: str


class MemoryClassifier(Base):
    model: Model = None
    system_prompt: str
    existing_memories: List[Memory]

    def update_model(self) -> None:
        if self.model is None:
            self.model = Ollama()

    def get_system_message(self) -> Message:
        system_prompt_lines = [
            '你的任务是确定用户的消息是否包含值得记住的信息，以便将来进行对话。',
            '这包括可以个性化与用户正在进行的交互的详细信息，例如：\n- 个人事实：姓名、年龄、职业、地点、兴趣、偏好等。\n',
            '-用户分享的重大生活事件或经历\n-有关用户当前状况、挑战或目标的重要上下文\n- 用户喜欢或不喜欢什么，他们的观点、信仰、价值观等。\n',
            '-提供对用户个性、观点或需求有价值见解的任何其他细节', '你的任务是决定用户输入是否包含任何值得记住的上述信息。',
            '如果用户输入包含任何值得记住以备将来对话的信息，请回答‘是’。', '如果输入不包含任何值得保存的重要细节，请回答“否”以忽略它。',
            '您还将收到一份现有记忆列表，以帮助您决定输入是新的还是已知的。', '如果与输入匹配的内存已经存在，请使用‘否’进行响应以保持原样。',
            '如果存在需要更新或删除的内存，请回答“是”以更新 / 删除它。', '您只能回答“是”或“否”。其他任何内容都不会被视为有效的回答。'
        ]
        if self.existing_memories and len(self.existing_memories) > 0:
            system_prompt_lines.extend([
                    '\nExisting memories:', '<existing_memories>\n'
                    + '\n'.join([f'  - {m.memory}' for m in self.existing_memories])
                    + '\n</existing_memories>'])
        return Message(role='system', content='\n'.join(system_prompt_lines))

    def run(self, message: str = None, **kwargs: Any) -> Optional[str]:
        print('*********** MemoryClassifier Start ***********')
        self.update_model()
        messages_for_model: List[Message] = [self.get_system_message()]
        user_prompt_message = Message(role='user', content=message, **kwargs) if message else None
        if user_prompt_message:
            messages_for_model += [user_prompt_message]
        response = self.model.response(messages=messages_for_model)
        print('*********** MemoryClassifier End ***********')
        return response.content

    async def arun(self, message: str = None, **kwargs: Any) -> Optional[str]:
        print('*********** Async MemoryClassifier Start ***********')
        self.update_model()
        messages_for_model: List[Message] = [self.get_system_message()]
        user_prompt_message = Message(role='user', content=message, **kwargs) if message else None
        if user_prompt_message:
            messages_for_model += [user_prompt_message]
        response = await self.model.aresponse(messages=messages_for_model)
        print('*********** Async MemoryClassifier End ***********')
        return response.content


class SessionSummary(Base):
    summary: str  # '会议总结。简明扼要，只关注重要信息。不要编造任何东西。'
    topics: List[str]  # '会议讨论的主题'


class MemorySummarizer(Base):
    model: Model = Ollama()
    use_structured_outputs: bool = False

    def update_model(self) -> None:
        if self.model is None:
            self.model = Ollama()
        if self.use_structured_outputs:
            self.model.response_format = SessionSummary
            self.model.structured_outputs = True
        else:
            self.model.response_format = {'type': 'json_object'}

    def get_system_message(self, messages_for_summarization: List[Dict[str, str]]) -> Message:
        system_prompt = textwrap.dedent('分析用户和助手之间的以下对话，并提取以下详细信息：\n-摘要（str）：提供会议的简明摘要，重点介绍有助于未来互动的重要信息。\n-Topics（可选[List[str]]）：列出会话中讨论的主题。\n请忽略任何琐碎的信息。\n对话：')
        conversation = []
        for message_pair in messages_for_summarization:
            conversation.append(f'User: {message_pair["user"]}')
            if 'assistant' in message_pair:
                conversation.append(f'Assistant: {message_pair["assistant"]}')
            elif 'model' in message_pair:
                conversation.append(f'Assistant: {message_pair["model"]}')
        system_prompt += '\n'.join(conversation)
        if not self.use_structured_outputs:
            system_prompt += '\n\nProvide your output as a JSON containing the following fields:'
            json_schema = SessionSummary.model_json_schema()
            response_model_properties = {}
            json_schema_properties = json_schema.get('properties')
            if json_schema_properties:
                for field_name, field_properties in json_schema_properties.items():
                    formatted_field_properties = {prop_name: prop_value
                        for prop_name, prop_value in field_properties.items()
                        if prop_name != 'title'}
                    response_model_properties[field_name] = formatted_field_properties
            if len(response_model_properties) > 0:
                system_prompt += '\n<json_fields>'
                system_prompt += f'\n{json.dumps([key for key in response_model_properties.keys() if key != "$defs"])}'
                system_prompt += '\n</json_fields>'
                system_prompt += '\nHere are the properties for each field:'
                system_prompt += '\n<json_field_properties>'
                system_prompt += f'\n{json.dumps(response_model_properties, indent=2)}'
                system_prompt += '\n</json_field_properties>'
            system_prompt += '\nStart your response with `{` and end it with `}`.'
            system_prompt += '\nYour output will be passed to json.loads() to convert it to a Python object.'
            system_prompt += '\nMake sure it only contains valid JSON.'
        return Message(role='system', content=system_prompt)

    def run(self, message_pairs: List[Tuple[Message, Message]], **kwargs: Any) -> Optional[SessionSummary]:
        print('*********** MemorySummarizer Start ***********')
        if message_pairs is None or len(message_pairs) == 0:
            print('No message pairs provided for summarization.')
            return None
        self.update_model()
        messages_for_summarization: List[Dict[str, str]] = []
        for message_pair in message_pairs:
            user_message, assistant_message = message_pair
            messages_for_summarization.append({user_message.role: user_message.get_content_string(), assistant_message.role: assistant_message.get_content_string()})
        messages_for_model: List[Message] = [
            self.get_system_message(messages_for_summarization), Message(role='user', content='Provide the summary of the conversation.'), ]
        response = self.model.response(messages=messages_for_model)
        print('*********** MemorySummarizer End ***********')
        if self.use_structured_outputs and response.parsed and isinstance(response.parsed, SessionSummary):
            return response.parsed
        if isinstance(response.content, str):
            try:
                session_summary = None
                try:
                    session_summary = SessionSummary(**json.loads(response.content))
                except json.JSONDecodeError:
                    if response.content.startswith('```json'):
                        response.content = response.content.replace('```json\n', '').replace('\n```', '')
                        session_summary = SessionSummary(**json.loads(response.content))
                return session_summary
            except Exception as e:
                print(f'Failed to convert response to session_summary: {e}')
        return None

    async def arun(self, message_pairs: List[Tuple[Message, Message]], **kwargs: Any) -> Optional[SessionSummary]:
        print('*********** Async MemorySummarizer Start ***********')
        if message_pairs is None or len(message_pairs) == 0:
            print('No message pairs provided for summarization.')
            return None
        self.update_model()
        messages_for_summarization: List[Dict[str, str]] = []
        for message_pair in message_pairs:
            user_message, assistant_message = message_pair
            messages_for_summarization.append({user_message.role: user_message.get_content_string(), assistant_message.role: assistant_message.get_content_string()})
        messages_for_model: List[Message] = [
            self.get_system_message(messages_for_summarization), Message(role='user', content='Provide the summary of the conversation.'), ]
        response = await self.model.aresponse(messages=messages_for_model)
        print('*********** Async MemorySummarizer End ***********')
        if self.use_structured_outputs and response.parsed and isinstance(response.parsed, SessionSummary):
            return response.parsed
        if isinstance(response.content, str):
            try:
                session_summary = None
                try:
                    session_summary = SessionSummary(**json.loads(response.content))
                except json.JSONDecodeError:
                    if response.content.startswith('```json'):
                        response.content = response.content.replace('```json\n', '').replace('\n```', '')
                        session_summary = SessionSummary(**json.loads(response.content))
                return session_summary
            except Exception as e:
                print(f'Failed to convert response to session_summary: {e}')
        return None


class TeamRun:
    def __init__(self,  message: Message = None, member_runs: List[AgentRun] = None, response: TeamRunResponse = None):
        self.message = message
        self.member_runs = member_runs
        self.response = response

    def to_dict(self) -> Dict[str, Any]:
        message = self.message.to_dict() if self.message else None
        member_responses = [run.to_dict() for run in self.member_runs] if self.member_runs else None
        response = self.response.to_dict() if self.response else None
        return {'message': message, 'member_responses': member_responses, 'response': response}

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'TeamRun':
        message = Message.model_validate(data.get('message')) if data.get('message') else None
        member_runs = ([AgentRun.model_validate(run) for run in data.get('member_runs', [])] if data.get('member_runs') else None)
        response = TeamRunResponse.from_dict(data.get('response', {})) if data.get('response') else None
        return cls(message=message, member_runs=member_runs, response=response)


class TeamMemberInteraction:
    def __init__(self, member_name='', task='', response: RunResponse = None):
        self.member_name = member_name
        self.task = task
        self.response = response


class TeamContext:
    def __init__(self, member_interactions: List[TeamMemberInteraction] = None, text: str = None):
        self.member_interactions = member_interactions or []
        self.text = text


class AgentMemory(Base):
    runs: List[AgentRun]
    messages: List[Message]
    update_system_message_on_change: bool = False

    summary: SessionSummary = None
    create_session_summary: bool = False
    update_session_summary_after_run: bool = True
    summarizer: MemorySummarizer = None

    create_user_memories: bool = False
    update_user_memories_after_run: bool = True
    db: MemoryDb = None
    user_id: str
    memories: List[Memory] = None
    num_memories: int = 0
    classifier: MemoryClassifier = None
    manager: MemoryManager = None
    updating_memory: bool = False

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        _memory_dict = super().to_dict(include={'update_system_message_on_change', 'create_session_summary', 'update_session_summary_after_run', 'create_user_memories', 'update_user_memories_after_run', 'user_id', 'num_memories'})
        if self.summary:
            _memory_dict['summary'] = self.summary.to_dict()
        if self.memories:
            _memory_dict['memories'] = [memory.to_dict() for memory in self.memories]
        if self.messages:
            _memory_dict['messages'] = [message.to_dict() for message in self.messages]
        if self.runs:
            _memory_dict['runs'] = [run.to_dict() for run in self.runs]
        return _memory_dict

    def add_run(self, agent_run: AgentRun) -> None:
        self.runs.append(agent_run)
        print('Added AgentRun to AgentMemory')

    def add_system_message(self, message: Message, system_message_role: str = 'system') -> None:
        if len(self.messages) == 0:
            if message:
                self.messages.append(message)
        else:
            system_message_index = next((i for i, m in enumerate(self.messages) if m.role == system_message_role), None)
            if system_message_index:
                if self.messages[system_message_index].content != message.content and self.update_system_message_on_change:
                    print('Updating system message in memory with new content')
                    self.messages[system_message_index] = message
            else:
                self.messages.insert(0, message)

    def add_messages(self, messages: List[Message]) -> None:
        self.messages.extend(messages)
        print(f'Added {len(messages)} Messages to AgentMemory')

    def get_messages(self) -> List[Dict[str, Any]]:
        return [message.to_dict() for message in self.messages]

    def get_messages_from_last_n_runs(self, last_n: int = None, skip_role: str = None) -> List[Message]:
        if not self.runs:
            return []
        runs_to_process = self.runs if last_n is None else self.runs[-last_n:]
        messages_from_history = []
        for run in runs_to_process:
            if not (run.response and run.response.messages):
                continue
            for message in run.response.messages:
                if skip_role and message.role == skip_role:
                    continue
                if hasattr(message, 'from_history') and message.from_history:
                    continue
                messages_from_history.append(message)
        print(f'Getting messages from previous runs: {len(messages_from_history)}')
        return messages_from_history

    def get_message_pairs(self, user_role: str = 'user', assistant_role: List[str] = None) -> List[Tuple[Message, Message]]:
        if assistant_role is None:
            assistant_role = ['assistant', 'model', 'CHATBOT']
        runs_as_message_pairs: List[Tuple[Message, Message]] = []
        for run in self.runs:
            if run.response and run.response.messages:
                user_messages_from_run = None
                assistant_messages_from_run = None
                for message in run.response.messages:
                    if hasattr(message, 'from_history') and message.from_history:
                        continue
                    if message.role == user_role:
                        user_messages_from_run = message
                        break
                for message in run.response.messages[::-1]:
                    if hasattr(message, 'from_history') and message.from_history:
                        continue
                    if message.role in assistant_role:
                        assistant_messages_from_run = message
                        break
                if user_messages_from_run and assistant_messages_from_run:
                    runs_as_message_pairs.append((user_messages_from_run, assistant_messages_from_run))
        return runs_as_message_pairs

    def get_tool_calls(self, num_calls: int = None) -> List[Dict[str, Any]]:
        tool_calls = []
        for message in self.messages[::-1]:
            if message.tool_calls:
                for tool_call in message.tool_calls:
                    tool_calls.append(tool_call)
                    if num_calls and len(tool_calls) >= num_calls:
                        return tool_calls
        return tool_calls

    def load_user_memories(self) -> None:
        if self.db is None:
            return
        try:
            memory_rows = self.db.read_memories(user_id=self.user_id, limit=self.num_memories, sort='desc')
        except Exception as e:
            print(f'Error reading memory: {e}')
            return
        self.memories = []
        if memory_rows is None or len(memory_rows) == 0:
            return
        for row in memory_rows:
            try:
                self.memories.append(Memory.model_validate(row.memory))
            except Exception as e:
                print(f'Error loading memory: {e}')
                continue

    def should_update_memory(self, input: str) -> bool:
        if self.classifier is None:
            self.classifier = MemoryClassifier()
        self.classifier.existing_memories = self.memories
        classifier_response = self.classifier.run(input)
        if classifier_response == 'yes':
            return True
        return False

    async def ashould_update_memory(self, input: str) -> bool:
        if self.classifier is None:
            self.classifier = MemoryClassifier()
        self.classifier.existing_memories = self.memories
        classifier_response = await self.classifier.arun(input)
        if classifier_response == 'yes':
            return True
        return False

    def update_memory(self, input: str, force: bool = False) -> Optional[str]:
        if input is None or not isinstance(input, str):
            return 'Invalid message content'
        if self.db is None:
            print('MemoryDb not provided.')
            return 'Please provide a db to store memories'
        self.updating_memory = True
        should_update_memory = force or self.should_update_memory(input=input)
        print(f'Update memory: {should_update_memory}')
        if not should_update_memory:
            print('Memory update not required')
            return 'Memory update not required'
        if self.manager is None:
            self.manager = MemoryManager(user_id=self.user_id, db=self.db)
        else:
            self.manager.db = self.db
            self.manager.user_id = self.user_id
        response = self.manager.run(input)
        self.load_user_memories()
        self.updating_memory = False
        return response

    async def aupdate_memory(self, input: str, force: bool = False) -> Optional[str]:
        if input is None or not isinstance(input, str):
            return 'Invalid message content'
        if self.db is None:
            print('MemoryDb not provided.')
            return 'Please provide a db to store memories'
        self.updating_memory = True
        should_update_memory = force or await self.ashould_update_memory(input=input)
        print(f'Async update memory: {should_update_memory}')
        if not should_update_memory:
            print('Memory update not required')
            return 'Memory update not required'
        if self.manager is None:
            self.manager = MemoryManager(user_id=self.user_id, db=self.db)
        else:
            self.manager.db = self.db
            self.manager.user_id = self.user_id
        response = await self.manager.arun(input)
        self.load_user_memories()
        self.updating_memory = False
        return response

    def update_summary(self) -> Optional[SessionSummary]:
        self.updating_memory = True
        if self.summarizer is None:
            self.summarizer = MemorySummarizer()
        self.summary = self.summarizer.run(self.get_message_pairs())
        self.updating_memory = False
        return self.summary

    async def aupdate_summary(self) -> Optional[SessionSummary]:
        self.updating_memory = True
        if self.summarizer is None:
            self.summarizer = MemorySummarizer()
        self.summary = await self.summarizer.arun(self.get_message_pairs())
        self.updating_memory = False
        return self.summary

    def clear(self) -> None:
        self.runs = []
        self.messages = []
        self.summary = None
        self.memories = None

    def deep_copy(self) -> 'AgentMemory':
        copied_obj = self.__class__(**self.to_dict())
        for field_name, field_value in self.__dict__.items():
            if field_name not in ['db', 'classifier', 'manager', 'summarizer']:
                try:
                    setattr(copied_obj, field_name, deepcopy(field_value))
                except Exception as e:
                    print(f'Failed to deepcopy field: {field_name} - {e}')
                    setattr(copied_obj, field_name, field_value)
        copied_obj.db = self.db
        copied_obj.classifier = self.classifier
        copied_obj.manager = self.manager
        copied_obj.summarizer = self.summarizer
        return copied_obj


class TeamMemory(Base):
    runs: List[TeamRun]
    messages: List[Message]
    update_system_message_on_change = True
    team_context: TeamContext = None
    create_user_memories = False
    update_user_memories_after_run = True
    db: MemoryDb = None
    user_id: str = ''
    memories: List[Memory]
    num_memories: int = 0
    classifier: MemoryClassifier = None
    manager: MemoryManager = None
    updating_memory = False

    def to_dict(self, no_none=True, include: list = None) -> Dict[str, Any]:
        _memory_dict = {}
        for key, value in self.__dict__.items():
            if value and key in ['update_system_message_on_change', 'create_user_memories', 'update_user_memories_after_run', 'user_id', 'num_memories']:
                _memory_dict[key] = value
        if self.messages:
            _memory_dict['messages'] = [message.to_dict() for message in self.messages]
        if self.memories:
            _memory_dict['memories'] = [memory.to_dict() for memory in self.memories]
        if self.runs:
            _memory_dict['runs'] = [run.to_dict() for run in self.runs]
        return _memory_dict

    def add_interaction_to_team_context(self, member_name: str, task: str, run_response: RunResponse) -> None:
        if self.team_context is None:
            self.team_context = TeamContext()
        self.team_context.member_interactions.append(TeamMemberInteraction(member_name=member_name, task=task, response=run_response))
        print(f'Updated team context with member name: {member_name}')

    def set_team_context_text(self, text: str) -> None:
        if self.team_context:
            self.team_context.text = text
        else:
            self.team_context = TeamContext(text=text)

    def get_team_context_str(self) -> str:
        if self.team_context and self.team_context.text:
            return f'<team context>\n{self.team_context.text}\n</team context>\n'
        return ''

    def get_team_member_interactions_str(self) -> str:
        team_member_interactions_str = ''
        if self.team_context and self.team_context.member_interactions:
            team_member_interactions_str += '<member interactions>\n'
            for interaction in self.team_context.member_interactions:
                team_member_interactions_str += f'Member: {interaction.member_name}\n'
                team_member_interactions_str += f'Task: {interaction.task}\n'
                team_member_interactions_str += f"Response: {interaction.response.to_dict().get('content', '')}\n"
                team_member_interactions_str += '\n'
            team_member_interactions_str += '</member interactions>\n'
        return team_member_interactions_str

    def get_team_context_images(self) -> List[ImageArtifact]:
        images = []
        if self.team_context and self.team_context.member_interactions:
            for interaction in self.team_context.member_interactions:
                if interaction.response.images:
                    images.extend(interaction.response.images)
        return images

    def get_team_context_videos(self) -> List[VideoArtifact]:
        videos = []
        if self.team_context and self.team_context.member_interactions:
            for interaction in self.team_context.member_interactions:
                if interaction.response.videos:
                    videos.extend(interaction.response.videos)
        return videos

    def get_team_context_audio(self) -> List[AudioArtifact]:
        audio = []
        if self.team_context and self.team_context.member_interactions:
            for interaction in self.team_context.member_interactions:
                if interaction.response.audio:
                    audio.extend(interaction.response.audio)
        return audio

    def add_team_run(self, team_run: TeamRun) -> None:
        self.runs.append(team_run)
        print('Added TeamRun to TeamMemory')

    def add_system_message(self, message: Message, system_message_role: str = 'system') -> None:
        if len(self.messages) == 0:
            if message:
                self.messages.append(message)
        else:
            system_message_index = next((i for i, m in enumerate(self.messages) if m.role == system_message_role), None)
            if system_message_index:
                if self.messages[system_message_index].content != message.content and self.update_system_message_on_change:
                    print('Updating system message in memory with new content')
                    self.messages[system_message_index] = message
            else:
                self.messages.insert(0, message)

    def add_messages(self, messages: List[Message]) -> None:
        self.messages.extend(messages)
        print(f'Added {len(messages)} Messages to TeamMemory')

    def get_messages(self) -> List[Dict[str, Any]]:
        return [message.to_dict() for message in self.messages]

    def get_messages_from_last_n_runs(self, last_n: int = None, skip_role: str = None) -> List[Message]:
        if not self.runs:
            return []
        runs_to_process = self.runs if last_n is None else self.runs[-last_n:]
        messages_from_history = []
        for run in runs_to_process:
            if not (run.response and run.response.messages):
                continue
            for message in run.response.messages:
                if skip_role and message.role == skip_role:
                    continue
                if hasattr(message, 'from_history') and message.from_history:
                    continue
                messages_from_history.append(message)
        print(f'Getting messages from previous runs: {len(messages_from_history)}')
        return messages_from_history

    def get_all_messages(self) -> List[Tuple[Message, Message]]:
        assistant_role = ['assistant', 'model', 'CHATBOT']
        runs_as_message_pairs: List[Tuple[Message, Message]] = []
        for run in self.runs:
            if run.response and run.response.messages:
                user_message_from_run = None
                assistant_message_from_run = None
                for message in run.response.messages:
                    if message.role == 'user':
                        user_message_from_run = message
                        break
                for message in run.response.messages[::-1]:
                    if message.role in assistant_role:
                        assistant_message_from_run = message
                        break
                if user_message_from_run and assistant_message_from_run:
                    runs_as_message_pairs.append((user_message_from_run, assistant_message_from_run))
        return runs_as_message_pairs

    def load_user_memories(self) -> None:
        if self.db is None:
            return
        try:
            memory_rows = self.db.read_memories(user_id=self.user_id, limit=self.num_memories, sort='desc')
        except Exception as e:
            print(f'Error reading memory: {e}')
            return
        self.memories = []
        if memory_rows is None or len(memory_rows) == 0:
            return
        for row in memory_rows:
            try:
                self.memories.append(Memory.model_validate(row.memory))
            except Exception as e:
                print(f'Error loading memory: {e}')
                continue

    def should_update_memory(self, input: str) -> bool:
        if self.classifier is None:
            self.classifier = MemoryClassifier()
        self.classifier.existing_memories = self.memories
        classifier_response = self.classifier.run(input)
        if classifier_response == 'yes':
            return True
        return False

    async def ashould_update_memory(self, input: str) -> bool:
        if self.classifier is None:
            self.classifier = MemoryClassifier()
        self.classifier.existing_memories = self.memories
        classifier_response = await self.classifier.arun(input)
        if classifier_response == 'yes':
            return True
        return False

    def update_memory(self, input: str, force: bool = False) -> Optional[str]:
        if input is None or not isinstance(input, str):
            return 'Invalid message content'
        if self.db is None:
            print('MemoryDb not provided.')
            return 'Please provide a db to store memories'
        self.updating_memory = True
        should_update_memory = force or self.should_update_memory(input=input)
        print(f'Update memory: {should_update_memory}')
        if not should_update_memory:
            print('Memory update not required')
            return 'Memory update not required'
        if self.manager is None:
            self.manager = MemoryManager(user_id=self.user_id, db=self.db)
        else:
            self.manager.db = self.db
            self.manager.user_id = self.user_id
        response = self.manager.run(input)
        self.load_user_memories()
        self.updating_memory = False
        return response

    async def aupdate_memory(self, input: str, force: bool = False) -> Optional[str]:
        if input is None or not isinstance(input, str):
            return 'Invalid message content'
        if self.db is None:
            print('MemoryDb not provided.')
            return 'Please provide a db to store memories'
        self.updating_memory = True
        should_update_memory = force or await self.ashould_update_memory(input=input)
        print(f'Async update memory: {should_update_memory}')
        if not should_update_memory:
            print('Memory update not required')
            return 'Memory update not required'
        if self.manager is None:
            self.manager = MemoryManager(user_id=self.user_id, db=self.db)
        else:
            self.manager.db = self.db
            self.manager.user_id = self.user_id
        response = await self.manager.arun(input)
        self.load_user_memories()
        self.updating_memory = False
        return response

    def deep_copy(self) -> 'TeamMemory':
        copied_obj = self.__class__(**self.to_dict())
        for field_name, field_value in self.__dict__.items():
            if field_name not in ['db', 'classifier', 'manager']:
                try:
                    setattr(copied_obj, field_name, deepcopy(field_value))
                except Exception as e:
                    print(f'Failed to deepcopy field: {field_name} - {e}')
                    setattr(copied_obj, field_name, field_value)
        copied_obj.db = self.db
        copied_obj.classifier = self.classifier
        copied_obj.manager = self.manager
        return copied_obj


class WorkflowRun(Base):
    input: Dict[str, Any]
    response: RunResponse = None


class WorkflowMemory(Base):
    runs: List[WorkflowRun]

    def add_run(self, workflow_run: WorkflowRun) -> None:
        self.runs.append(workflow_run)
        print('已将WorkflowRun添加到WorkflowMemory')

    def clear(self) -> None:
        self.runs = []

    def deep_copy(self, *, update: Dict[str, Any] = None) -> 'WorkflowMemory':
        new_memory = self.model_copy(deep=True, update=update)
        new_memory.clear()
        return new_memory


class Embedder:
    def __init__(self, model='llama3.1:8b', host='', timeout=1000, options: Any = None, client_kwargs: dict = None):
        self.model = model
        self.options = options
        client_kwargs = client_kwargs or {}
        client_kwargs['host'] = host
        client_kwargs['timeout'] = timeout
        self.ollama_client = ollama.Client(**{k: v for k, v in client_kwargs.items() if v})

    def __call__(self, input: str) -> List[float]:
        kwargs = {'options': self.options} if self.options else {}
        response = self.ollama_client.embed(input=input, model=self.model, **kwargs)
        embedding = []
        if response and 'embeddings' in response:
            embeddings = response['embeddings']
            if isinstance(embeddings, list) and len(embeddings) > 0 and isinstance(embeddings[0], list):
                embedding = embeddings[0]
            elif isinstance(embeddings, list) and all(isinstance(x, (int, float)) for x in embeddings):
                embedding = embeddings
        return embedding


class Document:
    def __init__(self, content: str, id='', name='', meta_data: dict = None):
        self.id = id
        self.name = name
        self.content = content
        self.meta_data = meta_data or {}

    def embed(self, embedder: Embedder) -> List[float]:
        return embedder(self.content)

    def to_dict(self) -> Dict[str, Any]:
        return {'name': self.name, 'meta_data': self.meta_data, 'content': self.content}


class Reader:
    def __init__(self, chunk_size=5000, overlap=0, separators: List[str] = None):
        if overlap >= chunk_size:
            chunk_size, overlap = overlap, chunk_size
        self.overlap = overlap
        self.chunk_size = chunk_size
        self.separators = separators or ['\n\n', '\n', '\r', '\r\n', '\n\r', '\t', '。', ' ']

    def chunk_document(self, document: Document) -> List[Document]:
        if len(document.content) <= self.chunk_size:
            return [document]
        chunks: List[Document] = []
        start = 0
        chunk_meta_data = document.meta_data
        chunk_number = 1
        content = re.sub(r'\v+', '\v', re.sub(r'\f+', '\f', re.sub(r'\r+', '\r', re.sub(r'\t+', '\t', re.sub(r'\s+', ' ', document.content)))))
        while start < len(content):
            end = min(start + self.chunk_size, len(content))
            if end < len(content):
                for sep in self.separators:
                    last_sep = content[start:end].rfind(sep)
                    if last_sep != -1:
                        end = start + last_sep + 1
                        break
            chunk = content[start:end]
            meta_data = chunk_meta_data.copy()
            meta_data['chunk'] = chunk_number
            chunk_id = None
            if document.id:
                chunk_id = f'{document.id}_{chunk_number}'
            chunk_number += 1
            meta_data['chunk_size'] = len(chunk)
            chunks.append(Document(id=chunk_id, name=document.name, meta_data=meta_data, content=chunk))
            new_start = end - self.overlap
            if new_start <= start:
                new_start = min(len(content), start + max(1, self.chunk_size // 10))
            start = new_start
        return chunks

    def read(self, urls: list) -> Iterator[List[Document]]:
        for u in urls:
            if u.endswith('.csv'):
                yield self.read_csv(u)
            elif u.endswith(('.doc', '.docx')):
                yield self.read_docx(u)
            elif u.endswith('.pdf'):
                yield self.read_pdf(u)
            elif u.endswith(('.txt', '.md', '.py', '.js', '.vue')):
                yield self.read_text(u)
            elif u.endswith(('.html', '.xml')):
                yield self.read_url(u)
            elif u.startswith('http'):
                yield self.read_website(u, max_depth=3, max_links=10)

    def read_csv(self, url: str) -> List[Document]:
        if url.startswith('http'):
            response = requests.get(url)
            filename = url.split('/')[-1]
            file_obj = io.BytesIO(response.content)
            csv_reader = csv.reader(file_obj, delimiter=',', quotechar='"')
            file_obj.close()
        else:
            file = pathlib.Path(url)
            filename = pathlib.Path(file.name).stem
            with file.open(newline='', mode='r', encoding='utf-8') as csvfile:
                csv_reader = csv.reader(csvfile, delimiter=',', quotechar='"')
        documents = [Document(name=filename, id=filename, content='\n'.join([', '.join(row) for row in csv_reader]))]
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents

    def read_docx(self, url: str) -> List[Document]:
        docx_document = docx.Document(url)
        doc_name = pathlib.Path(url).stem
        doc_content = '\n\n'.join([para.text for para in docx_document.paragraphs])
        documents = [Document(name=doc_name, id=doc_name, content=doc_content)]
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents

    def read_pdf(self, url: str) -> List[Document]:
        doc_name = url.split('/')[-1].split('.')[0].replace(' ', '_')
        documents = []
        doc_reader = pypdf.PdfReader(io.BytesIO(requests.get(url).content) if url.startswith('http') else url)
        for page_number, page in enumerate(doc_reader.pages, start=1):
            documents.append(Document(name=doc_name, id=f'{doc_name}_{page_number}', meta_data={'url': url, 'page': page_number},
                                      content=page.extract_text()))
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents

    def read_text(self, url: str) -> List[Document]:
        with open(url, 'r', encoding='utf8') as f:
            file_contents = f.read()
            file_name = url.split('/')[-1]
            documents = [Document(name=file_name, id=file_name, content=file_contents)]
            return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents

    def read_url(self, url: str) -> List[Document]:
        response = requests.get(url)
        doc_name = url.split('//')[-1].strip('/').replace('/', '_').replace(' ', '_')
        document = Document(name=doc_name, id=doc_name, meta_data={'url': url}, content=response.text)
        return self.chunk_document(document) if self.chunk_size > 0 else [document]

    def read_website(self, url: str, max_depth=3, max_links=10):
        print(f'Reading: {url}')
        num_links = 0
        crawler_result: Dict[str, str] = {}
        primary_domain = '.'.join(url.split('/')[2].split('.')[1:])
        urls_to_crawl = [(url, 1)]
        visited = set()

        while urls_to_crawl:
            current_url, current_depth = urls_to_crawl.pop(0)
            if current_url in visited or not current_url.split('/')[2].endswith(
                    primary_domain) or current_depth > max_depth or num_links >= max_links:
                continue
            visited.add(current_url)
            time.sleep(random.uniform(1, 3))
            try:
                print(f'Crawling: {current_url}')
                response = requests.get(current_url, timeout=10)
                tree = etree.HTML(response.content)
                for tag in ['article', 'main', 'content', 'main-content', 'post-content']:
                    if tag in ['article', 'main']:
                        elements = tree.xpath(f'//{tag}')
                    else:
                        elements = tree.xpath(f'//*[@class="{tag}"]')
                    if elements:
                        element = elements[0]
                        crawler_result[current_url] = ''.join(element.itertext()).strip().replace('\n', ' ')
                        num_links += 1
                        break
                for link in tree.xpath('//a[@href]'):
                    href_str = link.get('href')
                    if href_str.startswith('http'):
                        full_url = href_str
                    else:
                        full_url = '/'.join([current_url.rstrip('/'), href_str.lstrip('.').lstrip('/')])
                    if full_url.split('/')[2].endswith(primary_domain) and not any(
                            full_url.endswith(ext) for ext in ['.pdf', '.jpg', '.png']):
                        if full_url not in visited and (full_url, current_depth + 1) not in urls_to_crawl:
                            urls_to_crawl.append((full_url, current_depth + 1))
            except Exception as e:
                print(f'Failed to crawl: {current_url}: {e}')
                pass
        documents = [
            Document(name=url, id=str(crawled_url), meta_data={'url': str(crawled_url)}, content=crawled_content) for
            crawled_url, crawled_content in crawler_result.items()]
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents


class Knowledge:
    """ 知识库 支持文档 csv,doc,docx,pdf,txt,html,website"""
    def __init__(self, database='local', path: str = 'static/chromadb', **kwargs):
        self.dbname: str = database
        self.embedder = Embedder()
        self.client = chromadb.PersistentClient(path=path, **kwargs)
        self.collection = self.client.get_or_create_collection(self.dbname,  metadata={'hnsw:space': 'cosine'}, embedding_function=self.embedder)

    def exists(self, doc: Union[str, Document]) -> bool:
        if isinstance(doc, Document):
            return doc.content.replace('\x00', '\ufffd') in self.collection.get().get('documents', [])
        return bool(self.collection.get(where={"url": doc})['ids'])

    def insert(self, documents: List[Document]):
        print(f'插入 {len(documents)} 个文档')
        ids: List = []
        docs: List = []
        docs_embeddings: List = []
        docs_metadata: List = []
        for document in documents:
            cleaned_content = document.content.replace('\x00', '\ufffd')
            doc_id = hashlib.md5(cleaned_content.encode()).hexdigest()
            docs_embeddings.append(document.embed(embedder=self.embedder))
            docs.append(cleaned_content)
            ids.append(doc_id)
            docs_metadata.append(document.meta_data)
            print(f'插入文档: {document.id} | {document.name} | {document.meta_data}')
        if len(docs) > 0:
            self.collection.add(ids=ids, embeddings=docs_embeddings, documents=docs, metadatas=docs_metadata)

    def upsert(self, documents: List[Document]):
        print(f'更新插入 {len(documents)} 个文档')
        ids: List = []
        docs: List = []
        docs_embeddings: List = []
        docs_metadata: List = []
        for document in documents:
            cleaned_content = document.content.replace('\x00', '\ufffd')
            doc_id = hashlib.md5(cleaned_content.encode()).hexdigest()
            docs_embeddings.append(document.embed(embedder=self.embedder))
            docs.append(cleaned_content)
            ids.append(doc_id)
            docs_metadata.append(document.meta_data)
            print(f'更新插入: {document.id} | {document.name} | {document.meta_data}')
        if len(docs) > 0:
            self.collection.upsert(ids=ids, embeddings=docs_embeddings, documents=docs, metadatas=docs_metadata)

    def search(self, query: str, limit: int = 5) -> List[Document]:
        result = self.collection.query(query_embeddings=self.embedder(query), n_results=limit, include=['metadatas', 'documents', 'embeddings', 'distances', 'uris'])
        search_results: List[Document] = []
        ids = result.get('ids', [[]])[0]
        metadata = result.get('metadatas', [{}])[0]
        documents = result.get('documents', [[]])[0]
        embeddings = result.get('embeddings')[0]
        embeddings = [e.tolist() if hasattr(e, 'tolist') else e for e in embeddings]
        distances = result.get('distances', [[]])[0]
        for idx, distance in enumerate(distances):
            metadata[idx]['distances'] = distance
        for idx, (id_, metadata, document) in enumerate(zip(ids, metadata, documents)):
            search_results.append(Document(id=id_, meta_data=metadata, content=document, embedding=embeddings[idx]))
        return search_results

    def delete(self):
        self.client.delete_collection(name=self.dbname)

    def load(self, docs: List[Union[Document, str]] = None, recreate=False, upsert=False, skip_existing=True):
        if recreate:
            self.delete()
        urls = [u for u in docs if isinstance(u, str) and not self.exists(u)]
        docs = [d for d in docs if isinstance(d, Document)] + list(Reader().read(urls))
        for documents in docs:
            if upsert:
                self.upsert(documents)
            else:
                self.insert([doc for doc in documents if not self.exists(doc)] if skip_existing else documents)


class AgentSession:
    def __init__(self, session_id: str, mode: Literal['agent', 'team', 'workflow'] = 'agent',  user_id: str = None, team_session_id: str = None, memory: Dict[str, Any] = None, session_data: Dict[str, Any] = None, extra_data: Dict[str, Any] = None, created_at: int = None, updated_at: int = None, agent_id: str = None, agent_data: Dict[str, Any] = None):
        self.mode = mode
        self.session_id = session_id
        self.user_id = user_id
        self.team_session_id = team_session_id
        self.memory = memory
        self.session_data = session_data
        self.extra_data = extra_data
        self.created_at = created_at
        self.updated_at = updated_at
        self.agent_id = agent_id
        self.agent_data = agent_data


class Storage:
    def __init__(self, table_name: str, schema='ai', db_url: str = None, db_engine: sqlalchemy.engine.Engine = None, auto_upgrade_schema=False):
        _engine: sqlalchemy.engine.Engine = db_engine
        if _engine is None and db_url:
            _engine = sqlalchemy.engine.create_engine(db_url, connect_args={'charset': 'utf8mb4'})
        if _engine is None:
            raise ValueError('Must provide either db_url or db_engine')
        self.table_name: str = table_name
        self.schema: str = schema
        self.db_url: str = db_url
        self.db_engine: sqlalchemy.engine.Engine = _engine
        self.metadata: sqlalchemy.schema.MetaData = sqlalchemy.schema.MetaData(schema=self.schema)
        self.auto_upgrade_schema: bool = auto_upgrade_schema
        self._schema_up_to_date: bool = False
        self.Session: sqlalchemy.orm.sessionmaker[sqlalchemy.orm.Session] = sqlalchemy.orm.sessionmaker(bind=self.db_engine)
        self.table: sqlalchemy.schema.Table = self.get_table()

    def get_table(self) -> sqlalchemy.schema.Table:
        from sqlalchemy.dialects import mysql
        common_columns = [
            sqlalchemy.schema.Column('session_id', mysql.TEXT, primary_key=True), sqlalchemy.schema.Column('user_id', mysql.TEXT),
            sqlalchemy.schema.Column('memory', mysql.JSON), sqlalchemy.schema.Column('session_data', mysql.JSON), sqlalchemy.schema.Column('extra_data', mysql.JSON),
            sqlalchemy.schema.Column('created_at', mysql.BIGINT), sqlalchemy.schema.Column('updated_at', mysql.BIGINT), ]
        specific_columns = [sqlalchemy.schema.Column('agent_id', mysql.TEXT), sqlalchemy.schema.Column('team_session_id', mysql.TEXT, nullable=True),
                            sqlalchemy.schema.Column('agent_data', mysql.JSON), ]
        table = sqlalchemy.schema.Table(self.table_name, self.metadata, *common_columns, *specific_columns, extend_existing=True,
                      schema=self.schema)
        return table

    def table_exists(self) -> bool:
        try:
            return sqlalchemy.inspection.inspect(self.db_engine).has_table(self.table.name, schema=self.schema)
        except Exception as e:
            print(e)

    def create(self) -> None:
        self.table = self.get_table()
        if not self.table_exists():
            print(f'Creating table: {self.table_name}\n')
            self.table.create(self.db_engine)

    def read(self, session_id: str, user_id: str = None) -> Optional[AgentSession]:
        with self.Session.begin() as sess:
            stmt = sqlalchemy.sql.expression.select(self.table).where(self.table.c.session_id == session_id)
            if user_id:
                stmt = stmt.where(self.table.c.user_id == user_id)
            try:
                existing_row = sess.execute(stmt).first()
                return AgentSession(**existing_row._mapping)
            except Exception as e:
                self.create()

    def get_all_sessions(self, user_id: str = None, entity_id: str = None) -> List[AgentSession]:
        sessions = []
        with self.Session.begin() as sess:
            stmt = sqlalchemy.sql.expression.select(self.table)
            if user_id:
                stmt = stmt.where(self.table.c.user_id == user_id)
            if entity_id:
                stmt = stmt.where(self.table.c.agent_id == entity_id)
            stmt = stmt.order_by(self.table.c.created_at.desc())
            rows = sess.execute(stmt).fetchall()
            for row in rows:
                if row.session_id:
                    _agent_session = AgentSession(**row._mapping)
                    if _agent_session:
                        sessions.append(_agent_session)
        return sessions

    def upgrade_schema(self) -> None:
        if not self.auto_upgrade_schema:
            print('Auto schema upgrade disabled. Skipping upgrade.')
            return
        try:
            if self.table_exists():
                with self.Session() as sess:
                    column_exists_query = sqlalchemy.sql.expression.text('''
                        SELECT 1 FROM information_schema.columns
                        WHERE table_schema = :schema AND table_name = :table
                        AND column_name = 'team_session_id'
                        ''')
                    column_exists = (sess.execute(column_exists_query, {'schema': self.schema, 'table': self.table_name}).scalar()
                       )
                    if not column_exists:
                        print(f'Adding "team_session_id" column to {self.schema}.{self.table_name}')
                        alter_table_query = sqlalchemy.sql.expression.text(f'ALTER TABLE {self.schema}.{self.table_name} ADD COLUMN team_session_id TEXT')
                        sess.execute(alter_table_query)
                        sess.commit()
                        self._schema_up_to_date = True
                        print('Schema upgrade completed successfully')
        except Exception as e:
            print(f'Error during schema upgrade: {e}')
            raise

    def upsert(self, session: AgentSession) -> Optional[AgentSession]:
        if self.auto_upgrade_schema and not self._schema_up_to_date:
            self.upgrade_schema()
        with self.Session.begin() as sess:
            upsert_sql = sqlalchemy.sql.expression.text(f'''
                INSERT INTO {self.schema}.{self.table_name}
                (session_id, agent_id, team_session_id, user_id, memory, agent_data, session_data, extra_data, created_at, updated_at)
                VALUES
                (:session_id, :agent_id, :team_session_id, :user_id, :memory, :agent_data, :session_data, :extra_data, UNIX_TIMESTAMP(), NULL)
                ON DUPLICATE KEY UPDATE
                    agent_id = VALUES(agent_id), team_session_id = VALUES(team_session_id), user_id = VALUES(user_id), memory = VALUES(memory), agent_data = VALUES(agent_data), session_data = VALUES(session_data), extra_data = VALUES(extra_data), updated_at = UNIX_TIMESTAMP();
                ''')
            try:
                sess.execute(upsert_sql, {'session_id': session.session_id, 'agent_id': session.agent_id, 'team_session_id': session.team_session_id, 'user_id': session.user_id, 'memory': json.dumps(session.memory, ensure_ascii=False)
                        if session.memory
                        else None, 'agent_data': json.dumps(session.agent_data, ensure_ascii=False)
                        if session.agent_data
                        else None, 'session_data': json.dumps(session.session_data, ensure_ascii=False)
                        if session.session_data
                        else None, 'extra_data': json.dumps(session.extra_data, ensure_ascii=False)
                        if session.extra_data
                        else None})
            except Exception as e:
                if not self.table_exists():
                    print(f'Table does not exist: {self.table.name}')
                    print('Creating table and retrying upsert')
                    self.create()
                    return self.upsert(session)
                else:
                    print(f'Exception upserting into table: {e}')
                    print('A table upgrade might be required, please review these docs for more information: https://agno.link/upgrade-schema')
                    return None
        return self.read(session_id=session.session_id)

    def delete_session(self, session_id: str = None):
        if session_id is None:
            print('No session_id provided for deletion.')
            return
        with self.Session() as sess, sess.begin():
            try:
                delete_stmt = self.table.delete().where(self.table.c.session_id == session_id)
                result = sess.execute(delete_stmt)
                if result.rowcount == 0:
                    print(f'No session found with session_id: {session_id}')
                else:
                    print(f'Successfully deleted session with session_id: {session_id}')
            except Exception as e:
                print(f'Error deleting session: {e}')
                raise

    def drop(self) -> None:
        if self.table_exists():
            print(f'Deleting table: {self.table_name}')
            self.table.drop(self.db_engine)

    def __deepcopy__(self, memo):
        cls = self.__class__
        copied_obj = cls.__new__(cls)
        memo[id(self)] = copied_obj
        for k, v in self.__dict__.items():
            if k in {'metadata', 'table'}:
                continue
            elif k in {'db_engine', 'Session'}:
                setattr(copied_obj, k, v)
            else:
                setattr(copied_obj, k, deepcopy(v, memo))
        copied_obj.metadata = sqlalchemy.schema.MetaData(schema=self.schema)
        copied_obj.table = copied_obj.get_table()
        return copied_obj


class Agent:
    def __init__(self, model: Model = None, name: str = None, agent_id: str = None, introduction: str = None,
                 user_id: str = None, session_id: str = None, session_name: str = None,
                 session_state: Dict[str, Any] = None, context: Dict[str, Any] = None, add_context=False,
                 resolve_context=True, memory: AgentMemory = None, add_history_to_messages=False,
                 num_history_responses: int = None, num_history_runs=3, knowledge: Knowledge = None,
                 add_references=False, retriever: Callable[..., List[Dict]] = None,
                 storage: Storage = None, extra_data: Dict[str, Any] = None, tools: List[Union[Toolkit, Callable, Function, Dict]] = None,
                 show_tool_calls=True, tool_call_limit: int = None, tool_choice: Union[str, Dict[str, Any]] = None, reasoning=False,
                 reasoning_model: Model = None, reasoning_agent: 'Agent' = None, reasoning_min_steps=1,
                 reasoning_max_steps=10, read_chat_history=False, search_knowledge=True, update_knowledge=False,
                 read_tool_call_history=False, system_message: Union[str, Callable, Message] = None, system_message_role: str = 'system',
                 create_default_system_message=True, description: str = None, goal: str = None, instructions: Union[str, List[str], Callable] = None,
                 expected_output: str = None, additional_context: str = None, markdown=False, add_name_to_instructions=False,
                 add_datetime_to_instructions=False, add_state_in_messages=False, add_messages: List[Union[Dict, Message]] = None,
                 user_message: Union[List, Dict, str, Callable, Message] = None, user_message_role='user', create_default_user_message=True,
                 retries=0, delay_between_retries=1, exponential_backoff=False, response_model: Type[Base] = None,
                 parse_response=True, structured_outputs=False, use_json_mode=False, save_response_to_file: str = None,
                 stream=False, stream_intermediate_steps=False, team: List['Agent'] = None, team_data: Dict[str, Any] = None,
                 role: str = None, respond_directly=False, add_transfer_instructions=True, team_response_separator='\n', debug_mode=False,
                 monitoring=False, telemetry=True):
        self.model = model
        self.name = name
        self.agent_id = agent_id
        self.introduction = introduction
        self.user_id = user_id
        self.session_id = session_id
        self.session_name = session_name
        self.session_state = session_state
        self.context = context
        self.add_context = add_context
        self.resolve_context = resolve_context
        self.memory = memory
        self.add_history_to_messages = add_history_to_messages
        self.num_history_responses = num_history_responses
        self.num_history_runs = num_history_responses or num_history_runs
        self.knowledge = knowledge
        self.add_references = add_references
        self.retriever = retriever
        self.storage = storage
        self.extra_data = extra_data
        self.tools = tools
        self.show_tool_calls = show_tool_calls
        self.tool_call_limit = tool_call_limit
        self.tool_choice = tool_choice
        self.reasoning = reasoning
        self.reasoning_model = reasoning_model
        self.reasoning_agent = reasoning_agent
        self.reasoning_min_steps = reasoning_min_steps
        self.reasoning_max_steps = reasoning_max_steps
        self.read_chat_history = read_chat_history
        self.search_knowledge = search_knowledge
        self.update_knowledge = update_knowledge
        self.read_tool_call_history = read_tool_call_history
        self.system_message = system_message
        self.system_message_role = system_message_role
        self.create_default_system_message = create_default_system_message
        self.description = description
        self.goal = goal
        self.instructions = instructions
        self.expected_output = expected_output
        self.additional_context = additional_context
        self.markdown = markdown
        self.add_name_to_instructions = add_name_to_instructions
        self.add_datetime_to_instructions = add_datetime_to_instructions
        self.add_state_in_messages = add_state_in_messages
        self.add_messages = add_messages
        self.user_message = user_message
        self.user_message_role = user_message_role
        self.create_default_user_message = create_default_user_message
        self.retries = retries
        self.delay_between_retries = delay_between_retries
        self.exponential_backoff = exponential_backoff
        self.response_model = response_model
        self.parse_response = parse_response
        self.structured_outputs = structured_outputs
        self.use_json_mode = use_json_mode
        self.save_response_to_file = save_response_to_file
        self.stream = stream
        self.stream_intermediate_steps = stream_intermediate_steps
        self.team = team
        self.team_data = team_data
        self.team_session_id: str = None
        self.role = role
        self.respond_directly = respond_directly
        self.add_transfer_instructions = add_transfer_instructions
        self.team_response_separator = team_response_separator
        self.debug_mode = debug_mode
        self.monitoring = monitoring
        self.telemetry = telemetry
        self.run_id: str = None
        self.run_input: Union[str, List, Dict, Message] = None
        self.run_messages: RunMessages = None
        self.run_response: RunResponse = None
        self.images: List[ImageArtifact] = None
        self.audio: List[AudioArtifact] = None
        self.videos: List[VideoArtifact] = None
        self.agent_session: AgentSession = None
        self._tools_for_model: List[Dict] = None
        self._functions_for_model: Dict[str, Function] = None
        if (self.reasoning or self.reasoning_model) and not self.reasoning_agent:
            self.reasoning_agent = Agent(model=reasoning_model or self.model.__class__(id=self.model.id),
                              description='你是一个细致、周到、有逻辑的推理代理，通过清晰、结构化、循序渐进的分析来解决复杂的问题',
                              instructions=textwrap.dedent(f'''
            步骤1-问题分析：\n-用你自己的话清楚地重述用户的任务，以确保完全理解。\n-明确指出需要哪些信息以及可能需要哪些工具或资源。
            第2步-分解和制定战略：\n-将问题分解为明确定义的子任务。\n-制定至少两种不同的策略或方法来解决问题，以确保彻底性。
            第3步-意图澄清和规划：\n-清楚地表达用户请求背后的意图。\n-从步骤2中选择最合适的策略，根据与用户意图和任务约束的一致性清楚地证明你的选择。\n-制定详细的分步行动计划，概述解决问题所需的行动顺序。
            步骤4-执行行动计划：
            对于每个计划步骤，记录：\n1.**标题**：概括步骤的简明标题。\n2.**行动**：以第一人称明确说明你的下一个行动（“我会……”）。\n3.**结果**：使用必要的工具执行行动，并提供结果的简明摘要。\n4.**推理**：清楚地解释你的理由，包括：
            -必要性：为什么需要采取这一行动。\n-注意事项：强调关键考虑因素、潜在挑战和缓解策略。\n-进展：这一步如何从逻辑上遵循或建立在之前的行动之上。\n-假设：明确说明所做的任何假设，并证明其有效性。
            5.**下一步行动**：从以下选项中明确选择下一步：
            -**继续**：如果需要进一步的步骤。\n-**验证**：当你得到一个潜在的答案时，表明它已经准备好进行验证。\n-**最终答案**：只有当您自信地验证了解决方案时。\n-**重置**：如果发现严重错误或不正确的结果，请立即重新开始分析。
            6.**置信度分数**：提供一个数字置信度分数（0.0-1.0），表明您对步骤的正确性及其结果的确定性。
            步骤5-验证（在最终确定答案之前必须进行）：
            -通过以下方式明确验证您的解决方案：\n-与替代方法进行交叉验证（在步骤2中开发）。\n-使用其他可用工具或方法独立确认准确性。\n-清楚地记录验证结果和所选验证方法背后的推理。\n-如果验证失败或出现差异，明确指出错误，重置分析，并相应地修改计划。
            第6步-提供最终答案：
            -一旦经过彻底验证并充满信心，就可以清晰简洁地交付您的解决方案。\n-简要重述你的答案如何满足用户的初衷并解决所述任务。
            一般操作指南：
            -确保您的分析保持不变：
            -**完成**：解决任务的所有要素。\n-**全面**：探索不同的观点并预测潜在的结果。\n-**逻辑**：保持所有步骤之间的连贯性。\n-**可操作**：提出明确可执行的步骤和行动。\n-**富有洞察力**：在适用的情况下提供创新和独特的视角。
            -始终通过立即重置或修改步骤来明确处理错误和失误。\n-严格遵守最小{self.reasoning_min_steps}和最大{self.reasoning_max_steps}步数，以确保有效的任务解决。
            -主动毫不犹豫地执行必要的工具，清楚地记录工具的使用情况。'''), tools=tools, show_tool_calls=False,
                              response_model=ReasoningSteps, use_json_mode=self.use_json_mode, monitoring=self.monitoring,
                              telemetry=self.telemetry, debug_mode=self.debug_mode)
            self.reasoning_agent.model.show_tool_calls = False

    def set_agent_id(self) -> str:
        if self.agent_id is None:
            self.agent_id = str(uuid.uuid4())
        print(f'Agent ID: {self.agent_id}')
        return self.agent_id

    def set_session_id(self) -> str:
        if self.session_id is None or self.session_id == '':
            self.session_id = str(uuid.uuid4())
        print(f'Session ID: {self.session_id}')
        return self.session_id

    def set_debug(self) -> None:
        if self.debug_mode or os.getenv('AGNO_DEBUG', 'false').lower() == 'true':
            self.debug_mode = True

    def set_monitoring(self) -> None:
        monitor_env = os.getenv('AGNO_MONITOR')
        if monitor_env:
            self.monitoring = monitor_env.lower() == 'true'
        telemetry_env = os.getenv('AGNO_TELEMETRY')
        if telemetry_env:
            self.telemetry = telemetry_env.lower() == 'true'

    def initialize_agent(self) -> None:
        self.set_debug()
        self.set_agent_id()
        self.set_session_id()
        if self.memory is None:
            self.memory = AgentMemory()

    @property
    def is_streamable(self) -> bool:
        return self.response_model is None

    @property
    def has_team(self) -> bool:
        return self.team and len(self.team) > 0

    def _run(self, message: Union[str, List, Dict, Message] = None, *, stream: bool = False, audio: Sequence[Audio] = None, images: Sequence[Image] = None, videos: Sequence[Video] = None, files: Sequence[File] = None, messages: Sequence[Union[Dict, Message]] = None, stream_intermediate_steps: bool = False, **kwargs: Any) -> Iterator[RunResponse]:
        """运行代理并生成RunResponse。
        步骤：
        1.让代理人做好运行准备
        2.更新模型并解析上下文
        3.从存储中读取现有会话
        4.准备运行信息
        5.如果启用推理，则说明任务的原因
        6.通过生成RunStarted事件来启动Run
        7.从模型生成响应（包括运行函数调用）
        8.更新RunResponse
        9.更新代理内存
        10.计算会话度量
        11.将会话保存到存储
        12.如果设置了Save_response_to_file，则将输出保存到文件
        """
        self.initialize_agent()
        self.stream = self.stream or (stream and self.is_streamable)
        self.stream_intermediate_steps = self.stream_intermediate_steps or (stream_intermediate_steps and self.stream)
        self.run_id = str(uuid.uuid4())
        self.run_response = RunResponse(run_id=self.run_id, session_id=self.session_id, agent_id=self.agent_id)
        print(f'Agent Run Start: {self.run_response.run_id}')
        self.update_model()
        self.run_response.model = self.model.id if self.model else None
        if self.context and self.resolve_context:
            self.resolve_run_context()
        self.read_from_storage()
        run_messages: RunMessages = self.get_run_messages(message=message, audio=audio, images=images, videos=videos, files=files, messages=messages, **kwargs)
        if len(run_messages.messages) == 0:
            print('没有消息要发送到模型.')
        self.run_messages = run_messages
        if self.reasoning or self.reasoning_model:
            reasoning_generator = self.reason(run_messages=run_messages)
            if self.stream:
                yield from reasoning_generator
            else:
                collections.deque(reasoning_generator, maxlen=0)
        index_of_last_user_message = len(run_messages.messages)
        if self.stream_intermediate_steps:
            yield self.create_run_response('Run started', event=RunEvent.run_started)
        model_response: ModelResponse
        if self.stream:
            model_response = ModelResponse()
            for model_response_chunk in self.model.response_stream(messages=run_messages.messages):
                if model_response_chunk.event == ResponseEvent.res.value:
                    if model_response_chunk.content:
                        model_response.content = (model_response.content or '') + model_response_chunk.content
                        self.run_response.content = model_response.content
                    if model_response_chunk.thinking:
                        model_response.thinking = (model_response.thinking or '') + model_response_chunk.thinking
                        self.run_response.thinking = model_response.thinking
                    if model_response_chunk.redacted_thinking:
                        model_response.redacted_thinking = (model_response.redacted_thinking or '') + model_response_chunk.redacted_thinking
                        self.run_response.thinking = model_response.redacted_thinking
                    if model_response_chunk.citations:
                        self.run_response.citations = model_response_chunk.citations
                    if model_response_chunk.content or model_response_chunk.thinking or model_response_chunk.redacted_thinking or model_response_chunk.citations:
                        yield self.create_run_response(content=model_response_chunk.content, thinking=model_response_chunk.thinking, redacted_thinking=model_response_chunk.redacted_thinking, citations=model_response_chunk.citations, created_at=model_response_chunk.created_at)
                    if model_response_chunk.audio:
                        if model_response.audio is None:
                            model_response.audio = AudioResponse(id=str(uuid.uuid4()), content='', transcript='')
                        if model_response_chunk.audio.id:
                            model_response.audio.id = model_response_chunk.audio.id
                        if model_response_chunk.audio.content:
                            model_response.audio.content += model_response_chunk.audio.content
                        if model_response_chunk.audio.transcript:
                            model_response.audio.transcript += model_response_chunk.audio.transcript
                        if model_response_chunk.audio.expires_at:
                            model_response.audio.expires_at = model_response_chunk.audio.expires_at
                        if model_response_chunk.audio.mime_type:
                            model_response.audio.mime_type = model_response_chunk.audio.mime_type
                        model_response.audio.sample_rate = model_response_chunk.audio.sample_rate
                        model_response.audio.channels = model_response_chunk.audio.channels
                        self.run_response.response_audio = AudioResponse(id=model_response_chunk.audio.id, content=model_response_chunk.audio.content, transcript=model_response_chunk.audio.transcript, sample_rate=model_response_chunk.audio.sample_rate, channels=model_response_chunk.audio.channels)
                        self.run_response.created_at = model_response_chunk.created_at
                        yield self.run_response
                    if model_response_chunk.image:
                        self.add_image(model_response_chunk.image)
                        yield self.run_response
                elif model_response_chunk.event == ResponseEvent.start.value:
                    tool_calls_list = model_response_chunk.tool_calls
                    if tool_calls_list:
                        if self.run_response.tools is None:
                            self.run_response.tools = tool_calls_list
                        else:
                            self.run_response.tools.extend(tool_calls_list)
                        self.run_response.formatted_tool_calls = format_tool_calls(self.run_response.tools)
                    if self.stream_intermediate_steps:
                        yield self.create_run_response(content=model_response_chunk.content, event=RunEvent.tool_call_started)
                elif model_response_chunk.event == ResponseEvent.end.value:
                    tool_calls_list = model_response_chunk.tool_calls
                    if tool_calls_list:
                        if self.run_response.tools:
                            tool_call_index_map = {tc['tool_call_id']: i for i, tc in enumerate(self.run_response.tools) if tc.get('tool_call_id')}
                            for tool_call_dict in tool_calls_list:
                                tool_call_id = tool_call_dict.get('tool_call_id')
                                index = tool_call_index_map.get(tool_call_id)
                                if index:
                                    self.run_response.tools[index] = tool_call_dict
                        else:
                            self.run_response.tools = tool_calls_list
                    if self.stream_intermediate_steps:
                        yield self.create_run_response(content=model_response_chunk.content, event=RunEvent.tool_call_completed)
        else:
            model_response = self.model.response(messages=run_messages.messages)
            if model_response.tool_calls:
                self.run_response.formatted_tool_calls = format_tool_calls(model_response.tool_calls)
            if self.response_model and model_response.parsed:
                if self.model.structured_outputs:
                    self.run_response.content = model_response.parsed
                    self.run_response.content_type = self.response_model.__name__
            else:
                self.run_response.content = model_response.content
            if model_response.thinking:
                self.run_response.thinking = model_response.thinking
            if model_response.redacted_thinking:
                if self.run_response.thinking is None:
                    self.run_response.thinking = model_response.redacted_thinking
                else:
                    self.run_response.thinking += model_response.redacted_thinking
            if model_response.citations:
                self.run_response.citations = model_response.citations
            if model_response.tool_calls:
                if self.run_response.tools is None:
                    self.run_response.tools = model_response.tool_calls
                else:
                    self.run_response.tools.extend(model_response.tool_calls)
            if model_response.audio:
                self.run_response.response_audio = model_response.audio
            if model_response.image:
                self.add_image(model_response.image)
            self.run_response.messages = run_messages.messages
            self.run_response.created_at = model_response.created_at
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        self.run_response.messages = messages_for_run_response
        self.run_response.metrics = self.aggregate_metrics_from_messages(messages_for_run_response)
        if self.stream and model_response.audio:
            self.run_response.response_audio = model_response.audio
        if run_messages.system_message:
            self.memory.add_system_message(run_messages.system_message, system_message_role=self.system_message_role)
        messages_for_memory: List[Message] = ([run_messages.user_message] if run_messages.user_message else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        if self.stream_intermediate_steps:
            yield self.create_run_response(content='Memory updated', event=RunEvent.updating_memory)
        agent_run = AgentRun(response=self.run_response)
        agent_run.message = run_messages.user_message
        if self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message:
            self.memory.update_memory(input=run_messages.user_message.get_content_string())
        if messages and len(messages) > 0:
            for _im in messages:
                mp = None
                if isinstance(_im, Message):
                    mp = _im
                elif isinstance(_im, dict):
                    try:
                        mp = Message(**_im)
                    except Exception as e:
                        print(f'消息验证失败: {e}')
                else:
                    print(f'不支持的消息类型: {type(_im)}')
                    continue
                if mp:
                    if agent_run.messages is None:
                        agent_run.messages = []
                    agent_run.messages.append(mp)
                    if self.memory.create_user_memories and self.memory.update_user_memories_after_run:
                        self.memory.update_memory(input=mp.get_content_string())
                else:
                    print('无法将消息添加到内存中')
        self.memory.add_run(agent_run)
        if self.memory.create_session_summary and self.memory.update_session_summary_after_run:
            self.memory.update_summary()
        self.write_to_storage()
        self.save_run_response_to_file(message=message)
        if message:
            if isinstance(message, str):
                self.run_input = message
            elif isinstance(message, Message):
                self.run_input = message.to_dict()
            else:
                self.run_input = message
        elif messages:
            self.run_input = [m.to_dict() if isinstance(m, Message) else m for m in messages]
        self.set_monitoring()
        if self.telemetry or self.monitoring:
            try:
                run_data = self._create_run_data()
                agent_session: AgentSession = self.agent_session or self.get_agent_session()
            except Exception as e:
                print(f'Could not create agent event: {e}')
        print(f'Agent Run End: {self.run_response.run_id}')
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=self.run_response.content, event=RunEvent.run_completed)
        if not self.stream:
            yield self.run_response

    def run(self, message: Union[str, List, Dict, Message] = None, *, stream: bool = None, audio: Sequence[Audio] = None, images: Sequence[Image] = None, videos: Sequence[Video] = None, files: Sequence[File] = None, messages: Sequence[Union[Dict, Message]] = None, stream_intermediate_steps: bool = False, retries: int = None, **kwargs: Any) -> Union[RunResponse, Iterator[RunResponse]]:
        if retries is None:
            retries = self.retries
        if stream is None:
            stream = False if self.stream is None else self.stream
        last_exception = None
        num_attempts = retries + 1
        for attempt in range(num_attempts):
            if self.response_model and self.parse_response:
                print('Setting stream=False as response_model is set')
                self.stream = False
                run_response: RunResponse = next(self._run(message=message, stream=False, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs))
                if isinstance(run_response.content, self.response_model):
                    return run_response
                if isinstance(run_response.content, str):
                    try:
                        structured_output = parse_response_model_str(run_response.content, self.response_model)
                        if structured_output:
                            run_response.content = structured_output
                            run_response.content_type = self.response_model.__name__
                            if self.run_response:
                                self.run_response.content = structured_output
                                self.run_response.content_type = self.response_model.__name__
                        else:
                            print('无法将响应转换为response_model')
                    except Exception as e:
                        print(f'无法将响应转换为输出模型: {e}')
                else:
                    print('出了点问题。运行响应内容不是字符串')
                return run_response
            else:
                if stream and self.is_streamable:
                    resp = self._run(message=message, stream=True, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs)
                    return resp
                else:
                    resp = self._run(message=message, stream=False, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs)
                    return next(resp)
        if last_exception:
            print(f'尝试{num_attempts}次后失败。上次错误使用 {last_exception.model_name}({last_exception.model_id})')
            raise last_exception
        else:
            raise Exception(f'{num_attempts}次后失败')

    async def _arun(self, message: Union[str, List, Dict, Message] = None, *, stream: bool = False, audio: Sequence[Audio] = None, images: Sequence[Image] = None, videos: Sequence[Video] = None, files: Sequence[File] = None, messages: Sequence[Union[Dict, Message]] = None, stream_intermediate_steps: bool = False, **kwargs: Any) -> AsyncIterator[RunResponse]:
        self.initialize_agent()
        self.stream = self.stream or (stream and self.is_streamable)
        self.stream_intermediate_steps = self.stream_intermediate_steps or (stream_intermediate_steps and self.stream)
        self.run_id = str(uuid.uuid4())
        self.run_response = RunResponse(run_id=self.run_id, session_id=self.session_id, agent_id=self.agent_id)
        print(f'Async Agent Run Start: {self.run_response.run_id}')
        self.update_model(async_mode=True)
        self.run_response.model = self.model.id if self.model else None
        if self.context and self.resolve_context:
            self.resolve_run_context()
        self.read_from_storage()
        run_messages: RunMessages = self.get_run_messages(message=message, audio=audio, images=images, videos=videos, files=files, messages=messages, **kwargs)
        if len(run_messages.messages) == 0:
            print('没有消息要发送到模型。')
        self.run_messages = run_messages
        if self.reasoning or self.reasoning_model:
            areason_generator = self.areason(run_messages=run_messages)
            if self.stream:
                async for item in areason_generator:
                    yield item
            else:
                async for _ in areason_generator:
                    pass
        index_of_last_user_message = len(run_messages.messages)
        if self.stream_intermediate_steps:
            yield self.create_run_response('Run started', event=RunEvent.run_started)
        model_response: ModelResponse
        if stream and self.is_streamable:
            model_response = ModelResponse(content='')
            model_response_stream = self.model.aresponse_stream(messages=run_messages.messages)
            async for model_response_chunk in model_response_stream:
                if model_response_chunk.event == ResponseEvent.res.value:
                    if model_response_chunk.content:
                        model_response.content = (model_response.content or '') + model_response_chunk.content
                        self.run_response.content = model_response.content
                    if model_response_chunk.thinking:
                        model_response.thinking = (model_response.thinking or '') + model_response_chunk.thinking
                        self.run_response.thinking = model_response.thinking
                    if model_response_chunk.redacted_thinking:
                        model_response.redacted_thinking = (model_response.redacted_thinking or '') + model_response_chunk.redacted_thinking
                        self.run_response.thinking = model_response.redacted_thinking
                    if model_response_chunk.citations:
                        self.run_response.citations = model_response_chunk.citations
                    if model_response_chunk.content or model_response_chunk.thinking or model_response_chunk.redacted_thinking or model_response_chunk.citations:
                        yield self.create_run_response(content=model_response_chunk.content, thinking=model_response_chunk.thinking, redacted_thinking=model_response_chunk.redacted_thinking, citations=model_response_chunk.citations, created_at=model_response_chunk.created_at)
                    if model_response_chunk.audio:
                        if model_response.audio is None:
                            model_response.audio = AudioResponse(id=str(uuid.uuid4()), content='', transcript='')
                        if model_response_chunk.audio.id:
                            model_response.audio.id = model_response_chunk.audio.id
                        if model_response_chunk.audio.content:
                            model_response.audio.content += model_response_chunk.audio.content
                        if model_response_chunk.audio.transcript:
                            model_response.audio.transcript += model_response_chunk.audio.transcript
                        if model_response_chunk.audio.expires_at:
                            model_response.audio.expires_at = model_response_chunk.audio.expires_at
                        if model_response_chunk.audio.mime_type:
                            model_response.audio.mime_type = model_response_chunk.audio.mime_type
                        model_response.audio.sample_rate = model_response_chunk.audio.sample_rate
                        model_response.audio.channels = model_response_chunk.audio.channels
                        self.run_response.response_audio = AudioResponse(id=model_response_chunk.audio.id, content=model_response_chunk.audio.content, transcript=model_response_chunk.audio.transcript, sample_rate=model_response_chunk.audio.sample_rate, channels=model_response_chunk.audio.channels)
                        self.run_response.created_at = model_response_chunk.created_at
                        yield self.run_response
                    if model_response_chunk.image:
                        self.add_image(model_response_chunk.image)
                        yield self.run_response
                elif model_response_chunk.event == ResponseEvent.start.value:
                    tool_calls_list = model_response_chunk.tool_calls
                    if tool_calls_list:
                        if self.run_response.tools is None:
                            self.run_response.tools = tool_calls_list
                        else:
                            self.run_response.tools.extend(tool_calls_list)
                        self.run_response.formatted_tool_calls = format_tool_calls(self.run_response.tools)
                    if self.stream_intermediate_steps:
                        yield self.create_run_response(content=model_response_chunk.content, event=RunEvent.tool_call_started)
                elif model_response_chunk.event == ResponseEvent.end.value:
                    tool_calls_list = model_response_chunk.tool_calls
                    if tool_calls_list:
                        if self.run_response.tools:
                            tool_call_index_map = {tc['tool_call_id']: i
                                for i, tc in enumerate(self.run_response.tools)
                                if tc.get('tool_call_id')}
                            for tool_call_dict in tool_calls_list:
                                tool_call_id = (tool_call_dict['tool_call_id'] if 'tool_call_id' in tool_call_dict else None)
                                index = tool_call_index_map.get(tool_call_id)
                                if index:
                                    self.run_response.tools[index] = tool_call_dict
                        else:
                            self.run_response.tools = tool_calls_list
                    if self.stream_intermediate_steps:
                        yield self.create_run_response(content=model_response_chunk.content, event=RunEvent.tool_call_completed)
        else:
            model_response = await self.model.aresponse(messages=run_messages.messages)
            if model_response.tool_calls:
                self.run_response.formatted_tool_calls = format_tool_calls(model_response.tool_calls)
            if self.response_model and model_response.parsed:
                if self.model.structured_outputs:
                    self.run_response.content = model_response.parsed
                    self.run_response.content_type = self.response_model.__name__
            else:
                self.run_response.content = model_response.content
            if model_response.thinking:
                self.run_response.thinking = model_response.thinking
            if model_response.redacted_thinking:
                if self.run_response.thinking is None:
                    self.run_response.thinking = model_response.redacted_thinking
                else:
                    self.run_response.thinking += model_response.redacted_thinking
            if model_response.citations:
                self.run_response.citations = model_response.citations
            if model_response.tool_calls:
                if self.run_response.tools is None:
                    self.run_response.tools = model_response.tool_calls
                else:
                    self.run_response.tools.extend(model_response.tool_calls)
            if model_response.audio:
                self.run_response.response_audio = model_response.audio
            if model_response.image:
                self.add_image(model_response.image)
            self.run_response.messages = run_messages.messages
            self.run_response.created_at = model_response.created_at
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        self.run_response.messages = messages_for_run_response
        self.run_response.metrics = self.aggregate_metrics_from_messages(messages_for_run_response)
        if self.stream and model_response.audio:
            self.run_response.response_audio = model_response.audio
        if run_messages.system_message:
            self.memory.add_system_message(run_messages.system_message, system_message_role=self.system_message_role)
        messages_for_memory: List[Message] = ([run_messages.user_message] if run_messages.user_message else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        if self.stream_intermediate_steps:
            yield self.create_run_response(content='Memory updated', event=RunEvent.updating_memory)
        agent_run = AgentRun(response=self.run_response)
        agent_run.message = run_messages.user_message
        if self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message:
            await self.memory.aupdate_memory(input=run_messages.user_message.get_content_string())
        if messages and len(messages) > 0:
            for _im in messages:
                mp = None
                if isinstance(_im, Message):
                    mp = _im
                elif isinstance(_im, dict):
                    try:
                        mp = Message(**_im)
                    except Exception as e:
                        print(f'验证消息失败: {e}')
                else:
                    print(f'不支持的消息类型: {type(_im)}')
                    continue
                if mp:
                    if agent_run.messages is None:
                        agent_run.messages = []
                    agent_run.messages.append(mp)
                    if self.memory.create_user_memories and self.memory.update_user_memories_after_run:
                        await self.memory.aupdate_memory(input=mp.get_content_string())
                else:
                    print('无法将消息添加到内存中')
        self.memory.add_run(agent_run)
        if self.memory.create_session_summary and self.memory.update_session_summary_after_run:
            await self.memory.aupdate_summary()
        self.write_to_storage()
        self.save_run_response_to_file(message=message)
        if message:
            if isinstance(message, str):
                self.run_input = message
            elif isinstance(message, Message):
                self.run_input = message.to_dict()
            else:
                self.run_input = message
        elif messages:
            self.run_input = [m.to_dict() if isinstance(m, Message) else m for m in messages]
        self.set_monitoring()
        if self.telemetry or self.monitoring:
            try:
                run_data = self._create_run_data()
                agent_session: AgentSession = self.agent_session or self.get_agent_session()
            except Exception as e:
                print(f'Could not create agent event: {e}')

        print(f'Agent Run End: {self.run_response.run_id}')
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=self.run_response.content, event=RunEvent.run_completed)
        if not self.stream:
            yield self.run_response

    async def arun(self, message: Union[str, List, Dict, Message] = None, *, stream: bool = None, audio: Sequence[Audio] = None, images: Sequence[Image] = None, videos: Sequence[Video] = None, files: Sequence[File] = None, messages: Sequence[Union[Dict, Message]] = None, stream_intermediate_steps: bool = False, retries: int = None, **kwargs: Any) -> Any:
        if retries is None:
            retries = self.retries
        if stream is None:
            stream = False if self.stream is None else self.stream
        last_exception = None
        num_attempts = retries + 1
        for attempt in range(num_attempts):
            print(f'Attempt {attempt + 1}/{num_attempts}')
            if self.response_model and self.parse_response:
                print('Setting stream=False as response_model is set')
                run_response = await self._arun(message=message, stream=False, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs).__anext__()
                if isinstance(run_response.content, self.response_model):
                    return run_response
                if isinstance(run_response.content, str):
                    try:
                        structured_output = parse_response_model_str(run_response.content, self.response_model)
                        if structured_output:
                            run_response.content = structured_output
                            run_response.content_type = self.response_model.__name__
                            if self.run_response:
                                self.run_response.content = structured_output
                                self.run_response.content_type = self.response_model.__name__
                        else:
                            print('无法将响应转换为response_model')
                    except Exception as e:
                        print(f'无法将响应转换为输出模型: {e}')
                else:
                    print('出了点问题。运行响应内容不是字符串')
                return run_response
            else:
                if stream and self.is_streamable:
                    resp = self._arun(message=message, stream=True, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs)
                    return resp
                else:
                    resp = self._arun(message=message, stream=False, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs)
                    return await resp.__anext__()
        if last_exception:
            print(f'Failed after {num_attempts} attempts. Last error using {last_exception.model_name}({last_exception.model_id})')
            raise last_exception
        else:
            raise Exception(f'{num_attempts}次后失败')

    def create_run_response(self, content: Any = None, *, thinking: str = None, redacted_thinking: str = None, event: RunEvent = RunEvent.run_response, content_type: str = None, created_at: int = None, citations: Citations = None) -> RunResponse:
        thinking_combined = (thinking or '') + (redacted_thinking or '')
        rr = RunResponse(run_id=self.run_id, session_id=self.session_id, agent_id=self.agent_id, content=content, thinking=thinking_combined if thinking_combined else None, tools=self.run_response.tools, audio=self.run_response.audio, images=self.run_response.images, videos=self.run_response.videos, citations=citations or self.run_response.citations, response_audio=self.run_response.response_audio, model=self.run_response.model, messages=self.run_response.messages, extra_data=self.run_response.extra_data, event=event.value)
        if content_type:
            rr.content_type = content_type
        if created_at:
            rr.created_at = created_at
        return rr

    def get_tools(self, async_mode: bool = False) -> Optional[List[Union[Toolkit, Callable, Function, Dict]]]:
        agent_tools: List[Union[Toolkit, Callable, Function, Dict]] = []
        if self.tools:
            for tool in self.tools:
                agent_tools.append(tool)
        if self.read_chat_history:
            agent_tools.append(self.get_chat_history)
        if self.read_tool_call_history:
            agent_tools.append(self.get_tool_call_history)
        if self.memory and self.memory.create_user_memories:
            agent_tools.append(self.update_memory)
        if self.knowledge or self.retriever:
            if self.search_knowledge:
                if async_mode:
                    agent_tools.append(self.async_search_knowledge_base)
                else:
                    agent_tools.append(self.search_knowledge_base)
            if self.update_knowledge:
                agent_tools.append(self.add_to_knowledge)
        if self.has_team and self.team:
            for agent_index, agent in enumerate(self.team):
                agent_tools.append(self.get_transfer_function(agent, agent_index))
        return agent_tools

    def add_tools_to_model(self, model: Model, async_mode: bool = False) -> None:
        if self._functions_for_model is None or self._tools_for_model is None:
            agent_tools = self.get_tools(async_mode=async_mode)
            if agent_tools and len(agent_tools) > 0:
                print('Processing tools for model')
                strict = False
                if self.response_model and (self.structured_outputs or (not self.use_json_mode)) and model.supports_native_structured_outputs:
                    strict = True
                self._tools_for_model = []
                self._functions_for_model = {}
                for tool in agent_tools:
                    if isinstance(tool, Dict):
                        self._tools_for_model.append(tool)
                        print(f'Included builtin tool {tool}')
                    elif isinstance(tool, Toolkit):
                        for name, func in tool.functions.items():
                            if name not in self._functions_for_model:
                                func._agent = self
                                func.process_entrypoint(strict=strict)
                                if strict:
                                    func.strict = True
                                self._functions_for_model[name] = func
                                self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                                print(f'Included function {name} from {tool.name}')
                    elif isinstance(tool, Function):
                        if tool.name not in self._functions_for_model:
                            tool._agent = self
                            tool.process_entrypoint(strict=strict)
                            if strict and tool.strict is None:
                                tool.strict = True
                            self._functions_for_model[tool.name] = tool
                            self._tools_for_model.append({'type': 'function', 'function': tool.to_dict()})
                            print(f'Included function {tool.name}')
                    elif callable(tool):
                        try:
                            function_name = tool.__name__
                            if function_name not in self._functions_for_model:
                                func = Function.from_callable(tool, strict=strict)
                                func._agent = self
                                if strict:
                                    func.strict = True
                                self._functions_for_model[func.name] = func
                                self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                                print(f'Included function {func.name}')
                        except Exception as e:
                            print(f'Could not add function {tool}: {e}')
                model.set_tools(tools=self._tools_for_model)
                model.set_functions(functions=self._functions_for_model)

    def update_model(self, async_mode: bool = False) -> None:
        if self.model is None:
            self.model = Ollama()
        if self.response_model is None:
            self.model.response_format = None
        else:
            json_response_format = {'type': 'json_object'}
            if self.model.supports_native_structured_outputs:
                if (not self.use_json_mode) or self.structured_outputs:
                    print('将Model.response_format设置为Agent.response_Model')
                    self.model.response_format = self.response_model
                    self.model.structured_outputs = True
                else:
                    print('模型支持本机结构化输出，但未启用。改用JSON模式。')
                    self.model.response_format = json_response_format
                    self.model.structured_outputs = False
            elif self.model.supports_json_schema_outputs:
                if self.use_json_mode or not self.structured_outputs:
                    print('将Model.response_format设置为JSON响应模式')
                    self.model.response_format = {'type': 'json_schema', 'json_schema': {'name': self.response_model.__name__, 'schema': self.response_model.model_json_schema()}}
                else:
                    self.model.response_format = None
                self.model.structured_outputs = False
            else:
                print('模型不支持结构化或JSON模式输出')
                self.model.response_format = (json_response_format if (self.use_json_mode or not self.structured_outputs) else None)
                self.model.structured_outputs = False
        self.add_tools_to_model(model=self.model, async_mode=async_mode)
        if self.show_tool_calls:
            self.model.show_tool_calls = self.show_tool_calls
        if self.tool_choice:
            self.model.tool_choice = self.tool_choice
        if self.tool_call_limit:
            self.model.tool_call_limit = self.tool_call_limit

    def resolve_run_context(self) -> None:
        print('Resolving context')
        if self.context:
            if isinstance(self.context, dict):
                for ctx_key, ctx_value in self.context.items():
                    if callable(ctx_value):
                        try:
                            sig = inspect.signature(ctx_value)
                            if 'agent' in sig.parameters:
                                resolved_ctx_value = ctx_value(agent=self)
                            else:
                                resolved_ctx_value = ctx_value()
                            if resolved_ctx_value:
                                self.context[ctx_key] = resolved_ctx_value
                        except Exception as e:
                            print(f'Failed to resolve context for {ctx_key}: {e}')
                    else:
                        self.context[ctx_key] = ctx_value
            else:
                print('Context is not a dict')

    def load_user_memories(self) -> None:
        if self.memory and self.memory.create_user_memories:
            if self.user_id and self.memory.user_id is None:
                self.memory.user_id = self.user_id
            self.memory.load_user_memories()
            if self.user_id:
                print(f'Memories loaded for user: {self.user_id}')
            else:
                print('Memories loaded')

    def get_agent_data(self) -> Dict[str, Any]:
        agent_data: Dict[str, Any] = {}
        if self.name:
            agent_data['name'] = self.name
        if self.agent_id:
            agent_data['agent_id'] = self.agent_id
        if self.model:
            agent_data['model'] = self.model.to_dict()
        return agent_data

    def get_session_data(self) -> Dict[str, Any]:
        session_data: Dict[str, Any] = {}
        if self.session_name:
            session_data['session_name'] = self.session_name
        if self.session_state and len(self.session_state) > 0:
            session_data['session_state'] = self.session_state
        if self.team_data:
            session_data['team_data'] = self.team_data
        if self.images:
            session_data['images'] = [img.to_dict() for img in self.images]
        if self.videos:
            session_data['videos'] = [vid.to_dict() for vid in self.videos]
        if self.audio:
            session_data['audio'] = [aud.to_dict() for aud in self.audio]
        return session_data

    def get_agent_session(self) -> AgentSession:
        return AgentSession(session_id=self.session_id, agent_id=self.agent_id, user_id=self.user_id, team_session_id=self.team_session_id, memory=self.memory.to_dict() if self.memory else None, agent_data=self.get_agent_data(), session_data=self.get_session_data(), extra_data=self.extra_data, created_at=int(time.time()))

    def load_agent_session(self, session: AgentSession):
        if self.agent_id is None and session.agent_id:
            self.agent_id = session.agent_id
        if self.user_id is None and session.user_id:
            self.user_id = session.user_id
        if self.session_id is None and session.session_id:
            self.session_id = session.session_id
        if session.agent_data:
            if self.name is None and 'name' in session.agent_data:
                self.name = session.agent_data.get('name')
        if session.session_data:
            if self.session_name is None and 'session_name' in session.session_data:
                self.session_name = session.session_data.get('session_name')
            if 'session_state' in session.session_data:
                session_state_from_db = session.session_data.get('session_state')
                if session_state_from_db and isinstance(session_state_from_db, dict) and len(session_state_from_db) > 0:
                    if self.session_state and len(self.session_state) > 0:
                        merge_dictionaries(session_state_from_db, self.session_state)
                    self.session_state = session_state_from_db
            if 'images' in session.session_data:
                images_from_db = session.session_data.get('images')
                if images_from_db and isinstance(images_from_db, list):
                    if self.images is None:
                        self.images = []
                    self.images.extend([ImageArtifact.model_validate(img) for img in images_from_db])
            if 'videos' in session.session_data:
                videos_from_db = session.session_data.get('videos')
                if videos_from_db and isinstance(videos_from_db, list):
                    if self.videos is None:
                        self.videos = []
                    self.videos.extend([VideoArtifact.model_validate(vid) for vid in videos_from_db])
            if 'audio' in session.session_data:
                audio_from_db = session.session_data.get('audio')
                if audio_from_db and isinstance(audio_from_db, list):
                    if self.audio is None:
                        self.audio = []
                    self.audio.extend([AudioArtifact.model_validate(aud) for aud in audio_from_db])
        if session.extra_data:
            if self.extra_data:
                merge_dictionaries(session.extra_data, self.extra_data)
            self.extra_data = session.extra_data
        if self.memory is None:
            self.memory = session.memory
        if not isinstance(self.memory, AgentMemory):
            if isinstance(self.memory, dict):
                self.memory = AgentMemory(**self.memory)
            else:
                raise TypeError(f'预期内存为dict或AgentMemory，但实际得到{type(self.memory)}')
        if session.memory:
            try:
                if 'runs' in session.memory:
                    try:
                        self.memory.runs = [AgentRun.model_validate(m) for m in session.memory['runs']]
                    except Exception as e:
                        print(f'无法从内存中加载运行: {e}')
                if 'messages' in session.memory:
                    try:
                        self.memory.messages = [Message.model_validate(m) for m in session.memory['messages']]
                    except Exception as e:
                        print(f'无法从内存中加载消息: {e}')
                if 'summary' in session.memory:
                    try:
                        self.memory.summary = SessionSummary.model_validate(session.memory['summary'])
                    except Exception as e:
                        print(f'无法从内存中加载会话摘要: {e}')
                if 'memories' in session.memory:
                    try:
                        self.memory.memories = [Memory.model_validate(m) for m in session.memory['memories']]
                    except Exception as e:
                        print(f'加载用户内存失败: {e}')
            except Exception as e:
                print(f'未能加载AgentMemory: {e}')
        print(f'-*- AgentSession loaded: {session.session_id}')

    def read_from_storage(self) -> Optional[AgentSession]:
        if self.storage and self.session_id:
            self.agent_session = self.storage.read(session_id=self.session_id)
            if self.agent_session:
                self.load_agent_session(session=self.agent_session)
            self.load_user_memories()
        return self.agent_session

    def write_to_storage(self) -> Optional[AgentSession]:
        if self.storage:
            self.agent_session = self.storage.upsert(session=self.get_agent_session())
        return self.agent_session

    def add_introduction(self, introduction: str) -> None:
        if introduction:
            if len(self.memory.runs) == 0:
                self.memory.add_run(AgentRun(response=RunResponse(content=introduction, messages=[Message(role=self.model.assistant_message_role, content=introduction)])))

    def load_session(self, force: bool = False) -> Optional[str]:
        if self.agent_session and not force:
            if self.session_id and self.agent_session.session_id == self.session_id:
                return self.agent_session.session_id
        if self.storage:
            print(f'Reading AgentSession: {self.session_id}')
            self.read_from_storage()
            if self.agent_session is None:
                print('-*- Creating new AgentSession')
                if self.agent_id is None or self.session_id is None:
                    self.initialize_agent()
                if self.introduction:
                    self.add_introduction(self.introduction)
                self.write_to_storage()
                if self.agent_session is None:
                    raise Exception('在存储中创建新代理会话失败')
                print(f'-*- Created AgentSession: {self.agent_session.session_id}')
                self._log_agent_session()
        return self.session_id

    def new_session(self) -> None:
        self.agent_session = None
        if self.model:
            self.model.clear()
        if self.memory:
            self.memory.clear()
        self.session_id = str(uuid.uuid4())
        self.load_session(force=True)

    def get_json_output_prompt(self) -> str:
        json_output_prompt = '以包含以下字段的JSON格式提供输出:'
        if self.response_model:
            if isinstance(self.response_model, str):
                json_output_prompt += '\n<json_fields>'
                json_output_prompt += f'\n{self.response_model}'
                json_output_prompt += '\n</json_fields>'
            elif isinstance(self.response_model, list):
                json_output_prompt += '\n<json_fields>'
                json_output_prompt += f'\n{json.dumps(self.response_model)}'
                json_output_prompt += '\n</json_fields>'
            elif issubclass(self.response_model, Base):
                json_schema = self.response_model.model_json_schema()
                if json_schema:
                    response_model_properties = {}
                    json_schema_properties = json_schema.get('properties')
                    if json_schema_properties:
                        for field_name, field_properties in json_schema_properties.items():
                            formatted_field_properties = {prop_name: prop_value
                                for prop_name, prop_value in field_properties.items()
                                if prop_name != 'title'}
                            if 'allOf' in formatted_field_properties:
                                ref = formatted_field_properties['allOf'][0].get('$ref', '')
                                if ref.startswith('#/$defs/'):
                                    enum_name = ref.split('/')[-1]
                                    formatted_field_properties['enum_type'] = enum_name
                            response_model_properties[field_name] = formatted_field_properties
                    json_schema_defs = json_schema.get('$defs')
                    if json_schema_defs:
                        response_model_properties['$defs'] = {}
                        for def_name, def_properties in json_schema_defs.items():
                            if 'enum' in def_properties:
                                response_model_properties['$defs'][def_name] = {'type': 'string', 'enum': def_properties['enum'], 'description': def_properties.get('description', '')}
                            else:
                                def_fields = def_properties.get('properties')
                                formatted_def_properties = {}
                                if def_fields:
                                    for field_name, field_properties in def_fields.items():
                                        formatted_field_properties = {prop_name: prop_value
                                            for prop_name, prop_value in field_properties.items()
                                            if prop_name != 'title'}
                                        formatted_def_properties[field_name] = formatted_field_properties
                                if len(formatted_def_properties) > 0:
                                    response_model_properties['$defs'][def_name] = formatted_def_properties
                    if len(response_model_properties) > 0:
                        json_output_prompt += '\n<json_fields>'
                        json_output_prompt += (f'\n{json.dumps([key for key in response_model_properties.keys() if key != "$defs"])}')
                        json_output_prompt += '\n</json_fields>'
                        json_output_prompt += '\n\nHere are the properties for each field:'
                        json_output_prompt += '\n<json_field_properties>'
                        json_output_prompt += f'\n{json.dumps(response_model_properties, indent=2)}'
                        json_output_prompt += '\n</json_field_properties>'
            else:
                print(f'无法为构建json架构{self.response_model}')
        else:
            json_output_prompt += 'Provide the output as JSON.'
        json_output_prompt += '\nStart your response with `{` and end it with `}`.'
        json_output_prompt += '\nYour output will be passed to json.loads() to convert it to a Python object.'
        json_output_prompt += '\nMake sure it only contains valid JSON.'
        return json_output_prompt

    def format_message_with_state_variables(self, msg: Any) -> Any:
        if not isinstance(msg, str):
            return msg
        format_variables = collections.ChainMap(self.session_state or {}, self.context or {}, self.extra_data or {}, {'user_id': self.user_id} if self.user_id else {})
        return SafeFormatter().format(msg, **format_variables)

    def get_system_message(self) -> Optional[Message]:
        if self.system_message:
            if isinstance(self.system_message, Message):
                return self.system_message
            sys_message_content: str = ''
            if isinstance(self.system_message, str):
                sys_message_content = self.system_message
            elif callable(self.system_message):
                sys_message_content = self.system_message(agent=self)
                if not isinstance(sys_message_content, str):
                    raise Exception('system_message must return a string')
            if self.add_state_in_messages:
                sys_message_content = self.format_message_with_state_variables(sys_message_content)
            if self.response_model and self.model and (self.model.supports_native_structured_outputs and (self.use_json_mode or self.structured_outputs is False)):
                sys_message_content += f'\n{self.get_json_output_prompt()}'
            return Message(role=self.system_message_role, content=sys_message_content)
        if not self.create_default_system_message:
            return None
        if self.model is None:
            raise Exception('model not set')
        instructions: List[str] = []
        if self.instructions:
            _instructions = self.instructions
            if callable(self.instructions):
                _instructions = self.instructions(agent=self)
            if isinstance(_instructions, str):
                instructions.append(_instructions)
            elif isinstance(_instructions, list):
                instructions.extend(_instructions)
        _model_instructions = self.model.get_instructions_for_model()
        if _model_instructions:
            instructions.extend(_model_instructions)
        additional_information: List[str] = []
        if self.markdown and self.response_model is None:
            additional_information.append('Use markdown to format your answers.')
        if self.add_datetime_to_instructions:
            additional_information.append(f'The current time is {datetime.now()}')
        if self.name and self.add_name_to_instructions:
            additional_information.append(f'Your name is: {self.name}.')
        system_message_content: str = ''
        if self.description:
            system_message_content += f'{self.description}\n\n'
        if self.goal:
            system_message_content += f'<your_goal>\n{self.goal}\n</your_goal>\n\n'
        if self.role:
            system_message_content += f'<your_role>\n{self.role}\n</your_role>\n\n'
        if self.has_team and self.add_transfer_instructions:
            system_message_content += ('<agent_team>\n'
                '您是AI代理团队的负责人：\n-您可以直接响应，也可以将任务转移给团队中的其他代理，具体取决于他们可用的工具。\n-如果将任务转移给另一个代理，请确保包括：\n-task_description（str）：任务的清晰描述。\n-expected_output（str）：预期输出。\n-additional_information（str）：有助于代理完成任务的其他信息。\n-在响应用户之前，您必须始终验证其他代理的输出。\n如果你对结果不满意，可以重新分配任务。\n'
                '</agent_team>\n\n')
        if len(instructions) > 0:
            system_message_content += '<instructions>'
            if len(instructions) > 1:
                for _upi in instructions:
                    system_message_content += f'\n- {_upi}'
            else:
                system_message_content += '\n' + instructions[0]
            system_message_content += '\n</instructions>\n\n'
        if len(additional_information) > 0:
            system_message_content += '<additional_information>'
            for _ai in additional_information:
                system_message_content += f'\n- {_ai}'
            system_message_content += '\n</additional_information>\n\n'
        if self.add_state_in_messages:
            system_message_content = self.format_message_with_state_variables(system_message_content)
        system_message_from_model = self.model.get_system_message_for_model()
        if system_message_from_model:
            system_message_content += system_message_from_model
        if self.expected_output:
            system_message_content += f'<expected_output>\n{self.expected_output.strip()}\n</expected_output>\n\n'
        if self.additional_context:
            system_message_content += f'{self.additional_context.strip()}\n'
        if self.has_team and self.add_transfer_instructions:
            system_message_content += f'<transfer_instructions>\n{self.get_transfer_instructions().strip()}\n</transfer_instructions>\n\n'
        if self.memory:
            if self.memory.create_user_memories:
                if self.memory.memories and len(self.memory.memories) > 0:
                    system_message_content += '您可以访问以前与用户交互的记忆，以便使用:\n\n'
                    system_message_content += '<memories_from_previous_interactions>'
                    for _memory in self.memory.memories:
                        system_message_content += f'\n- {_memory.memory}'
                    system_message_content += '\n</memories_from_previous_interactions>\n\n'
                    system_message_content += '注意：此信息来自之前的互动，可能会在本次对话中更新\n你应该总是更喜欢这次谈话中的信息，而不是过去的记忆。\n\n'
                else:
                    system_message_content += '你有能力保留之前与用户互动的记忆但尚未与用户进行任何交互。\n如果用户询问以前的记忆，你可以让他们知道你对用户没有任何记忆，因为你还没有任何互动。\n\n'
                system_message_content += '您可以使用`update_memory`工具添加新内存。\n如果使用`update_memory`工具，请记住将响应传递给用户。\n\n'
            if self.memory.create_session_summary:
                if self.memory.summary:
                    system_message_content += 'Here is a brief summary of your previous interactions if it helps:\n\n'
                    system_message_content += '<summary_of_previous_interactions>\n'
                    system_message_content += str(self.memory.summary)
                    system_message_content += '\n</summary_of_previous_interactions>\n\n'
                    system_message_content += '注意：此信息来自以前的交互，可能已经过时。你应该总是更喜欢这次谈话中的信息，而不是过去的总结。\n\n'
        if self.response_model and not (self.model.supports_native_structured_outputs and (not self.use_json_mode or self.structured_outputs is True)):
            system_message_content += f'{self.get_json_output_prompt()}'
        return Message(role=self.system_message_role, content=system_message_content.strip()) if system_message_content else None

    def get_user_message(self, *, message: Union[str, List], audio: Sequence[Audio] = None, images: Sequence[Image] = None, videos: Sequence[Video] = None, files: Sequence[File] = None, **kwargs: Any) -> Optional[Message]:
        references = None
        if self.add_references and message:
            message_str: str
            if isinstance(message, str):
                message_str = message
            elif callable(message):
                message_str = message(agent=self)
            else:
                raise Exception('当add_references为True时，消息必须是字符串或可调用的')
            retrieval_timer = Timer()
            retrieval_timer.start()
            docs_from_knowledge = self.get_relevant_docs_from_knowledge(query=message_str, **kwargs)
            if docs_from_knowledge:
                references = MessageReferences(query=message_str, references=docs_from_knowledge, time=round(retrieval_timer.elapsed, 4))
                if self.run_response.extra_data is None:
                    self.run_response.extra_data = RunResponseExtraData()
                if self.run_response.extra_data.references is None:
                    self.run_response.extra_data.references = []
                self.run_response.extra_data.references.append(references)
            retrieval_timer.stop()
            print(f'Time to get references: {retrieval_timer.elapsed:.4f}s')
        if self.user_message:
            if isinstance(self.user_message, Message):
                return self.user_message
            user_message_content = self.user_message
            if callable(self.user_message):
                user_message_kwargs = {'agent': self, 'message': message, 'references': references}
                user_message_content = self.user_message(**user_message_kwargs)
                if not isinstance(user_message_content, str):
                    raise Exception('user_message must return a string')
            if self.add_state_in_messages:
                user_message_content = self.format_message_with_state_variables(user_message_content)
            return Message(role=self.user_message_role, content=user_message_content, audio=audio, images=images, videos=videos, files=files, **kwargs)
        if not self.create_default_user_message or isinstance(message, list):
            return Message(role=self.user_message_role, content=message, images=images, audio=audio, videos=videos, files=files, **kwargs)
        if message is None:
            return None
        user_msg_content = message
        if self.add_state_in_messages:
            user_msg_content = self.format_message_with_state_variables(message)
        if self.add_references and references and references.references and len(references.references) > 0:
            user_msg_content += '\n\n如果有帮助，请使用知识库中的以下参考文献:\n'
            user_msg_content += '<references>\n'
            user_msg_content += self.convert_documents_to_string(references.references) + '\n'
            user_msg_content += '</references>'
        if self.add_context and self.context:
            user_msg_content += '\n\n<context>\n'
            user_msg_content += self.convert_context_to_string(self.context) + '\n'
            user_msg_content += '</context>'
        return Message(role=self.user_message_role, content=user_msg_content, audio=audio, images=images, videos=videos, files=files, **kwargs)

    def get_run_messages(self, *, message: Union[str, List, Dict, Message] = None, audio: Sequence[Audio] = None, images: Sequence[Image] = None, videos: Sequence[Video] = None, files: Sequence[File] = None, messages: Sequence[Union[Dict, Message]] = None, **kwargs: Any) -> RunMessages:
        """此函数返回具有以下属性的RunMessages对象：
        -system_message：此运行的系统消息
        -user_message：此运行的用户消息
        -messages：要发送到模型的消息列表
        要构建RunMessages对象，请执行以下操作：
        1.将系统消息添加到run_message
        2.在run_message中添加额外的消息（如果提供）
        3.向run_message添加历史记录
        4.将用户消息添加到run_message
        5.在run_message中添加消息（如果提供）
        返回：
        具有以下属性的RunMessages对象：
        -system_message：此运行的系统消息
        -user_message：此运行的用户消息
        -messages：要发送到模型的所有消息的列表
        典型用法：
        run_messages=self.get_run_messages（message=消息，audio=音频，image=图像，video=视频，file=文件，messages=消息，**kwargs）
        """
        run_messages = RunMessages()
        system_message = self.get_system_message()
        if system_message:
            run_messages.system_message = system_message
            run_messages.messages.append(system_message)
        if self.add_messages:
            messages_to_add_to_run_response: List[Message] = []
            if run_messages.extra_messages is None:
                run_messages.extra_messages = []
            for _m in self.add_messages:
                if isinstance(_m, Message):
                    messages_to_add_to_run_response.append(_m)
                    run_messages.messages.append(_m)
                    run_messages.extra_messages.append(_m)
                elif isinstance(_m, dict):
                    try:
                        _m_parsed = Message.model_validate(_m)
                        messages_to_add_to_run_response.append(_m_parsed)
                        run_messages.messages.append(_m_parsed)
                        run_messages.extra_messages.append(_m_parsed)
                    except Exception as e:
                        print(f'验证消息失败: {e}')
            if len(messages_to_add_to_run_response) > 0:
                print(f'Adding {len(messages_to_add_to_run_response)} extra messages')
                if self.run_response.extra_data is None:
                    self.run_response.extra_data = RunResponseExtraData(add_messages=messages_to_add_to_run_response)
                else:
                    if self.run_response.extra_data.add_messages is None:
                        self.run_response.extra_data.add_messages = messages_to_add_to_run_response
                    else:
                        self.run_response.extra_data.add_messages.extend(messages_to_add_to_run_response)
        if self.add_history_to_messages:
            history: List[Message] = self.memory.get_messages_from_last_n_runs(last_n=self.num_history_runs, skip_role=self.system_message_role)
            if len(history) > 0:
                history_copy = [deepcopy(msg) for msg in history]
                for _msg in history_copy:
                    _msg.from_history = True
                print(f'Adding {len(history_copy)} messages from history')
                run_messages.messages += history_copy
        user_message: Message = None
        if message is None or isinstance(message, str) or isinstance(message, list):
            user_message = self.get_user_message(message=message, audio=audio, images=images, videos=videos, files=files, **kwargs)
        elif isinstance(message, Message):
            user_message = message
        elif isinstance(message, dict):
            try:
                user_message = Message.model_validate(message)
            except Exception as e:
                print(f'验证消息失败: {e}')
        if user_message:
            run_messages.user_message = user_message
            run_messages.messages.append(user_message)
        if messages and len(messages) > 0:
            for _m in messages:
                if isinstance(_m, Message):
                    run_messages.messages.append(_m)
                    if run_messages.extra_messages is None:
                        run_messages.extra_messages = []
                    run_messages.extra_messages.append(_m)
                elif isinstance(_m, dict):
                    try:
                        run_messages.messages.append(Message.model_validate(_m))
                        if run_messages.extra_messages is None:
                            run_messages.extra_messages = []
                        run_messages.extra_messages.append(Message.model_validate(_m))
                    except Exception as e:
                        print(f'验证消息失败: {e}')
        return run_messages

    def get_transfer_function(self, member_agent: 'Agent', index: int) -> Function:
        def _transfer_task_to_agent(task_description: str, expected_output: str, additional_information: str = None) -> Iterator[str]:
            if member_agent.team_data is None:
                member_agent.team_data = {}
            member_agent.team_data['leader_session_id'] = self.session_id
            member_agent.team_data['leader_agent_id'] = self.agent_id
            member_agent.team_data['leader_run_id'] = self.run_id
            member_agent_task = f'{task_description}\n\n<expected_output>\n{expected_output}\n</expected_output>'
            try:
                if additional_information and additional_information.strip() != '':
                    member_agent_task += f'\n\n<additional_information>\n{additional_information}\n</additional_information>'
            except Exception as e:
                print(f'Failed to add additional information to the member agent: {e}')
            member_agent_session_id = member_agent.session_id
            member_agent_agent_id = member_agent.agent_id
            member_agent_info = {'session_id': member_agent_session_id, 'agent_id': member_agent_agent_id}
            if self.team_data is None:
                self.team_data = {}
            if 'members' not in self.team_data:
                self.team_data['members'] = [member_agent_info]
            else:
                if member_agent_info not in self.team_data['members']:
                    self.team_data['members'].append(member_agent_info)
            if self.stream and member_agent.is_streamable:
                member_agent_run_response_stream = member_agent.run(member_agent_task, stream=True)
                for member_agent_run_response_chunk in member_agent_run_response_stream:
                    yield member_agent_run_response_chunk.content
            else:
                member_agent_run_response: RunResponse = member_agent.run(member_agent_task, stream=False)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), Base):
                    try:
                        yield json.dumps(member_agent_run_response.content.to_dict())
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            yield self.team_response_separator
        agent_name = member_agent.name if member_agent.name else f'agent_{index}'
        agent_name = ''.join(c for c in agent_name if c.isalnum() or c in '_- ').strip()
        agent_name = agent_name.lower().replace(' ', '_')
        if member_agent.name is None:
            member_agent.name = agent_name
        strict = True if (member_agent.response_model and member_agent.model) else False
        transfer_function = Function.from_callable(_transfer_task_to_agent, strict=strict)
        transfer_function.strict = strict
        transfer_function.name = f'transfer_task_to_{agent_name}'
        transfer_function.description = textwrap.dedent(f'''使用此功能将任务转移到{agent_name},您必须清晰简洁地描述代理应该完成的任务和预期的输出。
            Args:
            task_description(str):对代理应该完成的任务的清晰简洁的描述。
            expected_output(str):代理的预期输出。
            additional_information(可选[str]):帮助代理完成任务的其他信息。
            return:
            str：委托任务的结果。''')
        if member_agent.respond_directly:
            transfer_function.show_result = True
            transfer_function.stop_after_tool_call = True
        return transfer_function

    def get_transfer_instructions(self) -> str:
        if self.team and len(self.team) > 0:
            transfer_instructions = '您可以将任务转移到团队中的以下代理:\n'
            for agent_index, agent in enumerate(self.team):
                transfer_instructions += f'\nAgent {agent_index + 1}:\n'
                if agent.name:
                    transfer_instructions += f'Name: {agent.name}\n'
                if agent.role:
                    transfer_instructions += f'Role: {agent.role}\n'
                if agent.tools:
                    _tools = []
                    for _tool in agent.tools:
                        if isinstance(_tool, Toolkit):
                            _tools.extend(list(_tool.functions.keys()))
                        elif isinstance(_tool, Function):
                            _tools.append(_tool.name)
                        elif callable(_tool):
                            _tools.append(_tool.__name__)
                    transfer_instructions += f'Available tools: {", ".join(_tools)}\n'
            return transfer_instructions
        return ''

    def get_relevant_docs_from_knowledge(self, query: str, num_documents: int = None, **kwargs) -> Optional[List[Dict[str, Any]]]:
        if self.retriever and callable(self.retriever):
            try:
                sig = inspect.signature(self.retriever)
                retriever_kwargs: Dict[str, Any] = {}
                if 'agent' in sig.parameters:
                    retriever_kwargs = {'agent': self}
                retriever_kwargs.update({'query': query, 'num_documents': num_documents, **kwargs})
                return self.retriever(**retriever_kwargs)
            except Exception as e:
                print(f'Retriever failed: {e}')
                return None
        if self.knowledge is None:
            return None
        relevant_docs: List[Document] = self.knowledge.search(query=query, num_documents=num_documents, **kwargs)
        if len(relevant_docs) == 0:
            return None
        return [doc.to_dict() for doc in relevant_docs]
    
    async def aget_relevant_docs_from_knowledge(self, query: str, num_documents: int = None, **kwargs) -> Optional[List[Dict[str, Any]]]:
        if self.retriever and callable(self.retriever):
            try:
                sig = inspect.signature(self.retriever)
                retriever_kwargs: Dict[str, Any] = {}
                if 'agent' in sig.parameters:
                    retriever_kwargs = {'agent': self}
                retriever_kwargs.update({'query': query, 'num_documents': num_documents, **kwargs})
                return self.retriever(**retriever_kwargs)
            except Exception as e:
                print(f'Retriever failed: {e}')
                return None
        if self.knowledge is None or self.knowledge.vector_db is None:
            return None
        relevant_docs: List[Document] = self.knowledge.search(query=query, num_documents=num_documents, **kwargs)
        if len(relevant_docs) == 0:
            return None
        return [doc.to_dict() for doc in relevant_docs]

    def convert_documents_to_string(self, docs: List[Dict[str, Any]]) -> str:
        if docs is None or len(docs) == 0:
            return ''
        return json.dumps(docs, indent=2)

    def convert_context_to_string(self, context: Dict[str, Any]) -> str:
        if context is None:
            return ''
        try:
            return json.dumps(context, indent=2, default=str)
        except (TypeError, ValueError, OverflowError) as e:
            print(f'Failed to convert context to JSON: {e}')
            sanitized_context = {}
            for key, value in context.items():
                try:
                    json.dumps({key: value}, default=str)
                    sanitized_context[key] = value
                except Exception:
                    sanitized_context[key] = str(value)
            try:
                return json.dumps(sanitized_context, indent=2)
            except Exception as e:
                print(f'未能将经过净化的上下文转换为JSON: {e}')
                return str(context)

    def save_run_response_to_file(self, message: Union[str, List, Dict, Message] = None) -> None:
        if self.save_response_to_file and self.run_response:
            message_str = None
            if message:
                if isinstance(message, str):
                    message_str = message
                else:
                    print('输出文件名中未使用消息：消息不是字符串')
            try:
                fn = self.save_response_to_file.format(name=self.name, session_id=self.session_id, user_id=self.user_id, message=message_str, run_id=self.run_id)
                fn_path = pathlib.Path(fn)
                if not fn_path.parent.exists():
                    fn_path.parent.mkdir(parents=True, exist_ok=True)
                if isinstance(self.run_response.content, str):
                    fn_path.write_text(self.run_response.content)
                else:
                    fn_path.write_text(json.dumps(self.run_response.content, indent=2))
            except Exception as e:
                print(f'未能将输出保存到文件: {e}')

    def update_run_response_with_reasoning(self, reasoning_steps: List[ReasoningStep], reasoning_agent_messages: List[Message]) -> None:
        if self.run_response.extra_data is None:
            self.run_response.extra_data = RunResponseExtraData()
        extra_data = self.run_response.extra_data
        if extra_data.reasoning_steps is None:
            extra_data.reasoning_steps = reasoning_steps
        else:
            extra_data.reasoning_steps.extend(reasoning_steps)
        if extra_data.reasoning_messages is None:
            extra_data.reasoning_messages = reasoning_agent_messages
        else:
            extra_data.reasoning_messages.extend(reasoning_agent_messages)

    def aggregate_metrics_from_messages(self, messages: List[Message]) -> Dict[str, Any]:
        aggregated_metrics: Dict[str, Any] = collections.defaultdict(list)
        assistant_message_role = self.model.assistant_message_role if self.model else 'assistant'
        for m in messages:
            if m.role == assistant_message_role and m.metrics:
                for k, v in m.metrics.__dict__.items():
                    if k == 'timer':
                        continue
                    if v:
                        aggregated_metrics[k].append(v)
        if aggregated_metrics:
            aggregated_metrics = dict(aggregated_metrics)
        return aggregated_metrics

    def rename(self, name: str) -> None:
        self.read_from_storage()
        self.name = name
        self.write_to_storage()
        self._log_agent_session()

    def rename_session(self, session_name: str) -> None:
        self.read_from_storage()
        self.session_name = session_name
        self.write_to_storage()
        self._log_agent_session()

    def generate_session_name(self) -> str:
        if self.model is None:
            raise Exception('Model not set')
        gen_session_name_prompt = 'Conversation\n'
        messages_for_generating_session_name = []
        try:
            message_pairs = self.memory.get_message_pairs()
            for message_pair in message_pairs[:3]:
                messages_for_generating_session_name.append(message_pair[0])
                messages_for_generating_session_name.append(message_pair[1])
        except Exception as e:
            print(f'Failed to generate name: {e}')
        for message in messages_for_generating_session_name:
            gen_session_name_prompt += f'{message.role.upper()}: {message.content}\n'
        gen_session_name_prompt += '\n\nConversation Name: '
        system_message = Message(role=self.system_message_role, content='请为这次对话提供一个合适的名称，最多5个字。')
        user_message = Message(role=self.user_message_role, content=gen_session_name_prompt)
        generate_name_messages = [system_message, user_message]
        generated_name = self.model.response(messages=generate_name_messages)
        content = generated_name.content
        if content is None:
            print('生成的名称为空。再试一次.')
            return self.generate_session_name()
        if len(content.split()) > 15:
            print('生成的名称太长。再试一次.')
            return self.generate_session_name()
        return content.replace('"', '').strip()

    def auto_rename_session(self) -> None:
        self.read_from_storage()
        generated_session_name = self.generate_session_name()
        print(f'Generated Session Name: {generated_session_name}')
        self.session_name = generated_session_name
        self.write_to_storage()
        self._log_agent_session()

    def delete_session(self, session_id: str):
        if self.storage is None:
            return
        self.storage.delete_session(session_id=session_id)

    def add_image(self, image: ImageArtifact) -> None:
        if self.images is None:
            self.images = []
        self.images.append(image)
        if self.run_response:
            if self.run_response.images is None:
                self.run_response.images = []
            self.run_response.images.append(image)

    def add_video(self, video: VideoArtifact) -> None:
        if self.videos is None:
            self.videos = []
        self.videos.append(video)
        if self.run_response:
            if self.run_response.videos is None:
                self.run_response.videos = []
            self.run_response.videos.append(video)

    def add_audio(self, audio: AudioArtifact) -> None:
        if self.audio is None:
            self.audio = []
        self.audio.append(audio)
        if self.run_response:
            if self.run_response.audio is None:
                self.run_response.audio = []
            self.run_response.audio.append(audio)

    def get_images(self) -> Optional[List[ImageArtifact]]:
        return self.images

    def get_videos(self) -> Optional[List[VideoArtifact]]:
        return self.videos

    def get_audio(self) -> Optional[List[AudioArtifact]]:
        return self.audio

    def reason(self, run_messages: RunMessages) -> Iterator[RunResponse]:
        if self.stream_intermediate_steps:
            yield self.create_run_response(content='Reasoning started', event=RunEvent.reasoning_started)
        for message in run_messages.get_input_messages():
            if message.role == 'developer':
                message.role = 'system'
        reasoning_content: str = ''
        reasoning_agent_response: RunResponse = self.reasoning_agent.run(messages=run_messages.get_input_messages())
        if reasoning_agent_response.messages:
            for msg in reasoning_agent_response.messages:
                if msg.reasoning_content:
                    reasoning_content = msg.reasoning_content
                    break
        ds_reasoning_message = Message(role='assistant', content=f'<thinking>\n{reasoning_content}\n</thinking>', reasoning_content=reasoning_content)
        if ds_reasoning_message is None:
            print('推理错误。推理反应为无，继续常规会话...')
            return
        run_messages.messages.append(ds_reasoning_message)
        self.update_run_response_with_reasoning(reasoning_steps=[ReasoningStep(result=ds_reasoning_message.content)], reasoning_agent_messages=[ds_reasoning_message])
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=ReasoningSteps(reasoning_steps=[ReasoningStep(result=ds_reasoning_message.content)]), event=RunEvent.reasoning_completed)
        reasoning_agent: Agent = self.reasoning_agent
        reasoning_agent.show_tool_calls = False
        reasoning_agent.model.show_tool_calls = False
        step_count = 1
        reasoning_messages: List[Message] = []
        all_reasoning_steps: List[ReasoningStep] = []
        print('Starting Reasoning')
        while step_count < self.reasoning_max_steps:
            print(f'Step {step_count}')
            step_count += 1
            try:
                reasoning_agent_response: RunResponse = reasoning_agent.run(messages=run_messages.get_input_messages())
                if reasoning_agent_response.content is None or reasoning_agent_response.messages is None:
                    print('推理错误。推理响应为空，继续常规会话...')
                    break
                if reasoning_agent_response.content.reasoning_steps is None:
                    print('推理错误。推理步骤为空，继续常规会话...')
                    break
                reasoning_steps: List[ReasoningStep] = reasoning_agent_response.content.reasoning_steps
                all_reasoning_steps.extend(reasoning_steps)
                if self.stream_intermediate_steps:
                    for reasoning_step in reasoning_steps:
                        yield self.create_run_response(content=reasoning_step, content_type=reasoning_step.__class__.__name__, event=RunEvent.reasoning_step)
                first_assistant_index = next((i for i, m in enumerate(reasoning_agent_response.messages) if m.role == 'assistant'), len(reasoning_agent_response.messages))
                reasoning_messages = reasoning_agent_response.messages[first_assistant_index:]
                self.update_run_response_with_reasoning(reasoning_steps=reasoning_steps, reasoning_agent_messages=reasoning_agent_response.messages)
                if reasoning_steps[-1].over:
                    break
            except Exception as e:
                print(f'Reasoning error: {e}')
                break
        print(f'Total Reasoning steps: {len(all_reasoning_steps)}\nReasoning finished')
        run_messages.messages.append(Message(role='assistant',
                                             content='I have worked through this problem in-depth, running all necessary tools and have included my raw, step by step research. ',
                                             add_to_agent_memory=False))
        for message in reasoning_messages:
            message.add_to_agent_memory = False
        run_messages.messages.extend(reasoning_messages)
        run_messages.messages.append(Message(role='assistant',
                                             content='Now I will summarize my reasoning and provide a final answer. I will skip any tool calls already executed and steps that are not relevant to the final answer.',
                                             add_to_agent_memory=False))
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=ReasoningSteps(reasoning_steps=all_reasoning_steps), content_type=ReasoningSteps.__class__.__name__, event=RunEvent.reasoning_completed)
                
    async def areason(self, run_messages: RunMessages) -> Any:
        if self.stream_intermediate_steps:
            yield self.create_run_response(content='Reasoning started', event=RunEvent.reasoning_started)
        for message in run_messages.get_input_messages():
            if message.role == 'developer':
                message.role = 'system'
        reasoning_content: str = ''
        reasoning_agent_response: RunResponse = await self.reasoning_agent.arun(messages=run_messages.get_input_messages())
        for msg in reasoning_agent_response.messages:
            if msg.reasoning_content:
                reasoning_content = msg.reasoning_content
                break
        ds_reasoning_message = Message(role='assistant', content=f'<thinking>\n{reasoning_content}\n</thinking>', reasoning_content=reasoning_content)
        if ds_reasoning_message is None:
            print('推理错误。推理反应为无，继续常规会话...')
            return
        run_messages.messages.append(ds_reasoning_message)
        self.update_run_response_with_reasoning(reasoning_steps=[ReasoningStep(result=ds_reasoning_message.content)], reasoning_agent_messages=[ds_reasoning_message])
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=ReasoningSteps(reasoning_steps=[ReasoningStep(result=ds_reasoning_message.content)]), event=RunEvent.reasoning_completed)
        reasoning_agent: Agent = self.reasoning_agent
        reasoning_agent.show_tool_calls = False
        reasoning_agent.model.show_tool_calls = False
        step_count = 1
        reasoning_messages: List[Message] = []
        all_reasoning_steps: List[ReasoningStep] = []
        print('Starting Reasoning')
        while step_count < self.reasoning_max_steps:
            print(f'Step {step_count}')
            step_count += 1
            try:
                reasoning_agent_response: RunResponse = await reasoning_agent.arun(messages=run_messages.get_input_messages())
                if reasoning_agent_response.content is None or reasoning_agent_response.messages is None:
                    print('推理错误。推理响应为空，继续常规会话...')
                    break
                if reasoning_agent_response.content.reasoning_steps is None:
                    print('推理错误。推理步骤为空，继续常规会话n...')
                    break
                reasoning_steps: List[ReasoningStep] = reasoning_agent_response.content.reasoning_steps
                all_reasoning_steps.extend(reasoning_steps)
                if self.stream_intermediate_steps:
                    for reasoning_step in reasoning_steps:
                        yield self.create_run_response(content=reasoning_step, content_type=reasoning_step.__class__.__name__, event=RunEvent.reasoning_step)
                first_assistant_index = next((i for i, m in enumerate(reasoning_agent_response.messages) if m.role == 'assistant'), len(reasoning_agent_response.messages))
                reasoning_messages = reasoning_agent_response.messages[first_assistant_index:]
                self.update_run_response_with_reasoning(reasoning_steps=reasoning_steps, reasoning_agent_messages=reasoning_agent_response.messages)
                if reasoning_steps[-1].over:
                    break
            except Exception as e:
                print(f'Reasoning error: {e}')
                break
        print(f'Total Reasoning steps: {len(all_reasoning_steps)}')
        print('Reasoning finished')
        run_messages.messages.append(Message(role='assistant',
                                             content='I have worked through this problem in-depth, running all necessary tools and have included my raw, step by step research. ',
                                             add_to_agent_memory=False))
        for message in reasoning_messages:
            message.add_to_agent_memory = False
        run_messages.messages.extend(reasoning_messages)
        run_messages.messages.append(Message(role='assistant',
                                             content='Now I will summarize my reasoning and provide a final answer. I will skip any tool calls already executed and steps that are not relevant to the final answer.',
                                             add_to_agent_memory=False))
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=ReasoningSteps(reasoning_steps=all_reasoning_steps), content_type=ReasoningSteps.__class__.__name__, event=RunEvent.reasoning_completed)

    def get_chat_history(self, num_chats: int = None) -> str:
        history: List[Dict[str, Any]] = []
        all_chats = self.memory.get_message_pairs()
        if len(all_chats) == 0:
            return ''
        chats_added = 0
        for chat in all_chats[::-1]:
            history.insert(0, chat[1].to_dict())
            history.insert(0, chat[0].to_dict())
            chats_added += 1
            if num_chats and chats_added >= num_chats:
                break
        return json.dumps(history)

    def get_tool_call_history(self, num_calls: int = 3) -> str:
        tool_calls = self.memory.get_tool_calls(num_calls)
        if len(tool_calls) == 0:
            return ''
        print(f'tool_calls: {tool_calls}')
        return json.dumps(tool_calls)

    def search_knowledge_base(self, query: str) -> str:
        retrieval_timer = Timer()
        retrieval_timer.start()
        docs_from_knowledge = self.get_relevant_docs_from_knowledge(query=query)
        if docs_from_knowledge:
            references = MessageReferences(query=query, references=docs_from_knowledge, time=round(retrieval_timer.elapsed, 4))
            if self.run_response.extra_data is None:
                self.run_response.extra_data = RunResponseExtraData()
            if self.run_response.extra_data.references is None:
                self.run_response.extra_data.references = []
            self.run_response.extra_data.references.append(references)
        retrieval_timer.stop()
        print(f'Time to get references: {retrieval_timer.elapsed:.4f}s')
        if docs_from_knowledge is None:
            return 'No documents found'
        return self.convert_documents_to_string(docs_from_knowledge)
    
    async def async_search_knowledge_base(self, query: str) -> str:
        retrieval_timer = Timer()
        retrieval_timer.start()
        docs_from_knowledge = await self.aget_relevant_docs_from_knowledge(query=query)
        if docs_from_knowledge:
            references = MessageReferences(query=query, references=docs_from_knowledge, time=round(retrieval_timer.elapsed, 4))
            if self.run_response.extra_data is None:
                self.run_response.extra_data = RunResponseExtraData()
            if self.run_response.extra_data.references is None:
                self.run_response.extra_data.references = []
            self.run_response.extra_data.references.append(references)
        retrieval_timer.stop()
        print(f'Time to get references: {retrieval_timer.elapsed:.4f}s')
        if docs_from_knowledge is None:
            return 'No documents found'
        return self.convert_documents_to_string(docs_from_knowledge)

    def add_to_knowledge(self, query: str, result: str) -> str:
        if self.knowledge is None:
            return 'Knowledge base not available'
        document_name = self.name
        if document_name is None:
            document_name = query.replace(' ', '_').replace('?', '').replace('!', '').replace('.', '')
        document_content = json.dumps({'query': query, 'result': result})
        print(f'将文档添加到知识库: {document_name}: {document_content}')
        self.knowledge.load([Document(name=document_name, content=document_content)])
        return '已成功添加到知识库'

    def update_memory(self, task: str) -> str:
        try:
            return self.memory.update_memory(input=task, force=True) or '内存更新成功'
        except Exception as e:
            return f'内存更新失败: {e}'

    def _log_agent_session(self):
        if not (self.telemetry or self.monitoring):
            return
        try:
            agent_session: AgentSession = self.agent_session or self.get_agent_session()
        except Exception as e:
            print(f'Could not create agent monitor: {e}')

    def _create_run_data(self) -> Dict[str, Any]:
        run_response_format = 'text'
        if self.response_model:
            run_response_format = 'json'
        elif self.markdown:
            run_response_format = 'markdown'
        functions = {}
        if self.model and self.model._functions:
            functions = {f_name: func.to_dict() for f_name, func in self.model._functions.items() if isinstance(func, Function)}
        run_data: Dict[str, Any] = {'functions': functions, 'metrics': self.run_response.metrics}
        if self.monitoring:
            run_data.update({'run_input': self.run_input, 'run_response': self.run_response.to_dict(), 'run_response_format': run_response_format})
        return run_data

    def print_response(self, message: Union[List, Dict, str, Message] = None, *, messages: List[Union[Dict, Message]] = None, audio: Sequence[Audio] = None, images: Sequence[Image] = None, videos: Sequence[Video] = None, files: Sequence[File] = None, stream: bool = False, markdown: bool = False, show_message: bool = True, show_reasoning: bool = True, show_full_reasoning: bool = False, console: Any = None, tags_to_include_in_markdown: Set[str] = {'think', 'thinking'}, **kwargs: Any) -> None:
        if markdown:
            self.markdown = True
        if self.response_model:
            self.markdown = False
        with Live(console=console) as live_log:
            status = Status('Thinking...', spinner='aesthetic', speed=0.4, refresh_per_second=10)
            live_log.update(status)
            response_timer = Timer()
            response_timer.start()
            panels = [status]
            if message and show_message:
                message_content = get_text_from_message(message)
                message_panel = create_panel(content=Text(message_content, style='green'), title='Message', border_style='cyan')
                panels.append(message_panel)
                live_log.update(Group(*panels))
            run_response = self.run(message=message, messages=messages, audio=audio, images=images, videos=videos, files=files, stream=False, **kwargs)
            response_timer.stop()
            reasoning_steps = []
            if isinstance(run_response, RunResponse) and run_response.extra_data and run_response.extra_data.reasoning_steps:
                reasoning_steps = run_response.extra_data.reasoning_steps
            if len(reasoning_steps) > 0 and show_reasoning:
                for i, step in enumerate(reasoning_steps, 1):
                    step_content = Text.assemble()
                    if step.title:
                        step_content.append(f'{step.title}\n', 'bold')
                    if step.action:
                        step_content.append(f'{step.action}\n', 'dim')
                    if step.result:
                        step_content.append(Text.from_markup(step.result, style='dim'))
                    if show_full_reasoning:
                        if step.reasoning:
                            step_content.append(Text.from_markup(f'\n[bold]Reasoning:[/bold] {step.reasoning}', style='dim'))
                        if step.confidence:
                            step_content.append(Text.from_markup(f'\n[bold]Confidence:[/bold] {step.confidence}', style='dim'))
                    reasoning_panel = create_panel(content=step_content, title=f'Reasoning step {i}', border_style='green')
                    panels.append(reasoning_panel)
                live_log.update(Group(*panels))
            if isinstance(run_response, RunResponse) and run_response.thinking:
                thinking_panel = create_panel(content=Text(run_response.thinking), title=f'Thinking ({response_timer.elapsed:.1f}s)', border_style='green')
                panels.append(thinking_panel)
                live_log.update(Group(*panels))
            if self.show_tool_calls and isinstance(run_response, RunResponse) and run_response.formatted_tool_calls:
                tool_calls_content = Text()
                for tool_call in run_response.formatted_tool_calls:
                    tool_calls_content.append(f'• {tool_call}\n')
                tool_calls_panel = create_panel(content=tool_calls_content.plain.rstrip(), title='Tool Calls', border_style='yellow')
                panels.append(tool_calls_panel)
                live_log.update(Group(*panels))
            response_content_batch: Union[str, JSON, Markdown] = ''
            if isinstance(run_response, RunResponse):
                if isinstance(run_response.content, str):
                    if self.markdown:
                        escaped_content = run_response.content
                        for tag in tags_to_include_in_markdown:
                            escaped_content = escaped_content.replace(f'<{tag}>', f'&lt;{tag}&gt;')
                            escaped_content = escaped_content.replace(f'</{tag}>', f'&lt;/{tag}&gt;')
                        response_content_batch = Markdown(escaped_content)
                    else:
                        response_content_batch = run_response.get_content_as_string(indent=4)
                elif self.response_model and isinstance(run_response.content, Base):
                    try:
                        response_content_batch = JSON(json.dumps(run_response.content.to_dict()), indent=2)
                    except Exception as e:
                        print(f'Failed to convert response to JSON: {e}')
                else:
                    try:
                        response_content_batch = JSON(json.dumps(run_response.content), indent=4)
                    except Exception as e:
                        print(f'Failed to convert response to JSON: {e}')
            response_panel = create_panel(content=response_content_batch, title=f'Response ({response_timer.elapsed:.1f}s)', border_style='blue')
            panels.append(response_panel)
            if isinstance(run_response, RunResponse) and run_response.citations and run_response.citations.urls:
                md_content = '\n'.join(f'{i + 1}. [{citation.title or citation.url}]({citation.url})'
                    for i, citation in enumerate(run_response.citations.urls)
                    if citation.url)
                if md_content:
                    citations_panel = create_panel(content=Markdown(md_content), title='Citations', border_style='green')
                    panels.append(citations_panel)
                    live_log.update(Group(*panels))
            panels = [p for p in panels if not isinstance(p, Status)]
            live_log.update(Group(*panels))

    def cli_app(self, message: str = None, user: str = 'User', emoji: str = ':sunglasses:', stream: bool = False, markdown: bool = False, exit_on: List[str] = None, **kwargs: Any) -> None:
        if message:
            self.print_response(message=message, stream=stream, markdown=markdown, **kwargs)
        _exit_on = exit_on or ['exit', 'quit', 'bye']
        while True:
            message = Prompt.ask(f'[bold] {emoji} {user} [/bold]')
            if message in _exit_on:
                break
            self.print_response(message=message, stream=stream, markdown=markdown, **kwargs)


class Team:
    def __init__(self, members: List[Union[Agent, 'Team']],
                 mode: Literal['route', 'coordinate', 'collaborate'] = 'coordinate',
                 model: Model = None, name: str = None, team_id: str = None, user_id: str = None,
                 session_id: str = None, session_name: str = None, session_state: Dict[str, Any] = None,
                 add_state_in_messages=False, description: str = None,
                 instructions: Union[str, List[str], Callable] = None,
                 expected_output: str = None, success_criteria: str = None, markdown=False,
                 add_datetime_to_instructions=False,
                 context: Dict[str, Any] = None, add_context=False, enable_agentic_context=False,
                 share_member_interactions=False,
                 read_team_history=False, tools: List[Union[Toolkit, Callable, Function, Dict]] = None,
                 show_tool_calls=True,
                 tool_call_limit: int = None, tool_choice: Union[str, Dict[str, Any]] = None,
                 response_model: Type[Base] = None,
                 use_json_mode=False, parse_response=True, memory: TeamMemory = None,
                 enable_team_history=False,
                 num_of_interactions_from_history=3, storage: Storage = None,
                 extra_data: Dict[str, Any] = None, reasoning=False,
                 reasoning_model: Model = None, reasoning_min_steps=1, reasoning_max_steps=10,
                 debug_mode=False,
                 show_members_responses=False, monitoring=False, telemetry=True, role: str = None,
                 team_session_id: str = None):
        self.role = role
        self.team_session_id = team_session_id
        self.members = members
        self.mode = mode
        self.model = model
        self.name = name
        self.team_id = team_id
        self.user_id = user_id
        self.session_id = session_id
        self.session_name = session_name
        self.session_state = session_state
        self.add_state_in_messages = add_state_in_messages
        self.description = description
        self.instructions = instructions
        self.expected_output = expected_output
        self.markdown = markdown
        self.add_datetime_to_instructions = add_datetime_to_instructions
        self.success_criteria = success_criteria
        self.context = context
        self.add_context = add_context
        self.enable_agentic_context = enable_agentic_context
        self.share_member_interactions = share_member_interactions
        self.read_team_history = read_team_history
        self.tools = tools
        self.show_tool_calls = show_tool_calls
        self.tool_choice = tool_choice
        self.tool_call_limit = tool_call_limit
        self.response_model = response_model
        self.use_json_mode = use_json_mode
        self.parse_response = parse_response
        self.memory = memory
        self.enable_team_history = enable_team_history
        self.num_of_interactions_from_history = num_of_interactions_from_history
        self.storage = storage
        self.extra_data = extra_data
        self.reasoning = reasoning
        self.reasoning_model = reasoning_model
        self.reasoning_min_steps = reasoning_min_steps
        self.reasoning_max_steps = reasoning_max_steps
        self.debug_mode = debug_mode
        self.show_members_responses = show_members_responses
        self.monitoring = monitoring
        self.telemetry = telemetry
        self.run_id: str = None
        self.run_input: Union[str, List, Dict] = None
        self.run_messages: RunMessages = None
        self.run_response: TeamRunResponse = None
        self.images: List[ImageArtifact] = None
        self.audio: List[AudioArtifact] = None
        self.videos: List[VideoArtifact] = None
        self.team_session: AgentSession = None
        self._tools_for_model: List[Dict] = None
        self._functions_for_model: Dict[str, Function] = None
        self._member_response_model: Type[Base] = None
        self.reasoning_agent = None
        if (self.reasoning or self.reasoning_model) and not self.reasoning_agent:
            self.reasoning_agent = Agent(model=reasoning_model or self.model.__class__(id=self.model.id),
                                         description='你是一个细致、周到、有逻辑的推理代理，通过清晰、结构化、循序渐进的分析来解决复杂的问题',
                                         instructions=textwrap.dedent(f'''
            步骤1-问题分析：\n-用你自己的话清楚地重述用户的任务，以确保完全理解。\n-明确指出需要哪些信息以及可能需要哪些工具或资源。
            第2步-分解和制定战略：\n-将问题分解为明确定义的子任务。\n-制定至少两种不同的策略或方法来解决问题，以确保彻底性。
            第3步-意图澄清和规划：\n-清楚地表达用户请求背后的意图。\n-从步骤2中选择最合适的策略，根据与用户意图和任务约束的一致性清楚地证明你的选择。\n-制定详细的分步行动计划，概述解决问题所需的行动顺序。
            步骤4-执行行动计划：
            对于每个计划步骤，记录：\n1.**标题**：概括步骤的简明标题。\n2.**行动**：以第一人称明确说明你的下一个行动（“我会……”）。\n3.**结果**：使用必要的工具执行行动，并提供结果的简明摘要。\n4.**推理**：清楚地解释你的理由，包括：
            -必要性：为什么需要采取这一行动。\n-注意事项：强调关键考虑因素、潜在挑战和缓解策略。\n-进展：这一步如何从逻辑上遵循或建立在之前的行动之上。\n-假设：明确说明所做的任何假设，并证明其有效性。
            5.**下一步行动**：从以下选项中明确选择下一步：
            -**继续**：如果需要进一步的步骤。\n-**验证**：当你得到一个潜在的答案时，表明它已经准备好进行验证。\n-**最终答案**：只有当您自信地验证了解决方案时。\n-**重置**：如果发现严重错误或不正确的结果，请立即重新开始分析。
            6.**置信度分数**：提供一个数字置信度分数（0.0-1.0），表明您对步骤的正确性及其结果的确定性。
            步骤5-验证（在最终确定答案之前必须进行）：
            -通过以下方式明确验证您的解决方案：\n-与替代方法进行交叉验证（在步骤2中开发）。\n-使用其他可用工具或方法独立确认准确性。\n-清楚地记录验证结果和所选验证方法背后的推理。\n-如果验证失败或出现差异，明确指出错误，重置分析，并相应地修改计划。
            第6步-提供最终答案：
            -一旦经过彻底验证并充满信心，就可以清晰简洁地交付您的解决方案。\n-简要重述你的答案如何满足用户的初衷并解决所述任务。
            一般操作指南：
            -确保您的分析保持不变：
            -**完成**：解决任务的所有要素。\n-**全面**：探索不同的观点并预测潜在的结果。\n-**逻辑**：保持所有步骤之间的连贯性。\n-**可操作**：提出明确可执行的步骤和行动。\n-**富有洞察力**：在适用的情况下提供创新和独特的视角。
            -始终通过立即重置或修改步骤来明确处理错误和失误。\n-严格遵守最小{self.reasoning_min_steps}和最大{self.reasoning_max_steps}步数，以确保有效的任务解决。
            -主动毫不犹豫地执行必要的工具，清楚地记录工具的使用情况。'''), tools=tools, show_tool_calls=False,
                                         response_model=ReasoningSteps, use_json_mode=self.use_json_mode,
                                         monitoring=self.monitoring,
                                         telemetry=self.telemetry, debug_mode=self.debug_mode)
            self.reasoning_agent.model.show_tool_calls = False

    def _set_team_id(self) -> str:
        if self.team_id is None:
            self.team_id = str(uuid.uuid4())
        return self.team_id

    def _set_session_id(self) -> str:
        if self.session_id is None or self.session_id == '':
            self.session_id = str(uuid.uuid4())
        return self.session_id

    def _set_debug(self) -> None:
        if self.debug_mode or os.getenv('AGNO_DEBUG', 'false').lower() == 'true':
            self.debug_mode = True

    def _set_monitoring(self) -> None:
        monitor_env = os.getenv('AGNO_MONITOR')
        if monitor_env:
            self.monitoring = monitor_env.lower() == 'true'
        telemetry_env = os.getenv('AGNO_TELEMETRY')
        if telemetry_env:
            self.telemetry = telemetry_env.lower() == 'true'

    def _initialize_member(self, member: Union['Team', Agent]):
        if self.debug_mode:
            member.debug_mode = True
        if self.show_tool_calls:
            member.show_tool_calls = True
        if self.markdown:
            member.markdown = True
        member.team_session_id = self.session_id
        member.team_id = self.team_id
        if member.name is None and member.role is None:
            print('Team member name and role is undefined.')

    def _initialize_team(self) -> None:
        self._set_debug()
        self._set_monitoring()
        self._set_team_id()
        self._set_session_id()
        print(f'Team ID: {self.team_id}\nSession ID: {self.session_id}')
        for member in self.members:
            self._initialize_member(member)

    def run(self, message: Union[str, List, Dict, Message], *, stream: bool = False,
            stream_intermediate_steps: bool = False, retries=3, audio: Sequence[Audio] = None,
            images: Sequence[Image] = None, videos: Sequence[Video] = None,
            files: Sequence[File] = None, **kwargs: Any) -> Union[TeamRunResponse, Iterator[TeamRunResponse]]:
        self._initialize_team()
        show_tool_calls = self.show_tool_calls
        self.read_from_storage()
        if self.memory is None:
            self.memory = TeamMemory()
        if self.context:
            self._resolve_run_context()
        if self.response_model and self.parse_response:
            stream = False
            print('Disabling stream as response_model is set')
        self._configure_model(show_tool_calls=show_tool_calls)
        last_exception = None
        num_attempts = retries + 1
        for attempt in range(num_attempts):
            run_id = str(uuid.uuid4())
            self.run_id = run_id
            print(f'Team Run Start: {self.run_id}\nMode: "{self.mode}"')
            if message:
                if isinstance(message, str):
                    self.run_input = message
                elif isinstance(message, Message):
                    self.run_input = message.to_dict()
                else:
                    self.run_input = message
            _tools: List[Union[Toolkit, Callable, Function, Dict]] = []
            if self.tools:
                for tool in self.tools:
                    _tools.append(tool)
            if self.read_team_history:
                _tools.append(self.get_team_history)
            if self.mode == 'route':
                user_message = self._get_user_message(message, audio=audio, images=images, videos=videos, files=files)
                forward_task_func: Function = self.get_forward_task_function(message=user_message, stream=stream,
                                                                             async_mode=False, images=images,
                                                                             videos=videos, audio=audio, files=files)
                _tools.append(forward_task_func)
            elif self.mode == 'coordinate':
                _tools.append(
                    self.get_transfer_task_function(stream=stream, async_mode=False, images=images, videos=videos,
                                                    audio=audio, files=files))
                if self.enable_agentic_context:
                    _tools.append(self.set_team_context)
            elif self.mode == 'collaborate':
                run_member_agents_func = self.get_run_member_agents_function(stream=stream, async_mode=False,
                                                                             images=images, videos=videos, audio=audio,
                                                                             files=files)
                _tools.append(run_member_agents_func)
                if self.enable_agentic_context:
                    _tools.append(self.set_team_context)
            self._add_tools_to_model(self.model, tools=_tools)
            self.run_response = TeamRunResponse(run_id=self.run_id, session_id=self.session_id, team_id=self.team_id)
            self.run_response.model = self.model.id if self.model else None
            if self.mode == 'route':
                run_messages: RunMessages = self.get_run_messages(run_response=self.run_response, message=message,
                                                                  audio=audio, images=images, videos=videos,
                                                                  files=files, **kwargs)
            else:
                run_messages = self.get_run_messages(run_response=self.run_response, message=message, audio=audio,
                                                     images=images, videos=videos, files=files, **kwargs)
            if stream:
                resp = self._run_stream(run_response=self.run_response, run_messages=run_messages,
                                        stream_intermediate_steps=stream_intermediate_steps)
                return resp
            else:
                self._run(run_response=self.run_response, run_messages=run_messages)
                return self.run_response
        if last_exception:
            print(
                f'Failed after {num_attempts} attempts. Last error using {last_exception.model_name}({last_exception.model_id})')
            raise last_exception
        else:
            raise Exception(f'Failed after {num_attempts} attempts.')

    def _run(self, run_response: TeamRunResponse, run_messages: RunMessages) -> None:
        """运行团队并返回响应。
        步骤：
        1.如果启用推理，则说明任务的原因
        2.从模型中获取响应
        3.更新run_response
        4.更新团队记忆
        5.计算会话度量
        6.将会话保存到存储
        7.解析任何结构化输出
        8.记录团队跑步记录
        """
        if self.reasoning or self.reasoning_model:
            reasoning_generator = self._reason(run_response=run_response, run_messages=run_messages)
            collections.deque(reasoning_generator, maxlen=0)
        self.run_messages = run_messages
        index_of_last_user_message = len(run_messages.messages)
        model_response = self.model.response(messages=run_messages.messages)
        if self.response_model and not self.use_json_mode and model_response.parsed:
            run_response.content = model_response.parsed
            run_response.content_type = self.response_model.__name__
        else:
            if not run_response.content:
                run_response.content = model_response.content
            else:
                run_response.content += model_response.content
        if model_response.thinking:
            if not run_response.thinking:
                run_response.thinking = model_response.thinking
            else:
                run_response.thinking += model_response.thinking
        if model_response.citations:
            run_response.citations = model_response.citations
        if model_response.tool_calls:
            if run_response.tools is None:
                run_response.tools = model_response.tool_calls
            else:
                run_response.tools.extend(model_response.tool_calls)
        run_response.formatted_tool_calls = format_tool_calls(run_response.tools)
        if model_response.audio:
            run_response.response_audio = model_response.audio
        run_response.created_at = model_response.created_at
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        run_response.messages = messages_for_run_response
        run_response.metrics = self._aggregate_metrics_from_messages(messages_for_run_response)
        if run_messages.system_message:
            self.memory.add_system_message(run_messages.system_message, system_message_role='system')
        messages_for_memory: List[Message] = (
            [run_messages.user_message] if run_messages.user_message else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        team_run = TeamRun(response=run_response)
        team_run.message = run_messages.user_message
        if self.memory and self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message:
            self.memory.update_memory(input=run_messages.user_message.get_content_string())
        self.memory.add_team_run(team_run)
        self.write_to_storage()
        if self.response_model:
            if isinstance(run_response.content, str) and self.parse_response:
                try:
                    parsed_response_content = parse_response_model_str(run_response.content, self.response_model)
                    if parsed_response_content:
                        run_response.content = parsed_response_content
                        run_response.content_type = self.response_model.__name__
                    else:
                        print('Failed to convert response to response_model')
                except Exception as e:
                    print(f'Failed to convert response to output model: {e}')
            else:
                print('Something went wrong. Run response content is not a string')
        elif self._member_response_model:
            if isinstance(run_response.content, str):
                try:
                    parsed_response_content = parse_response_model_str(run_response.content,
                                                                       self._member_response_model)
                    if parsed_response_content:
                        run_response.content = parsed_response_content
                        run_response.content_type = self._member_response_model.__name__
                    else:
                        print('Failed to convert response to response_model')
                except Exception as e:
                    print(f'Failed to convert response to output model: {e}')
            else:
                print('Something went wrong. Run response content is not a string')
        self._log_team_run()
        print(f'Team Run End: {self.run_id}')

    def _run_stream(self, run_response: TeamRunResponse, run_messages: RunMessages,
                    stream_intermediate_steps: bool = False) -> Iterator[TeamRunResponse]:
        """运行Team并返回响应迭代器。
        步骤：
        1.如果启用推理，则说明任务的原因
        2.从模型中获取响应
        3.更新run_response
        4.更新团队记忆
        5.计算会话度量
        6.将会话保存到存储
        7.记录团队运行日志
        """
        if self.reasoning or self.reasoning_model:
            reasoning_generator = self._reason(run_response=run_response, run_messages=run_messages)
            yield from reasoning_generator
        self.run_messages = run_messages
        index_of_last_user_message = len(run_messages.messages)
        if stream_intermediate_steps:
            yield self._create_run_response(content='Run started', event=RunEvent.run_started)
        full_model_response = ModelResponse()
        model_stream = self.model.response_stream(messages=run_messages.messages)
        for model_response_chunk in model_stream:
            if model_response_chunk.event == ResponseEvent.res.value:
                should_yield = False
                if model_response_chunk.content:
                    if not full_model_response.content:
                        full_model_response.content = model_response_chunk.content
                    else:
                        full_model_response.content += model_response_chunk.content
                    should_yield = True
                if model_response_chunk.thinking:
                    if not full_model_response.thinking:
                        full_model_response.thinking = model_response_chunk.thinking
                    else:
                        full_model_response.thinking += model_response_chunk.thinking
                    should_yield = True
                if model_response_chunk.citations:
                    full_model_response.citations = model_response_chunk.citations
                    should_yield = True
                if model_response_chunk.audio:
                    if full_model_response.audio is None:
                        full_model_response.audio = AudioResponse(id=str(uuid.uuid4()), content='', transcript='')
                    if model_response_chunk.audio.id:
                        full_model_response.audio.id = model_response_chunk.audio.id
                    if model_response_chunk.audio.content:
                        full_model_response.audio.content += model_response_chunk.audio.content
                    if model_response_chunk.audio.transcript:
                        full_model_response.audio.transcript += model_response_chunk.audio.transcript
                    if model_response_chunk.audio.expires_at:
                        full_model_response.audio.expires_at = model_response_chunk.audio.expires_at
                    if model_response_chunk.audio.mime_type:
                        full_model_response.audio.mime_type = model_response_chunk.audio.mime_type
                    if model_response_chunk.audio.sample_rate:
                        full_model_response.audio.sample_rate = model_response_chunk.audio.sample_rate
                    if model_response_chunk.audio.channels:
                        full_model_response.audio.channels = model_response_chunk.audio.channels
                    should_yield = True
                if should_yield:
                    yield self._create_run_response(content=model_response_chunk.content,
                                                    thinking=model_response_chunk.thinking,
                                                    response_audio=model_response_chunk.audio,
                                                    citations=model_response_chunk.citations,
                                                    created_at=model_response_chunk.created_at)
            elif model_response_chunk.event == ResponseEvent.start.value:
                tool_calls_list = model_response_chunk.tool_calls
                if tool_calls_list:
                    if run_response.tools is None:
                        run_response.tools = tool_calls_list
                    else:
                        run_response.tools.extend(tool_calls_list)
                run_response.formatted_tool_calls = format_tool_calls(run_response.tools)
                if stream_intermediate_steps:
                    yield self._create_run_response(content=model_response_chunk.content,
                                                    event=RunEvent.tool_call_started, from_run_response=run_response)
            elif model_response_chunk.event == ResponseEvent.end.value:
                tool_calls_list = model_response_chunk.tool_calls
                if tool_calls_list:
                    if run_response.tools:
                        tool_call_index_map = {tc['tool_call_id']: i
                                               for i, tc in enumerate(run_response.tools)
                                               if tc.get('tool_call_id')}
                        for tool_call_dict in tool_calls_list:
                            tool_call_id = tool_call_dict.get('tool_call_id')
                            index = tool_call_index_map.get(tool_call_id)
                            if index:
                                run_response.tools[index] = tool_call_dict
                    else:
                        run_response.tools = tool_calls_list
                    if stream_intermediate_steps:
                        yield self._create_run_response(content=model_response_chunk.content,
                                                        event=RunEvent.tool_call_completed,
                                                        from_run_response=run_response)
        run_response.created_at = full_model_response.created_at
        if full_model_response.content:
            run_response.content = full_model_response.content
        if full_model_response.thinking:
            run_response.thinking = full_model_response.thinking
        if full_model_response.audio:
            run_response.response_audio = full_model_response.audio
        if full_model_response.citations:
            run_response.citations = full_model_response.citations
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        run_response.messages = messages_for_run_response
        run_response.metrics = self._aggregate_metrics_from_messages(messages_for_run_response)
        if run_messages.system_message:
            self.memory.add_system_message(run_messages.system_message, system_message_role='system')
        messages_for_memory: List[Message] = (
            [run_messages.user_message] if run_messages.user_message else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        team_run = TeamRun(response=run_response)
        team_run.message = run_messages.user_message
        if self.memory and self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message:
            self.memory.update_memory(input=run_messages.user_message.get_content_string())
        self.memory.add_team_run(team_run)
        self.write_to_storage()
        self._log_team_run()
        if stream_intermediate_steps:
            yield self._create_run_response(from_run_response=run_response, event=RunEvent.run_completed)
        print(f'Team Run End: {self.run_id}')

    async def arun(self, message: Union[str, List, Dict, Message], *, stream: bool = False,
                   stream_intermediate_steps: bool = False, retries: int = None,
                   audio: Sequence[Audio] = None, images: Sequence[Image] = None,
                   videos: Sequence[Video] = None, files: Sequence[File] = None, **kwargs: Any) -> \
    Union[TeamRunResponse, AsyncIterator[TeamRunResponse]]:
        self._initialize_team()
        retries = retries or 3
        if retries < 1:
            raise ValueError('Retries must be at least 1')
        show_tool_calls = self.show_tool_calls
        self.read_from_storage()
        if self.memory is None:
            self.memory = TeamMemory()
        if self.context:
            self._resolve_run_context()
        if self.response_model and self.parse_response:
            stream = False
            print('Disabling stream as response_model is set')
        self._configure_model(show_tool_calls=show_tool_calls)
        last_exception = None
        num_attempts = retries + 1
        for attempt in range(num_attempts):
            run_id = str(uuid.uuid4())
            self.run_id = run_id
            print(f'Team Run Start: {self.run_id}')
            print(f'Mode: "{self.mode}"')
            if message:
                if isinstance(message, str):
                    self.run_input = message
                elif isinstance(message, Message):
                    self.run_input = message.to_dict()
                else:
                    self.run_input = message
            _tools: List[Union[Function, Callable, Toolkit, Dict]] = []
            if self.tools:
                for tool in self.tools:
                    _tools.append(tool)
            if self.read_team_history:
                _tools.append(self.get_team_history)
            if self.mode == 'route':
                user_message = self._get_user_message(message, audio=audio, images=images, videos=videos, files=files)
                forward_task_func: Function = self.get_forward_task_function(message=user_message, stream=stream,
                                                                             async_mode=True, images=images,
                                                                             videos=videos, audio=audio, files=files)
                _tools.append(forward_task_func)
                self.model.tool_choice = 'required'
            elif self.mode == 'coordinate':
                _tools.append(
                    self.get_transfer_task_function(stream=stream, async_mode=True, images=images, videos=videos,
                                                    audio=audio, files=files))
                self.model.tool_choice = 'auto'
                if self.enable_agentic_context:
                    _tools.append(self.set_team_context)
            elif self.mode == 'collaborate':
                run_member_agents_func = self.get_run_member_agents_function(stream=stream, async_mode=True,
                                                                             images=images, videos=videos, audio=audio,
                                                                             files=files)
                _tools.append(run_member_agents_func)
                self.model.tool_choice = 'auto'
                if self.enable_agentic_context:
                    _tools.append(self.set_team_context)
            self._add_tools_to_model(self.model, tools=_tools)
            self.run_response = TeamRunResponse(run_id=self.run_id, session_id=self.session_id, team_id=self.team_id)
            self.run_response.model = self.model.id if self.model else None
            if self.mode == 'route':
                run_messages: RunMessages = self.get_run_messages(run_response=self.run_response, message=message,
                                                                  audio=audio, images=images, videos=videos,
                                                                  files=files, **kwargs)
            else:
                run_messages = self.get_run_messages(run_response=self.run_response, message=message, audio=audio,
                                                     images=images, videos=videos, files=files, **kwargs)
            if stream:
                resp = self._arun_stream(run_response=self.run_response, run_messages=run_messages,
                                         stream_intermediate_steps=stream_intermediate_steps)
                return resp
            else:
                await self._arun(run_response=self.run_response, run_messages=run_messages)
                return self.run_response
        if last_exception:
            print(
                f'Failed after {num_attempts} attempts. Last error using {last_exception.model_name}({last_exception.model_id})')
            raise last_exception
        else:
            raise Exception(f'Failed after {num_attempts} attempts.')

    async def _arun(self, run_response: TeamRunResponse, run_messages: RunMessages) -> None:
        """运行团队并返回响应。
        步骤：
        1.如果启用推理，则说明任务的原因
        2.从模型中获取响应
        3.更新run_response
        4.更新团队记忆
        5.计算会话度量
        6.将会话保存到存储
        7.解析任何结构化输出
        8.记录团队运行记录
        """
        if self.reasoning or self.reasoning_model:
            reasoning_generator = self._areason(run_response=run_response, run_messages=run_messages)
            async for _ in reasoning_generator:
                pass
        self.run_messages = run_messages
        index_of_last_user_message = len(run_messages.messages)
        model_response = await self.model.aresponse(messages=run_messages.messages)
        if (self.response_model) and not self.use_json_mode and (model_response.parsed):
            run_response.content = model_response.parsed
            run_response.content_type = self.response_model.__name__
        else:
            if not run_response.content:
                run_response.content = model_response.content
            else:
                run_response.content += model_response.content
        if model_response.thinking:
            if not run_response.thinking:
                run_response.thinking = model_response.thinking
            else:
                run_response.thinking += model_response.thinking
        if model_response.citations:
            run_response.citations = model_response.citations
        if model_response.tool_calls:
            if run_response.tools is None:
                run_response.tools = model_response.tool_calls
            else:
                run_response.tools.extend(model_response.tool_calls)
        run_response.formatted_tool_calls = format_tool_calls(run_response.tools)
        if model_response.audio:
            run_response.response_audio = model_response.audio
        run_response.created_at = model_response.created_at
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        run_response.messages = messages_for_run_response
        run_response.metrics = self._aggregate_metrics_from_messages(messages_for_run_response)
        if run_messages.system_message:
            self.memory.add_system_message(run_messages.system_message, system_message_role='system')
        messages_for_memory: List[Message] = (
            [run_messages.user_message] if run_messages.user_message else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        team_run = TeamRun(response=run_response)
        team_run.message = run_messages.user_message
        if self.memory and self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message:
            await self.memory.aupdate_memory(input=run_messages.user_message.get_content_string())
        self.memory.add_team_run(team_run)
        self.write_to_storage()
        if self.response_model:
            if isinstance(run_response.content, str) and self.parse_response:
                try:
                    parsed_response_content = parse_response_model_str(run_response.content, self.response_model)
                    if parsed_response_content:
                        run_response.content = parsed_response_content
                        run_response.content_type = self.response_model.__name__
                    else:
                        print('Failed to convert response to response_model')
                except Exception as e:
                    print(f'Failed to convert response to output model: {e}')
            else:
                print('Something went wrong. Run response content is not a string')
        elif self._member_response_model:
            if isinstance(run_response.content, str):
                try:
                    parsed_response_content = parse_response_model_str(run_response.content,
                                                                       self._member_response_model)
                    if parsed_response_content:
                        run_response.content = parsed_response_content
                        run_response.content_type = self._member_response_model.__name__
                    else:
                        print('Failed to convert response to response_model')
                except Exception as e:
                    print(f'Failed to convert response to output model: {e}')
            else:
                print('Something went wrong. Run response content is not a string')
        await self._alog_team_run()
        print(f'Team Run End: {self.run_id}')

    async def _arun_stream(self, run_response: TeamRunResponse, run_messages: RunMessages,
                           stream_intermediate_steps: bool = False) -> AsyncIterator[TeamRunResponse]:
        """运行团队并返回响应。
        步骤：
        1.如果启用推理，则说明任务的原因
        2.从模型中获取响应
        3.更新run_response
        4.更新团队记忆
        5.计算会话度量
        6.将会话保存到存储
        7.记录团队运行日志
        """
        if self.reasoning or self.reasoning_model:
            reasoning_generator = self._areason(run_response=run_response, run_messages=run_messages)
            async for reasoning_response in reasoning_generator:
                yield reasoning_response
        self.run_messages = run_messages
        index_of_last_user_message = len(run_messages.messages)
        if stream_intermediate_steps:
            yield self._create_run_response(content='Run started', event=RunEvent.run_started)
        full_model_response = ModelResponse()
        model_stream = self.model.aresponse_stream(messages=run_messages.messages)
        async for model_response_chunk in model_stream:
            if model_response_chunk.event == ResponseEvent.res.value:
                should_yield = False
                if model_response_chunk.content:
                    if not full_model_response.content:
                        full_model_response.content = model_response_chunk.content
                    else:
                        full_model_response.content += model_response_chunk.content
                    should_yield = True
                if model_response_chunk.thinking:
                    if not full_model_response.thinking:
                        full_model_response.thinking = model_response_chunk.thinking
                    else:
                        full_model_response.thinking += model_response_chunk.thinking
                    should_yield = True
                if model_response_chunk.citations:
                    full_model_response.citations = model_response_chunk.citations
                    should_yield = True
                if model_response_chunk.audio:
                    if full_model_response.audio is None:
                        full_model_response.audio = AudioResponse(id=str(uuid.uuid4()), content='', transcript='')
                    if model_response_chunk.audio.id:
                        full_model_response.audio.id = model_response_chunk.audio.id
                    if model_response_chunk.audio.content:
                        full_model_response.audio.content += model_response_chunk.audio.content
                    if model_response_chunk.audio.transcript:
                        full_model_response.audio.transcript += model_response_chunk.audio.transcript
                    if model_response_chunk.audio.expires_at:
                        full_model_response.audio.expires_at = model_response_chunk.audio.expires_at
                    if model_response_chunk.audio.mime_type:
                        full_model_response.audio.mime_type = model_response_chunk.audio.mime_type
                    if model_response_chunk.audio.sample_rate:
                        full_model_response.audio.sample_rate = model_response_chunk.audio.sample_rate
                    if model_response_chunk.audio.channels:
                        full_model_response.audio.channels = model_response_chunk.audio.channels
                    should_yield = True
                if should_yield:
                    yield self._create_run_response(content=model_response_chunk.content,
                                                    thinking=model_response_chunk.thinking,
                                                    response_audio=model_response_chunk.audio,
                                                    citations=model_response_chunk.citations,
                                                    created_at=model_response_chunk.created_at)
            elif model_response_chunk.event == ResponseEvent.start.value:
                tool_calls_list = model_response_chunk.tool_calls
                if tool_calls_list:
                    if run_response.tools is None:
                        run_response.tools = tool_calls_list
                    else:
                        run_response.tools.extend(tool_calls_list)
                run_response.formatted_tool_calls = format_tool_calls(run_response.tools)
                if stream_intermediate_steps:
                    yield self._create_run_response(content=model_response_chunk.content,
                                                    event=RunEvent.tool_call_started, from_run_response=run_response)
            elif model_response_chunk.event == ResponseEvent.end.value:
                tool_calls_list = model_response_chunk.tool_calls
                if tool_calls_list:
                    if run_response.tools:
                        tool_call_index_map = {tc['tool_call_id']: i
                                               for i, tc in enumerate(run_response.tools)
                                               if tc.get('tool_call_id')}
                        for tool_call_dict in tool_calls_list:
                            tool_call_id = tool_call_dict.get('tool_call_id')
                            index = tool_call_index_map.get(tool_call_id)
                            if index:
                                run_response.tools[index] = tool_call_dict
                    else:
                        run_response.tools = tool_calls_list
                    if stream_intermediate_steps:
                        yield self._create_run_response(content=model_response_chunk.content,
                                                        event=RunEvent.tool_call_completed,
                                                        from_run_response=run_response)
        if (self.response_model) and not self.use_json_mode and (full_model_response.parsed):
            run_response.content = full_model_response.parsed
        run_response.created_at = full_model_response.created_at
        if full_model_response.content:
            run_response.content = full_model_response.content
        if full_model_response.thinking:
            run_response.thinking = full_model_response.thinking
        if full_model_response.audio:
            run_response.response_audio = full_model_response.audio
        if full_model_response.citations:
            run_response.citations = full_model_response.citations
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        run_response.messages = messages_for_run_response
        run_response.metrics = self._aggregate_metrics_from_messages(messages_for_run_response)
        if run_messages.system_message:
            self.memory.add_system_message(run_messages.system_message, system_message_role='system')
        messages_for_memory: List[Message] = (
            [run_messages.user_message] if run_messages.user_message else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        team_run = TeamRun(response=run_response)
        team_run.message = run_messages.user_message
        if self.memory and self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message:
            await self.memory.aupdate_memory(input=run_messages.user_message.get_content_string())
        self.memory.add_team_run(team_run)
        self.write_to_storage()
        await self._alog_team_run()
        if stream_intermediate_steps:
            yield self._create_run_response(from_run_response=run_response, event=RunEvent.run_completed)
        print(f'Team Run End: {self.run_id}')

    def print_response(self, message: Union[List, Dict, str, Message] = None, *, stream: bool = False,
                       stream_intermediate_steps: bool = False, show_message: bool = True, show_reasoning: bool = True,
                       show_reasoning_verbose: bool = False, console: Any = None,
                       tags_to_include_in_markdown: Set[str] = None, audio: Sequence[Audio] = None,
                       images: Sequence[Image] = None, videos: Sequence[Video] = None,
                       files: Sequence[File] = None, markdown: bool = None, **kwargs: Any) -> None:
        if not tags_to_include_in_markdown:
            tags_to_include_in_markdown = {'think', 'thinking'}
        if markdown is None:
            markdown = self.markdown
        else:
            self.markdown = markdown
        if not tags_to_include_in_markdown:
            tags_to_include_in_markdown = {'think', 'thinking'}
        with Live(console=console) as live_console:
            status = Status('Thinking...', spinner='aesthetic', speed=0.4, refresh_per_second=10)
            live_console.update(status)
            response_timer = Timer()
            response_timer.start()
            panels = [status]
            if message and show_message:
                message_content = get_text_from_message(message)
                message_panel = create_panel(content=Text(message_content, style='green'), title='Message',
                                             border_style='cyan')
                panels.append(message_panel)
                live_console.update(Group(*panels))
            run_response: TeamRunResponse = self.run(message=message, images=images, audio=audio, videos=videos,
                                                     files=files, stream=False, **kwargs)
            response_timer.stop()
            team_markdown = False
            member_markdown = {}
            if markdown:
                for member in self.members:
                    if isinstance(member, Agent) and member.agent_id:
                        member_markdown[member.agent_id] = True
                    if isinstance(member, Team) and member.team_id:
                        member_markdown[member.team_id] = True
                team_markdown = True
            if self.response_model:
                team_markdown = False
            for member in self.members:
                if member.response_model and isinstance(member, Agent) and member.agent_id:
                    member_markdown[member.agent_id] = False
                if member.response_model and isinstance(member, Team) and member.team_id:
                    member_markdown[member.team_id] = False
            reasoning_steps = []
            if isinstance(run_response,
                          TeamRunResponse) and run_response.extra_data and run_response.extra_data.reasoning_steps:
                reasoning_steps = run_response.extra_data.reasoning_steps
            if len(reasoning_steps) > 0 and show_reasoning:
                for i, step in enumerate(reasoning_steps, 1):
                    reasoning_panel = self._build_reasoning_step_panel(i, step, show_reasoning_verbose)
                    panels.append(reasoning_panel)
                live_console.update(Group(*panels))
            if isinstance(run_response, TeamRunResponse) and run_response.thinking:
                thinking_panel = create_panel(content=Text(run_response.thinking),
                                              title=f'Thinking ({response_timer.elapsed:.1f}s)', border_style='green')
                panels.append(thinking_panel)
                live_console.update(Group(*panels))
            if isinstance(run_response, TeamRunResponse):
                if self.show_members_responses:
                    for member_response in run_response.member_responses:
                        reasoning_steps = []
                        if isinstance(member_response,
                                      RunResponse) and member_response.extra_data and member_response.extra_data.reasoning_steps:
                            reasoning_steps.extend(member_response.extra_data.reasoning_steps)
                        if len(reasoning_steps) > 0 and show_reasoning:
                            for i, step in enumerate(reasoning_steps, 1):
                                member_reasoning_panel = self._build_reasoning_step_panel(i, step,
                                                                                          show_reasoning_verbose,
                                                                                          color='magenta')
                                panels.append(member_reasoning_panel)
                        if self.show_tool_calls and hasattr(member_response, 'tools') and member_response.tools:
                            member_name = None
                            if isinstance(member_response, RunResponse) and member_response.agent_id:
                                member_name = self._get_member_name(member_response.agent_id)
                            elif isinstance(member_response, TeamRunResponse) and member_response.team_id:
                                member_name = self._get_member_name(member_response.team_id)
                            if member_name:
                                formatted_calls = format_tool_calls(member_response.tools)
                                if formatted_calls:
                                    console_width = console.width if console else 80
                                    panel_width = console_width + 30
                                    lines = []
                                    for call in formatted_calls:
                                        wrapped_call = textwrap.fill(f'• {call}', width=panel_width,
                                                                     subsequent_indent='  ')
                                        lines.append(wrapped_call)
                                    tool_calls_text = '\n\n'.join(lines)
                                    member_tool_calls_panel = create_panel(content=tool_calls_text,
                                                                           title=f'{member_name} Tool Calls',
                                                                           border_style='yellow')
                                    panels.append(member_tool_calls_panel)
                                    live_console.update(Group(*panels))
                        show_markdown = False
                        if member_markdown:
                            if isinstance(member_response, RunResponse) and member_response.agent_id:
                                show_markdown = member_markdown.get(member_response.agent_id, False)
                            elif isinstance(member_response, TeamRunResponse) and member_response.team_id:
                                show_markdown = member_markdown.get(member_response.team_id, False)
                        member_response_content: Union[str, JSON, Markdown] = self._parse_response_content(
                            member_response, tags_to_include_in_markdown, show_markdown=show_markdown)
                        if isinstance(member_response, RunResponse) and member_response.agent_id:
                            member_response_panel = create_panel(content=member_response_content,
                                                                 title=f'{self._get_member_name(member_response.agent_id)} Response',
                                                                 border_style='magenta')
                        elif isinstance(member_response, TeamRunResponse) and member_response.team_id:
                            member_response_panel = create_panel(content=member_response_content,
                                                                 title=f'{self._get_member_name(member_response.team_id)} Response',
                                                                 border_style='magenta')
                        panels.append(member_response_panel)
                        if member_response.citations and member_response.citations.urls:
                            md_content = '\n'.join(
                                f'{i + 1}. [{citation.title or citation.url}]({citation.url})' for i, citation in
                                enumerate(member_response.citations.urls) if citation.url)
                            if md_content:
                                citations_panel = create_panel(content=Markdown(md_content), title='Citations',
                                                               border_style='magenta')
                                panels.append(citations_panel)
                    live_console.update(Group(*panels))
                if self.show_tool_calls and run_response.tools:
                    formatted_calls = format_tool_calls(run_response.tools)
                    if formatted_calls:
                        console_width = console.width if console else 80
                        panel_width = console_width + 30
                        lines = []
                        for call in formatted_calls:
                            wrapped_call = textwrap.fill(f'• {call}', width=panel_width, subsequent_indent='  ')
                            lines.append(wrapped_call)
                        tool_calls_text = '\n\n'.join(lines)
                        team_tool_calls_panel = create_panel(content=tool_calls_text, title='Team Tool Calls',
                                                             border_style='yellow')
                        panels.append(team_tool_calls_panel)
                        live_console.update(Group(*panels))
                response_content_batch: Union[str, JSON, Markdown] = self._parse_response_content(run_response,
                                                                                                  tags_to_include_in_markdown,
                                                                                                  show_markdown=team_markdown)
                response_panel = create_panel(content=response_content_batch,
                                              title=f'Response ({response_timer.elapsed:.1f}s)', border_style='blue')
                panels.append(response_panel)
                if run_response.citations and run_response.citations.urls:
                    md_content = '\n'.join(
                        f'{i + 1}. [{citation.title or citation.url}]({citation.url})' for i, citation in
                        enumerate(run_response.citations.urls) if citation.url)
                    if md_content:
                        citations_panel = create_panel(content=Markdown(md_content), title='Citations',
                                                       border_style='green')
                        panels.append(citations_panel)
            panels = [p for p in panels if not isinstance(p, Status)]
            live_console.update(Group(*panels))

    def _build_reasoning_step_panel(self, step_idx: int, step: ReasoningStep, show_reasoning_verbose: bool = False,
                                    color: str = 'green'):
        step_content = Text.assemble()
        if step.title:
            step_content.append(f'{step.title}\n', 'bold')
        if step.action:
            step_content.append(f'{step.action}\n', 'dim')
        if step.result:
            step_content.append(Text.from_markup(step.result, style='dim'))
        if show_reasoning_verbose:
            if step.reasoning:
                step_content.append(Text.from_markup(f'\n[bold]Reasoning:[/bold] {step.reasoning}', style='dim'))
            if step.confidence:
                step_content.append(Text.from_markup(f'\n[bold]Confidence:[/bold] {step.confidence}', style='dim'))
        return create_panel(content=step_content, title=f'Reasoning step {step_idx}', border_style=color)

    def _get_member_name(self, entity_id: str) -> str:
        for member in self.members:
            if isinstance(member, Agent):
                if member.agent_id == entity_id:
                    return member.name or entity_id
            elif isinstance(member, Team):
                if member.team_id == entity_id:
                    return member.name or entity_id
        return entity_id

    def _parse_response_content(self, run_response: Union[TeamRunResponse, RunResponse],
                                tags_to_include_in_markdown: Set[str], show_markdown: bool = True) -> Any:
        if isinstance(run_response.content, str):
            if show_markdown:
                escaped_content = run_response.content
                for tag in tags_to_include_in_markdown:
                    escaped_content = escaped_content.replace(f'<{tag}>', f'&lt;{tag}&gt;')
                    escaped_content = escaped_content.replace(f'</{tag}>', f'&lt;/{tag}&gt;')
                return Markdown(escaped_content)
            else:
                return run_response.get_content_as_string(indent=4)
        elif isinstance(run_response.content, Base):
            try:
                return JSON(json.dumps(run_response.content.to_dict()), indent=2)
            except Exception as e:
                print(f'Failed to convert response to JSON: {e}')
        else:
            try:
                return JSON(json.dumps(run_response.content), indent=4)
            except Exception as e:
                print(f'Failed to convert response to JSON: {e}')

    def cli_app(self, message: str = None, user: str = 'User', emoji: str = ':sunglasses:',
                stream: bool = False, markdown: bool = False, exit_on: List[str] = None,
                **kwargs: Any) -> None:
        if message:
            self.print_response(message=message, stream=stream, markdown=markdown, **kwargs)
        _exit_on = exit_on or ['exit', 'quit', 'bye']
        while True:
            message = Prompt.ask(f'[bold] {emoji} {user} [/bold]')
            if message in _exit_on:
                break
            self.print_response(message=message, stream=stream, markdown=markdown, **kwargs)

    def _aggregate_metrics_from_messages(self, messages: List[Message]) -> Dict[str, Any]:
        aggregated_metrics: Dict[str, Any] = collections.defaultdict(list)
        assistant_message_role = self.model.assistant_message_role if self.model else 'assistant'
        for m in messages:
            if m.role == assistant_message_role and m.metrics:
                for k, v in m.metrics.__dict__.items():
                    if k == 'timer':
                        continue
                    if v:
                        aggregated_metrics[k].append(v)
        if aggregated_metrics:
            aggregated_metrics = dict(aggregated_metrics)
        return aggregated_metrics

    def _reason(self, run_response: TeamRunResponse, run_messages: RunMessages,
                stream_intermediate_steps: bool = False) -> Iterator[TeamRunResponse]:
        if stream_intermediate_steps:
            yield self._create_run_response(from_run_response=run_response, content='Reasoning started',
                                            event=RunEvent.reasoning_started)
        reasoning_agent = self.reasoning_agent
        for message in run_messages.get_input_messages():
            if message.role == 'developer':
                message.role = 'system'
        reasoning_content: str = ''
        reasoning_agent_response: RunResponse = reasoning_agent.run(messages=run_messages.get_input_messages())
        if reasoning_agent_response.messages:
            for msg in reasoning_agent_response.messages:
                if msg.reasoning_content:
                    reasoning_content = msg.reasoning_content
                    break
        reasoning_message = Message(role='assistant', content=f'<thinking>\n{reasoning_content}\n</thinking>',
                       reasoning_content=reasoning_content)
        if reasoning_message is None:
            print('Reasoning error. Reasoning response is None, continuing regular session...')
            return
        if reasoning_message:
            run_messages.messages.append(reasoning_message)
            update_run_response_with_reasoning(run_response=run_response,
                                               reasoning_steps=[ReasoningStep(result=reasoning_message.content)],
                                               reasoning_agent_messages=[reasoning_message])
        if stream_intermediate_steps:
            yield self._create_run_response(
                content=ReasoningSteps(reasoning_steps=[ReasoningStep(result=reasoning_message.content)]),
                content_type=ReasoningSteps.__class__.__name__, event=RunEvent.reasoning_completed)

    async def _areason(self, run_response: TeamRunResponse, run_messages: RunMessages,
                       stream_intermediate_steps: bool = False) -> AsyncIterator[TeamRunResponse]:
        if stream_intermediate_steps:
            yield self._create_run_response(from_run_response=run_response, content='Reasoning started',
                                            event=RunEvent.reasoning_started)
        reasoning_agent = self.reasoning_agent
        for message in run_messages.get_input_messages():
            if message.role == 'developer':
                message.role = 'system'
        reasoning_content: str = ''
        reasoning_agent_response: RunResponse = await reasoning_agent.arun(messages=run_messages.get_input_messages())
        for msg in reasoning_agent_response.messages:
            if msg.reasoning_content:
                reasoning_content = msg.reasoning_content
                break
        reasoning_message = Message(role='assistant', content=f'<thinking>\n{reasoning_content}\n</thinking>',
                       reasoning_content=reasoning_content)
        if reasoning_message is None:
            print('Reasoning error. Reasoning response is None, continuing regular session...')
            return
        if reasoning_message:
            run_messages.messages.append(reasoning_message)
            update_run_response_with_reasoning(run_response=run_response,
                                               reasoning_steps=[ReasoningStep(result=reasoning_message.content)],
                                               reasoning_agent_messages=[reasoning_message])
        if stream_intermediate_steps:
            yield self._create_run_response(
                content=ReasoningSteps(reasoning_steps=[ReasoningStep(result=reasoning_message.content)]),
                content_type=ReasoningSteps.__class__.__name__, event=RunEvent.reasoning_completed)

    def _create_run_response(self, content: Any = None, content_type: str = None,
                             thinking: str = None, event: RunEvent = RunEvent.run_response,
                             tools: List[Dict[str, Any]] = None, audio: List[AudioArtifact] = None,
                             images: List[ImageArtifact] = None, videos: List[VideoArtifact] = None,
                             response_audio: AudioResponse = None, citations: Citations = None,
                             model: str = None, messages: List[Message] = None,
                             created_at: int = None,
                             from_run_response: TeamRunResponse = None) -> TeamRunResponse:
        extra_data = None
        member_responses = None
        formatted_tool_calls = None
        if from_run_response:
            content = from_run_response.content
            content_type = from_run_response.content_type
            tools = from_run_response.tools
            audio = from_run_response.audio
            images = from_run_response.images
            videos = from_run_response.videos
            response_audio = from_run_response.response_audio
            model = from_run_response.model
            messages = from_run_response.messages
            extra_data = from_run_response.extra_data
            member_responses = from_run_response.member_responses
            citations = from_run_response.citations
            tools = from_run_response.tools
            formatted_tool_calls = from_run_response.formatted_tool_calls
        rr = TeamRunResponse(run_id=self.run_id, session_id=self.session_id, team_id=self.team_id, content=content,
                             thinking=thinking, tools=tools, audio=audio, images=images, videos=videos,
                             response_audio=response_audio, citations=citations, model=model, messages=messages,
                             extra_data=extra_data, event=event.value)
        if formatted_tool_calls:
            rr.formatted_tool_calls = formatted_tool_calls
        if member_responses:
            rr.member_responses = member_responses
        if content_type:
            rr.content_type = content_type
        if created_at:
            rr.created_at = created_at
        return rr

    def _resolve_run_context(self) -> None:
        print('Resolving context')
        if self.context:
            if isinstance(self.context, dict):
                for ctx_key, ctx_value in self.context.items():
                    if callable(ctx_value):
                        try:
                            sig = inspect.signature(ctx_value)
                            if 'agent' in sig.parameters:
                                resolved_ctx_value = ctx_value(agent=self)
                            else:
                                resolved_ctx_value = ctx_value()
                            if resolved_ctx_value:
                                self.context[ctx_key] = resolved_ctx_value
                        except Exception as e:
                            print(f'Failed to resolve context for {ctx_key}: {e}')
                    else:
                        self.context[ctx_key] = ctx_value
            else:
                print('Context is not a dict')

    def _configure_model(self, show_tool_calls: bool = False) -> None:
        if self.model is None:
            print('Setting default model to OpenAI Chat')
            self.model = Ollama()
        if self.response_model is None:
            self.model.response_format = None
        else:
            json_response_format = {'type': 'json_object'}
            if self.model.supports_native_structured_outputs:
                if self.use_json_mode:
                    self.model.response_format = json_response_format
                    self.model.structured_outputs = False
                else:
                    print('Setting Model.response_format to Agent.response_model')
                    self.model.response_format = self.response_model
                    self.model.structured_outputs = True
            elif self.model.supports_json_schema_outputs:
                if self.use_json_mode:
                    print('Setting Model.response_format to JSON response mode')
                    self.model.response_format = {'type': 'json_schema',
                                                  'json_schema': {'name': self.response_model.__name__,
                                                                  'schema': self.response_model.model_json_schema()}}
                else:
                    self.model.response_format = None
                self.model.structured_outputs = False
            else:
                self.model.response_format = json_response_format if self.use_json_mode else None
                self.model.structured_outputs = False
        self.model.show_tool_calls = show_tool_calls
        if self.tool_choice:
            self.model.tool_choice = self.tool_choice
        if self.tool_call_limit:
            self.model.tool_call_limit = self.tool_call_limit

    def _add_tools_to_model(self, model: Model, tools: List[Union[Function, Callable, Toolkit, Dict]]) -> None:
        self._functions_for_model = {}
        self._tools_for_model = []
        if len(tools) > 0:
            print('Processing tools for model')
            strict = False
            if self.response_model and not self.use_json_mode and model.supports_native_structured_outputs:
                strict = True
            for tool in tools:
                if isinstance(tool, Dict):
                    self._tools_for_model.append(tool)
                    print(f'Included builtin tool {tool}')
                elif isinstance(tool, Toolkit):
                    for name, func in tool.functions.items():
                        if name not in self._functions_for_model:
                            func._agent = self
                            func.process_entrypoint(strict=strict)
                            if strict:
                                func.strict = True
                            self._functions_for_model[name] = func
                            self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                            print(f'Included function {name} from {tool.name}')
                elif isinstance(tool, Function):
                    if tool.name not in self._functions_for_model:
                        tool._agent = self
                        tool.process_entrypoint(strict=strict)
                        if strict and tool.strict is None:
                            tool.strict = True
                        self._functions_for_model[tool.name] = tool
                        self._tools_for_model.append({'type': 'function', 'function': tool.to_dict()})
                        print(f'Included function {tool.name}')
                elif callable(tool):
                    try:
                        func = Function.from_callable(tool, strict=strict)
                        func._agent = self
                        if strict:
                            func.strict = True
                        self._functions_for_model[func.name] = func
                        self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                        print(f'Included function {func.name}')
                    except Exception as e:
                        print(f'Could not add function {tool}: {e}')
            model.set_tools(tools=self._tools_for_model)
            model.set_functions(functions=self._functions_for_model)

    def get_members_system_message_content(self, indent: int = 0) -> str:
        system_message_content = ''
        for idx, member in enumerate(self.members):
            if isinstance(member, Team):
                system_message_content += f'{indent * " "} - Team: {member.name}\n'
                if member.members:
                    system_message_content += member.get_members_system_message_content(indent=indent + 2)
            else:
                system_message_content += f'{indent * " "} - Agent {idx + 1}:\n'
                if member.name:
                    system_message_content += f'{indent * " "}   - Name: {member.name}\n'
                if member.role:
                    system_message_content += f'{indent * " "}   - Role: {member.role}\n'
                if member.description:
                    system_message_content += f'{indent * " "}   - Description: {member.description}\n'
                if member.tools:
                    system_message_content += f'{indent * " "}   - Available tools:\n'
                    tool_name_and_description = []
                    for _tool in member.tools:
                        if isinstance(_tool, Toolkit):
                            for _func in _tool.functions.values():
                                if _func.entrypoint:
                                    tool_name_and_description.append(
                                        (_func.name, Function.get_entrypoint_docstring(_func.entrypoint)))
                        elif isinstance(_tool, Function) and _tool.entrypoint:
                            tool_name_and_description.append(
                                (_tool.name, Function.get_entrypoint_docstring(_tool.entrypoint)))
                        elif callable(_tool):
                            tool_name_and_description.append((_tool.__name__, Function.get_entrypoint_docstring(_tool)))
                    for _tool_name, _tool_description in tool_name_and_description:
                        system_message_content += f'{indent * " "}    - {_tool_name}: {_tool_description}\n'
        return system_message_content

    def get_system_message(self, audio: Sequence[Audio] = None, images: Sequence[Image] = None,
                           videos: Sequence[Video] = None, files: Sequence[File] = None) -> Optional[Message]:
        instructions: List[str] = []
        if self.instructions:
            _instructions = self.instructions
            if callable(self.instructions):
                _instructions = self.instructions(agent=self)
            if isinstance(_instructions, str):
                instructions.append(_instructions)
            elif isinstance(_instructions, list):
                instructions.extend(_instructions)
        _model_instructions = self.model.get_instructions_for_model()
        if _model_instructions:
            instructions.extend(_model_instructions)
        additional_information: List[str] = []
        if self.markdown and self.response_model is None:
            additional_information.append('Use markdown to format your answers.')
        if self.add_datetime_to_instructions:
            additional_information.append(f'The current time is {datetime.now()}')
        system_message_content: str = ''
        if self.mode == 'coordinate':
            system_message_content += '您是AI代理团队和可能的子团队的领导者:\n'
            system_message_content += self.get_members_system_message_content()
            system_message_content += (
                '''\n-您可以直接响应，也可以将任务转移给团队中的其他代理，具体取决于他们可用的工具及其角色。\n-如果将任务转移给另一个代理，请确保包括：\n-agent_name（str）：要将任务传输到的代理的名称。\n-task_description（str）：任务的清晰描述。\n-expected_output（str）：预期输出。\n-您可以同时将任务传递给多个成员。\n-在响应用户之前，您必须始终验证其他代理的输出。\n-评估其他代理人的反应。如果你觉得任务已经完成，你可以停下来回应用户。\n如果你对结果不满意，可以重新分配任务。\n''')
        elif self.mode == 'route':
            system_message_content += '您是AI代理团队和可能的子团队的领导者:\n'
            system_message_content += self.get_members_system_message_content()
            system_message_content += '-您充当用户请求的路由器。您必须选择正确的代理来转发用户的请求。这应该是完成任务可能性最高的代理。\n-将任务转发给另一个代理时，请确保包括：\n-agent_name（str）：要将任务传输到的代理的名称。\n-expected_output（str）：预期输出。\n你应该尽最大努力把任务交给一个代理人。\n-如果用户请求需要它（例如，如果他们要求多个东西），您可以一次转发给多个代理。\n'
        elif self.mode == 'collaborate':
            system_message_content += '您正在领导一个由代理和可能的子团队组成的协作团队:\n'
            system_message_content += self.get_members_system_message_content()
            system_message_content += '-对于团队中的所有代理，只调用run_member_agent一次。\n-考虑其他代理的所有回复，并评估任务是否已完成。\n如果你觉得任务已经完成，你可以停下来回应用户。\n'
            system_message_content += '\n'
        if self.enable_agentic_context:
            system_message_content += '你可以而且应该更新团队的背景。使用`set_team_text`工具更新共享团队上下文.\n'
        if self.name:
            system_message_content += f'你是: {self.name}.\n\n'
        if self.success_criteria:
            system_message_content += f'<success_criteria>\n如果满足以下标准，该团队将被视为成功: {self.success_criteria}\n满足条件后停止团队运行.\n</success_criteria>\n\n'
        if self.description:
            system_message_content += f'<description>\n{self.description}\n</description>\n\n'
        if len(instructions) > 0:
            system_message_content += '<instructions>'
            if len(instructions) > 1:
                for _upi in instructions:
                    system_message_content += f'\n- {_upi}'
            else:
                system_message_content += '\n' + instructions[0]
            system_message_content += '\n</instructions>\n\n'
        if len(additional_information) > 0:
            system_message_content += '<additional_information>'
            for _ai in additional_information:
                system_message_content += f'\n- {_ai}'
            system_message_content += '\n</additional_information>\n\n'
        if audio or images or videos or files:
            system_message_content += '<attached_media>\n'
            system_message_content += 'You have the following media attached to your message:\n'
            if audio and len(audio) > 0:
                system_message_content += ' - Audio\n'
            if images and len(images) > 0:
                system_message_content += ' - Images\n'
            if videos and len(videos) > 0:
                system_message_content += ' - Videos\n'
            if files and len(files) > 0:
                system_message_content += ' - Files\n'
            system_message_content += '</attached_media>\n\n'
        if self.add_state_in_messages:
            system_message_content = self._format_message_with_state_variables(system_message_content)
        system_message_from_model = self.model.get_system_message_for_model()
        if system_message_from_model:
            system_message_content += system_message_from_model
        if self.expected_output:
            system_message_content += f'<expected_output>\n{self.expected_output.strip()}\n</expected_output>\n\n'
        if self.response_model and self.use_json_mode and self.model and self.model.supports_native_structured_outputs:
            system_message_content += f'{self._get_json_output_prompt()}'
        return Message(role='system', content=system_message_content.strip())

    def get_run_messages(self, *, run_response: TeamRunResponse, message: Union[str, List, Dict, Message],
                         audio: Sequence[Audio] = None, images: Sequence[Image] = None,
                         videos: Sequence[Video] = None, files: Sequence[File] = None,
                         **kwargs: Any) -> RunMessages:
        """此函数返回具有以下属性的RunMessages对象：
        -system_message：此运行的系统消息
        -user_message：此运行的用户消息
        -messages：要发送到模型的消息列表
        要构建RunMessages对象，请执行以下操作：
        1.将系统消息添加到run_message
        2.向run_message添加历史记录
        3.将用户消息添加到run_message
        """
        run_messages = RunMessages()
        system_message = self.get_system_message(images=images, audio=audio, videos=videos, files=files)
        if system_message:
            run_messages.system_message = system_message
            run_messages.messages.append(system_message)
        if self.enable_team_history:
            history: List[Message] = self.memory.get_messages_from_last_n_runs(
                last_n=self.num_of_interactions_from_history, skip_role='system')
            if len(history) > 0:
                history_copy = [deepcopy(msg) for msg in history]
                for _msg in history_copy:
                    _msg.from_history = True
                print(f'Adding {len(history_copy)} messages from history')
                run_messages.messages += history_copy
        user_message = self._get_user_message(message, audio=audio, images=images, videos=videos, files=files, **kwargs)
        if user_message:
            run_messages.user_message = user_message
            run_messages.messages.append(user_message)
        return run_messages

    def _get_user_message(self, message: Union[str, List, Dict, Message], audio: Sequence[Audio] = None,
                          images: Sequence[Image] = None, videos: Sequence[Video] = None,
                          files: Sequence[File] = None, **kwargs):
        user_message_content: str = ''
        if isinstance(message, str) or isinstance(message, list):
            if self.add_state_in_messages:
                if isinstance(message, str):
                    user_message_content = self._format_message_with_state_variables(message)
                elif isinstance(message, list):
                    user_message_content = '\n'.join(
                        [self._format_message_with_state_variables(msg) for msg in message])
            else:
                if isinstance(message, str):
                    user_message_content = message
                else:
                    user_message_content = '\n'.join(message)
            if self.add_context and self.context:
                user_message_content += '\n\n<context>\n'
                user_message_content += self._convert_context_to_string(self.context) + '\n'
                user_message_content += '</context>'
            return Message(role='user', content=user_message_content, audio=audio, images=images, videos=videos,
                           files=files, **kwargs)
        elif isinstance(message, Message):
            return message
        elif isinstance(message, dict):
            try:
                return Message.model_validate(message)
            except Exception as e:
                print(f'Failed to validate message: {e}')

    def _format_message_with_state_variables(self, message: str) -> Any:
        format_variables = collections.ChainMap(self.session_state or {}, self.context or {}, self.extra_data or {},
                                    {'user_id': self.user_id} if self.user_id else {})
        return SafeFormatter().format(message, **format_variables)

    def _convert_context_to_string(self, context: Dict[str, Any]) -> str:
        if context is None:
            return ''
        try:
            return json.dumps(context, indent=2, default=str)
        except (TypeError, ValueError, OverflowError) as e:
            print(f'Failed to convert context to JSON: {e}')
            sanitized_context = {}
            for key, value in context.items():
                try:
                    json.dumps({key: value}, default=str)
                    sanitized_context[key] = value
                except Exception as e:
                    print(f'Failed to serialize to JSON: {e}')
                    sanitized_context[key] = str(value)
            try:
                return json.dumps(sanitized_context, indent=2)
            except Exception as e:
                print(f'Failed to convert sanitized context to JSON: {e}')
                return str(context)

    def _get_json_output_prompt(self) -> str:
        json_output_prompt = 'Provide your output as a JSON containing the following fields:'
        if self.response_model:
            if isinstance(self.response_model, str):
                json_output_prompt += '\n<json_fields>'
                json_output_prompt += f'\n{self.response_model}'
                json_output_prompt += '\n</json_fields>'
            elif isinstance(self.response_model, list):
                json_output_prompt += '\n<json_fields>'
                json_output_prompt += f'\n{json.dumps(self.response_model)}'
                json_output_prompt += '\n</json_fields>'
            elif issubclass(self.response_model, Base):
                json_schema = self.response_model.model_json_schema()
                if json_schema:
                    response_model_properties = {}
                    json_schema_properties = json_schema.get('properties')
                    if json_schema_properties:
                        for field_name, field_properties in json_schema_properties.items():
                            formatted_field_properties = {prop_name: prop_value
                                                          for prop_name, prop_value in field_properties.items()
                                                          if prop_name != 'title'}
                            response_model_properties[field_name] = formatted_field_properties
                    json_schema_defs = json_schema.get('$defs')
                    if json_schema_defs:
                        response_model_properties['$defs'] = {}
                        for def_name, def_properties in json_schema_defs.items():
                            def_fields = def_properties.get('properties')
                            formatted_def_properties = {}
                            if def_fields:
                                for field_name, field_properties in def_fields.items():
                                    formatted_field_properties = {prop_name: prop_value
                                                                  for prop_name, prop_value in field_properties.items()
                                                                  if prop_name != 'title'}
                                    formatted_def_properties[field_name] = formatted_field_properties
                            if len(formatted_def_properties) > 0:
                                response_model_properties['$defs'][def_name] = formatted_def_properties
                    if len(response_model_properties) > 0:
                        json_output_prompt += '\n<json_fields>'
                        json_output_prompt += (
                            f'\n{json.dumps([key for key in response_model_properties.keys() if key != "$defs"])}')
                        json_output_prompt += '\n</json_fields>'
                        json_output_prompt += '\n\nHere are the properties for each field:'
                        json_output_prompt += '\n<json_field_properties>'
                        json_output_prompt += f'\n{json.dumps(response_model_properties, indent=2)}'
                        json_output_prompt += '\n</json_field_properties>'
            else:
                print(f'Could not build json schema for {self.response_model}')
        else:
            json_output_prompt += 'Provide the output as JSON.'
        json_output_prompt += '\nStart your response with `{` and end it with `}`.'
        json_output_prompt += '\nYour output will be passed to json.loads() to convert it to a Python object.'
        json_output_prompt += '\nMake sure it only contains valid JSON.'
        return json_output_prompt

    def _update_team_state(self, run_response: Union[TeamRunResponse, RunResponse]) -> None:
        if run_response.images:
            if self.images is None:
                self.images = []
            self.images.extend(run_response.images)
        if run_response.videos:
            if self.videos is None:
                self.videos = []
            self.videos.extend(run_response.videos)
        if run_response.audio:
            if self.audio is None:
                self.audio = []
            self.audio.extend(run_response.audio)

    def get_team_history(self, num_chats: int = None) -> str:
        history: List[Dict[str, Any]] = []
        all_chats = self.memory.get_all_messages()
        if len(all_chats) == 0:
            return ''
        chats_added = 0
        for chat in all_chats[::-1]:
            history.insert(0, chat[1].to_dict())
            history.insert(0, chat[0].to_dict())
            chats_added += 1
            if num_chats and chats_added >= num_chats:
                break
        return json.dumps(history)

    def set_team_context(self, state: Union[str, dict]) -> str:
        if isinstance(state, str):
            self.memory.set_team_context_text(state)
        elif isinstance(state, dict):
            self.memory.set_team_context_text(json.dumps(state))
        msg = f'Current team context: {self.memory.get_team_context_str()}'
        print(msg)
        return msg

    def get_run_member_agents_function(self, stream: bool = False, async_mode: bool = False,
                                       images: List[Image] = None, videos: List[Video] = None,
                                       audio: List[Audio] = None,
                                       files: List[File] = None) -> Function:
        if not images:
            images = []
        if not videos:
            videos = []
        if not audio:
            audio = []
        if not files:
            files = []

        def run_member_agents(task_description: str, expected_output: str = None) -> Iterator[str]:
            team_context_str = None
            if self.enable_agentic_context:
                team_context_str = self.memory.get_team_context_str()
            team_member_interactions_str = None
            if self.share_member_interactions:
                team_member_interactions_str = self.memory.get_team_member_interactions_str()
                if context_images := self.memory.get_team_context_images():
                    images.extend([Image.from_artifact(img) for img in context_images])
                if context_videos := self.memory.get_team_context_videos():
                    videos.extend([Video.from_artifact(vid) for vid in context_videos])
                if context_audio := self.memory.get_team_context_audio():
                    audio.extend([Audio.from_artifact(aud) for aud in context_audio])
            member_agent_task = '您是协作完成任务的代理团队的成员.'
            if expected_output:
                member_agent_task += f'\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if team_context_str:
                member_agent_task += f'\n\n{team_context_str}'
            if team_member_interactions_str:
                member_agent_task += f'\n\n{team_member_interactions_str}'
            member_agent_task += f'\n\n<task>\n{task_description}\n</task>'
            for member_agent_index, member_agent in enumerate(self.members):
                if stream:
                    member_agent_run_response_stream = member_agent.run(member_agent_task, images=images, videos=videos,
                                                                        audio=audio, files=files, stream=True)
                    for member_agent_run_response_chunk in member_agent_run_response_stream:
                        check_if_run_cancelled(member_agent_run_response_chunk)
                        yield member_agent_run_response_chunk.content or ''
                else:
                    member_agent_run_response = member_agent.run(member_agent_task, images=images, videos=videos,
                                                                 audio=audio, files=files, stream=False)
                    check_if_run_cancelled(member_agent_run_response)
                    if member_agent_run_response.content is None:
                        yield 'No response from the member agent.'
                    elif isinstance(member_agent_run_response.content, str):
                        yield member_agent_run_response.content
                    elif issubclass(type(member_agent_run_response.content), Base):
                        try:
                            yield json.dumps(member_agent_run_response.content.to_dict())
                        except Exception as e:
                            yield str(e)
                    else:
                        try:
                            yield json.dumps(member_agent_run_response.content, indent=2)
                        except Exception as e:
                            yield str(e)
                member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
                self.memory.add_interaction_to_team_context(member_name=member_name, task=task_description,
                                                            run_response=member_agent.run_response)
                self.run_response.add_member_run(member_agent.run_response)
                self._update_team_state(member_agent.run_response)

        async def arun_member_agents(task_description: str, expected_output: str = None) -> AsyncIterator[
            str]:
            team_context_str = None
            if self.enable_agentic_context:
                team_context_str = self.memory.get_team_context_str()
            team_member_interactions_str = None
            if self.share_member_interactions:
                team_member_interactions_str = self.memory.get_team_member_interactions_str()
                if context_images := self.memory.get_team_context_images():
                    images.extend([Image.from_artifact(img) for img in context_images])
                if context_videos := self.memory.get_team_context_videos():
                    videos.extend([Video.from_artifact(vid) for vid in context_videos])
                if context_audio := self.memory.get_team_context_audio():
                    audio.extend([Audio.from_artifact(aud) for aud in context_audio])
            member_agent_task = '您是协作完成任务的代理团队的成员.'
            if expected_output:
                member_agent_task += f'\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if team_context_str:
                member_agent_task += f'\n\n{team_context_str}'
            if team_member_interactions_str:
                member_agent_task += f'\n\n{team_member_interactions_str}'
            member_agent_task += f'\n\n<task>\n{task_description}\n</task>'
            tasks = []
            for member_agent_index, member_agent in enumerate(self.members):
                current_agent = member_agent
                current_index = member_agent_index

                async def run_member_agent(agent=current_agent, idx=current_index) -> str:
                    response = await agent.arun(member_agent_task, images=images, videos=videos, audio=audio,
                                                files=files, stream=False)
                    check_if_run_cancelled(response)
                    member_name = agent.name if agent.name else f'agent_{idx}'
                    self.memory.add_interaction_to_team_context(member_name=member_name, task=task_description,
                                                                run_response=agent.run_response)
                    self.run_response.add_member_run(agent.run_response)
                    self._update_team_state(agent.run_response)
                    if response.content is None:
                        return f'Agent {member_name}: No response from the member agent.'
                    elif isinstance(response.content, str):
                        return f'Agent {member_name}: {response.content}'
                    elif issubclass(type(response.content), Base):
                        try:
                            return f'Agent {member_name}: {json.dumps(response.content.to_dict())}'
                        except Exception as e:
                            return f'Agent {member_name}: Error - {str(e)}'
                    else:
                        try:
                            return f'Agent {member_name}: {json.dumps(response.content, indent=2)}'
                        except Exception as e:
                            return f'Agent {member_name}: Error - {str(e)}'

                tasks.append(run_member_agent)
            results = await asyncio.gather(*[task() for task in tasks])
            for result in results:
                yield result

        if async_mode:
            run_member_agents_function = arun_member_agents
        else:
            run_member_agents_function = run_member_agents
        run_member_agents_func = Function.from_callable(run_member_agents_function, strict=True)
        return run_member_agents_func

    def get_transfer_task_function(self, stream: bool = False, async_mode: bool = False,
                                   images: List[Image] = None, videos: List[Video] = None,
                                   audio: List[Audio] = None, files: List[File] = None) -> Function:
        if not images:
            images = []
        if not videos:
            videos = []
        if not audio:
            audio = []
        if not files:
            files = []

        def transfer_task_to_member(agent_name: str, task_description: str, expected_output: str) -> Iterator[str]:
            result = self._find_member_by_name(agent_name)
            if result is None:
                yield f'Agent with name {agent_name} not found in the team or any subteams. Please choose the correct agent from the list of agents.'
                return
            member_agent_index, member_agent = result
            self._initialize_member(member_agent)
            team_context_str = None
            if self.enable_agentic_context:
                team_context_str = self.memory.get_team_context_str()
            team_member_interactions_str = None
            if self.share_member_interactions:
                team_member_interactions_str = self.memory.get_team_member_interactions_str()
                if context_images := self.memory.get_team_context_images():
                    images.extend([Image.from_artifact(img) for img in context_images])
                if context_videos := self.memory.get_team_context_videos():
                    videos.extend([Video.from_artifact(vid) for vid in context_videos])
                if context_audio := self.memory.get_team_context_audio():
                    audio.extend([Audio.from_artifact(aud) for aud in context_audio])
            member_agent_task = f'你是一个特工团队的成员。你的目标是完成以下任务:\n\n{task_description}\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if team_context_str:
                member_agent_task += f'\n\n{team_context_str}'
            if team_member_interactions_str:
                member_agent_task += f'\n\n{team_member_interactions_str}'
            if stream:
                member_agent_run_response_stream = member_agent.run(member_agent_task, images=images, videos=videos,
                                                                    audio=audio, files=files, stream=True)
                for member_agent_run_response_chunk in member_agent_run_response_stream:
                    check_if_run_cancelled(member_agent_run_response_chunk)
                    yield member_agent_run_response_chunk.content or ''
            else:
                member_agent_run_response = member_agent.run(member_agent_task, images=images, videos=videos,
                                                             audio=audio, files=files, stream=False)
                check_if_run_cancelled(member_agent_run_response)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), Base):
                    try:
                        yield json.dumps(member_agent_run_response.content.to_dict())
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
            self.memory.add_interaction_to_team_context(member_name=member_name, task=task_description,
                                                        run_response=member_agent.run_response)
            self.run_response.add_member_run(member_agent.run_response)
            self._update_team_state(member_agent.run_response)

        async def atransfer_task_to_member(agent_name: str, task_description: str, expected_output: str) -> \
        AsyncIterator[str]:
            result = self._find_member_by_name(agent_name)
            if result is None:
                yield f'Agent with name {agent_name} not found in the team or any subteams. Please choose the correct agent from the list of agents.'
                return
            member_agent_index, member_agent = result
            self._initialize_member(member_agent)
            team_context_str = None
            if self.enable_agentic_context:
                team_context_str = self.memory.get_team_context_str()
            team_member_interactions_str = None
            if self.share_member_interactions:
                team_member_interactions_str = self.memory.get_team_member_interactions_str()
                if context_images := self.memory.get_team_context_images():
                    images.extend([Image.from_artifact(img) for img in context_images])
                if context_videos := self.memory.get_team_context_videos():
                    videos.extend([Video.from_artifact(vid) for vid in context_videos])
                if context_audio := self.memory.get_team_context_audio():
                    audio.extend([Audio.from_artifact(aud) for aud in context_audio])
            member_agent_task = f'你是一个特工团队的成员。你的目标是完成以下任务:\n\n{task_description}\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if team_context_str:
                member_agent_task += f'\n\n{team_context_str}'
            if team_member_interactions_str:
                member_agent_task += f'\n\n{team_member_interactions_str}'
            if stream:
                member_agent_run_response_stream = await member_agent.arun(member_agent_task, images=images,
                                                                           videos=videos, audio=audio, files=files,
                                                                           stream=True)
                async for member_agent_run_response_chunk in member_agent_run_response_stream:
                    check_if_run_cancelled(member_agent_run_response_chunk)
                    yield member_agent_run_response_chunk.content or ''
            else:
                member_agent_run_response = await member_agent.arun(member_agent_task, images=images, videos=videos,
                                                                    audio=audio, files=files, stream=False)
                check_if_run_cancelled(member_agent_run_response)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), Base):
                    try:
                        yield json.dumps(member_agent_run_response.content.to_dict())
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
            self.memory.add_interaction_to_team_context(member_name=member_name, task=task_description,
                                                        run_response=member_agent.run_response)
            self.run_response.add_member_run(member_agent.run_response)
            self._update_team_state(member_agent.run_response)

        if async_mode:
            transfer_function = atransfer_task_to_member
        else:
            transfer_function = transfer_task_to_member
        transfer_func = Function.from_callable(transfer_function, strict=True)
        return transfer_func

    def _find_member_by_name(self, agent_name: str):
        for i, member in enumerate(self.members):
            if member.name == agent_name:
                return (i, member)
            if isinstance(member, Team):
                result = member._find_member_by_name(agent_name)
                if result:
                    return (i, member)
        return None

    def get_forward_task_function(self, message: Message, stream: bool = False, async_mode: bool = False,
                                  images: Sequence[Image] = None, videos: Sequence[Video] = None,
                                  audio: Sequence[Audio] = None,
                                  files: Sequence[File] = None) -> Function:
        if not images:
            images = []
        if not videos:
            videos = []
        if not audio:
            audio = []
        if not files:
            files = []

        def forward_task_to_member(agent_name: str, expected_output: str = None) -> Iterator[str]:
            self._member_response_model = None
            result = self._find_member_by_name(agent_name)
            if result is None:
                yield f'Agent with name {agent_name} not found in the team or any subteams. Please choose the correct agent from the list of agents.'
                return
            member_agent_index, member_agent = result
            self._initialize_member(member_agent)
            if member_agent.response_model:
                self._member_response_model = member_agent.response_model
            member_agent_task = message.get_content_string()
            if expected_output:
                member_agent_task += f'\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if stream:
                member_agent_run_response_stream = member_agent.run(member_agent_task, images=images, videos=videos,
                                                                    audio=audio, files=files, stream=True)
                for member_agent_run_response_chunk in member_agent_run_response_stream:
                    yield member_agent_run_response_chunk.content or ''
            else:
                member_agent_run_response = member_agent.run(member_agent_task, images=images, videos=videos,
                                                             audio=audio, files=files, stream=False)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), Base):
                    try:
                        yield json.dumps(member_agent_run_response.content.to_dict())
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
            self.memory.add_interaction_to_team_context(member_name=member_name, task=message.get_content_string(),
                                                        run_response=member_agent.run_response)
            self.run_response.add_member_run(member_agent.run_response)
            self._update_team_state(member_agent.run_response)

        async def aforward_task_to_member(agent_name: str, expected_output: str = None) -> AsyncIterator[str]:
            self._member_response_model = None
            result = self._find_member_by_name(agent_name)
            if result is None:
                yield f'Agent with name {agent_name} not found in the team or any subteams. Please choose the correct agent from the list of agents.'
                return
            member_agent_index, member_agent = result
            self._initialize_member(member_agent)
            if member_agent.response_model:
                self._member_response_model = member_agent.response_model
            member_agent_task = message.get_content_string()
            if expected_output:
                member_agent_task += f'\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if stream:
                member_agent_run_response_stream = await member_agent.arun(member_agent_task, images=images,
                                                                           videos=videos, audio=audio, files=files,
                                                                           stream=True)
                async for member_agent_run_response_chunk in member_agent_run_response_stream:
                    check_if_run_cancelled(member_agent_run_response_chunk)
                    yield member_agent_run_response_chunk.content or ''
            else:
                member_agent_run_response = await member_agent.arun(member_agent_task, images=images, videos=videos,
                                                                    audio=audio, files=files, stream=False)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), Base):
                    try:
                        yield json.dumps(member_agent_run_response.content.to_dict())
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
            self.memory.add_interaction_to_team_context(member_name=member_name, task=message.get_content_string(),
                                                        run_response=member_agent.run_response)
            self.run_response.add_member_run(member_agent.run_response)
            self._update_team_state(member_agent.run_response)

        if async_mode:
            forward_function = aforward_task_to_member
        else:
            forward_function = forward_task_to_member
        forward_func = Function.from_callable(forward_function, strict=True)
        forward_func.stop_after_tool_call = True
        forward_func.show_result = True
        return forward_func

    def load_user_memories(self) -> None:
        if self.memory and self.memory.create_user_memories:
            if self.user_id:
                self.memory.user_id = self.user_id
            self.memory.load_user_memories()
            if self.user_id:
                print(f'Memories loaded for user: {self.user_id}')
            else:
                print('Memories loaded')

    def read_from_storage(self) -> Optional[AgentSession]:
        if self.storage and self.session_id:
            self.team_session = self.storage.read(session_id=self.session_id)
            if self.team_session:
                self.load_team_session(session=self.team_session)
            self.load_user_memories()
        return self.team_session

    def write_to_storage(self) -> Optional[AgentSession]:
        if self.storage:
            self.team_session = self.storage.upsert(session=self._get_team_session())
        return self.team_session

    def rename_session(self, session_name: str) -> None:
        self.read_from_storage()
        self.session_name = session_name
        self.write_to_storage()
        self._log_team_session()

    def delete_session(self, session_id: str) -> None:
        if self.storage:
            self.storage.delete_session(session_id=session_id)

    def load_team_session(self, session: AgentSession):
        if self.team_id is None and session.team_id:
            self.team_id = session.team_id
        if self.user_id is None and session.user_id:
            self.user_id = session.user_id
        if self.session_id is None and session.session_id:
            self.session_id = session.session_id
        if session.team_data:
            if self.name is None and 'name' in session.team_data:
                self.name = session.team_data.get('name')
        if session.session_data:
            if self.session_name is None and 'session_name' in session.session_data:
                self.session_name = session.session_data.get('session_name')
            if 'session_state' in session.session_data:
                session_state_from_db = session.session_data.get('session_state')
                if session_state_from_db and isinstance(session_state_from_db, dict) and len(
                        session_state_from_db) > 0:
                    if self.session_state and len(self.session_state) > 0:
                        merge_dictionaries(session_state_from_db, self.session_state)
                    self.session_state = session_state_from_db
            if 'images' in session.session_data:
                images_from_db = session.session_data.get('images')
                if images_from_db and isinstance(images_from_db, list):
                    if self.images is None:
                        self.images = []
                    self.images.extend([ImageArtifact.model_validate(img) for img in images_from_db])
            if 'videos' in session.session_data:
                videos_from_db = session.session_data.get('videos')
                if videos_from_db and isinstance(videos_from_db, list):
                    if self.videos is None:
                        self.videos = []
                    self.videos.extend([VideoArtifact.model_validate(vid) for vid in videos_from_db])
            if 'audio' in session.session_data:
                audio_from_db = session.session_data.get('audio')
                if audio_from_db and isinstance(audio_from_db, list):
                    if self.audio is None:
                        self.audio = []
                    self.audio.extend([AudioArtifact.model_validate(aud) for aud in audio_from_db])
        if session.extra_data:
            if self.extra_data:
                merge_dictionaries(session.extra_data, self.extra_data)
            self.extra_data = session.extra_data
        if self.memory is None:
            self.memory = session.memory
        if not isinstance(self.memory, TeamMemory):
            if isinstance(self.memory, dict):
                self.memory = TeamMemory(**self.memory)
            elif self.memory:
                raise TypeError(f'Expected memory to be a dict or TeamMemory, but got {type(self.memory)}')
        if session.memory and self.memory:
            try:
                if 'runs' in session.memory:
                    try:
                        self.memory.runs = [TeamRun.from_dict(m) for m in session.memory['runs']]
                    except Exception as e:
                        print(f'Failed to load runs from memory: {e}')
                if 'messages' in session.memory:
                    try:
                        self.memory.messages = [Message.model_validate(m) for m in session.memory['messages']]
                    except Exception as e:
                        print(f'Failed to load messages from memory: {e}')
                if 'memories' in session.memory:
                    try:
                        self.memory.memories = [Memory.model_validate(m) for m in session.memory['memories']]
                    except Exception as e:
                        print(f'Failed to load user memories: {e}')
            except Exception as e:
                print(f'Failed to load AgentMemory: {e}')
        print(f'-*- AgentSession loaded: {session.session_id}')

    def _create_run_data(self) -> Dict[str, Any]:
        run_response_format = 'text'
        if self.response_model:
            run_response_format = 'json'
        elif self.markdown:
            run_response_format = 'markdown'
        functions = {}
        if self.model and self.model._functions:
            functions = {f_name: func.to_dict() for f_name, func in self.model._functions.items() if
                         isinstance(func, Function)}
        run_data: Dict[str, Any] = {'functions': functions, 'metrics': self.run_response.metrics}
        if self.monitoring:
            run_data.update({'run_input': self.run_input, 'run_response': self.run_response.to_dict(),
                             'run_response_format': run_response_format})
        return run_data

    def _get_team_data(self) -> Dict[str, Any]:
        team_data: Dict[str, Any] = {}
        if self.name:
            team_data['name'] = self.name
        if self.team_id:
            team_data['team_id'] = self.team_id
        if self.model:
            team_data['model'] = self.model.to_dict()
        return team_data

    def _get_session_data(self) -> Dict[str, Any]:
        session_data: Dict[str, Any] = {}
        if self.session_name:
            session_data['session_name'] = self.session_name
        if self.session_state and len(self.session_state) > 0:
            session_data['session_state'] = self.session_state
        if self.images:
            session_data['images'] = [img.to_dict() for img in self.images]
        if self.videos:
            session_data['videos'] = [vid.to_dict() for vid in self.videos]
        if self.audio:
            session_data['audio'] = [aud.to_dict() for aud in self.audio]
        return session_data

    def _get_team_session(self) -> AgentSession:
        return AgentSession(mode='team', session_id=self.session_id, agent_id=self.team_id, user_id=self.user_id,
                           agent_session_id=self.team_session_id,
                           memory=self.memory.to_dict() if self.memory else None,
                           team_data=self._get_team_data(), session_data=self._get_session_data(),
                           extra_data=self.extra_data, created_at=int(time.time()))

    def _log_team_run(self) -> None:
        if not self.telemetry and not self.monitoring:
            return
        try:
            run_data = self._create_run_data()
            team_session = self.team_session or self._get_team_session()
        except Exception as e:
            print(f'Could not create team event: {e}')

    async def _alog_team_run(self) -> None:
        if not self.telemetry and not self.monitoring:
            return
        try:
            run_data = self._create_run_data()
            team_session = self.team_session or self._get_team_session()
        except Exception as e:
            print(f'Could not create team event: {e}')

    def _log_team_session(self):
        if not (self.telemetry or self.monitoring):
            return
        try:
            team_session = self.team_session or self._get_team_session()
        except Exception as e:
            print(f'Could not create team monitor: {e}')


class Workflow:
    """工作流"""
    def __init__(self, name: str = None, workflow_id: str = None, description: str = None, user_id: str = None, session_id: str = None, session_name: str = None, session_state: Dict[str, Any] = None, memory: WorkflowMemory = None, extra_data: Dict[str, Any] = None, debug_mode: bool = False, monitoring: bool = False, telemetry: bool = True):
        self.name = name or self.__class__.__name__
        self.workflow_id = workflow_id
        self.description = description or self.__class__.__doc__
        self.user_id = user_id
        self.session_id = session_id
        self.session_name = session_name
        self.session_state: Dict[str, Any] = session_state or {}
        self.memory = memory
        self.extra_data = extra_data
        self.debug_mode = debug_mode
        self.monitoring = monitoring
        self.telemetry = telemetry
        self.run_id = None
        self.run_input = None
        self.run_response = None
        self.images = None
        self.videos = None
        self.audio = None
        self.workflow_session: AgentSession = None
        self._subclass_run: Callable = None
        self._run_parameters: Dict[str, Any] = None
        self._run_return_type: str = None
        self.update_run_method()
        for field_name, value in self.__class__.__dict__.items():
            if isinstance(value, Agent):
                value.session_id = self.session_id

    def run(self, **kwargs: Any):
        print(f'{self.__class__.__name__}.run() method not implemented.')
        return

    def run_workflow(self, **kwargs: Any):
        self.set_debug()
        self.set_workflow_id()
        self.set_session_id()
        self.initialize_memory()
        self.run_id = str(uuid.uuid4())
        self.run_input = kwargs
        self.run_response = RunResponse(run_id=self.run_id, session_id=self.session_id, workflow_id=self.workflow_id)
        self.update_agent_session_ids()
        print(f'*********** Workflow Run Start: {self.run_id} ***********')
        try:
            result = self._subclass_run(**kwargs)
        except Exception as e:
            print(f'Workflow.run() failed: {e}')
            raise e
        if isinstance(result, (types.GeneratorType, collections.abc.Iterator)):
            self.run_response.content = ''
            def result_generator():
                for item in result:
                    if isinstance(item, RunResponse):
                        item.run_id = self.run_id
                        item.session_id = self.session_id
                        item.workflow_id = self.workflow_id
                        if item.content and isinstance(item.content, str):
                            self.run_response.content += item.content
                    else:
                        print(f'Workflow.run() should only yield RunResponse objects, got: {type(item)}')
                    yield item
                self.memory.add_run(WorkflowRun(input=self.run_input, response=self.run_response))
                print(f'*********** Workflow Run End: {self.run_id} ***********')
            return result_generator()
        elif isinstance(result, RunResponse):
            result.run_id = self.run_id
            result.session_id = self.session_id
            result.workflow_id = self.workflow_id
            if result.content and isinstance(result.content, str):
                self.run_response.content = result.content
            self.memory.add_run(WorkflowRun(input=self.run_input, response=self.run_response))
            print(f'*********** Workflow Run End: {self.run_id} ***********')
            return result
        else:
            print(f'Workflow.run() should only return RunResponse objects, got: {type(result)}')
            return None

    def set_workflow_id(self) -> str:
        if self.workflow_id is None:
            self.workflow_id = str(uuid.uuid4())
        print(f'*********** Workflow ID: {self.workflow_id} ***********')
        return self.workflow_id

    def set_session_id(self) -> str:
        if self.session_id is None:
            self.session_id = str(uuid.uuid4())
        print(f'*********** Session ID: {self.session_id} ***********')
        return self.session_id

    def set_debug(self) -> None:
        if self.debug_mode or os.getenv('AGNO_DEBUG', 'false').lower() == 'true':
            self.debug_mode = True
            print('Debug logs enabled')

    def set_monitoring(self) -> None:
        if self.monitoring or os.getenv('AGNO_MONITOR', 'false').lower() == 'true':
            self.monitoring = True
        else:
            self.monitoring = False
        if self.telemetry or os.getenv('AGNO_TELEMETRY', 'true').lower() == 'true':
            self.telemetry = True
        else:
            self.telemetry = False

    def initialize_memory(self) -> None:
        if self.memory is None:
            self.memory = WorkflowMemory()

    def update_run_method(self):
        if self.__class__.run is not Workflow.run:
            self._subclass_run = self.__class__.run.__get__(self)
            sig = inspect.signature(self.__class__.run)
            self._run_parameters = {param_name: {'name': param_name, 'default': param.default.default
                    if hasattr(param.default, '__class__') and param.default.__class__.__name__ == 'FieldInfo'
                    else (param.default if param.default is not inspect.Parameter.empty else None), 'annotation': (param.annotation.__name__
                        if hasattr(param.annotation, '__name__')
                        else (str(param.annotation).replace('typing.Optional[', '').replace(']', '')
                            if 'typing.Optional' in str(param.annotation)
                            else str(param.annotation)))
                    if param.annotation is not inspect.Parameter.empty
                    else None, 'required': param.default is inspect.Parameter.empty}
                for param_name, param in sig.parameters.items()
                if param_name != 'self'}
            return_annotation = sig.return_annotation
            self._run_return_type = (return_annotation.__name__
                if return_annotation is not inspect.Signature.empty and hasattr(return_annotation, '__name__')
                else str(return_annotation)
                if return_annotation is not inspect.Signature.empty
                else None)
            object.__setattr__(self, 'run', self.run_workflow.__get__(self))
        else:
            self._subclass_run = self.run
            self._run_parameters = {}
            self._run_return_type = None

    def update_agent_session_ids(self):
        for field_name, value in self.__class__.__dict__.items():
            if isinstance(value, Agent):
                field_value = getattr(self, field_name)
                field_value.session_id = self.session_id

    def get_workflow_data(self) -> Dict[str, Any]:
        workflow_data: Dict[str, Any] = {}
        if self.name:
            workflow_data['name'] = self.name
        if self.workflow_id:
            workflow_data['workflow_id'] = self.workflow_id
        if self.description:
            workflow_data['description'] = self.description
        return workflow_data

    def get_session_data(self) -> Dict[str, Any]:
        session_data: Dict[str, Any] = {}
        if self.session_name:
            session_data['session_name'] = self.session_name
        if self.session_state and len(self.session_state) > 0:
            session_data['session_state'] = self.session_state
        if self.images:
            session_data['images'] = [img.to_dict() for img in self.images]
        if self.videos:
            session_data['videos'] = [vid.to_dict() for vid in self.videos]
        if self.audio:
            session_data['audio'] = [aud.to_dict() for aud in self.audio]
        return session_data

    def get_workflow_session(self) -> AgentSession:
        return AgentSession(mode='workflow', session_id=self.session_id, workflow_id=self.workflow_id, user_id=self.user_id, memory=self.memory.to_dict() if self.memory else None, workflow_data=self.get_workflow_data(), session_data=self.get_session_data(), extra_data=self.extra_data)

    def load_workflow_session(self, session: AgentSession):
        if self.workflow_id is None and session.workflow_id:
            self.workflow_id = session.workflow_id
        if self.user_id is None and session.user_id:
            self.user_id = session.user_id
        if self.session_id is None and session.session_id:
            self.session_id = session.session_id
        if session.workflow_data:
            if self.name is None and 'name' in session.workflow_data:
                self.name = session.workflow_data.get('name')
        if session.session_data:
            if self.session_name is None and 'session_name' in session.session_data:
                self.session_name = session.session_data.get('session_name')
            if 'session_state' in session.session_data:
                session_state_from_db = session.session_data.get('session_state')
                if session_state_from_db and isinstance(session_state_from_db, dict) and len(session_state_from_db) > 0:
                    if len(self.session_state) > 0:
                        merge_dictionaries(session_state_from_db, self.session_state)
                    self.session_state = session_state_from_db
            if 'images' in session.session_data:
                images_from_db = session.session_data.get('images')
                if images_from_db and isinstance(images_from_db, list):
                    if self.images is None:
                        self.images = []
                    self.images.extend([ImageArtifact.model_validate(img) for img in images_from_db])
            if 'videos' in session.session_data:
                videos_from_db = session.session_data.get('videos')
                if videos_from_db and isinstance(videos_from_db, list):
                    if self.videos is None:
                        self.videos = []
                    self.videos.extend([VideoArtifact.model_validate(vid) for vid in videos_from_db])
            if 'audio' in session.session_data:
                audio_from_db = session.session_data.get('audio')
                if audio_from_db and isinstance(audio_from_db, list):
                    if self.audio is None:
                        self.audio = []
                    self.audio.extend([AudioArtifact.model_validate(aud) for aud in audio_from_db])
        if session.extra_data:
            if self.extra_data:
                merge_dictionaries(session.extra_data, self.extra_data)
            self.extra_data = session.extra_data
        if session.memory:
            if self.memory is None:
                self.memory = WorkflowMemory()
            try:
                if 'runs' in session.memory:
                    try:
                        self.memory.runs = [WorkflowRun(**m) for m in session.memory['runs']]
                    except Exception as e:
                        print(f'Failed to load runs from memory: {e}')
            except Exception as e:
                print(f'Failed to load WorkflowMemory: {e}')
        print(f'-*- agent session loaded: {session.session_id}')

    def load_session(self, force: bool = False) -> Optional[str]:
        if self.workflow_session and not force:
            if self.session_id and self.workflow_session.session_id == self.session_id:
                return self.workflow_session.session_id
        return self.session_id

    def new_session(self) -> None:
        self.workflow_session = None
        self.session_id = str(uuid.uuid4())
        self.load_session(force=True)

    def log_workflow_session(self):
        print(f'*********** Logging : {self.session_id} ***********')

    def rename(self, name: str) -> None:
        self.name = name
        self.log_workflow_session()

    def rename_session(self, session_name: str):
        self.session_name = session_name
        self.log_workflow_session()


def check_if_run_cancelled(run_response: Union[RunResponse, TeamRunResponse]):
    if run_response.event == RunEvent.run_cancelled:
        raise KeyboardInterrupt()


def merge_dictionaries(a: Dict[str, Any], b: Dict[str, Any]):
    for key in b:
        if key in a and isinstance(a[key], dict) and isinstance(b[key], dict):
            merge_dictionaries(a[key], b[key])
        else:
            a[key] = b[key]


def create_panel(content, title, border_style='blue'):
    return Panel(content, title=title, title_align='left', border_style=border_style, box=HEAVY, expand=True, padding=(1, 1))


def update_run_response_with_reasoning(run_response: Union[RunResponse, TeamRunResponse], reasoning_steps: List[ReasoningStep], reasoning_agent_messages: List[Message]) -> None:
    if run_response.extra_data is None:
        run_response.extra_data = RunResponseExtraData()
    if run_response.extra_data.reasoning_steps is None:
        run_response.extra_data.reasoning_steps = reasoning_steps
    else:
        run_response.extra_data.reasoning_steps.extend(reasoning_steps)
    if run_response.extra_data.reasoning_messages is None:
        run_response.extra_data.reasoning_messages = reasoning_agent_messages
    else:
        run_response.extra_data.reasoning_messages.extend(reasoning_agent_messages)


def format_tool_calls(tool_calls: List[Dict[str, Any]]) -> List[str]:
    formatted_tool_calls = []
    for tool_call in tool_calls:
        if 'tool_name' in tool_call and 'tool_args' in tool_call:
            tool_name = tool_call['tool_name']
            args_str = ''
            if 'tool_args' in tool_call and tool_call['tool_args']:
                args_str = ', '.join(f'{k}={v}' for k, v in tool_call['tool_args'].items())
            formatted_tool_calls.append(f'{tool_name}({args_str})')
    return formatted_tool_calls


def parse_response_model_str(content: str, response_model: Type[Base]) -> Optional[Base]:
    structured_output = None
    try:
        structured_output = response_model(**json.loads(content))
    except json.JSONDecodeError:
        content = content
        if '```json' in content:
            content = content.split('```json')[-1].split('```')[0].strip()
        elif '```' in content:
            content = content.split('```')[1].strip()
        content = re.sub(r'[*_`#]', '', content)
        content = content.replace('\n', ' ').replace('\r', '')
        content = re.sub(r'[\x00-\x1F\x7F]', '', content)
        def escape_quotes_in_values(match):
            key = match.group(1)
            value = match.group(2)
            escaped_value = value.replace('"', '\\"')
            return f'"{key}": {escaped_value}'
        content = re.sub(r'"(?P<key>[^"]+)"\s*:\s*"(?P<value>.*?)(?="\s*(?:,|\}))', escape_quotes_in_values, content)
        try:
            structured_output = response_model(**json.loads(content))
        except json.JSONDecodeError as e:
            print(f'Failed to parse cleaned JSON: {e}')
            try:
                data = json.loads(content)
                structured_output = response_model.model_validate(data)
            except json.JSONDecodeError as e:
                print(f'Failed to parse as Python dict: {e}')
    return structured_output


def get_text_from_message(message: Union[List, Dict, str, Message]) -> str:
    if isinstance(message, str):
        return message
    if isinstance(message, list):
        text_messages = []
        if len(message) == 0:
            return ''
        if 'type' in message[0]:
            for m in message:
                m_type = m.get('type')
                if m_type and isinstance(m_type, str):
                    m_value = m.get(m_type)
                    if m_value and isinstance(m_value, str):
                        if m_type == 'text':
                            text_messages.append(m_value)
                        if m_type == 'image_url':
                            text_messages.append(f'Image: {m_value}')
                        else:
                            text_messages.append(f'{m_type}: {m_value}')
        elif 'role' in message[0]:
            for m in message:
                m_role = m.get('role')
                if m_role and isinstance(m_role, str):
                    m_content = m.get('content')
                    if m_content and isinstance(m_content, str):
                        if m_role == 'user':
                            text_messages.append(m_content)
        if len(text_messages) > 0:
            return '\n'.join(text_messages)
    if isinstance(message, dict):
        if 'content' in message:
            return get_text_from_message(message['content'])
    if isinstance(message, Message) and message.content:
        return get_text_from_message(message.content)
    return ''
