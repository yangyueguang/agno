import random
import time
import csv
import io
import os
import re
import json
import httpx
from io import BytesIO
from pathlib import Path
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup, Tag
from pypdf import PdfReader as DocumentReader
from pypdf.errors import PdfStreamError
from docx import Document as DocxDocument
from typing import Any, Dict, List, Optional, Tuple, IO, Union, Set
from ollama import Client as OllamaClient


class Embedder:
    def __init__(self, id: str = 'llama3.1:8b', dimensions: int = 4096, host: Optional[str] = None,
                 timeout: Optional[Any] = None, options: Optional[Any] = None,
                 client_kwargs: Optional[Dict[str, Any]] = None, ollama_client: Optional[OllamaClient] = None):
        self.id = id
        self.dimensions = dimensions
        self.host = host
        self.timeout = timeout
        self.options = options
        self.client_kwargs = client_kwargs
        self.ollama_client = ollama_client

    @property
    def client(self) -> OllamaClient:
        if self.ollama_client:
            return self.ollama_client
        _ollama_params: Dict[str, Any] = {'host': self.host, 'timeout': self.timeout}
        _ollama_params = {k: v for k, v in _ollama_params.items() if v is not None}
        if self.client_kwargs:
            _ollama_params.update(self.client_kwargs)
        self.ollama_client = OllamaClient(**_ollama_params)
        return self.ollama_client

    def _response(self, text: str) -> Dict[str, Any]:
        kwargs: Dict[str, Any] = {}
        if self.options is not None:
            kwargs['options'] = self.options
        response = self.client.embed(input=text, model=self.id, **kwargs)
        if response and 'embeddings' in response:
            embeddings = response['embeddings']
            if isinstance(embeddings, list) and len(embeddings) > 0 and isinstance(embeddings[0], list):
                return {'embeddings': embeddings[0]}
            elif isinstance(embeddings, list) and all(isinstance(x, (int, float)) for x in embeddings):
                return {'embeddings': embeddings}
        return {'embeddings': []}

    def get_embedding(self, text: str) -> List[float]:
        try:
            response = self._response(text=text)
            embedding = response.get('embeddings', [])
            if len(embedding) != self.dimensions:
                print(f'Expected embedding dimension {self.dimensions}, but got {len(embedding)}')
                return []
            return embedding
        except Exception as e:
            print(e)
            return []

    def get_embedding_and_usage(self, text: str) -> Tuple[List[float], Optional[Dict]]:
        embedding = self.get_embedding(text=text)
        usage = None
        return embedding, usage


class Document:
    def __init__(self, content: str, id: str = None, name: str = None, meta_data: Dict[str, Any] = None, embedder: Embedder = None, embedding: List[float] = None, usage: Dict[str, Any] = None, reranking_score: float = None):
        self.content = content
        self.id = id
        self.name = name
        self.meta_data = meta_data or {}
        self.embedder = embedder
        self.embedding = embedding
        self.usage = usage
        self.reranking_score = reranking_score

    def embed(self, embedder: Optional[Embedder] = None) -> None:
        _embedder = embedder or self.embedder
        if _embedder is None:
            raise ValueError('No embedder provided')
        self.embedding, self.usage = _embedder.get_embedding_and_usage(self.content)

    def to_dict(self) -> Dict[str, Any]:
        fields = {'name', 'meta_data', 'content'}
        return {field: getattr(self, field)
            for field in fields
            if getattr(self, field) is not None or field == 'content'}

    @classmethod
    def from_dict(cls, document: Dict[str, Any]) -> 'Document':
        return cls(**document)

    @classmethod
    def from_json(cls, document: str) -> 'Document':
        return cls(**json.loads(document))


class ChunkingStrategy:
    def __init__(self, chunk_size=5000, overlap=0, separators: List[str] = None):
        if overlap >= chunk_size:
            chunk_size, overlap = overlap, chunk_size
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.separators = separators or ['\n\n', '\n', '。', ' ']

    def chunk(self, document: Document) -> List[Document]:
        if len(document.content) <= self.chunk_size:
            return [document]
        chunks: List[Document] = []
        start = 0
        chunk_meta_data = document.meta_data
        chunk_number = 1
        content = re.sub(r'\v+', '\v', re.sub(r'\f+', '\f', re.sub(r'\r+', '\r', re.sub(r'\t+', '\t', re.sub(r'\s+', ' ', document.content)))))
        while start < len(content):
            end = min(start + self.chunk_size, len(content))
            if end < len(content):
                for sep in self.separators:
                    last_sep = content[start:end].rfind(sep)
                    if last_sep != -1:
                        end = start + last_sep + 1
                        break
            chunk = content[start:end]
            meta_data = chunk_meta_data.copy()
            meta_data['chunk'] = chunk_number
            chunk_id = None
            if document.id:
                chunk_id = f'{document.id}_{chunk_number}'
            chunk_number += 1
            meta_data['chunk_size'] = len(chunk)
            chunks.append(Document(id=chunk_id, name=document.name, meta_data=meta_data, content=chunk))
            new_start = end - self.overlap
            if new_start <= start:
                new_start = min(len(content), start + max(1, self.chunk_size // 10))
            start = new_start
        return chunks


class Reader:
    def __init__(self, chunk: bool = True, chunk_size: int = 3000, separators: List[str] = None, chunking_strategy: ChunkingStrategy = None):
        self.chunk = chunk
        self.chunk_size = chunk_size
        self.separators = separators or ['\n', '\n\n', '\r', '\r\n', '\n\r', '\t', ' ', '  ']
        self.chunking_strategy = chunking_strategy or ChunkingStrategy()

    def read(self, obj: Any) -> List[Document]:
        raise NotImplementedError

    def chunk_document(self, document: Document) -> List[Document]:
        return self.chunking_strategy.chunk(document)


class CSVReader(Reader):
    def read(self, file: str, delimiter: str = ',', quotechar: str = '"') -> List[Document]:
        file = Path(file)
        csv_name = Path(file.name).stem
        with file.open(newline='', mode='r', encoding='utf-8') as csvfile:
            csv_reader = csv.reader(csvfile, delimiter=delimiter, quotechar=quotechar)
            documents = [Document(name=csv_name, id=csv_name, content='\n'.join([', '.join(row) for row in csv_reader]))]
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk else documents


class CSVUrlReader(Reader):
    def read(self, url: str) -> List[Document]:
        def fetch_with_retry(url: str, max_retries: int = 3, backoff_factor: int = 2):
            for attempt in range(max_retries):
                try:
                    response = httpx.get(url)
                    response.raise_for_status()
                    return response
                except httpx.RequestError as e:
                    if attempt == max_retries - 1:
                        print(f'Failed to fetch {url} after {max_retries} attempts: {e}')
                        raise
                    wait_time = backoff_factor ** attempt
                    print(f'Request failed (attempt {attempt + 1}), retrying in {wait_time} seconds...')
                    time.sleep(wait_time)
                except httpx.HTTPStatusError as e:
                    print(f'HTTP error for {url}: {e.response.status_code} - {e.response.text}')
                    raise
            raise httpx.RequestError(f'Failed to fetch {url} after {max_retries} attempts')

        response = fetch_with_retry(url)
        parsed_url = urlparse(url)
        filename = os.path.basename(parsed_url.path) or 'data.csv'
        file_obj = io.BytesIO(response.content)
        file_obj.name = filename
        documents = CSVReader().read(file=file_obj)
        file_obj.close()
        return documents


class DocxReader(Reader):
    def read(self, file: Union[Path, io.BytesIO]) -> List[Document]:
        if isinstance(file, Path):
            print(f'Reading: {file}')
            docx_document = DocxDocument(str(file))
            doc_name = file.stem
        else:
            print(f'Reading uploaded file: {file.name}')
            docx_document = DocxDocument(file)
            doc_name = file.name.split('.')[0]
        doc_content = '\n\n'.join([para.text for para in docx_document.paragraphs])
        documents = [
            Document(name=doc_name, id=doc_name, content=doc_content)
        ]
        if self.chunk:
            chunked_documents = []
            for document in documents:
                chunked_documents.extend(self.chunk_document(document))
            return chunked_documents
        return documents


class PDFReader(Reader):
    def read(self, pdf: Union[str, Path, IO[Any]]) -> List[Document]:
        try:
            if isinstance(pdf, str):
                doc_name = pdf.split('/')[-1].split('.')[0].replace(' ', '_')
            else:
                doc_name = pdf.name.split('.')[0]
        except Exception:
            doc_name = 'pdf'
        print(f'Reading: {doc_name}')
        try:
            doc_reader = DocumentReader(pdf)
        except PdfStreamError as e:
            print(f'Error reading PDF: {e}')
            return []
        documents = []
        for page_number, page in enumerate(doc_reader.pages, start=1):
            documents.append(Document(name=doc_name, id=f'{doc_name}_{page_number}', meta_data={'page': page_number}, content=page.extract_text()))
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk else documents


class PDFUrlReader(Reader):
    def read(self, url: str) -> List[Document]:
        if not url:
            raise ValueError('No url provided')
        print(f'Reading: {url}')
        for attempt in range(3):
            try:
                response = httpx.get(url)
                break
            except httpx.RequestError as e:
                if attempt == 2:
                    print(f'Failed to fetch PDF after 3 attempts: {e}')
                    raise
                wait_time = 2**attempt
                print(f'Request failed, retrying in {wait_time} seconds...')
                time.sleep(wait_time)
        try:
            response.raise_for_status()
        except httpx.HTTPStatusError as e:
            print(f'HTTP error occurred: {e.response.status_code} - {e.response.text}')
            raise
        doc_name = url.split('/')[-1].split('.')[0].replace('/', '_').replace(' ', '_')
        doc_reader = DocumentReader(BytesIO(response.content))
        documents = []
        for page_number, page in enumerate(doc_reader.pages, start=1):
            documents.append(Document(name=doc_name, id=f'{doc_name}_{page_number}', meta_data={'page': page_number}, content=page.extract_text()))
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk else documents


class TextReader(Reader):
    def read(self, file: Union[Path, IO[Any]]) -> List[Document]:
        try:
            if isinstance(file, Path):
                if not file.exists():
                    raise FileNotFoundError(f'Could not find file: {file}')
                print(f'Reading: {file}')
                file_name = file.stem
                file_contents = file.read_text('utf-8')
            else:
                print(f'Reading uploaded file: {file.name}')
                file_name = file.name.split('.')[0]
                file.seek(0)
                file_contents = file.read().decode('utf-8')
            documents = [
                Document(name=file_name, id=file_name, content=file_contents)
            ]
            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents
            return documents
        except Exception as e:
            print(f'Error reading: {file}: {e}')
            return []


class URLReader(Reader):
    def read(self, url: str) -> List[Document]:
        if not url:
            raise ValueError('No url provided')
        print(f'Reading: {url}')
        for attempt in range(3):
            try:
                response = httpx.get(url)
                break
            except httpx.RequestError as e:
                if attempt == 2:
                    print(f'Failed to fetch PDF after 3 attempts: {e}')
                    raise
                wait_time = 2**attempt
                print(f'Request failed, retrying in {wait_time} seconds...')
                time.sleep(wait_time)
        try:
            print(f'Status: {response.status_code}')
            print(f'Content size: {len(response.content)} bytes')
        except Exception:
            pass
        try:
            response.raise_for_status()
        except httpx.HTTPStatusError as e:
            print(f'HTTP error occurred: {e.response.status_code} - {e.response.text}')
            raise
        parsed_url = urlparse(url)
        doc_name = parsed_url.path.strip('/').replace('/', '_').replace(' ', '_')
        if not doc_name:
            doc_name = parsed_url.netloc
        document = Document(name=doc_name, id=doc_name, meta_data={'url': url}, content=response.text)
        if self.chunk:
            return self.chunk_document(document)
        return [document]


class WebsiteReader(Reader):
    def __init__(self, chunk: bool = True, chunk_size: int = 3000, separators: List[str] = None, chunking_strategy: ChunkingStrategy = None, max_depth=3, max_links=10, _visited: Set[str] = None, _urls_to_crawl: List[Tuple[str, int]] = None):
        super().__init__(chunk, chunk_size, separators, chunking_strategy)
        self.max_depth = max_depth
        self.max_links = max_links
        self._visited = _visited or set()
        self._urls_to_crawl = _urls_to_crawl or []

    def read(self, url: str) -> List[Document]:
        print(f'Reading: {url}')
        num_links = 0
        crawler_result: Dict[str, str] = {}
        domain_parts = urlparse(url).netloc.split('.')
        primary_domain = '.'.join(domain_parts[-2:])
        self._urls_to_crawl.append((url, 1))
        while self._urls_to_crawl:
            current_url, current_depth = self._urls_to_crawl.pop(0)
            if current_url in self._visited or not urlparse(current_url).netloc.endswith(primary_domain) or current_depth > self.max_depth or num_links >= self.max_links:
                continue
            self._visited.add(current_url)
            time.sleep(random.uniform(1, 3))
            try:
                print(f'Crawling: {current_url}')
                response = httpx.get(current_url, timeout=10)
                response.raise_for_status()
                soup = BeautifulSoup(response.content, 'html.parser')
                for tag in ['article', 'main', 'content', 'main-content', 'post-content']:
                    if element := soup.find(tag) if tag in ['article', 'main'] else soup.find(class_=tag):
                        crawler_result[current_url] = element.get_text(strip=True, separator=' ')
                        num_links += 1
                        break
                for link in soup.find_all('a', href=True):
                    if not isinstance(link, Tag):
                        continue
                    href_str = str(link['href'])
                    full_url = urljoin(current_url, href_str)
                    if not isinstance(full_url, str):
                        continue
                    parsed_url = urlparse(full_url)
                    if parsed_url.netloc.endswith(primary_domain) and not any(parsed_url.path.endswith(ext) for ext in ['.pdf', '.jpg', '.png']):
                        full_url_str = str(full_url)
                        if full_url_str not in self._visited and (full_url_str, current_depth + 1) not in self._urls_to_crawl:
                            self._urls_to_crawl.append((full_url_str, current_depth + 1))
            except Exception as e:
                print(f'Failed to crawl: {current_url}: {e}')
                pass
        documents = []
        for crawled_url, crawled_content in crawler_result.items():
            if self.chunk:
                documents.extend(self.chunk_document(Document(name=url, id=str(crawled_url), meta_data={'url': str(crawled_url)}, content=crawled_content)))
            else:
                documents.append(Document(name=url, id=str(crawled_url), meta_data={'url': str(crawled_url)}, content=crawled_content))
        return documents
