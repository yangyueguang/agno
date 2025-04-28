import re
import os
import json
import time
import yaml
import string
import chromadb
import asyncio
import textwrap
from rich.box import HEAVY
from rich.panel import Panel
from rich.prompt import Prompt
from rich.console import Group
from rich.json import JSON
from rich.live import Live
from rich.markdown import Markdown
from rich.status import Status
from rich.text import Text
from pathlib import Path
from copy import copy, deepcopy
from datetime import datetime
from time import time
from uuid import uuid4
from collections import ChainMap, defaultdict, deque
from pydantic import BaseModel, ValidationError
from typing import Any, AsyncIterator, Callable, Dict, Iterator, List, Literal, Optional, Sequence, Set, Type, Union, cast, Mapping
from agno.models import Ollama, Audio, AudioArtifact, AudioResponse, File, Image, ImageArtifact, Video, VideoArtifact, Model, Citations, Message, MessageReferences, MessageMetrics, ModelResponse, ModelResponseEvent, Timer, Function, Toolkit
from agno.memory import TeamMemory, TeamRun, Memory, SessionSummary, AgentMemory, AgentRun, RunEvent, RunResponse, RunResponseExtraData, TeamRunResponse, RunMessages, NextAction, ReasoningStep, ReasoningSteps, get_deepseek_reasoning, get_openai_reasoning, aget_deepseek_reasoning, get_next_action, update_messages_with_reasoning, aget_openai_reasoning
from hashlib import md5
import random
import csv
import io
import inspect
import requests
from sqlalchemy.dialects import mysql
from sqlalchemy.engine import Engine, create_engine
from sqlalchemy.engine.row import Row
from sqlalchemy.inspection import inspect as sqlinspect
from sqlalchemy.orm import sessionmaker, Session
from sqlalchemy.schema import Column, MetaData, Table
from sqlalchemy.sql.expression import select, text
from dataclasses import dataclass
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup, Tag
from pypdf import PdfReader as DocumentReader
from docx import Document as DocxDocument
from ollama import Client as OllamaClient
import collections.abc
from types import GeneratorType
from agno.memory import WorkflowMemory, WorkflowRun


def check_if_run_cancelled(run_response: Union[RunResponse, TeamRunResponse]):
    if run_response.event == RunEvent.run_cancelled:
        raise KeyboardInterrupt()


class SafeFormatter(string.Formatter):
    def get_value(self, key, args, kwargs):
        if key not in kwargs:
            return f'{key}'
        return kwargs[key]

    def format_field(self, value, format_spec):
        try:
            return super().format_field(value, format_spec)
        except ValueError:
            return f'{{{value}:{format_spec}}}'


class Embedder:
    def __init__(self, model='llama3.1:8b', host='', timeout=0, options: Any = None, client_kwargs: dict = None):
        self.model = model
        self.options = options
        client_kwargs = client_kwargs or {}
        client_kwargs['host'] = host
        client_kwargs['timeout'] = timeout
        self.ollama_client = OllamaClient(**{k: v for k, v in client_kwargs.items() if v})

    def __call__(self, input: str) -> List[float]:
        kwargs = {'options': self.options} if self.options else {}
        response = self.ollama_client.embed(input=input, model=self.model, **kwargs)
        embedding = []
        if response and 'embeddings' in response:
            embeddings = response['embeddings']
            if isinstance(embeddings, list) and len(embeddings) > 0 and isinstance(embeddings[0], list):
                embedding = embeddings[0]
            elif isinstance(embeddings, list) and all(isinstance(x, (int, float)) for x in embeddings):
                embedding = embeddings
        return embedding


class Document:
    def __init__(self, content: str, id='', name='', meta_data: dict = None):
        self.id = id
        self.name = name
        self.content = content
        self.meta_data = meta_data or {}

    def embed(self, embedder: Embedder) -> List[float]:
        return embedder(self.content)

    def to_dict(self) -> Dict[str, Any]:
        return {'name': self.name, 'meta_data': self.meta_data, 'content': self.content}


class Reader:
    def __init__(self, chunk_size=5000, overlap=0, separators: List[str] = None):
        if overlap >= chunk_size:
            chunk_size, overlap = overlap, chunk_size
        self.overlap = overlap
        self.chunk_size = chunk_size
        self.separators = separators or ['\n\n', '\n', '\r', '\r\n', '\n\r', '\t', '。', ' ']

    def chunk_document(self, document: Document) -> List[Document]:
        if len(document.content) <= self.chunk_size:
            return [document]
        chunks: List[Document] = []
        start = 0
        chunk_meta_data = document.meta_data
        chunk_number = 1
        content = re.sub(r'\v+', '\v', re.sub(r'\f+', '\f', re.sub(r'\r+', '\r', re.sub(r'\t+', '\t', re.sub(r'\s+', ' ', document.content)))))
        while start < len(content):
            end = min(start + self.chunk_size, len(content))
            if end < len(content):
                for sep in self.separators:
                    last_sep = content[start:end].rfind(sep)
                    if last_sep != -1:
                        end = start + last_sep + 1
                        break
            chunk = content[start:end]
            meta_data = chunk_meta_data.copy()
            meta_data['chunk'] = chunk_number
            chunk_id = None
            if document.id:
                chunk_id = f'{document.id}_{chunk_number}'
            chunk_number += 1
            meta_data['chunk_size'] = len(chunk)
            chunks.append(Document(id=chunk_id, name=document.name, meta_data=meta_data, content=chunk))
            new_start = end - self.overlap
            if new_start <= start:
                new_start = min(len(content), start + max(1, self.chunk_size // 10))
            start = new_start
        return chunks

    def read(self, urls: list) -> Iterator[List[Document]]:
        for u in urls:
            if u.endswith('.csv'):
                yield self.read_csv(u)
            elif u.endswith(('.doc', '.docx')):
                yield self.read_docx(u)
            elif u.endswith('.pdf'):
                yield self.read_pdf(u)
            elif u.endswith(('.txt', '.md', '.py', '.js', '.vue')):
                yield self.read_text(u)
            elif u.endswith(('.html', '.xml')):
                yield self.read_url(u)
            elif u.startswith('http'):
                yield self.read_website(u, max_depth=3, max_links=10)

    def read_csv(self, url: str) -> List[Document]:
        if url.startswith('http'):
            response = requests.get(url)
            filename = os.path.basename(urlparse(url).path) or 'data.csv'
            file_obj = io.BytesIO(response.content)
            csv_reader = csv.reader(file_obj, delimiter=',', quotechar='"')
            file_obj.close()
        else:
            file = Path(url)
            filename = Path(file.name).stem
            with file.open(newline='', mode='r', encoding='utf-8') as csvfile:
                csv_reader = csv.reader(csvfile, delimiter=',', quotechar='"')
        documents = [Document(name=filename, id=filename, content='\n'.join([', '.join(row) for row in csv_reader]))]
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents

    def read_docx(self, url: str) -> List[Document]:
        docx_document = DocxDocument(url)
        doc_name = Path(url).stem
        doc_content = '\n\n'.join([para.text for para in docx_document.paragraphs])
        documents = [Document(name=doc_name, id=doc_name, content=doc_content)]
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents

    def read_pdf(self, url: str) -> List[Document]:
        doc_name = url.split('/')[-1].split('.')[0].replace(' ', '_')
        if url.startswith('http'):
            response = requests.get(url)
            doc_reader = DocumentReader(io.BytesIO(response.content))
        else:
            doc_reader = DocumentReader(url)
        documents = []
        for page_number, page in enumerate(doc_reader.pages, start=1):
            documents.append(Document(name=doc_name, id=f'{doc_name}_{page_number}', meta_data={'url': url, 'page': page_number},
                                      content=page.extract_text()))
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents

    def read_text(self, url: str) -> List[Document]:
        with open(url, 'r', encoding='utf8') as f:
            file_contents = f.read()
            file_name = url.split('/')[-1]
            documents = [Document(name=file_name, id=file_name, content=file_contents)]
            return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents

    def read_url(self, url: str) -> List[Document]:
        response = requests.get(url)
        parsed_url = urlparse(url)
        doc_name = parsed_url.path.strip('/').replace('/', '_').replace(' ', '_') or parsed_url.netloc
        document = Document(name=doc_name, id=doc_name, meta_data={'url': url}, content=response.text)
        return self.chunk_document(document) if self.chunk_size > 0 else [document]

    def read_website(self, url: str, max_depth=3, max_links=10) -> List[Document]:
        print(f'Reading: {url}')
        num_links = 0
        crawler_result: Dict[str, str] = {}
        domain_parts = urlparse(url).netloc.split('.')
        primary_domain = '.'.join(domain_parts[-2:])
        urls_to_crawl = [(url, 1)]
        visited = set()
        while urls_to_crawl:
            current_url, current_depth = urls_to_crawl.pop(0)
            if current_url in visited or not urlparse(current_url).netloc.endswith(
                    primary_domain) or current_depth > max_depth or num_links >= max_links:
                continue
            visited.add(current_url)
            time.sleep(random.uniform(1, 3))
            try:
                print(f'Crawling: {current_url}')
                response = requests.get(current_url, timeout=10)
                soup = BeautifulSoup(response.content, 'html.parser')
                for tag in ['article', 'main', 'content', 'main-content', 'post-content']:
                    if element := soup.find(tag) if tag in ['article', 'main'] else soup.find(class_=tag):
                        crawler_result[current_url] = element.get_text(strip=True, separator=' ')
                        num_links += 1
                        break
                for link in soup.find_all('a', href=True):
                    if not isinstance(link, Tag):
                        continue
                    href_str = str(link['href'])
                    full_url = urljoin(current_url, href_str)
                    if not isinstance(full_url, str):
                        continue
                    parsed_url = urlparse(full_url)
                    if parsed_url.netloc.endswith(primary_domain) and not any(
                            parsed_url.path.endswith(ext) for ext in ['.pdf', '.jpg', '.png']):
                        full_url_str = str(full_url)
                        if full_url_str not in visited and (full_url_str, current_depth + 1) not in urls_to_crawl:
                            urls_to_crawl.append((full_url_str, current_depth + 1))
            except Exception as e:
                print(f'Failed to crawl: {current_url}: {e}')
                pass
        documents = [
            Document(name=url, id=str(crawled_url), meta_data={'url': str(crawled_url)}, content=crawled_content) for
            crawled_url, crawled_content in crawler_result.items()]
        return [i for document in documents for i in self.chunk_document(document)] if self.chunk_size > 0 else documents


class Knowledge:
    """ 知识库 支持文档 csv,doc,docx,pdf,txt,html,website"""
    def __init__(self, database='local', path: str = 'static/chromadb', **kwargs):
        self.dbname: str = database
        self.embedder = Embedder()
        self.client = chromadb.PersistentClient(path=path, **kwargs)
        self.collection = self.client.get_or_create_collection(self.dbname,  metadata={'hnsw:space': 'cosine'}, embedding_function=self.embedder)

    def exists(self, doc: Union[str, Document]) -> bool:
        if isinstance(doc, Document):
            return doc.content.replace('\x00', '\ufffd') in self.collection.get().get('documents', [])
        return bool(self.collection.get(where={"url": doc})['ids'])

    def insert(self, documents: List[Document]):
        print(f'插入 {len(documents)} 个文档')
        ids: List = []
        docs: List = []
        docs_embeddings: List = []
        docs_metadata: List = []
        for document in documents:
            cleaned_content = document.content.replace('\x00', '\ufffd')
            doc_id = md5(cleaned_content.encode()).hexdigest()
            docs_embeddings.append(document.embed(embedder=self.embedder))
            docs.append(cleaned_content)
            ids.append(doc_id)
            docs_metadata.append(document.meta_data)
            print(f'插入文档: {document.id} | {document.name} | {document.meta_data}')
        if len(docs) > 0:
            self.collection.add(ids=ids, embeddings=docs_embeddings, documents=docs, metadatas=docs_metadata)

    def upsert(self, documents: List[Document]):
        print(f'更新插入 {len(documents)} 个文档')
        ids: List = []
        docs: List = []
        docs_embeddings: List = []
        docs_metadata: List = []
        for document in documents:
            cleaned_content = document.content.replace('\x00', '\ufffd')
            doc_id = md5(cleaned_content.encode()).hexdigest()
            docs_embeddings.append(document.embed(embedder=self.embedder))
            docs.append(cleaned_content)
            ids.append(doc_id)
            docs_metadata.append(document.meta_data)
            print(f'更新插入: {document.id} | {document.name} | {document.meta_data}')
        if len(docs) > 0:
            self.collection.upsert(ids=ids, embeddings=docs_embeddings, documents=docs, metadatas=docs_metadata)

    def search(self, query: str, limit: int = 5) -> List[Document]:
        result = self.collection.query(query_embeddings=self.embedder(query), n_results=limit, include=['metadatas', 'documents', 'embeddings', 'distances', 'uris'])
        search_results: List[Document] = []
        ids = result.get('ids', [[]])[0]
        metadata = result.get('metadatas', [{}])[0]
        documents = result.get('documents', [[]])[0]
        embeddings = result.get('embeddings')[0]
        embeddings = [e.tolist() if hasattr(e, 'tolist') else e for e in embeddings]
        distances = result.get('distances', [[]])[0]
        for idx, distance in enumerate(distances):
            metadata[idx]['distances'] = distance
        for idx, (id_, metadata, document) in enumerate(zip(ids, metadata, documents)):
            search_results.append(Document(id=id_, meta_data=metadata, content=document, embedding=embeddings[idx]))
        return search_results

    def delete(self):
        self.client.delete_collection(name=self.dbname)

    def load(self, docs: List[Union[Document, str]] = None, recreate=False, upsert=False, skip_existing=True):
        if recreate:
            self.delete()
        urls = [u for u in docs if isinstance(u, str) and not self.exists(u)]
        docs = [d for d in docs if isinstance(d, Document)] + list(Reader().read(urls))
        for documents in docs:
            if upsert:
                self.upsert(documents)
            else:
                self.insert([doc for doc in documents if not self.exists(doc)] if skip_existing else documents)


@dataclass
class AgentSession:
    session_id: str
    user_id: Optional[str] = None
    team_session_id: Optional[str] = None
    memory: Optional[Dict[str, Any]] = None
    session_data: Optional[Dict[str, Any]] = None
    extra_data: Optional[Dict[str, Any]] = None
    created_at: Optional[int] = None
    updated_at: Optional[int] = None
    agent_id: Optional[str] = None
    agent_data: Optional[Dict[str, Any]] = None

    def to_dict(self) -> Dict[str, Any]:
        return self.__dict__

    def telemetry_data(self) -> Dict[str, Any]:
        return {'model': self.agent_data.get('model') if self.agent_data else None, 'created_at': self.created_at, 'updated_at': self.updated_at}

    @classmethod
    def from_dict(cls, data: Mapping[str, Any]):
        if data is None or data.get('session_id') is None:
            print('AgentSession is missing session_id')
            return None
        return cls(session_id=data.get('session_id'), agent_id=data.get('agent_id'), team_session_id=data.get('team_session_id'), user_id=data.get('user_id'), memory=data.get('memory'), agent_data=data.get('agent_data'), session_data=data.get('session_data'), extra_data=data.get('extra_data'), created_at=data.get('created_at'), updated_at=data.get('updated_at'))


@dataclass
class TeamSession:
    session_id: str
    team_session_id: Optional[str] = None
    team_id: Optional[str] = None
    user_id: Optional[str] = None
    memory: Optional[Dict[str, Any]] = None
    team_data: Optional[Dict[str, Any]] = None
    session_data: Optional[Dict[str, Any]] = None
    extra_data: Optional[Dict[str, Any]] = None
    created_at: Optional[int] = None
    updated_at: Optional[int] = None

    def to_dict(self) -> Dict[str, Any]:
        return self.__dict__

    def telemetry_data(self) -> Dict[str, Any]:
        return {'model': self.team_data.get('model') if self.team_data else None, 'created_at': self.created_at, 'updated_at': self.updated_at}

    @classmethod
    def from_dict(cls, data: Mapping[str, Any]):
        if data is None or data.get('session_id') is None:
            print('AgentSession is missing session_id')
            return None
        return cls(session_id=data.get('session_id'), team_id=data.get('team_id'), team_session_id=data.get('team_session_id'), user_id=data.get('user_id'), memory=data.get('memory'), team_data=data.get('team_data'), session_data=data.get('session_data'), extra_data=data.get('extra_data'), created_at=data.get('created_at'), updated_at=data.get('updated_at'))


@dataclass
class WorkflowSession:
    session_id: str
    user_id: Optional[str] = None
    memory: Optional[Dict[str, Any]] = None
    session_data: Optional[Dict[str, Any]] = None
    extra_data: Optional[Dict[str, Any]] = None
    created_at: Optional[int] = None
    updated_at: Optional[int] = None
    workflow_id: Optional[str] = None
    workflow_data: Optional[Dict[str, Any]] = None

    def to_dict(self) -> Dict[str, Any]:
        return self.__dict__

    def monitoring_data(self) -> Dict[str, Any]:
        return self.to_dict()

    def telemetry_data(self) -> Dict[str, Any]:
        return {'created_at': self.created_at, 'updated_at': self.updated_at}

    @classmethod
    def from_dict(cls, data: Mapping[str, Any]):
        if data is None or data.get('session_id') is None:
            print('WorkflowSession is missing session_id')
            return None
        return cls(session_id=data.get('session_id'), workflow_id=data.get('workflow_id'), user_id=data.get('user_id'), memory=data.get('memory'), workflow_data=data.get('workflow_data'), session_data=data.get('session_data'), extra_data=data.get('extra_data'), created_at=data.get('created_at'), updated_at=data.get('updated_at'))


class Storage:
    def __init__(self, table_name: str, schema: Optional[str] = 'ai', db_url: Optional[str] = None, db_engine: Optional[Engine] = None, schema_version: int = 1, auto_upgrade_schema: bool = False, mode: Optional[Literal['agent', 'team', 'workflow']] = 'agent'):
        self._mode = mode
        _engine: Optional[Engine] = db_engine
        if _engine is None and db_url is not None:
            _engine = create_engine(db_url, connect_args={'charset': 'utf8mb4'})
        if _engine is None:
            raise ValueError('Must provide either db_url or db_engine')
        self.table_name: str = table_name
        self.schema: Optional[str] = schema
        self.db_url: Optional[str] = db_url
        self.db_engine: Engine = _engine
        self.metadata: MetaData = MetaData(schema=self.schema)
        self.schema_version: int = schema_version
        self.auto_upgrade_schema: bool = auto_upgrade_schema
        self._schema_up_to_date: bool = False
        self.Session: sessionmaker[Session] = sessionmaker(bind=self.db_engine)
        self.table: Table = self.get_table()

    @property
    def mode(self) -> Literal['agent', 'team', 'workflow']:
        return self._mode

    @mode.setter
    def mode(self, value: Optional[Literal['agent', 'team', 'workflow']]) -> None:
        self._mode = 'agent' if value is None else value
        if value is not None:
            self.table = self.get_table()

    def get_table_v1(self) -> Table:
        common_columns = [
            Column('session_id', mysql.TEXT, primary_key=True), Column('user_id', mysql.TEXT), Column('memory', mysql.JSON), Column('session_data', mysql.JSON), Column('extra_data', mysql.JSON), Column('created_at', mysql.BIGINT), Column('updated_at', mysql.BIGINT), ]
        specific_columns = []
        if self.mode == 'agent':
            specific_columns = [
                Column('agent_id', mysql.TEXT), Column('team_session_id', mysql.TEXT, nullable=True), Column('agent_data', mysql.JSON), ]
        elif self.mode == 'team':
            specific_columns = [
                Column('team_id', mysql.TEXT), Column('team_session_id', mysql.TEXT, nullable=True), Column('team_data', mysql.JSON), ]
        elif self.mode == 'workflow':
            specific_columns = [
                Column('workflow_id', mysql.TEXT), Column('workflow_data', mysql.JSON), ]
        table = Table(self.table_name, self.metadata, *common_columns, *specific_columns, extend_existing=True, schema=self.schema)
        return table

    def get_table(self) -> Table:
        if self.schema_version == 1:
            return self.get_table_v1()
        else:
            raise ValueError(f'Unsupported schema version: {self.schema_version}')

    def table_exists(self) -> bool:
        print(f'Checking if table exists: {self.table.name}')
        try:
            return sqlinspect(self.db_engine).has_table(self.table.name, schema=self.schema)
        except Exception as e:
            print(e)
            return False

    def create(self) -> None:
        self.table = self.get_table()
        if not self.table_exists():
            print(f'Creating table: {self.table_name}\n')
            self.table.create(self.db_engine)

    def _read(self, session: Session, session_id: str, user_id: Optional[str] = None) -> Optional[Row[Any]]:
        stmt = select(self.table).where(self.table.c.session_id == session_id)
        if user_id is not None:
            stmt = stmt.where(self.table.c.user_id == user_id)
        try:
            return session.execute(stmt).first()
        except Exception as e:
            print(f'Exception reading from table: {e}')
            print(f'Table does not exist: {self.table.name}')
            print(f'Creating table: {self.table_name}')
            self.create()
        return None

    def read(self, session_id: str, user_id: Optional[str] = None) -> Optional[Union[AgentSession, TeamSession, WorkflowSession]]:
        with self.Session.begin() as sess:
            existing_row: Optional[Row[Any]] = self._read(session=sess, session_id=session_id, user_id=user_id)
            if existing_row is not None:
                if self.mode == 'agent':
                    return AgentSession.from_dict(existing_row._mapping)
                elif self.mode == 'team':
                    return TeamSession.from_dict(existing_row._mapping)
                elif self.mode == 'workflow':
                    return WorkflowSession.from_dict(existing_row._mapping)
            return None

    def get_all_session_ids(self, user_id: Optional[str] = None, entity_id: Optional[str] = None) -> List[str]:
        session_ids: List[str] = []
        try:
            with self.Session.begin() as sess:
                stmt = select(self.table)
                if user_id is not None:
                    stmt = stmt.where(self.table.c.user_id == user_id)
                if entity_id is not None:
                    if self.mode == 'agent':
                        stmt = stmt.where(self.table.c.agent_id == entity_id)
                    elif self.mode == 'team':
                        stmt = stmt.where(self.table.c.team_id == entity_id)
                    elif self.mode == 'workflow':
                        stmt = stmt.where(self.table.c.workflow_id == entity_id)
                stmt = stmt.order_by(self.table.c.created_at.desc())
                rows = sess.execute(stmt).fetchall()
                for row in rows:
                    if row is not None and row.session_id is not None:
                        session_ids.append(row.session_id)
        except Exception as e:
            print(f'An unexpected error occurred: {str(e)}')
        return session_ids

    def get_all_sessions(self, user_id: Optional[str] = None, entity_id: Optional[str] = None) -> List[Union[AgentSession, TeamSession, WorkflowSession]]:
        sessions = []
        try:
            with self.Session.begin() as sess:
                stmt = select(self.table)
                if user_id is not None:
                    stmt = stmt.where(self.table.c.user_id == user_id)
                if entity_id is not None:
                    if self.mode == 'agent':
                        stmt = stmt.where(self.table.c.agent_id == entity_id)
                    elif self.mode == 'team':
                        stmt = stmt.where(self.table.c.team_id == entity_id)
                    elif self.mode == 'workflow':
                        stmt = stmt.where(self.table.c.workflow_id == entity_id)
                stmt = stmt.order_by(self.table.c.created_at.desc())
                rows = sess.execute(stmt).fetchall()
                for row in rows:
                    if row.session_id is not None:
                        if self.mode == 'agent':
                            _agent_session = AgentSession.from_dict(row._mapping)
                            if _agent_session is not None:
                                sessions.append(_agent_session)
                        elif self.mode == 'team':
                            _team_session = TeamSession.from_dict(row._mapping)
                            if _team_session is not None:
                                sessions.append(_team_session)
                        elif self.mode == 'workflow':
                            _workflow_session = WorkflowSession.from_dict(row._mapping)
                            if _workflow_session is not None:
                                sessions.append(_workflow_session)
        except Exception:
            print(f'Table does not exist: {self.table.name}')
        return sessions

    def upgrade_schema(self) -> None:
        if not self.auto_upgrade_schema:
            print('Auto schema upgrade disabled. Skipping upgrade.')
            return
        try:
            if self.mode == 'agent' and self.table_exists():
                with self.Session() as sess:
                    column_exists_query = text('''
                        SELECT 1 FROM information_schema.columns
                        WHERE table_schema = :schema AND table_name = :table
                        AND column_name = 'team_session_id'
                        ''')
                    column_exists = (sess.execute(column_exists_query, {'schema': self.schema, 'table': self.table_name}).scalar()
                        is not None)
                    if not column_exists:
                        print(f'Adding "team_session_id" column to {self.schema}.{self.table_name}')
                        alter_table_query = text(f'ALTER TABLE {self.schema}.{self.table_name} ADD COLUMN team_session_id TEXT')
                        sess.execute(alter_table_query)
                        sess.commit()
                        self._schema_up_to_date = True
                        print('Schema upgrade completed successfully')
        except Exception as e:
            print(f'Error during schema upgrade: {e}')
            raise

    def upsert(self, session: Union[AgentSession, TeamSession, WorkflowSession]) -> Optional[Union[AgentSession, TeamSession, WorkflowSession]]:
        if self.auto_upgrade_schema and not self._schema_up_to_date:
            self.upgrade_schema()
        with self.Session.begin() as sess:
            if self.mode == 'agent':
                upsert_sql = text(f'''
                    INSERT INTO {self.schema}.{self.table_name}
                    (session_id, agent_id, team_session_id, user_id, memory, agent_data, session_data, extra_data, created_at, updated_at)
                    VALUES
                    (:session_id, :agent_id, :team_session_id, :user_id, :memory, :agent_data, :session_data, :extra_data, UNIX_TIMESTAMP(), NULL)
                    ON DUPLICATE KEY UPDATE
                        agent_id = VALUES(agent_id), team_session_id = VALUES(team_session_id), user_id = VALUES(user_id), memory = VALUES(memory), agent_data = VALUES(agent_data), session_data = VALUES(session_data), extra_data = VALUES(extra_data), updated_at = UNIX_TIMESTAMP();
                    ''')
            elif self.mode == 'team':
                upsert_sql = text(f'''
                    INSERT INTO {self.schema}.{self.table_name}
                    (session_id, team_id, user_id, team_session_id, memory, team_data, session_data, extra_data, created_at, updated_at)
                    VALUES
                    (:session_id, :team_id, :user_id, :team_session_id, :memory, :team_data, :session_data, :extra_data, UNIX_TIMESTAMP(), NULL)
                    ON DUPLICATE KEY UPDATE
                        team_id = VALUES(team_id), team_session_id = VALUES(team_session_id), user_id = VALUES(user_id), memory = VALUES(memory), team_data = VALUES(team_data), session_data = VALUES(session_data), extra_data = VALUES(extra_data), updated_at = UNIX_TIMESTAMP();
                    ''')
            elif self.mode == 'workflow':
                upsert_sql = text(f'''
                    INSERT INTO {self.schema}.{self.table_name}
                    (session_id, workflow_id, user_id, memory, workflow_data, session_data, extra_data, created_at, updated_at)
                    VALUES
                    (:session_id, :workflow_id, :user_id, :memory, :workflow_data, :session_data, :extra_data, UNIX_TIMESTAMP(), NULL)
                    ON DUPLICATE KEY UPDATE
                        workflow_id = VALUES(workflow_id), user_id = VALUES(user_id), memory = VALUES(memory), workflow_data = VALUES(workflow_data), session_data = VALUES(session_data), extra_data = VALUES(extra_data), updated_at = UNIX_TIMESTAMP();
                    ''')
            try:
                if self.mode == 'agent':
                    sess.execute(upsert_sql, {'session_id': session.session_id, 'agent_id': session.agent_id, 'team_session_id': session.team_session_id, 'user_id': session.user_id, 'memory': json.dumps(session.memory, ensure_ascii=False)
                            if session.memory is not None
                            else None, 'agent_data': json.dumps(session.agent_data, ensure_ascii=False)
                            if session.agent_data is not None
                            else None, 'session_data': json.dumps(session.session_data, ensure_ascii=False)
                            if session.session_data is not None
                            else None, 'extra_data': json.dumps(session.extra_data, ensure_ascii=False)
                            if session.extra_data is not None
                            else None})
                elif self.mode == 'team':
                    sess.execute(upsert_sql, {'session_id': session.session_id, 'team_id': session.team_id, 'user_id': session.user_id, 'team_session_id': session.team_session_id, 'memory': json.dumps(session.memory, ensure_ascii=False)
                            if session.memory is not None
                            else None, 'team_data': json.dumps(session.team_data, ensure_ascii=False)
                            if session.team_data is not None
                            else None, 'session_data': json.dumps(session.session_data, ensure_ascii=False)
                            if session.session_data is not None
                            else None, 'extra_data': json.dumps(session.extra_data, ensure_ascii=False)
                            if session.extra_data is not None
                            else None})
                elif self.mode == 'workflow':
                    sess.execute(upsert_sql, {'session_id': session.session_id, 'workflow_id': session.workflow_id, 'user_id': session.user_id, 'memory': json.dumps(session.memory, ensure_ascii=False)
                            if session.memory is not None
                            else None, 'workflow_data': json.dumps(session.workflow_data, ensure_ascii=False)
                            if session.workflow_data is not None
                            else None, 'session_data': json.dumps(session.session_data, ensure_ascii=False)
                            if session.session_data is not None
                            else None, 'extra_data': json.dumps(session.extra_data, ensure_ascii=False)
                            if session.extra_data is not None
                            else None})
            except Exception as e:
                if not self.table_exists():
                    print(f'Table does not exist: {self.table.name}')
                    print('Creating table and retrying upsert')
                    self.create()
                    return self.upsert(session)
                else:
                    print(f'Exception upserting into table: {e}')
                    print('A table upgrade might be required, please review these docs for more information: https://agno.link/upgrade-schema')
                    return None
        return self.read(session_id=session.session_id)

    def delete_session(self, session_id: Optional[str] = None):
        if session_id is None:
            print('No session_id provided for deletion.')
            return
        with self.Session() as sess, sess.begin():
            try:
                delete_stmt = self.table.delete().where(self.table.c.session_id == session_id)
                result = sess.execute(delete_stmt)
                if result.rowcount == 0:
                    print(f'No session found with session_id: {session_id}')
                else:
                    print(f'Successfully deleted session with session_id: {session_id}')
            except Exception as e:
                print(f'Error deleting session: {e}')
                raise

    def drop(self) -> None:
        if self.table_exists():
            print(f'Deleting table: {self.table_name}')
            self.table.drop(self.db_engine)

    def __deepcopy__(self, memo):
        cls = self.__class__
        copied_obj = cls.__new__(cls)
        memo[id(self)] = copied_obj
        for k, v in self.__dict__.items():
            if k in {'metadata', 'table'}:
                continue
            elif k in {'db_engine', 'Session'}:
                setattr(copied_obj, k, v)
            else:
                setattr(copied_obj, k, deepcopy(v, memo))
        copied_obj.metadata = MetaData(schema=self.schema)
        copied_obj.table = copied_obj.get_table()
        return copied_obj


class Agent:
    def __init__(self, model: Model = None, name: str = None, agent_id: str = None, introduction: str = None,
                 user_id: str = None, session_id: str = None, session_name: str = None,
                 session_state: Optional[Dict[str, Any]] = None, context: Optional[Dict[str, Any]] = None, add_context=False,
                 resolve_context=True, memory: Optional[AgentMemory] = None, add_history_to_messages=False,
                 num_history_responses: int = None, num_history_runs=3, knowledge: Optional[Knowledge] = None,
                 add_references=False, retriever: Optional[Callable[..., Optional[List[Dict]]]] = None, references_format: Literal['json', 'yaml'] = 'json',
                 storage: Optional[Storage] = None, extra_data: Optional[Dict[str, Any]] = None, tools: Optional[List[Union[Toolkit, Callable, Function, Dict]]] = None,
                 show_tool_calls=True, tool_call_limit: Optional[int] = None, tool_choice: Optional[Union[str, Dict[str, Any]]] = None, reasoning=False,
                 reasoning_model: Optional[Model] = None, reasoning_agent: Optional['Agent'] = None, reasoning_min_steps=1,
                 reasoning_max_steps=10, read_chat_history=False, search_knowledge=True, update_knowledge=False,
                 read_tool_call_history=False, system_message: Optional[Union[str, Callable, Message]] = None, system_message_role: str = 'system',
                 create_default_system_message=True, description: str = None, goal: str = None, instructions: Optional[Union[str, List[str], Callable]] = None,
                 expected_output: str = None, additional_context: str = None, markdown=False, add_name_to_instructions=False,
                 add_datetime_to_instructions=False, add_state_in_messages=False, add_messages: Optional[List[Union[Dict, Message]]] = None,
                 user_message: Optional[Union[List, Dict, str, Callable, Message]] = None, user_message_role='user', create_default_user_message=True,
                 retries=0, delay_between_retries=1, exponential_backoff=False, response_model: Optional[Type[BaseModel]] = None,
                 parse_response=True, structured_outputs=False, use_json_mode=False, save_response_to_file: str = None,
                 stream=False, stream_intermediate_steps=False, team: Optional[List['Agent']] = None, team_data: Optional[Dict[str, Any]] = None,
                 role: str = None, respond_directly=False, add_transfer_instructions=True, team_response_separator='\n', debug_mode=False,
                 monitoring=False, telemetry=True):
        self.model = model
        self.name = name
        self.agent_id = agent_id
        self.introduction = introduction
        self.user_id = user_id
        self.session_id = session_id
        self.session_name = session_name
        self.session_state = session_state
        self.context = context
        self.add_context = add_context
        self.resolve_context = resolve_context
        self.memory = memory
        self.add_history_to_messages = add_history_to_messages
        self.num_history_responses = num_history_responses
        self.num_history_runs = num_history_responses or num_history_runs
        self.knowledge = knowledge
        self.add_references = add_references
        self.retriever = retriever
        self.references_format = references_format
        self.storage = storage
        self.extra_data = extra_data
        self.tools = tools
        self.show_tool_calls = show_tool_calls
        self.tool_call_limit = tool_call_limit
        self.tool_choice = tool_choice
        self.reasoning = reasoning
        self.reasoning_model = reasoning_model
        self.reasoning_agent = reasoning_agent
        self.reasoning_min_steps = reasoning_min_steps
        self.reasoning_max_steps = reasoning_max_steps
        self.read_chat_history = read_chat_history
        self.search_knowledge = search_knowledge
        self.update_knowledge = update_knowledge
        self.read_tool_call_history = read_tool_call_history
        self.system_message = system_message
        self.system_message_role = system_message_role
        self.create_default_system_message = create_default_system_message
        self.description = description
        self.goal = goal
        self.instructions = instructions
        self.expected_output = expected_output
        self.additional_context = additional_context
        self.markdown = markdown
        self.add_name_to_instructions = add_name_to_instructions
        self.add_datetime_to_instructions = add_datetime_to_instructions
        self.add_state_in_messages = add_state_in_messages
        self.add_messages = add_messages
        self.user_message = user_message
        self.user_message_role = user_message_role
        self.create_default_user_message = create_default_user_message
        self.retries = retries
        self.delay_between_retries = delay_between_retries
        self.exponential_backoff = exponential_backoff
        self.response_model = response_model
        self.parse_response = parse_response
        self.structured_outputs = structured_outputs
        self.use_json_mode = use_json_mode
        self.save_response_to_file = save_response_to_file
        self.stream = stream
        self.stream_intermediate_steps = stream_intermediate_steps
        self.team = team
        self.team_data = team_data
        self.team_session_id: Optional[str] = None
        self.role = role
        self.respond_directly = respond_directly
        self.add_transfer_instructions = add_transfer_instructions
        self.team_response_separator = team_response_separator
        self.debug_mode = debug_mode
        self.monitoring = monitoring
        self.telemetry = telemetry
        self.run_id: Optional[str] = None
        self.run_input: Optional[Union[str, List, Dict, Message]] = None
        self.run_messages: Optional[RunMessages] = None
        self.run_response: Optional[RunResponse] = None
        self.images: Optional[List[ImageArtifact]] = None
        self.audio: Optional[List[AudioArtifact]] = None
        self.videos: Optional[List[VideoArtifact]] = None
        self.agent_session: Optional[AgentSession] = None
        self._tools_for_model: Optional[List[Dict]] = None
        self._functions_for_model: Optional[Dict[str, Function]] = None
        if (self.reasoning or self.reasoning_model) and not self.reasoning_agent:
            self.reasoning_agent = Agent(model=reasoning_model or self.model.__class__(id=self.model.id),
                              description='你是一个细致、周到、有逻辑的推理代理，通过清晰、结构化、循序渐进的分析来解决复杂的问题',
                              instructions=textwrap.dedent(f'''
            步骤1-问题分析：\n-用你自己的话清楚地重述用户的任务，以确保完全理解。\n-明确指出需要哪些信息以及可能需要哪些工具或资源。
            第2步-分解和制定战略：\n-将问题分解为明确定义的子任务。\n-制定至少两种不同的策略或方法来解决问题，以确保彻底性。
            第3步-意图澄清和规划：\n-清楚地表达用户请求背后的意图。\n-从步骤2中选择最合适的策略，根据与用户意图和任务约束的一致性清楚地证明你的选择。\n-制定详细的分步行动计划，概述解决问题所需的行动顺序。
            步骤4-执行行动计划：
            对于每个计划步骤，记录：\n1.**标题**：概括步骤的简明标题。\n2.**行动**：以第一人称明确说明你的下一个行动（“我会……”）。\n3.**结果**：使用必要的工具执行行动，并提供结果的简明摘要。\n4.**推理**：清楚地解释你的理由，包括：
            -必要性：为什么需要采取这一行动。\n-注意事项：强调关键考虑因素、潜在挑战和缓解策略。\n-进展：这一步如何从逻辑上遵循或建立在之前的行动之上。\n-假设：明确说明所做的任何假设，并证明其有效性。
            5.**下一步行动**：从以下选项中明确选择下一步：
            -**继续**：如果需要进一步的步骤。\n-**验证**：当你得到一个潜在的答案时，表明它已经准备好进行验证。\n-**最终答案**：只有当您自信地验证了解决方案时。\n-**重置**：如果发现严重错误或不正确的结果，请立即重新开始分析。
            6.**置信度分数**：提供一个数字置信度分数（0.0-1.0），表明您对步骤的正确性及其结果的确定性。
            步骤5-验证（在最终确定答案之前必须进行）：
            -通过以下方式明确验证您的解决方案：\n-与替代方法进行交叉验证（在步骤2中开发）。\n-使用其他可用工具或方法独立确认准确性。\n-清楚地记录验证结果和所选验证方法背后的推理。\n-如果验证失败或出现差异，明确指出错误，重置分析，并相应地修改计划。
            第6步-提供最终答案：
            -一旦经过彻底验证并充满信心，就可以清晰简洁地交付您的解决方案。\n-简要重述你的答案如何满足用户的初衷并解决所述任务。
            一般操作指南：
            -确保您的分析保持不变：
            -**完成**：解决任务的所有要素。\n-**全面**：探索不同的观点并预测潜在的结果。\n-**逻辑**：保持所有步骤之间的连贯性。\n-**可操作**：提出明确可执行的步骤和行动。\n-**富有洞察力**：在适用的情况下提供创新和独特的视角。
            -始终通过立即重置或修改步骤来明确处理错误和失误。\n-严格遵守最小{self.reasoning_min_steps}和最大{self.reasoning_max_steps}步数，以确保有效的任务解决。
            -主动毫不犹豫地执行必要的工具，清楚地记录工具的使用情况。'''), tools=tools, show_tool_calls=False,
                              response_model=ReasoningSteps, use_json_mode=self.use_json_mode, monitoring=self.monitoring,
                              telemetry=self.telemetry, debug_mode=self.debug_mode)
            self.reasoning_agent.model.show_tool_calls = False

    def set_agent_id(self) -> str:
        if self.agent_id is None:
            self.agent_id = str(uuid4())
        print(f'Agent ID: {self.agent_id}')
        return self.agent_id

    def set_session_id(self) -> str:
        if self.session_id is None or self.session_id == '':
            self.session_id = str(uuid4())
        print(f'Session ID: {self.session_id}')
        return self.session_id

    def set_debug(self) -> None:
        if self.debug_mode or os.getenv('AGNO_DEBUG', 'false').lower() == 'true':
            self.debug_mode = True

    def set_storage_mode(self):
        if self.storage is not None:
            if self.storage.mode in ['workflow', 'team']:
                print(f'您不应该以多种模式使用存储。当前模式为 {self.storage.mode}.')
            self.storage.mode = 'agent'

    def set_monitoring(self) -> None:
        monitor_env = os.getenv('AGNO_MONITOR')
        if monitor_env is not None:
            self.monitoring = monitor_env.lower() == 'true'
        telemetry_env = os.getenv('AGNO_TELEMETRY')
        if telemetry_env is not None:
            self.telemetry = telemetry_env.lower() == 'true'

    def initialize_agent(self) -> None:
        self.set_storage_mode()
        self.set_debug()
        self.set_agent_id()
        self.set_session_id()
        if self.memory is None:
            self.memory = AgentMemory()

    @property
    def is_streamable(self) -> bool:
        return self.response_model is None

    @property
    def has_team(self) -> bool:
        return self.team is not None and len(self.team) > 0

    def _run(self, message: Optional[Union[str, List, Dict, Message]] = None, *, stream: bool = False, audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None, messages: Optional[Sequence[Union[Dict, Message]]] = None, stream_intermediate_steps: bool = False, **kwargs: Any) -> Iterator[RunResponse]:
        """运行代理并生成RunResponse。
        步骤：
        1.让代理人做好运行准备
        2.更新模型并解析上下文
        3.从存储中读取现有会话
        4.准备运行信息
        5.如果启用推理，则说明任务的原因
        6.通过生成RunStarted事件来启动Run
        7.从模型生成响应（包括运行函数调用）
        8.更新RunResponse
        9.更新代理内存
        10.计算会话度量
        11.将会话保存到存储
        12.如果设置了Save_response_to_file，则将输出保存到文件
        """
        self.initialize_agent()
        self.memory = cast(AgentMemory, self.memory)
        self.stream = self.stream or (stream and self.is_streamable)
        self.stream_intermediate_steps = self.stream_intermediate_steps or (stream_intermediate_steps and self.stream)
        self.run_id = str(uuid4())
        self.run_response = RunResponse(run_id=self.run_id, session_id=self.session_id, agent_id=self.agent_id)
        print(f'Agent Run Start: {self.run_response.run_id}')
        self.update_model()
        self.run_response.model = self.model.id if self.model is not None else None
        if self.context is not None and self.resolve_context:
            self.resolve_run_context()
        self.read_from_storage()
        run_messages: RunMessages = self.get_run_messages(message=message, audio=audio, images=images, videos=videos, files=files, messages=messages, **kwargs)
        if len(run_messages.messages) == 0:
            print('没有消息要发送到模型.')
        self.run_messages = run_messages
        if self.reasoning or self.reasoning_model is not None:
            reasoning_generator = self.reason(run_messages=run_messages)
            if self.stream:
                yield from reasoning_generator
            else:
                deque(reasoning_generator, maxlen=0)
        index_of_last_user_message = len(run_messages.messages)
        if self.stream_intermediate_steps:
            yield self.create_run_response('Run started', event=RunEvent.run_started)
        model_response: ModelResponse
        self.model = cast(Model, self.model)
        if self.stream:
            model_response = ModelResponse()
            for model_response_chunk in self.model.response_stream(messages=run_messages.messages):
                if model_response_chunk.event == ModelResponseEvent.assistant_response.value:
                    if model_response_chunk.content is not None:
                        model_response.content = (model_response.content or '') + model_response_chunk.content
                        self.run_response.content = model_response.content
                    if model_response_chunk.thinking is not None:
                        model_response.thinking = (model_response.thinking or '') + model_response_chunk.thinking
                        self.run_response.thinking = model_response.thinking
                    if model_response_chunk.redacted_thinking is not None:
                        model_response.redacted_thinking = (model_response.redacted_thinking or '') + model_response_chunk.redacted_thinking
                        self.run_response.thinking = model_response.redacted_thinking
                    if model_response_chunk.citations is not None:
                        self.run_response.citations = model_response_chunk.citations
                    if model_response_chunk.content is not None or model_response_chunk.thinking is not None or model_response_chunk.redacted_thinking is not None or model_response_chunk.citations is not None:
                        yield self.create_run_response(content=model_response_chunk.content, thinking=model_response_chunk.thinking, redacted_thinking=model_response_chunk.redacted_thinking, citations=model_response_chunk.citations, created_at=model_response_chunk.created_at)
                    if model_response_chunk.audio is not None:
                        if model_response.audio is None:
                            model_response.audio = AudioResponse(id=str(uuid4()), content='', transcript='')
                        if model_response_chunk.audio.id is not None:
                            model_response.audio.id = model_response_chunk.audio.id
                        if model_response_chunk.audio.content is not None:
                            model_response.audio.content += model_response_chunk.audio.content
                        if model_response_chunk.audio.transcript is not None:
                            model_response.audio.transcript += model_response_chunk.audio.transcript
                        if model_response_chunk.audio.expires_at is not None:
                            model_response.audio.expires_at = model_response_chunk.audio.expires_at
                        if model_response_chunk.audio.mime_type is not None:
                            model_response.audio.mime_type = model_response_chunk.audio.mime_type
                        model_response.audio.sample_rate = model_response_chunk.audio.sample_rate
                        model_response.audio.channels = model_response_chunk.audio.channels
                        self.run_response.response_audio = AudioResponse(id=model_response_chunk.audio.id, content=model_response_chunk.audio.content, transcript=model_response_chunk.audio.transcript, sample_rate=model_response_chunk.audio.sample_rate, channels=model_response_chunk.audio.channels)
                        self.run_response.created_at = model_response_chunk.created_at
                        yield self.run_response
                    if model_response_chunk.image is not None:
                        self.add_image(model_response_chunk.image)
                        yield self.run_response
                elif model_response_chunk.event == ModelResponseEvent.tool_call_started.value:
                    tool_calls_list = model_response_chunk.tool_calls
                    if tool_calls_list is not None:
                        if self.run_response.tools is None:
                            self.run_response.tools = tool_calls_list
                        else:
                            self.run_response.tools.extend(tool_calls_list)
                        self.run_response.formatted_tool_calls = format_tool_calls(self.run_response.tools)
                    if self.stream_intermediate_steps:
                        yield self.create_run_response(content=model_response_chunk.content, event=RunEvent.tool_call_started)
                elif model_response_chunk.event == ModelResponseEvent.tool_call_completed.value:
                    tool_calls_list = model_response_chunk.tool_calls
                    if tool_calls_list is not None:
                        if self.run_response.tools:
                            tool_call_index_map = {tc['tool_call_id']: i for i, tc in enumerate(self.run_response.tools) if tc.get('tool_call_id')}
                            for tool_call_dict in tool_calls_list:
                                tool_call_id = tool_call_dict.get('tool_call_id')
                                index = tool_call_index_map.get(tool_call_id)
                                if index is not None:
                                    self.run_response.tools[index] = tool_call_dict
                        else:
                            self.run_response.tools = tool_calls_list
                    if self.stream_intermediate_steps:
                        yield self.create_run_response(content=model_response_chunk.content, event=RunEvent.tool_call_completed)
        else:
            model_response = self.model.response(messages=run_messages.messages)
            if model_response.tool_calls:
                self.run_response.formatted_tool_calls = format_tool_calls(model_response.tool_calls)
            if self.response_model is not None and model_response.parsed is not None:
                if self.model.structured_outputs:
                    self.run_response.content = model_response.parsed
                    self.run_response.content_type = self.response_model.__name__
            else:
                self.run_response.content = model_response.content
            if model_response.thinking is not None:
                self.run_response.thinking = model_response.thinking
            if model_response.redacted_thinking is not None:
                if self.run_response.thinking is None:
                    self.run_response.thinking = model_response.redacted_thinking
                else:
                    self.run_response.thinking += model_response.redacted_thinking
            if model_response.citations is not None:
                self.run_response.citations = model_response.citations
            if model_response.tool_calls is not None:
                if self.run_response.tools is None:
                    self.run_response.tools = model_response.tool_calls
                else:
                    self.run_response.tools.extend(model_response.tool_calls)
            if model_response.audio is not None:
                self.run_response.response_audio = model_response.audio
            if model_response.image is not None:
                self.add_image(model_response.image)
            self.run_response.messages = run_messages.messages
            self.run_response.created_at = model_response.created_at
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        self.run_response.messages = messages_for_run_response
        self.run_response.metrics = self.aggregate_metrics_from_messages(messages_for_run_response)
        if self.stream and model_response.audio is not None:
            self.run_response.response_audio = model_response.audio
        if run_messages.system_message is not None:
            self.memory.add_system_message(run_messages.system_message, system_message_role=self.system_message_role)
        messages_for_memory: List[Message] = ([run_messages.user_message] if run_messages.user_message is not None else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        if self.stream_intermediate_steps:
            yield self.create_run_response(content='Memory updated', event=RunEvent.updating_memory)
        agent_run = AgentRun(response=self.run_response)
        agent_run.message = run_messages.user_message
        if self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message is not None:
            self.memory.update_memory(input=run_messages.user_message.get_content_string())
        if messages is not None and len(messages) > 0:
            for _im in messages:
                mp = None
                if isinstance(_im, Message):
                    mp = _im
                elif isinstance(_im, dict):
                    try:
                        mp = Message(**_im)
                    except Exception as e:
                        print(f'消息验证失败: {e}')
                else:
                    print(f'不支持的消息类型: {type(_im)}')
                    continue
                if mp:
                    if agent_run.messages is None:
                        agent_run.messages = []
                    agent_run.messages.append(mp)
                    if self.memory.create_user_memories and self.memory.update_user_memories_after_run:
                        self.memory.update_memory(input=mp.get_content_string())
                else:
                    print('无法将消息添加到内存中')
        self.memory.add_run(agent_run)
        if self.memory.create_session_summary and self.memory.update_session_summary_after_run:
            self.memory.update_summary()
        self.write_to_storage()
        self.save_run_response_to_file(message=message)
        if message is not None:
            if isinstance(message, str):
                self.run_input = message
            elif isinstance(message, Message):
                self.run_input = message.to_dict()
            else:
                self.run_input = message
        elif messages is not None:
            self.run_input = [m.to_dict() if isinstance(m, Message) else m for m in messages]
        self.set_monitoring()
        if self.telemetry or self.monitoring:
            try:
                run_data = self._create_run_data()
                agent_session: AgentSession = self.agent_session or self.get_agent_session()
            except Exception as e:
                print(f'Could not create agent event: {e}')
        print(f'Agent Run End: {self.run_response.run_id}')
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=self.run_response.content, event=RunEvent.run_completed)
        if not self.stream:
            yield self.run_response

    def run(self, message: Optional[Union[str, List, Dict, Message]] = None, *, stream: Optional[bool] = None, audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None, messages: Optional[Sequence[Union[Dict, Message]]] = None, stream_intermediate_steps: bool = False, retries: Optional[int] = None, **kwargs: Any) -> Union[RunResponse, Iterator[RunResponse]]:
        if retries is None:
            retries = self.retries
        if stream is None:
            stream = False if self.stream is None else self.stream
        last_exception = None
        num_attempts = retries + 1
        for attempt in range(num_attempts):
            if self.response_model is not None and self.parse_response:
                print('Setting stream=False as response_model is set')
                self.stream = False
                run_response: RunResponse = next(self._run(message=message, stream=False, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs))
                if isinstance(run_response.content, self.response_model):
                    return run_response
                if isinstance(run_response.content, str):
                    try:
                        structured_output = parse_response_model_str(run_response.content, self.response_model)
                        if structured_output is not None:
                            run_response.content = structured_output
                            run_response.content_type = self.response_model.__name__
                            if self.run_response is not None:
                                self.run_response.content = structured_output
                                self.run_response.content_type = self.response_model.__name__
                        else:
                            print('无法将响应转换为response_model')
                    except Exception as e:
                        print(f'无法将响应转换为输出模型: {e}')
                else:
                    print('出了点问题。运行响应内容不是字符串')
                return run_response
            else:
                if stream and self.is_streamable:
                    resp = self._run(message=message, stream=True, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs)
                    return resp
                else:
                    resp = self._run(message=message, stream=False, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs)
                    return next(resp)
        if last_exception is not None:
            print(f'尝试{num_attempts}次后失败。上次错误使用 {last_exception.model_name}({last_exception.model_id})')
            raise last_exception
        else:
            raise Exception(f'{num_attempts}次后失败')

    async def _arun(self, message: Optional[Union[str, List, Dict, Message]] = None, *, stream: bool = False, audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None, messages: Optional[Sequence[Union[Dict, Message]]] = None, stream_intermediate_steps: bool = False, **kwargs: Any) -> AsyncIterator[RunResponse]:
        self.initialize_agent()
        self.memory = cast(AgentMemory, self.memory)
        self.stream = self.stream or (stream and self.is_streamable)
        self.stream_intermediate_steps = self.stream_intermediate_steps or (stream_intermediate_steps and self.stream)
        self.run_id = str(uuid4())
        self.run_response = RunResponse(run_id=self.run_id, session_id=self.session_id, agent_id=self.agent_id)
        print(f'Async Agent Run Start: {self.run_response.run_id}')
        self.update_model(async_mode=True)
        self.run_response.model = self.model.id if self.model is not None else None
        if self.context is not None and self.resolve_context:
            self.resolve_run_context()
        self.read_from_storage()
        run_messages: RunMessages = self.get_run_messages(message=message, audio=audio, images=images, videos=videos, files=files, messages=messages, **kwargs)
        if len(run_messages.messages) == 0:
            print('没有消息要发送到模型。')
        self.run_messages = run_messages
        if self.reasoning or self.reasoning_model is not None:
            areason_generator = self.areason(run_messages=run_messages)
            if self.stream:
                async for item in areason_generator:
                    yield item
            else:
                async for _ in areason_generator:
                    pass
        index_of_last_user_message = len(run_messages.messages)
        if self.stream_intermediate_steps:
            yield self.create_run_response('Run started', event=RunEvent.run_started)
        model_response: ModelResponse
        self.model = cast(Model, self.model)
        if stream and self.is_streamable:
            model_response = ModelResponse(content='')
            model_response_stream = self.model.aresponse_stream(messages=run_messages.messages)
            async for model_response_chunk in model_response_stream:
                if model_response_chunk.event == ModelResponseEvent.assistant_response.value:
                    if model_response_chunk.content is not None:
                        model_response.content = (model_response.content or '') + model_response_chunk.content
                        self.run_response.content = model_response.content
                    if model_response_chunk.thinking is not None:
                        model_response.thinking = (model_response.thinking or '') + model_response_chunk.thinking
                        self.run_response.thinking = model_response.thinking
                    if model_response_chunk.redacted_thinking is not None:
                        model_response.redacted_thinking = (model_response.redacted_thinking or '') + model_response_chunk.redacted_thinking
                        self.run_response.thinking = model_response.redacted_thinking
                    if model_response_chunk.citations is not None:
                        self.run_response.citations = model_response_chunk.citations
                    if model_response_chunk.content is not None or model_response_chunk.thinking is not None or model_response_chunk.redacted_thinking is not None or model_response_chunk.citations is not None:
                        yield self.create_run_response(content=model_response_chunk.content, thinking=model_response_chunk.thinking, redacted_thinking=model_response_chunk.redacted_thinking, citations=model_response_chunk.citations, created_at=model_response_chunk.created_at)
                    if model_response_chunk.audio is not None:
                        if model_response.audio is None:
                            model_response.audio = AudioResponse(id=str(uuid4()), content='', transcript='')
                        if model_response_chunk.audio.id is not None:
                            model_response.audio.id = model_response_chunk.audio.id
                        if model_response_chunk.audio.content is not None:
                            model_response.audio.content += model_response_chunk.audio.content
                        if model_response_chunk.audio.transcript is not None:
                            model_response.audio.transcript += model_response_chunk.audio.transcript
                        if model_response_chunk.audio.expires_at is not None:
                            model_response.audio.expires_at = model_response_chunk.audio.expires_at
                        if model_response_chunk.audio.mime_type is not None:
                            model_response.audio.mime_type = model_response_chunk.audio.mime_type
                        model_response.audio.sample_rate = model_response_chunk.audio.sample_rate
                        model_response.audio.channels = model_response_chunk.audio.channels
                        self.run_response.response_audio = AudioResponse(id=model_response_chunk.audio.id, content=model_response_chunk.audio.content, transcript=model_response_chunk.audio.transcript, sample_rate=model_response_chunk.audio.sample_rate, channels=model_response_chunk.audio.channels)
                        self.run_response.created_at = model_response_chunk.created_at
                        yield self.run_response
                    if model_response_chunk.image is not None:
                        self.add_image(model_response_chunk.image)
                        yield self.run_response
                elif model_response_chunk.event == ModelResponseEvent.tool_call_started.value:
                    tool_calls_list = model_response_chunk.tool_calls
                    if tool_calls_list is not None:
                        if self.run_response.tools is None:
                            self.run_response.tools = tool_calls_list
                        else:
                            self.run_response.tools.extend(tool_calls_list)
                        self.run_response.formatted_tool_calls = format_tool_calls(self.run_response.tools)
                    if self.stream_intermediate_steps:
                        yield self.create_run_response(content=model_response_chunk.content, event=RunEvent.tool_call_started)
                elif model_response_chunk.event == ModelResponseEvent.tool_call_completed.value:
                    tool_calls_list = model_response_chunk.tool_calls
                    if tool_calls_list is not None:
                        if self.run_response.tools:
                            tool_call_index_map = {tc['tool_call_id']: i
                                for i, tc in enumerate(self.run_response.tools)
                                if tc.get('tool_call_id') is not None}
                            for tool_call_dict in tool_calls_list:
                                tool_call_id = (tool_call_dict['tool_call_id'] if 'tool_call_id' in tool_call_dict else None)
                                index = tool_call_index_map.get(tool_call_id)
                                if index is not None:
                                    self.run_response.tools[index] = tool_call_dict
                        else:
                            self.run_response.tools = tool_calls_list
                    if self.stream_intermediate_steps:
                        yield self.create_run_response(content=model_response_chunk.content, event=RunEvent.tool_call_completed)
        else:
            model_response = await self.model.aresponse(messages=run_messages.messages)
            if model_response.tool_calls:
                self.run_response.formatted_tool_calls = format_tool_calls(model_response.tool_calls)
            if self.response_model is not None and model_response.parsed is not None:
                if self.model.structured_outputs:
                    self.run_response.content = model_response.parsed
                    self.run_response.content_type = self.response_model.__name__
            else:
                self.run_response.content = model_response.content
            if model_response.thinking is not None:
                self.run_response.thinking = model_response.thinking
            if model_response.redacted_thinking is not None:
                if self.run_response.thinking is None:
                    self.run_response.thinking = model_response.redacted_thinking
                else:
                    self.run_response.thinking += model_response.redacted_thinking
            if model_response.citations is not None:
                self.run_response.citations = model_response.citations
            if model_response.tool_calls is not None:
                if self.run_response.tools is None:
                    self.run_response.tools = model_response.tool_calls
                else:
                    self.run_response.tools.extend(model_response.tool_calls)
            if model_response.audio is not None:
                self.run_response.response_audio = model_response.audio
            if model_response.image is not None:
                self.add_image(model_response.image)
            self.run_response.messages = run_messages.messages
            self.run_response.created_at = model_response.created_at
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        self.run_response.messages = messages_for_run_response
        self.run_response.metrics = self.aggregate_metrics_from_messages(messages_for_run_response)
        if self.stream and model_response.audio is not None:
            self.run_response.response_audio = model_response.audio
        if run_messages.system_message is not None:
            self.memory.add_system_message(run_messages.system_message, system_message_role=self.system_message_role)
        messages_for_memory: List[Message] = ([run_messages.user_message] if run_messages.user_message is not None else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        if self.stream_intermediate_steps:
            yield self.create_run_response(content='Memory updated', event=RunEvent.updating_memory)
        agent_run = AgentRun(response=self.run_response)
        agent_run.message = run_messages.user_message
        if self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message is not None:
            await self.memory.aupdate_memory(input=run_messages.user_message.get_content_string())
        if messages is not None and len(messages) > 0:
            for _im in messages:
                mp = None
                if isinstance(_im, Message):
                    mp = _im
                elif isinstance(_im, dict):
                    try:
                        mp = Message(**_im)
                    except Exception as e:
                        print(f'验证消息失败: {e}')
                else:
                    print(f'不支持的消息类型: {type(_im)}')
                    continue
                if mp:
                    if agent_run.messages is None:
                        agent_run.messages = []
                    agent_run.messages.append(mp)
                    if self.memory.create_user_memories and self.memory.update_user_memories_after_run:
                        await self.memory.aupdate_memory(input=mp.get_content_string())
                else:
                    print('无法将消息添加到内存中')
        self.memory.add_run(agent_run)
        if self.memory.create_session_summary and self.memory.update_session_summary_after_run:
            await self.memory.aupdate_summary()
        self.write_to_storage()
        self.save_run_response_to_file(message=message)
        if message is not None:
            if isinstance(message, str):
                self.run_input = message
            elif isinstance(message, Message):
                self.run_input = message.to_dict()
            else:
                self.run_input = message
        elif messages is not None:
            self.run_input = [m.to_dict() if isinstance(m, Message) else m for m in messages]
        self.set_monitoring()
        if self.telemetry or self.monitoring:
            try:
                run_data = self._create_run_data()
                agent_session: AgentSession = self.agent_session or self.get_agent_session()
            except Exception as e:
                print(f'Could not create agent event: {e}')

        print(f'Agent Run End: {self.run_response.run_id}')
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=self.run_response.content, event=RunEvent.run_completed)
        if not self.stream:
            yield self.run_response

    async def arun(self, message: Optional[Union[str, List, Dict, Message]] = None, *, stream: Optional[bool] = None, audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None, messages: Optional[Sequence[Union[Dict, Message]]] = None, stream_intermediate_steps: bool = False, retries: Optional[int] = None, **kwargs: Any) -> Any:
        if retries is None:
            retries = self.retries
        if stream is None:
            stream = False if self.stream is None else self.stream
        last_exception = None
        num_attempts = retries + 1
        for attempt in range(num_attempts):
            print(f'Attempt {attempt + 1}/{num_attempts}')
            if self.response_model is not None and self.parse_response:
                print('Setting stream=False as response_model is set')
                run_response = await self._arun(message=message, stream=False, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs).__anext__()
                if isinstance(run_response.content, self.response_model):
                    return run_response
                if isinstance(run_response.content, str):
                    try:
                        structured_output = parse_response_model_str(run_response.content, self.response_model)
                        if structured_output is not None:
                            run_response.content = structured_output
                            run_response.content_type = self.response_model.__name__
                            if self.run_response is not None:
                                self.run_response.content = structured_output
                                self.run_response.content_type = self.response_model.__name__
                        else:
                            print('无法将响应转换为response_model')
                    except Exception as e:
                        print(f'无法将响应转换为输出模型: {e}')
                else:
                    print('出了点问题。运行响应内容不是字符串')
                return run_response
            else:
                if stream and self.is_streamable:
                    resp = self._arun(message=message, stream=True, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs)
                    return resp
                else:
                    resp = self._arun(message=message, stream=False, audio=audio, images=images, videos=videos, files=files, messages=messages, stream_intermediate_steps=stream_intermediate_steps, **kwargs)
                    return await resp.__anext__()
        if last_exception is not None:
            print(f'Failed after {num_attempts} attempts. Last error using {last_exception.model_name}({last_exception.model_id})')
            raise last_exception
        else:
            raise Exception(f'{num_attempts}次后失败')

    def create_run_response(self, content: Optional[Any] = None, *, thinking: Optional[str] = None, redacted_thinking: Optional[str] = None, event: RunEvent = RunEvent.run_response, content_type: Optional[str] = None, created_at: Optional[int] = None, citations: Optional[Citations] = None) -> RunResponse:
        self.run_response = cast(RunResponse, self.run_response)
        thinking_combined = (thinking or '') + (redacted_thinking or '')
        rr = RunResponse(run_id=self.run_id, session_id=self.session_id, agent_id=self.agent_id, content=content, thinking=thinking_combined if thinking_combined else None, tools=self.run_response.tools, audio=self.run_response.audio, images=self.run_response.images, videos=self.run_response.videos, citations=citations or self.run_response.citations, response_audio=self.run_response.response_audio, model=self.run_response.model, messages=self.run_response.messages, extra_data=self.run_response.extra_data, event=event.value)
        if content_type is not None:
            rr.content_type = content_type
        if created_at is not None:
            rr.created_at = created_at
        return rr

    def get_tools(self, async_mode: bool = False) -> Optional[List[Union[Toolkit, Callable, Function, Dict]]]:
        self.memory = cast(AgentMemory, self.memory)
        agent_tools: List[Union[Toolkit, Callable, Function, Dict]] = []
        if self.tools is not None:
            for tool in self.tools:
                agent_tools.append(tool)
        if self.read_chat_history:
            agent_tools.append(self.get_chat_history)
        if self.read_tool_call_history:
            agent_tools.append(self.get_tool_call_history)
        if self.memory and self.memory.create_user_memories:
            agent_tools.append(self.update_memory)
        if self.knowledge is not None or self.retriever is not None:
            if self.search_knowledge:
                if async_mode:
                    agent_tools.append(self.async_search_knowledge_base)
                else:
                    agent_tools.append(self.search_knowledge_base)
            if self.update_knowledge:
                agent_tools.append(self.add_to_knowledge)
        if self.has_team and self.team is not None:
            for agent_index, agent in enumerate(self.team):
                agent_tools.append(self.get_transfer_function(agent, agent_index))
        return agent_tools

    def add_tools_to_model(self, model: Model, async_mode: bool = False) -> None:
        if self._functions_for_model is None or self._tools_for_model is None:
            agent_tools = self.get_tools(async_mode=async_mode)
            if agent_tools is not None and len(agent_tools) > 0:
                print('Processing tools for model')
                strict = False
                if self.response_model is not None and (self.structured_outputs or (not self.use_json_mode)) and model.supports_native_structured_outputs:
                    strict = True
                self._tools_for_model = []
                self._functions_for_model = {}
                for tool in agent_tools:
                    if isinstance(tool, Dict):
                        self._tools_for_model.append(tool)
                        print(f'Included builtin tool {tool}')
                    elif isinstance(tool, Toolkit):
                        for name, func in tool.functions.items():
                            if name not in self._functions_for_model:
                                func._agent = self
                                func.process_entrypoint(strict=strict)
                                if strict:
                                    func.strict = True
                                self._functions_for_model[name] = func
                                self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                                print(f'Included function {name} from {tool.name}')
                    elif isinstance(tool, Function):
                        if tool.name not in self._functions_for_model:
                            tool._agent = self
                            tool.process_entrypoint(strict=strict)
                            if strict and tool.strict is None:
                                tool.strict = True
                            self._functions_for_model[tool.name] = tool
                            self._tools_for_model.append({'type': 'function', 'function': tool.to_dict()})
                            print(f'Included function {tool.name}')
                    elif callable(tool):
                        try:
                            function_name = tool.__name__
                            if function_name not in self._functions_for_model:
                                func = Function.from_callable(tool, strict=strict)
                                func._agent = self
                                if strict:
                                    func.strict = True
                                self._functions_for_model[func.name] = func
                                self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                                print(f'Included function {func.name}')
                        except Exception as e:
                            print(f'Could not add function {tool}: {e}')
                model.set_tools(tools=self._tools_for_model)
                model.set_functions(functions=self._functions_for_model)

    def update_model(self, async_mode: bool = False) -> None:
        if self.model is None:
            self.model = Ollama()
        if self.response_model is None:
            self.model.response_format = None
        else:
            json_response_format = {'type': 'json_object'}
            if self.model.supports_native_structured_outputs:
                if (not self.use_json_mode) or self.structured_outputs:
                    print('将Model.response_format设置为Agent.response_Model')
                    self.model.response_format = self.response_model
                    self.model.structured_outputs = True
                else:
                    print('模型支持本机结构化输出，但未启用。改用JSON模式。')
                    self.model.response_format = json_response_format
                    self.model.structured_outputs = False
            elif self.model.supports_json_schema_outputs:
                if self.use_json_mode or not self.structured_outputs:
                    print('将Model.response_format设置为JSON响应模式')
                    self.model.response_format = {'type': 'json_schema', 'json_schema': {'name': self.response_model.__name__, 'schema': self.response_model.model_json_schema()}}
                else:
                    self.model.response_format = None
                self.model.structured_outputs = False
            else:
                print('模型不支持结构化或JSON模式输出')
                self.model.response_format = (json_response_format if (self.use_json_mode or not self.structured_outputs) else None)
                self.model.structured_outputs = False
        self.add_tools_to_model(model=self.model, async_mode=async_mode)
        if self.show_tool_calls is not None:
            self.model.show_tool_calls = self.show_tool_calls
        if self.tool_choice is not None:
            self.model.tool_choice = self.tool_choice
        if self.tool_call_limit is not None:
            self.model.tool_call_limit = self.tool_call_limit

    def resolve_run_context(self) -> None:
        print('Resolving context')
        if self.context is not None:
            if isinstance(self.context, dict):
                for ctx_key, ctx_value in self.context.items():
                    if callable(ctx_value):
                        try:
                            sig = inspect.signature(ctx_value)
                            if 'agent' in sig.parameters:
                                resolved_ctx_value = ctx_value(agent=self)
                            else:
                                resolved_ctx_value = ctx_value()
                            if resolved_ctx_value is not None:
                                self.context[ctx_key] = resolved_ctx_value
                        except Exception as e:
                            print(f'Failed to resolve context for {ctx_key}: {e}')
                    else:
                        self.context[ctx_key] = ctx_value
            else:
                print('Context is not a dict')

    def load_user_memories(self) -> None:
        self.memory = cast(AgentMemory, self.memory)
        if self.memory and self.memory.create_user_memories:
            if self.user_id is not None and self.memory.user_id is None:
                self.memory.user_id = self.user_id
            self.memory.load_user_memories()
            if self.user_id is not None:
                print(f'Memories loaded for user: {self.user_id}')
            else:
                print('Memories loaded')

    def get_agent_data(self) -> Dict[str, Any]:
        agent_data: Dict[str, Any] = {}
        if self.name is not None:
            agent_data['name'] = self.name
        if self.agent_id is not None:
            agent_data['agent_id'] = self.agent_id
        if self.model is not None:
            agent_data['model'] = self.model.to_dict()
        return agent_data

    def get_session_data(self) -> Dict[str, Any]:
        session_data: Dict[str, Any] = {}
        if self.session_name is not None:
            session_data['session_name'] = self.session_name
        if self.session_state is not None and len(self.session_state) > 0:
            session_data['session_state'] = self.session_state
        if self.team_data is not None:
            session_data['team_data'] = self.team_data
        if self.images is not None:
            session_data['images'] = [img.model_dump() for img in self.images]
        if self.videos is not None:
            session_data['videos'] = [vid.model_dump() for vid in self.videos]
        if self.audio is not None:
            session_data['audio'] = [aud.model_dump() for aud in self.audio]
        return session_data

    def get_agent_session(self) -> AgentSession:
        self.memory = cast(AgentMemory, self.memory)
        self.session_id = cast(str, self.session_id)
        self.team_session_id = cast(str, self.team_session_id)
        self.agent_id = cast(str, self.agent_id)
        return AgentSession(session_id=self.session_id, agent_id=self.agent_id, user_id=self.user_id, team_session_id=self.team_session_id, memory=self.memory.to_dict() if self.memory is not None else None, agent_data=self.get_agent_data(), session_data=self.get_session_data(), extra_data=self.extra_data, created_at=int(time()))

    def load_agent_session(self, session: AgentSession):
        if self.agent_id is None and session.agent_id is not None:
            self.agent_id = session.agent_id
        if self.user_id is None and session.user_id is not None:
            self.user_id = session.user_id
        if self.session_id is None and session.session_id is not None:
            self.session_id = session.session_id
        if session.agent_data is not None:
            if self.name is None and 'name' in session.agent_data:
                self.name = session.agent_data.get('name')
        if session.session_data is not None:
            if self.session_name is None and 'session_name' in session.session_data:
                self.session_name = session.session_data.get('session_name')
            if 'session_state' in session.session_data:
                session_state_from_db = session.session_data.get('session_state')
                if session_state_from_db is not None and isinstance(session_state_from_db, dict) and len(session_state_from_db) > 0:
                    if self.session_state is not None and len(self.session_state) > 0:
                        merge_dictionaries(session_state_from_db, self.session_state)
                    self.session_state = session_state_from_db
            if 'images' in session.session_data:
                images_from_db = session.session_data.get('images')
                if images_from_db is not None and isinstance(images_from_db, list):
                    if self.images is None:
                        self.images = []
                    self.images.extend([ImageArtifact.model_validate(img) for img in images_from_db])
            if 'videos' in session.session_data:
                videos_from_db = session.session_data.get('videos')
                if videos_from_db is not None and isinstance(videos_from_db, list):
                    if self.videos is None:
                        self.videos = []
                    self.videos.extend([VideoArtifact.model_validate(vid) for vid in videos_from_db])
            if 'audio' in session.session_data:
                audio_from_db = session.session_data.get('audio')
                if audio_from_db is not None and isinstance(audio_from_db, list):
                    if self.audio is None:
                        self.audio = []
                    self.audio.extend([AudioArtifact.model_validate(aud) for aud in audio_from_db])
        if session.extra_data is not None:
            if self.extra_data is not None:
                merge_dictionaries(session.extra_data, self.extra_data)
            self.extra_data = session.extra_data
        if self.memory is None:
            self.memory = session.memory
        if not isinstance(self.memory, AgentMemory):
            if isinstance(self.memory, dict):
                self.memory = AgentMemory(**self.memory)
            else:
                raise TypeError(f'预期内存为dict或AgentMemory，但实际得到{type(self.memory)}')
        if session.memory is not None:
            try:
                if 'runs' in session.memory:
                    try:
                        self.memory.runs = [AgentRun.model_validate(m) for m in session.memory['runs']]
                    except Exception as e:
                        print(f'无法从内存中加载运行: {e}')
                if 'messages' in session.memory:
                    try:
                        self.memory.messages = [Message.model_validate(m) for m in session.memory['messages']]
                    except Exception as e:
                        print(f'无法从内存中加载消息: {e}')
                if 'summary' in session.memory:
                    try:
                        self.memory.summary = SessionSummary.model_validate(session.memory['summary'])
                    except Exception as e:
                        print(f'无法从内存中加载会话摘要: {e}')
                if 'memories' in session.memory:
                    try:
                        self.memory.memories = [Memory.model_validate(m) for m in session.memory['memories']]
                    except Exception as e:
                        print(f'加载用户内存失败: {e}')
            except Exception as e:
                print(f'未能加载AgentMemory: {e}')
        print(f'-*- AgentSession loaded: {session.session_id}')

    def read_from_storage(self) -> Optional[AgentSession]:
        if self.storage is not None and self.session_id is not None:
            self.agent_session = cast(AgentSession, self.storage.read(session_id=self.session_id))
            if self.agent_session is not None:
                self.load_agent_session(session=self.agent_session)
            self.load_user_memories()
        return self.agent_session

    def write_to_storage(self) -> Optional[AgentSession]:
        if self.storage is not None:
            self.agent_session = cast(AgentSession, self.storage.upsert(session=self.get_agent_session()))
        return self.agent_session

    def add_introduction(self, introduction: str) -> None:
        self.memory = cast(AgentMemory, self.memory)
        if introduction is not None:
            if len(self.memory.runs) == 0:
                self.memory.add_run(AgentRun(response=RunResponse(content=introduction, messages=[Message(role=self.model.assistant_message_role, content=introduction)])))

    def load_session(self, force: bool = False) -> Optional[str]:
        if self.agent_session is not None and not force:
            if self.session_id is not None and self.agent_session.session_id == self.session_id:
                return self.agent_session.session_id
        if self.storage is not None:
            print(f'Reading AgentSession: {self.session_id}')
            self.read_from_storage()
            if self.agent_session is None:
                print('-*- Creating new AgentSession')
                if self.agent_id is None or self.session_id is None:
                    self.initialize_agent()
                if self.introduction is not None:
                    self.add_introduction(self.introduction)
                self.write_to_storage()
                if self.agent_session is None:
                    raise Exception('在存储中创建新代理会话失败')
                print(f'-*- Created AgentSession: {self.agent_session.session_id}')
                self._log_agent_session()
        return self.session_id

    def new_session(self) -> None:
        self.agent_session = None
        if self.model is not None:
            self.model.clear()
        if self.memory is not None:
            self.memory.clear()
        self.session_id = str(uuid4())
        self.load_session(force=True)

    def get_json_output_prompt(self) -> str:
        json_output_prompt = '以包含以下字段的JSON格式提供输出:'
        if self.response_model is not None:
            if isinstance(self.response_model, str):
                json_output_prompt += '\n<json_fields>'
                json_output_prompt += f'\n{self.response_model}'
                json_output_prompt += '\n</json_fields>'
            elif isinstance(self.response_model, list):
                json_output_prompt += '\n<json_fields>'
                json_output_prompt += f'\n{json.dumps(self.response_model)}'
                json_output_prompt += '\n</json_fields>'
            elif issubclass(self.response_model, BaseModel):
                json_schema = self.response_model.model_json_schema()
                if json_schema is not None:
                    response_model_properties = {}
                    json_schema_properties = json_schema.get('properties')
                    if json_schema_properties is not None:
                        for field_name, field_properties in json_schema_properties.items():
                            formatted_field_properties = {prop_name: prop_value
                                for prop_name, prop_value in field_properties.items()
                                if prop_name != 'title'}
                            if 'allOf' in formatted_field_properties:
                                ref = formatted_field_properties['allOf'][0].get('$ref', '')
                                if ref.startswith('#/$defs/'):
                                    enum_name = ref.split('/')[-1]
                                    formatted_field_properties['enum_type'] = enum_name
                            response_model_properties[field_name] = formatted_field_properties
                    json_schema_defs = json_schema.get('$defs')
                    if json_schema_defs is not None:
                        response_model_properties['$defs'] = {}
                        for def_name, def_properties in json_schema_defs.items():
                            if 'enum' in def_properties:
                                response_model_properties['$defs'][def_name] = {'type': 'string', 'enum': def_properties['enum'], 'description': def_properties.get('description', '')}
                            else:
                                def_fields = def_properties.get('properties')
                                formatted_def_properties = {}
                                if def_fields is not None:
                                    for field_name, field_properties in def_fields.items():
                                        formatted_field_properties = {prop_name: prop_value
                                            for prop_name, prop_value in field_properties.items()
                                            if prop_name != 'title'}
                                        formatted_def_properties[field_name] = formatted_field_properties
                                if len(formatted_def_properties) > 0:
                                    response_model_properties['$defs'][def_name] = formatted_def_properties
                    if len(response_model_properties) > 0:
                        json_output_prompt += '\n<json_fields>'
                        json_output_prompt += (f'\n{json.dumps([key for key in response_model_properties.keys() if key != "$defs"])}')
                        json_output_prompt += '\n</json_fields>'
                        json_output_prompt += '\n\nHere are the properties for each field:'
                        json_output_prompt += '\n<json_field_properties>'
                        json_output_prompt += f'\n{json.dumps(response_model_properties, indent=2)}'
                        json_output_prompt += '\n</json_field_properties>'
            else:
                print(f'无法为构建json架构{self.response_model}')
        else:
            json_output_prompt += 'Provide the output as JSON.'
        json_output_prompt += '\nStart your response with `{` and end it with `}`.'
        json_output_prompt += '\nYour output will be passed to json.loads() to convert it to a Python object.'
        json_output_prompt += '\nMake sure it only contains valid JSON.'
        return json_output_prompt

    def format_message_with_state_variables(self, msg: Any) -> Any:
        if not isinstance(msg, str):
            return msg
        format_variables = ChainMap(self.session_state or {}, self.context or {}, self.extra_data or {}, {'user_id': self.user_id} if self.user_id is not None else {})
        return SafeFormatter().format(msg, **format_variables)

    def get_system_message(self) -> Optional[Message]:
        self.memory = cast(AgentMemory, self.memory)
        if self.system_message is not None:
            if isinstance(self.system_message, Message):
                return self.system_message
            sys_message_content: str = ''
            if isinstance(self.system_message, str):
                sys_message_content = self.system_message
            elif callable(self.system_message):
                sys_message_content = self.system_message(agent=self)
                if not isinstance(sys_message_content, str):
                    raise Exception('system_message must return a string')
            if self.add_state_in_messages:
                sys_message_content = self.format_message_with_state_variables(sys_message_content)
            if self.response_model is not None and self.model and (self.model.supports_native_structured_outputs and (self.use_json_mode or self.structured_outputs is False)):
                sys_message_content += f'\n{self.get_json_output_prompt()}'
            return Message(role=self.system_message_role, content=sys_message_content)
        if not self.create_default_system_message:
            return None
        if self.model is None:
            raise Exception('model not set')
        instructions: List[str] = []
        if self.instructions is not None:
            _instructions = self.instructions
            if callable(self.instructions):
                _instructions = self.instructions(agent=self)
            if isinstance(_instructions, str):
                instructions.append(_instructions)
            elif isinstance(_instructions, list):
                instructions.extend(_instructions)
        _model_instructions = self.model.get_instructions_for_model()
        if _model_instructions is not None:
            instructions.extend(_model_instructions)
        additional_information: List[str] = []
        if self.markdown and self.response_model is None:
            additional_information.append('Use markdown to format your answers.')
        if self.add_datetime_to_instructions:
            additional_information.append(f'The current time is {datetime.now()}')
        if self.name is not None and self.add_name_to_instructions:
            additional_information.append(f'Your name is: {self.name}.')
        system_message_content: str = ''
        if self.description is not None:
            system_message_content += f'{self.description}\n\n'
        if self.goal is not None:
            system_message_content += f'<your_goal>\n{self.goal}\n</your_goal>\n\n'
        if self.role is not None:
            system_message_content += f'<your_role>\n{self.role}\n</your_role>\n\n'
        if self.has_team and self.add_transfer_instructions:
            system_message_content += ('<agent_team>\n'
                '您是AI代理团队的负责人：\n-您可以直接响应，也可以将任务转移给团队中的其他代理，具体取决于他们可用的工具。\n-如果将任务转移给另一个代理，请确保包括：\n-task_description（str）：任务的清晰描述。\n-expected_output（str）：预期输出。\n-additional_information（str）：有助于代理完成任务的其他信息。\n-在响应用户之前，您必须始终验证其他代理的输出。\n如果你对结果不满意，可以重新分配任务。\n'
                '</agent_team>\n\n')
        if len(instructions) > 0:
            system_message_content += '<instructions>'
            if len(instructions) > 1:
                for _upi in instructions:
                    system_message_content += f'\n- {_upi}'
            else:
                system_message_content += '\n' + instructions[0]
            system_message_content += '\n</instructions>\n\n'
        if len(additional_information) > 0:
            system_message_content += '<additional_information>'
            for _ai in additional_information:
                system_message_content += f'\n- {_ai}'
            system_message_content += '\n</additional_information>\n\n'
        if self.add_state_in_messages:
            system_message_content = self.format_message_with_state_variables(system_message_content)
        system_message_from_model = self.model.get_system_message_for_model()
        if system_message_from_model is not None:
            system_message_content += system_message_from_model
        if self.expected_output is not None:
            system_message_content += f'<expected_output>\n{self.expected_output.strip()}\n</expected_output>\n\n'
        if self.additional_context is not None:
            system_message_content += f'{self.additional_context.strip()}\n'
        if self.has_team and self.add_transfer_instructions:
            system_message_content += f'<transfer_instructions>\n{self.get_transfer_instructions().strip()}\n</transfer_instructions>\n\n'
        if self.memory:
            if self.memory.create_user_memories:
                if self.memory.memories and len(self.memory.memories) > 0:
                    system_message_content += '您可以访问以前与用户交互的记忆，以便使用:\n\n'
                    system_message_content += '<memories_from_previous_interactions>'
                    for _memory in self.memory.memories:
                        system_message_content += f'\n- {_memory.memory}'
                    system_message_content += '\n</memories_from_previous_interactions>\n\n'
                    system_message_content += '注意：此信息来自之前的互动，可能会在本次对话中更新\n你应该总是更喜欢这次谈话中的信息，而不是过去的记忆。\n\n'
                else:
                    system_message_content += '你有能力保留之前与用户互动的记忆但尚未与用户进行任何交互。\n如果用户询问以前的记忆，你可以让他们知道你对用户没有任何记忆，因为你还没有任何互动。\n\n'
                system_message_content += '您可以使用`update_memory`工具添加新内存。\n如果使用`update_memory`工具，请记住将响应传递给用户。\n\n'
            if self.memory.create_session_summary:
                if self.memory.summary is not None:
                    system_message_content += 'Here is a brief summary of your previous interactions if it helps:\n\n'
                    system_message_content += '<summary_of_previous_interactions>\n'
                    system_message_content += str(self.memory.summary)
                    system_message_content += '\n</summary_of_previous_interactions>\n\n'
                    system_message_content += '注意：此信息来自以前的交互，可能已经过时。你应该总是更喜欢这次谈话中的信息，而不是过去的总结。\n\n'
        if self.response_model is not None and not (self.model.supports_native_structured_outputs and (not self.use_json_mode or self.structured_outputs is True)):
            system_message_content += f'{self.get_json_output_prompt()}'
        return Message(role=self.system_message_role, content=system_message_content.strip()) if system_message_content else None

    def get_user_message(self, *, message: Optional[Union[str, List]], audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None, **kwargs: Any) -> Optional[Message]:
        references = None
        self.run_response = cast(RunResponse, self.run_response)
        if self.add_references and message:
            message_str: str
            if isinstance(message, str):
                message_str = message
            elif callable(message):
                message_str = message(agent=self)
            else:
                raise Exception('当add_references为True时，消息必须是字符串或可调用的')
            retrieval_timer = Timer()
            retrieval_timer.start()
            docs_from_knowledge = self.get_relevant_docs_from_knowledge(query=message_str, **kwargs)
            if docs_from_knowledge is not None:
                references = MessageReferences(query=message_str, references=docs_from_knowledge, time=round(retrieval_timer.elapsed, 4))
                if self.run_response.extra_data is None:
                    self.run_response.extra_data = RunResponseExtraData()
                if self.run_response.extra_data.references is None:
                    self.run_response.extra_data.references = []
                self.run_response.extra_data.references.append(references)
            retrieval_timer.stop()
            print(f'Time to get references: {retrieval_timer.elapsed:.4f}s')
        if self.user_message is not None:
            if isinstance(self.user_message, Message):
                return self.user_message
            user_message_content = self.user_message
            if callable(self.user_message):
                user_message_kwargs = {'agent': self, 'message': message, 'references': references}
                user_message_content = self.user_message(**user_message_kwargs)
                if not isinstance(user_message_content, str):
                    raise Exception('user_message must return a string')
            if self.add_state_in_messages:
                user_message_content = self.format_message_with_state_variables(user_message_content)
            return Message(role=self.user_message_role, content=user_message_content, audio=audio, images=images, videos=videos, files=files, **kwargs)
        if not self.create_default_user_message or isinstance(message, list):
            return Message(role=self.user_message_role, content=message, images=images, audio=audio, videos=videos, files=files, **kwargs)
        if message is None:
            return None
        user_msg_content = message
        if self.add_state_in_messages:
            user_msg_content = self.format_message_with_state_variables(message)
        if self.add_references and references is not None and references.references is not None and len(references.references) > 0:
            user_msg_content += '\n\n如果有帮助，请使用知识库中的以下参考文献:\n'
            user_msg_content += '<references>\n'
            user_msg_content += self.convert_documents_to_string(references.references) + '\n'
            user_msg_content += '</references>'
        if self.add_context and self.context is not None:
            user_msg_content += '\n\n<context>\n'
            user_msg_content += self.convert_context_to_string(self.context) + '\n'
            user_msg_content += '</context>'
        return Message(role=self.user_message_role, content=user_msg_content, audio=audio, images=images, videos=videos, files=files, **kwargs)

    def get_run_messages(self, *, message: Optional[Union[str, List, Dict, Message]] = None, audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None, messages: Optional[Sequence[Union[Dict, Message]]] = None, **kwargs: Any) -> RunMessages:
        """此函数返回具有以下属性的RunMessages对象：
        -system_message：此运行的系统消息
        -user_message：此运行的用户消息
        -messages：要发送到模型的消息列表
        要构建RunMessages对象，请执行以下操作：
        1.将系统消息添加到run_message
        2.在run_message中添加额外的消息（如果提供）
        3.向run_message添加历史记录
        4.将用户消息添加到run_message
        5.在run_message中添加消息（如果提供）
        返回：
        具有以下属性的RunMessages对象：
        -system_message：此运行的系统消息
        -user_message：此运行的用户消息
        -messages：要发送到模型的所有消息的列表
        典型用法：
        run_messages=self.get_run_messages（message=消息，audio=音频，image=图像，video=视频，file=文件，messages=消息，**kwargs）
        """
        run_messages = RunMessages()
        self.memory = cast(AgentMemory, self.memory)
        self.run_response = cast(RunResponse, self.run_response)
        system_message = self.get_system_message()
        if system_message is not None:
            run_messages.system_message = system_message
            run_messages.messages.append(system_message)
        if self.add_messages is not None:
            messages_to_add_to_run_response: List[Message] = []
            if run_messages.extra_messages is None:
                run_messages.extra_messages = []
            for _m in self.add_messages:
                if isinstance(_m, Message):
                    messages_to_add_to_run_response.append(_m)
                    run_messages.messages.append(_m)
                    run_messages.extra_messages.append(_m)
                elif isinstance(_m, dict):
                    try:
                        _m_parsed = Message.model_validate(_m)
                        messages_to_add_to_run_response.append(_m_parsed)
                        run_messages.messages.append(_m_parsed)
                        run_messages.extra_messages.append(_m_parsed)
                    except Exception as e:
                        print(f'验证消息失败: {e}')
            if len(messages_to_add_to_run_response) > 0:
                print(f'Adding {len(messages_to_add_to_run_response)} extra messages')
                if self.run_response.extra_data is None:
                    self.run_response.extra_data = RunResponseExtraData(add_messages=messages_to_add_to_run_response)
                else:
                    if self.run_response.extra_data.add_messages is None:
                        self.run_response.extra_data.add_messages = messages_to_add_to_run_response
                    else:
                        self.run_response.extra_data.add_messages.extend(messages_to_add_to_run_response)
        if self.add_history_to_messages:
            history: List[Message] = self.memory.get_messages_from_last_n_runs(last_n=self.num_history_runs, skip_role=self.system_message_role)
            if len(history) > 0:
                history_copy = [deepcopy(msg) for msg in history]
                for _msg in history_copy:
                    _msg.from_history = True
                print(f'Adding {len(history_copy)} messages from history')
                run_messages.messages += history_copy
        user_message: Optional[Message] = None
        if message is None or isinstance(message, str) or isinstance(message, list):
            user_message = self.get_user_message(message=message, audio=audio, images=images, videos=videos, files=files, **kwargs)
        elif isinstance(message, Message):
            user_message = message
        elif isinstance(message, dict):
            try:
                user_message = Message.model_validate(message)
            except Exception as e:
                print(f'验证消息失败: {e}')
        if user_message is not None:
            run_messages.user_message = user_message
            run_messages.messages.append(user_message)
        if messages is not None and len(messages) > 0:
            for _m in messages:
                if isinstance(_m, Message):
                    run_messages.messages.append(_m)
                    if run_messages.extra_messages is None:
                        run_messages.extra_messages = []
                    run_messages.extra_messages.append(_m)
                elif isinstance(_m, dict):
                    try:
                        run_messages.messages.append(Message.model_validate(_m))
                        if run_messages.extra_messages is None:
                            run_messages.extra_messages = []
                        run_messages.extra_messages.append(Message.model_validate(_m))
                    except Exception as e:
                        print(f'验证消息失败: {e}')
        return run_messages

    def get_transfer_function(self, member_agent: 'Agent', index: int) -> Function:
        def _transfer_task_to_agent(task_description: str, expected_output: str, additional_information: Optional[str] = None) -> Iterator[str]:
            if member_agent.team_data is None:
                member_agent.team_data = {}
            member_agent.team_data['leader_session_id'] = self.session_id
            member_agent.team_data['leader_agent_id'] = self.agent_id
            member_agent.team_data['leader_run_id'] = self.run_id
            member_agent_task = f'{task_description}\n\n<expected_output>\n{expected_output}\n</expected_output>'
            try:
                if additional_information is not None and additional_information.strip() != '':
                    member_agent_task += f'\n\n<additional_information>\n{additional_information}\n</additional_information>'
            except Exception as e:
                print(f'Failed to add additional information to the member agent: {e}')
            member_agent_session_id = member_agent.session_id
            member_agent_agent_id = member_agent.agent_id
            member_agent_info = {'session_id': member_agent_session_id, 'agent_id': member_agent_agent_id}
            if self.team_data is None:
                self.team_data = {}
            if 'members' not in self.team_data:
                self.team_data['members'] = [member_agent_info]
            else:
                if member_agent_info not in self.team_data['members']:
                    self.team_data['members'].append(member_agent_info)
            if self.stream and member_agent.is_streamable:
                member_agent_run_response_stream = member_agent.run(member_agent_task, stream=True)
                for member_agent_run_response_chunk in member_agent_run_response_stream:
                    yield member_agent_run_response_chunk.content
            else:
                member_agent_run_response: RunResponse = member_agent.run(member_agent_task, stream=False)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), BaseModel):
                    try:
                        yield member_agent_run_response.content.model_dump_json(indent=2)
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            yield self.team_response_separator
        agent_name = member_agent.name if member_agent.name else f'agent_{index}'
        agent_name = ''.join(c for c in agent_name if c.isalnum() or c in '_- ').strip()
        agent_name = agent_name.lower().replace(' ', '_')
        if member_agent.name is None:
            member_agent.name = agent_name
        strict = True if (member_agent.response_model is not None and member_agent.model is not None) else False
        transfer_function = Function.from_callable(_transfer_task_to_agent, strict=strict)
        transfer_function.strict = strict
        transfer_function.name = f'transfer_task_to_{agent_name}'
        transfer_function.description = textwrap.dedent(f'''使用此功能将任务转移到{agent_name},您必须清晰简洁地描述代理应该完成的任务和预期的输出。
            Args:
            task_description(str):对代理应该完成的任务的清晰简洁的描述。
            expected_output(str):代理的预期输出。
            additional_information(可选[str]):帮助代理完成任务的其他信息。
            return:
            str：委托任务的结果。''')
        if member_agent.respond_directly:
            transfer_function.show_result = True
            transfer_function.stop_after_tool_call = True
        return transfer_function

    def get_transfer_instructions(self) -> str:
        if self.team and len(self.team) > 0:
            transfer_instructions = '您可以将任务转移到团队中的以下代理:\n'
            for agent_index, agent in enumerate(self.team):
                transfer_instructions += f'\nAgent {agent_index + 1}:\n'
                if agent.name:
                    transfer_instructions += f'Name: {agent.name}\n'
                if agent.role:
                    transfer_instructions += f'Role: {agent.role}\n'
                if agent.tools is not None:
                    _tools = []
                    for _tool in agent.tools:
                        if isinstance(_tool, Toolkit):
                            _tools.extend(list(_tool.functions.keys()))
                        elif isinstance(_tool, Function):
                            _tools.append(_tool.name)
                        elif callable(_tool):
                            _tools.append(_tool.__name__)
                    transfer_instructions += f'Available tools: {", ".join(_tools)}\n'
            return transfer_instructions
        return ''

    def get_relevant_docs_from_knowledge(self, query: str, num_documents: Optional[int] = None, **kwargs) -> Optional[List[Dict[str, Any]]]:
        if self.retriever is not None and callable(self.retriever):
            try:
                sig = inspect.signature(self.retriever)
                retriever_kwargs: Dict[str, Any] = {}
                if 'agent' in sig.parameters:
                    retriever_kwargs = {'agent': self}
                retriever_kwargs.update({'query': query, 'num_documents': num_documents, **kwargs})
                return self.retriever(**retriever_kwargs)
            except Exception as e:
                print(f'Retriever failed: {e}')
                return None
        if self.knowledge is None:
            return None
        relevant_docs: List[Document] = self.knowledge.search(query=query, num_documents=num_documents, **kwargs)
        if len(relevant_docs) == 0:
            return None
        return [doc.to_dict() for doc in relevant_docs]
    
    async def aget_relevant_docs_from_knowledge(self, query: str, num_documents: Optional[int] = None, **kwargs) -> Optional[List[Dict[str, Any]]]:
        if self.retriever is not None and callable(self.retriever):
            try:
                sig = inspect.signature(self.retriever)
                retriever_kwargs: Dict[str, Any] = {}
                if 'agent' in sig.parameters:
                    retriever_kwargs = {'agent': self}
                retriever_kwargs.update({'query': query, 'num_documents': num_documents, **kwargs})
                return self.retriever(**retriever_kwargs)
            except Exception as e:
                print(f'Retriever failed: {e}')
                return None
        if self.knowledge is None or self.knowledge.vector_db is None:
            return None
        relevant_docs: List[Document] = self.knowledge.search(query=query, num_documents=num_documents, **kwargs)
        if len(relevant_docs) == 0:
            return None
        return [doc.to_dict() for doc in relevant_docs]

    def convert_documents_to_string(self, docs: List[Dict[str, Any]]) -> str:
        if docs is None or len(docs) == 0:
            return ''
        if self.references_format == 'yaml':
            return yaml.dump(docs)
        return json.dumps(docs, indent=2)

    def convert_context_to_string(self, context: Dict[str, Any]) -> str:
        if context is None:
            return ''
        try:
            return json.dumps(context, indent=2, default=str)
        except (TypeError, ValueError, OverflowError) as e:
            print(f'Failed to convert context to JSON: {e}')
            sanitized_context = {}
            for key, value in context.items():
                try:
                    json.dumps({key: value}, default=str)
                    sanitized_context[key] = value
                except Exception:
                    sanitized_context[key] = str(value)
            try:
                return json.dumps(sanitized_context, indent=2)
            except Exception as e:
                print(f'未能将经过净化的上下文转换为JSON: {e}')
                return str(context)

    def save_run_response_to_file(self, message: Optional[Union[str, List, Dict, Message]] = None) -> None:
        if self.save_response_to_file is not None and self.run_response is not None:
            message_str = None
            if message is not None:
                if isinstance(message, str):
                    message_str = message
                else:
                    print('输出文件名中未使用消息：消息不是字符串')
            try:
                fn = self.save_response_to_file.format(name=self.name, session_id=self.session_id, user_id=self.user_id, message=message_str, run_id=self.run_id)
                fn_path = Path(fn)
                if not fn_path.parent.exists():
                    fn_path.parent.mkdir(parents=True, exist_ok=True)
                if isinstance(self.run_response.content, str):
                    fn_path.write_text(self.run_response.content)
                else:
                    fn_path.write_text(json.dumps(self.run_response.content, indent=2))
            except Exception as e:
                print(f'未能将输出保存到文件: {e}')

    def update_run_response_with_reasoning(self, reasoning_steps: List[ReasoningStep], reasoning_agent_messages: List[Message]) -> None:
        self.run_response = cast(RunResponse, self.run_response)
        if self.run_response.extra_data is None:
            self.run_response.extra_data = RunResponseExtraData()
        extra_data = self.run_response.extra_data
        if extra_data.reasoning_steps is None:
            extra_data.reasoning_steps = reasoning_steps
        else:
            extra_data.reasoning_steps.extend(reasoning_steps)
        if extra_data.reasoning_messages is None:
            extra_data.reasoning_messages = reasoning_agent_messages
        else:
            extra_data.reasoning_messages.extend(reasoning_agent_messages)

    def aggregate_metrics_from_messages(self, messages: List[Message]) -> Dict[str, Any]:
        aggregated_metrics: Dict[str, Any] = defaultdict(list)
        assistant_message_role = self.model.assistant_message_role if self.model is not None else 'assistant'
        for m in messages:
            if m.role == assistant_message_role and m.metrics is not None:
                for k, v in m.metrics.__dict__.items():
                    if k == 'timer':
                        continue
                    if v is not None:
                        aggregated_metrics[k].append(v)
        if aggregated_metrics is not None:
            aggregated_metrics = dict(aggregated_metrics)
        return aggregated_metrics

    def rename(self, name: str) -> None:
        self.read_from_storage()
        self.name = name
        self.write_to_storage()
        self._log_agent_session()

    def rename_session(self, session_name: str) -> None:
        self.read_from_storage()
        self.session_name = session_name
        self.write_to_storage()
        self._log_agent_session()

    def generate_session_name(self) -> str:
        if self.model is None:
            raise Exception('Model not set')
        gen_session_name_prompt = 'Conversation\n'
        messages_for_generating_session_name = []
        self.memory = cast(AgentMemory, self.memory)
        try:
            message_pairs = self.memory.get_message_pairs()
            for message_pair in message_pairs[:3]:
                messages_for_generating_session_name.append(message_pair[0])
                messages_for_generating_session_name.append(message_pair[1])
        except Exception as e:
            print(f'Failed to generate name: {e}')
        for message in messages_for_generating_session_name:
            gen_session_name_prompt += f'{message.role.upper()}: {message.content}\n'
        gen_session_name_prompt += '\n\nConversation Name: '
        system_message = Message(role=self.system_message_role, content='请为这次对话提供一个合适的名称，最多5个字。')
        user_message = Message(role=self.user_message_role, content=gen_session_name_prompt)
        generate_name_messages = [system_message, user_message]
        generated_name = self.model.response(messages=generate_name_messages)
        content = generated_name.content
        if content is None:
            print('生成的名称为空。再试一次.')
            return self.generate_session_name()
        if len(content.split()) > 15:
            print('生成的名称太长。再试一次.')
            return self.generate_session_name()
        return content.replace('"', '').strip()

    def auto_rename_session(self) -> None:
        self.read_from_storage()
        generated_session_name = self.generate_session_name()
        print(f'Generated Session Name: {generated_session_name}')
        self.session_name = generated_session_name
        self.write_to_storage()
        self._log_agent_session()

    def delete_session(self, session_id: str):
        if self.storage is None:
            return
        self.storage.delete_session(session_id=session_id)

    def add_image(self, image: ImageArtifact) -> None:
        if self.images is None:
            self.images = []
        self.images.append(image)
        if self.run_response is not None:
            if self.run_response.images is None:
                self.run_response.images = []
            self.run_response.images.append(image)

    def add_video(self, video: VideoArtifact) -> None:
        if self.videos is None:
            self.videos = []
        self.videos.append(video)
        if self.run_response is not None:
            if self.run_response.videos is None:
                self.run_response.videos = []
            self.run_response.videos.append(video)

    def add_audio(self, audio: AudioArtifact) -> None:
        if self.audio is None:
            self.audio = []
        self.audio.append(audio)
        if self.run_response is not None:
            if self.run_response.audio is None:
                self.run_response.audio = []
            self.run_response.audio.append(audio)

    def get_images(self) -> Optional[List[ImageArtifact]]:
        return self.images

    def get_videos(self) -> Optional[List[VideoArtifact]]:
        return self.videos

    def get_audio(self) -> Optional[List[AudioArtifact]]:
        return self.audio

    def reason(self, run_messages: RunMessages) -> Iterator[RunResponse]:
        if self.stream_intermediate_steps:
            yield self.create_run_response(content='Reasoning started', event=RunEvent.reasoning_started)
        ds_reasoning_message: Optional[Message] = get_deepseek_reasoning(reasoning_agent=self.reasoning_agent, messages=run_messages.get_input_messages())
        if ds_reasoning_message is None:
            print('推理错误。推理反应为无，继续常规会话...')
            return
        run_messages.messages.append(ds_reasoning_message)
        self.update_run_response_with_reasoning(reasoning_steps=[ReasoningStep(result=ds_reasoning_message.content)], reasoning_agent_messages=[ds_reasoning_message])
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=ReasoningSteps(reasoning_steps=[ReasoningStep(result=ds_reasoning_message.content)]), event=RunEvent.reasoning_completed)
        reasoning_agent: Optional[Agent] = self.reasoning_agent
        reasoning_agent.show_tool_calls = False
        reasoning_agent.model.show_tool_calls = False
        step_count = 1
        next_action = NextAction.CONTINUE
        reasoning_messages: List[Message] = []
        all_reasoning_steps: List[ReasoningStep] = []
        print('Starting Reasoning')
        while next_action == NextAction.CONTINUE and step_count < self.reasoning_max_steps:
            print(f'Step {step_count}')
            step_count += 1
            try:
                reasoning_agent_response: RunResponse = reasoning_agent.run(messages=run_messages.get_input_messages())
                if reasoning_agent_response.content is None or reasoning_agent_response.messages is None:
                    print('推理错误。推理响应为空，继续常规会话...')
                    break
                if reasoning_agent_response.content.reasoning_steps is None:
                    print('推理错误。推理步骤为空，继续常规会话...')
                    break
                reasoning_steps: List[ReasoningStep] = reasoning_agent_response.content.reasoning_steps
                all_reasoning_steps.extend(reasoning_steps)
                if self.stream_intermediate_steps:
                    for reasoning_step in reasoning_steps:
                        yield self.create_run_response(content=reasoning_step, content_type=reasoning_step.__class__.__name__, event=RunEvent.reasoning_step)
                first_assistant_index = next((i for i, m in enumerate(reasoning_agent_response.messages) if m.role == 'assistant'), len(reasoning_agent_response.messages))
                reasoning_messages = reasoning_agent_response.messages[first_assistant_index:]
                self.update_run_response_with_reasoning(reasoning_steps=reasoning_steps, reasoning_agent_messages=reasoning_agent_response.messages)
                next_action = get_next_action(reasoning_steps[-1])
                if next_action == NextAction.FINAL_ANSWER:
                    break
            except Exception as e:
                print(f'Reasoning error: {e}')
                break
        print(f'Total Reasoning steps: {len(all_reasoning_steps)}\nReasoning finished')
        update_messages_with_reasoning(run_messages=run_messages, reasoning_messages=reasoning_messages)
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=ReasoningSteps(reasoning_steps=all_reasoning_steps), content_type=ReasoningSteps.__class__.__name__, event=RunEvent.reasoning_completed)
                
    async def areason(self, run_messages: RunMessages) -> Any:
        if self.stream_intermediate_steps:
            yield self.create_run_response(content='Reasoning started', event=RunEvent.reasoning_started)
        ds_reasoning_message: Optional[Message] = await aget_deepseek_reasoning(reasoning_agent=self.reasoning_agent, messages=run_messages.get_input_messages())
        if ds_reasoning_message is None:
            print('推理错误。推理反应为无，继续常规会话...')
            return
        run_messages.messages.append(ds_reasoning_message)
        self.update_run_response_with_reasoning(reasoning_steps=[ReasoningStep(result=ds_reasoning_message.content)], reasoning_agent_messages=[ds_reasoning_message])
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=ReasoningSteps(reasoning_steps=[ReasoningStep(result=ds_reasoning_message.content)]), event=RunEvent.reasoning_completed)
        reasoning_agent: Optional[Agent] = self.reasoning_agent
        reasoning_agent.show_tool_calls = False
        reasoning_agent.model.show_tool_calls = False
        step_count = 1
        next_action = NextAction.CONTINUE
        reasoning_messages: List[Message] = []
        all_reasoning_steps: List[ReasoningStep] = []
        print('Starting Reasoning')
        while next_action == NextAction.CONTINUE and step_count < self.reasoning_max_steps:
            print(f'Step {step_count}')
            step_count += 1
            try:
                reasoning_agent_response: RunResponse = await reasoning_agent.arun(messages=run_messages.get_input_messages())
                if reasoning_agent_response.content is None or reasoning_agent_response.messages is None:
                    print('推理错误。推理响应为空，继续常规会话...')
                    break
                if reasoning_agent_response.content.reasoning_steps is None:
                    print('推理错误。推理步骤为空，继续常规会话n...')
                    break
                reasoning_steps: List[ReasoningStep] = reasoning_agent_response.content.reasoning_steps
                all_reasoning_steps.extend(reasoning_steps)
                if self.stream_intermediate_steps:
                    for reasoning_step in reasoning_steps:
                        yield self.create_run_response(content=reasoning_step, content_type=reasoning_step.__class__.__name__, event=RunEvent.reasoning_step)
                first_assistant_index = next((i for i, m in enumerate(reasoning_agent_response.messages) if m.role == 'assistant'), len(reasoning_agent_response.messages))
                reasoning_messages = reasoning_agent_response.messages[first_assistant_index:]
                self.update_run_response_with_reasoning(reasoning_steps=reasoning_steps, reasoning_agent_messages=reasoning_agent_response.messages)
                next_action = get_next_action(reasoning_steps[-1])
                if next_action == NextAction.FINAL_ANSWER:
                    break
            except Exception as e:
                print(f'Reasoning error: {e}')
                break
        print(f'Total Reasoning steps: {len(all_reasoning_steps)}')
        print('Reasoning finished')
        update_messages_with_reasoning(run_messages=run_messages, reasoning_messages=reasoning_messages)
        if self.stream_intermediate_steps:
            yield self.create_run_response(content=ReasoningSteps(reasoning_steps=all_reasoning_steps), content_type=ReasoningSteps.__class__.__name__, event=RunEvent.reasoning_completed)

    def get_chat_history(self, num_chats: Optional[int] = None) -> str:
        history: List[Dict[str, Any]] = []
        self.memory = cast(AgentMemory, self.memory)
        all_chats = self.memory.get_message_pairs()
        if len(all_chats) == 0:
            return ''
        chats_added = 0
        for chat in all_chats[::-1]:
            history.insert(0, chat[1].to_dict())
            history.insert(0, chat[0].to_dict())
            chats_added += 1
            if num_chats is not None and chats_added >= num_chats:
                break
        return json.dumps(history)

    def get_tool_call_history(self, num_calls: int = 3) -> str:
        self.memory = cast(AgentMemory, self.memory)
        tool_calls = self.memory.get_tool_calls(num_calls)
        if len(tool_calls) == 0:
            return ''
        print(f'tool_calls: {tool_calls}')
        return json.dumps(tool_calls)

    def search_knowledge_base(self, query: str) -> str:
        self.run_response = cast(RunResponse, self.run_response)
        retrieval_timer = Timer()
        retrieval_timer.start()
        docs_from_knowledge = self.get_relevant_docs_from_knowledge(query=query)
        if docs_from_knowledge is not None:
            references = MessageReferences(query=query, references=docs_from_knowledge, time=round(retrieval_timer.elapsed, 4))
            if self.run_response.extra_data is None:
                self.run_response.extra_data = RunResponseExtraData()
            if self.run_response.extra_data.references is None:
                self.run_response.extra_data.references = []
            self.run_response.extra_data.references.append(references)
        retrieval_timer.stop()
        print(f'Time to get references: {retrieval_timer.elapsed:.4f}s')
        if docs_from_knowledge is None:
            return 'No documents found'
        return self.convert_documents_to_string(docs_from_knowledge)
    
    async def async_search_knowledge_base(self, query: str) -> str:
        self.run_response = cast(RunResponse, self.run_response)
        retrieval_timer = Timer()
        retrieval_timer.start()
        docs_from_knowledge = await self.aget_relevant_docs_from_knowledge(query=query)
        if docs_from_knowledge is not None:
            references = MessageReferences(query=query, references=docs_from_knowledge, time=round(retrieval_timer.elapsed, 4))
            if self.run_response.extra_data is None:
                self.run_response.extra_data = RunResponseExtraData()
            if self.run_response.extra_data.references is None:
                self.run_response.extra_data.references = []
            self.run_response.extra_data.references.append(references)
        retrieval_timer.stop()
        print(f'Time to get references: {retrieval_timer.elapsed:.4f}s')
        if docs_from_knowledge is None:
            return 'No documents found'
        return self.convert_documents_to_string(docs_from_knowledge)

    def add_to_knowledge(self, query: str, result: str) -> str:
        if self.knowledge is None:
            return 'Knowledge base not available'
        document_name = self.name
        if document_name is None:
            document_name = query.replace(' ', '_').replace('?', '').replace('!', '').replace('.', '')
        document_content = json.dumps({'query': query, 'result': result})
        print(f'将文档添加到知识库: {document_name}: {document_content}')
        self.knowledge.load([Document(name=document_name, content=document_content)])
        return '已成功添加到知识库'

    def update_memory(self, task: str) -> str:
        self.memory = cast(AgentMemory, self.memory)
        try:
            return self.memory.update_memory(input=task, force=True) or '内存更新成功'
        except Exception as e:
            return f'内存更新失败: {e}'

    def _log_agent_session(self):
        if not (self.telemetry or self.monitoring):
            return
        try:
            agent_session: AgentSession = self.agent_session or self.get_agent_session()
        except Exception as e:
            print(f'Could not create agent monitor: {e}')

    def _create_run_data(self) -> Dict[str, Any]:
        run_response_format = 'text'
        self.run_response = cast(RunResponse, self.run_response)
        if self.response_model is not None:
            run_response_format = 'json'
        elif self.markdown:
            run_response_format = 'markdown'
        functions = {}
        if self.model is not None and self.model._functions is not None:
            functions = {f_name: func.to_dict() for f_name, func in self.model._functions.items() if isinstance(func, Function)}
        run_data: Dict[str, Any] = {'functions': functions, 'metrics': self.run_response.metrics}
        if self.monitoring:
            run_data.update({'run_input': self.run_input, 'run_response': self.run_response.to_dict(), 'run_response_format': run_response_format})
        return run_data

    def print_response(self, message: Optional[Union[List, Dict, str, Message]] = None, *, messages: Optional[List[Union[Dict, Message]]] = None, audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None, stream: bool = False, markdown: bool = False, show_message: bool = True, show_reasoning: bool = True, show_full_reasoning: bool = False, console: Optional[Any] = None, tags_to_include_in_markdown: Set[str] = {'think', 'thinking'}, **kwargs: Any) -> None:
        if markdown:
            self.markdown = True
        if self.response_model is not None:
            self.markdown = False
        with Live(console=console) as live_log:
            status = Status('Thinking...', spinner='aesthetic', speed=0.4, refresh_per_second=10)
            live_log.update(status)
            response_timer = Timer()
            response_timer.start()
            panels = [status]
            if message and show_message:
                message_content = get_text_from_message(message)
                message_panel = create_panel(content=Text(message_content, style='green'), title='Message', border_style='cyan')
                panels.append(message_panel)
                live_log.update(Group(*panels))
            run_response = self.run(message=message, messages=messages, audio=audio, images=images, videos=videos, files=files, stream=False, **kwargs)
            response_timer.stop()
            reasoning_steps = []
            if isinstance(run_response, RunResponse) and run_response.extra_data is not None and run_response.extra_data.reasoning_steps is not None:
                reasoning_steps = run_response.extra_data.reasoning_steps
            if len(reasoning_steps) > 0 and show_reasoning:
                for i, step in enumerate(reasoning_steps, 1):
                    step_content = Text.assemble()
                    if step.title is not None:
                        step_content.append(f'{step.title}\n', 'bold')
                    if step.action is not None:
                        step_content.append(f'{step.action}\n', 'dim')
                    if step.result is not None:
                        step_content.append(Text.from_markup(step.result, style='dim'))
                    if show_full_reasoning:
                        if step.reasoning is not None:
                            step_content.append(Text.from_markup(f'\n[bold]Reasoning:[/bold] {step.reasoning}', style='dim'))
                        if step.confidence is not None:
                            step_content.append(Text.from_markup(f'\n[bold]Confidence:[/bold] {step.confidence}', style='dim'))
                    reasoning_panel = create_panel(content=step_content, title=f'Reasoning step {i}', border_style='green')
                    panels.append(reasoning_panel)
                live_log.update(Group(*panels))
            if isinstance(run_response, RunResponse) and run_response.thinking is not None:
                thinking_panel = create_panel(content=Text(run_response.thinking), title=f'Thinking ({response_timer.elapsed:.1f}s)', border_style='green')
                panels.append(thinking_panel)
                live_log.update(Group(*panels))
            if self.show_tool_calls and isinstance(run_response, RunResponse) and run_response.formatted_tool_calls:
                tool_calls_content = Text()
                for tool_call in run_response.formatted_tool_calls:
                    tool_calls_content.append(f'• {tool_call}\n')
                tool_calls_panel = create_panel(content=tool_calls_content.plain.rstrip(), title='Tool Calls', border_style='yellow')
                panels.append(tool_calls_panel)
                live_log.update(Group(*panels))
            response_content_batch: Union[str, JSON, Markdown] = ''
            if isinstance(run_response, RunResponse):
                if isinstance(run_response.content, str):
                    if self.markdown:
                        escaped_content = escape_markdown_tags(run_response.content, tags_to_include_in_markdown)
                        response_content_batch = Markdown(escaped_content)
                    else:
                        response_content_batch = run_response.get_content_as_string(indent=4)
                elif self.response_model is not None and isinstance(run_response.content, BaseModel):
                    try:
                        response_content_batch = JSON(run_response.content.model_dump_json(exclude_none=True), indent=2)
                    except Exception as e:
                        print(f'Failed to convert response to JSON: {e}')
                else:
                    try:
                        response_content_batch = JSON(json.dumps(run_response.content), indent=4)
                    except Exception as e:
                        print(f'Failed to convert response to JSON: {e}')
            response_panel = create_panel(content=response_content_batch, title=f'Response ({response_timer.elapsed:.1f}s)', border_style='blue')
            panels.append(response_panel)
            if isinstance(run_response, RunResponse) and run_response.citations is not None and run_response.citations.urls is not None:
                md_content = '\n'.join(f'{i + 1}. [{citation.title or citation.url}]({citation.url})'
                    for i, citation in enumerate(run_response.citations.urls)
                    if citation.url)
                if md_content:
                    citations_panel = create_panel(content=Markdown(md_content), title='Citations', border_style='green')
                    panels.append(citations_panel)
                    live_log.update(Group(*panels))
            panels = [p for p in panels if not isinstance(p, Status)]
            live_log.update(Group(*panels))

    def cli_app(self, message: Optional[str] = None, user: str = 'User', emoji: str = ':sunglasses:', stream: bool = False, markdown: bool = False, exit_on: Optional[List[str]] = None, **kwargs: Any) -> None:
        if message:
            self.print_response(message=message, stream=stream, markdown=markdown, **kwargs)
        _exit_on = exit_on or ['exit', 'quit', 'bye']
        while True:
            message = Prompt.ask(f'[bold] {emoji} {user} [/bold]')
            if message in _exit_on:
                break
            self.print_response(message=message, stream=stream, markdown=markdown, **kwargs)


class Team:
    def __init__(self, members: List[Union[Agent, 'Team']],
                 mode: Literal['route', 'coordinate', 'collaborate'] = 'coordinate',
                 model: Optional[Model] = None, name: str = None, team_id: str = None, user_id: str = None,
                 session_id: str = None, session_name: str = None, session_state: Optional[Dict[str, Any]] = None,
                 add_state_in_messages=False, description: str = None,
                 instructions: Optional[Union[str, List[str], Callable]] = None,
                 expected_output: str = None, success_criteria: str = None, markdown=False,
                 add_datetime_to_instructions=False,
                 context: Optional[Dict[str, Any]] = None, add_context=False, enable_agentic_context=False,
                 share_member_interactions=False,
                 read_team_history=False, tools: Optional[List[Union[Toolkit, Callable, Function, Dict]]] = None,
                 show_tool_calls=True,
                 tool_call_limit: int = None, tool_choice: Optional[Union[str, Dict[str, Any]]] = None,
                 response_model: Optional[Type[BaseModel]] = None,
                 use_json_mode=False, parse_response=True, memory: Optional[TeamMemory] = None,
                 enable_team_history=False,
                 num_of_interactions_from_history=3, storage: Storage = None,
                 extra_data: Optional[Dict[str, Any]] = None, reasoning=False,
                 reasoning_model: Optional[Model] = None, reasoning_min_steps=1, reasoning_max_steps=10,
                 debug_mode=False,
                 show_members_responses=False, monitoring=False, telemetry=True, role: str = None,
                 team_session_id: str = None):
        self.role = role
        self.team_session_id = team_session_id
        self.members = members
        self.mode = mode
        self.model = model
        self.name = name
        self.team_id = team_id
        self.user_id = user_id
        self.session_id = session_id
        self.session_name = session_name
        self.session_state = session_state
        self.add_state_in_messages = add_state_in_messages
        self.description = description
        self.instructions = instructions
        self.expected_output = expected_output
        self.markdown = markdown
        self.add_datetime_to_instructions = add_datetime_to_instructions
        self.success_criteria = success_criteria
        self.context = context
        self.add_context = add_context
        self.enable_agentic_context = enable_agentic_context
        self.share_member_interactions = share_member_interactions
        self.read_team_history = read_team_history
        self.tools = tools
        self.show_tool_calls = show_tool_calls
        self.tool_choice = tool_choice
        self.tool_call_limit = tool_call_limit
        self.response_model = response_model
        self.use_json_mode = use_json_mode
        self.parse_response = parse_response
        self.memory = memory
        self.enable_team_history = enable_team_history
        self.num_of_interactions_from_history = num_of_interactions_from_history
        self.storage = storage
        self.extra_data = extra_data
        self.reasoning = reasoning
        self.reasoning_model = reasoning_model
        self.reasoning_min_steps = reasoning_min_steps
        self.reasoning_max_steps = reasoning_max_steps
        self.debug_mode = debug_mode
        self.show_members_responses = show_members_responses
        self.monitoring = monitoring
        self.telemetry = telemetry
        self.run_id: Optional[str] = None
        self.run_input: Optional[Union[str, List, Dict]] = None
        self.run_messages: Optional[RunMessages] = None
        self.run_response: Optional[TeamRunResponse] = None
        self.images: Optional[List[ImageArtifact]] = None
        self.audio: Optional[List[AudioArtifact]] = None
        self.videos: Optional[List[VideoArtifact]] = None
        self.team_session: Optional[TeamSession] = None
        self._tools_for_model: Optional[List[Dict]] = None
        self._functions_for_model: Optional[Dict[str, Function]] = None
        self._member_response_model: Optional[Type[BaseModel]] = None
        self.reasoning_agent = None
        if (self.reasoning or self.reasoning_model) and not self.reasoning_agent:
            self.reasoning_agent = Agent(model=reasoning_model or self.model.__class__(id=self.model.id),
                                         description='你是一个细致、周到、有逻辑的推理代理，通过清晰、结构化、循序渐进的分析来解决复杂的问题',
                                         instructions=textwrap.dedent(f'''
            步骤1-问题分析：\n-用你自己的话清楚地重述用户的任务，以确保完全理解。\n-明确指出需要哪些信息以及可能需要哪些工具或资源。
            第2步-分解和制定战略：\n-将问题分解为明确定义的子任务。\n-制定至少两种不同的策略或方法来解决问题，以确保彻底性。
            第3步-意图澄清和规划：\n-清楚地表达用户请求背后的意图。\n-从步骤2中选择最合适的策略，根据与用户意图和任务约束的一致性清楚地证明你的选择。\n-制定详细的分步行动计划，概述解决问题所需的行动顺序。
            步骤4-执行行动计划：
            对于每个计划步骤，记录：\n1.**标题**：概括步骤的简明标题。\n2.**行动**：以第一人称明确说明你的下一个行动（“我会……”）。\n3.**结果**：使用必要的工具执行行动，并提供结果的简明摘要。\n4.**推理**：清楚地解释你的理由，包括：
            -必要性：为什么需要采取这一行动。\n-注意事项：强调关键考虑因素、潜在挑战和缓解策略。\n-进展：这一步如何从逻辑上遵循或建立在之前的行动之上。\n-假设：明确说明所做的任何假设，并证明其有效性。
            5.**下一步行动**：从以下选项中明确选择下一步：
            -**继续**：如果需要进一步的步骤。\n-**验证**：当你得到一个潜在的答案时，表明它已经准备好进行验证。\n-**最终答案**：只有当您自信地验证了解决方案时。\n-**重置**：如果发现严重错误或不正确的结果，请立即重新开始分析。
            6.**置信度分数**：提供一个数字置信度分数（0.0-1.0），表明您对步骤的正确性及其结果的确定性。
            步骤5-验证（在最终确定答案之前必须进行）：
            -通过以下方式明确验证您的解决方案：\n-与替代方法进行交叉验证（在步骤2中开发）。\n-使用其他可用工具或方法独立确认准确性。\n-清楚地记录验证结果和所选验证方法背后的推理。\n-如果验证失败或出现差异，明确指出错误，重置分析，并相应地修改计划。
            第6步-提供最终答案：
            -一旦经过彻底验证并充满信心，就可以清晰简洁地交付您的解决方案。\n-简要重述你的答案如何满足用户的初衷并解决所述任务。
            一般操作指南：
            -确保您的分析保持不变：
            -**完成**：解决任务的所有要素。\n-**全面**：探索不同的观点并预测潜在的结果。\n-**逻辑**：保持所有步骤之间的连贯性。\n-**可操作**：提出明确可执行的步骤和行动。\n-**富有洞察力**：在适用的情况下提供创新和独特的视角。
            -始终通过立即重置或修改步骤来明确处理错误和失误。\n-严格遵守最小{self.reasoning_min_steps}和最大{self.reasoning_max_steps}步数，以确保有效的任务解决。
            -主动毫不犹豫地执行必要的工具，清楚地记录工具的使用情况。'''), tools=tools, show_tool_calls=False,
                                         response_model=ReasoningSteps, use_json_mode=self.use_json_mode,
                                         monitoring=self.monitoring,
                                         telemetry=self.telemetry, debug_mode=self.debug_mode)
            self.reasoning_agent.model.show_tool_calls = False

    def _set_team_id(self) -> str:
        if self.team_id is None:
            self.team_id = str(uuid4())
        return self.team_id

    def _set_session_id(self) -> str:
        if self.session_id is None or self.session_id == '':
            self.session_id = str(uuid4())
        return self.session_id

    def _set_debug(self) -> None:
        if self.debug_mode or os.getenv('AGNO_DEBUG', 'false').lower() == 'true':
            self.debug_mode = True

    def _set_storage_mode(self):
        if self.storage is not None:
            self.storage.mode = 'team'

    def _set_monitoring(self) -> None:
        monitor_env = os.getenv('AGNO_MONITOR')
        if monitor_env is not None:
            self.monitoring = monitor_env.lower() == 'true'
        telemetry_env = os.getenv('AGNO_TELEMETRY')
        if telemetry_env is not None:
            self.telemetry = telemetry_env.lower() == 'true'

    def _initialize_member(self, member: Union['Team', Agent]):
        if self.debug_mode:
            member.debug_mode = True
        if self.show_tool_calls:
            member.show_tool_calls = True
        if self.markdown:
            member.markdown = True
        member.team_session_id = self.session_id
        member.team_id = self.team_id
        if member.name is None and member.role is None:
            print('Team member name and role is undefined.')

    def _initialize_team(self) -> None:
        self._set_storage_mode()
        self._set_debug()
        self._set_monitoring()
        self._set_team_id()
        self._set_session_id()
        print(f'Team ID: {self.team_id}\nSession ID: {self.session_id}')
        for member in self.members:
            self._initialize_member(member)

    def run(self, message: Union[str, List, Dict, Message], *, stream: bool = False,
            stream_intermediate_steps: bool = False, retries=3, audio: Optional[Sequence[Audio]] = None,
            images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None,
            files: Optional[Sequence[File]] = None, **kwargs: Any) -> Union[TeamRunResponse, Iterator[TeamRunResponse]]:
        self._initialize_team()
        show_tool_calls = self.show_tool_calls
        self.read_from_storage()
        if self.memory is None:
            self.memory = TeamMemory()
        if self.context is not None:
            self._resolve_run_context()
        if self.response_model is not None and self.parse_response:
            stream = False
            print('Disabling stream as response_model is set')
        self._configure_model(show_tool_calls=show_tool_calls)
        last_exception = None
        num_attempts = retries + 1
        for attempt in range(num_attempts):
            run_id = str(uuid4())
            self.run_id = run_id
            print(f'Team Run Start: {self.run_id}\nMode: "{self.mode}"')
            if message is not None:
                if isinstance(message, str):
                    self.run_input = message
                elif isinstance(message, Message):
                    self.run_input = message.to_dict()
                else:
                    self.run_input = message
            _tools: List[Union[Toolkit, Callable, Function, Dict]] = []
            if self.tools is not None:
                for tool in self.tools:
                    _tools.append(tool)
            if self.read_team_history:
                _tools.append(self.get_team_history)
            if self.mode == 'route':
                user_message = self._get_user_message(message, audio=audio, images=images, videos=videos, files=files)
                forward_task_func: Function = self.get_forward_task_function(message=user_message, stream=stream,
                                                                             async_mode=False, images=images,
                                                                             videos=videos, audio=audio, files=files)
                _tools.append(forward_task_func)
            elif self.mode == 'coordinate':
                _tools.append(
                    self.get_transfer_task_function(stream=stream, async_mode=False, images=images, videos=videos,
                                                    audio=audio, files=files))
                if self.enable_agentic_context:
                    _tools.append(self.set_team_context)
            elif self.mode == 'collaborate':
                run_member_agents_func = self.get_run_member_agents_function(stream=stream, async_mode=False,
                                                                             images=images, videos=videos, audio=audio,
                                                                             files=files)
                _tools.append(run_member_agents_func)
                if self.enable_agentic_context:
                    _tools.append(self.set_team_context)
            self._add_tools_to_model(self.model, tools=_tools)
            self.run_response = TeamRunResponse(run_id=self.run_id, session_id=self.session_id, team_id=self.team_id)
            self.run_response.model = self.model.id if self.model is not None else None
            if self.mode == 'route':
                run_messages: RunMessages = self.get_run_messages(run_response=self.run_response, message=message,
                                                                  audio=audio, images=images, videos=videos,
                                                                  files=files, **kwargs)
            else:
                run_messages = self.get_run_messages(run_response=self.run_response, message=message, audio=audio,
                                                     images=images, videos=videos, files=files, **kwargs)
            if stream:
                resp = self._run_stream(run_response=self.run_response, run_messages=run_messages,
                                        stream_intermediate_steps=stream_intermediate_steps)
                return resp
            else:
                self._run(run_response=self.run_response, run_messages=run_messages)
                return self.run_response
        if last_exception is not None:
            print(
                f'Failed after {num_attempts} attempts. Last error using {last_exception.model_name}({last_exception.model_id})')
            raise last_exception
        else:
            raise Exception(f'Failed after {num_attempts} attempts.')

    def _run(self, run_response: TeamRunResponse, run_messages: RunMessages) -> None:
        """运行团队并返回响应。
        步骤：
        1.如果启用推理，则说明任务的原因
        2.从模型中获取响应
        3.更新run_response
        4.更新团队记忆
        5.计算会话度量
        6.将会话保存到存储
        7.解析任何结构化输出
        8.记录团队跑步记录
        """
        self.memory = cast(TeamMemory, self.memory)
        self.model = cast(Model, self.model)
        if self.reasoning or self.reasoning_model is not None:
            reasoning_generator = self._reason(run_response=run_response, run_messages=run_messages)
            deque(reasoning_generator, maxlen=0)
        self.run_messages = run_messages
        index_of_last_user_message = len(run_messages.messages)
        model_response = self.model.response(messages=run_messages.messages)
        if (self.response_model is not None) and not self.use_json_mode and (model_response.parsed is not None):
            run_response.content = model_response.parsed
            run_response.content_type = self.response_model.__name__
        else:
            if not run_response.content:
                run_response.content = model_response.content
            else:
                run_response.content += model_response.content
        if model_response.thinking is not None:
            if not run_response.thinking:
                run_response.thinking = model_response.thinking
            else:
                run_response.thinking += model_response.thinking
        if model_response.citations is not None:
            run_response.citations = model_response.citations
        if model_response.tool_calls is not None:
            if run_response.tools is None:
                run_response.tools = model_response.tool_calls
            else:
                run_response.tools.extend(model_response.tool_calls)
        run_response.formatted_tool_calls = format_tool_calls(run_response.tools)
        if model_response.audio is not None:
            run_response.response_audio = model_response.audio
        run_response.created_at = model_response.created_at
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        run_response.messages = messages_for_run_response
        run_response.metrics = self._aggregate_metrics_from_messages(messages_for_run_response)
        if run_messages.system_message is not None:
            self.memory.add_system_message(run_messages.system_message, system_message_role='system')
        messages_for_memory: List[Message] = (
            [run_messages.user_message] if run_messages.user_message is not None else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        team_run = TeamRun(response=run_response)
        team_run.message = run_messages.user_message
        if self.memory is not None and self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message is not None:
            self.memory.update_memory(input=run_messages.user_message.get_content_string())
        self.memory.add_team_run(team_run)
        self.write_to_storage()
        if self.response_model is not None:
            if isinstance(run_response.content, str) and self.parse_response:
                try:
                    parsed_response_content = parse_response_model_str(run_response.content, self.response_model)
                    if parsed_response_content is not None:
                        run_response.content = parsed_response_content
                        run_response.content_type = self.response_model.__name__
                    else:
                        print('Failed to convert response to response_model')
                except Exception as e:
                    print(f'Failed to convert response to output model: {e}')
            else:
                print('Something went wrong. Run response content is not a string')
        elif self._member_response_model is not None:
            if isinstance(run_response.content, str):
                try:
                    parsed_response_content = parse_response_model_str(run_response.content,
                                                                       self._member_response_model)
                    if parsed_response_content is not None:
                        run_response.content = parsed_response_content
                        run_response.content_type = self._member_response_model.__name__
                    else:
                        print('Failed to convert response to response_model')
                except Exception as e:
                    print(f'Failed to convert response to output model: {e}')
            else:
                print('Something went wrong. Run response content is not a string')
        self._log_team_run()
        print(f'Team Run End: {self.run_id}')

    def _run_stream(self, run_response: TeamRunResponse, run_messages: RunMessages,
                    stream_intermediate_steps: bool = False) -> Iterator[TeamRunResponse]:
        """运行Team并返回响应迭代器。
        步骤：
        1.如果启用推理，则说明任务的原因
        2.从模型中获取响应
        3.更新run_response
        4.更新团队记忆
        5.计算会话度量
        6.将会话保存到存储
        7.记录团队运行日志
        """
        self.memory = cast(TeamMemory, self.memory)
        self.model = cast(Model, self.model)
        if self.reasoning or self.reasoning_model is not None:
            reasoning_generator = self._reason(run_response=run_response, run_messages=run_messages)
            yield from reasoning_generator
        self.run_messages = run_messages
        index_of_last_user_message = len(run_messages.messages)
        if stream_intermediate_steps:
            yield self._create_run_response(content='Run started', event=RunEvent.run_started)
        full_model_response = ModelResponse()
        model_stream = self.model.response_stream(messages=run_messages.messages)
        for model_response_chunk in model_stream:
            if model_response_chunk.event == ModelResponseEvent.assistant_response.value:
                should_yield = False
                if model_response_chunk.content is not None:
                    if not full_model_response.content:
                        full_model_response.content = model_response_chunk.content
                    else:
                        full_model_response.content += model_response_chunk.content
                    should_yield = True
                if model_response_chunk.thinking is not None:
                    if not full_model_response.thinking:
                        full_model_response.thinking = model_response_chunk.thinking
                    else:
                        full_model_response.thinking += model_response_chunk.thinking
                    should_yield = True
                if model_response_chunk.citations is not None:
                    full_model_response.citations = model_response_chunk.citations
                    should_yield = True
                if model_response_chunk.audio is not None:
                    if full_model_response.audio is None:
                        full_model_response.audio = AudioResponse(id=str(uuid4()), content='', transcript='')
                    if model_response_chunk.audio.id is not None:
                        full_model_response.audio.id = model_response_chunk.audio.id
                    if model_response_chunk.audio.content is not None:
                        full_model_response.audio.content += model_response_chunk.audio.content
                    if model_response_chunk.audio.transcript is not None:
                        full_model_response.audio.transcript += model_response_chunk.audio.transcript
                    if model_response_chunk.audio.expires_at is not None:
                        full_model_response.audio.expires_at = model_response_chunk.audio.expires_at
                    if model_response_chunk.audio.mime_type is not None:
                        full_model_response.audio.mime_type = model_response_chunk.audio.mime_type
                    if model_response_chunk.audio.sample_rate is not None:
                        full_model_response.audio.sample_rate = model_response_chunk.audio.sample_rate
                    if model_response_chunk.audio.channels is not None:
                        full_model_response.audio.channels = model_response_chunk.audio.channels
                    should_yield = True
                if should_yield:
                    yield self._create_run_response(content=model_response_chunk.content,
                                                    thinking=model_response_chunk.thinking,
                                                    response_audio=model_response_chunk.audio,
                                                    citations=model_response_chunk.citations,
                                                    created_at=model_response_chunk.created_at)
            elif model_response_chunk.event == ModelResponseEvent.tool_call_started.value:
                tool_calls_list = model_response_chunk.tool_calls
                if tool_calls_list is not None:
                    if run_response.tools is None:
                        run_response.tools = tool_calls_list
                    else:
                        run_response.tools.extend(tool_calls_list)
                run_response.formatted_tool_calls = format_tool_calls(run_response.tools)
                if stream_intermediate_steps:
                    yield self._create_run_response(content=model_response_chunk.content,
                                                    event=RunEvent.tool_call_started, from_run_response=run_response)
            elif model_response_chunk.event == ModelResponseEvent.tool_call_completed.value:
                tool_calls_list = model_response_chunk.tool_calls
                if tool_calls_list is not None:
                    if run_response.tools:
                        tool_call_index_map = {tc['tool_call_id']: i
                                               for i, tc in enumerate(run_response.tools)
                                               if tc.get('tool_call_id') is not None}
                        for tool_call_dict in tool_calls_list:
                            tool_call_id = tool_call_dict.get('tool_call_id')
                            index = tool_call_index_map.get(tool_call_id)
                            if index is not None:
                                run_response.tools[index] = tool_call_dict
                    else:
                        run_response.tools = tool_calls_list
                    if stream_intermediate_steps:
                        yield self._create_run_response(content=model_response_chunk.content,
                                                        event=RunEvent.tool_call_completed,
                                                        from_run_response=run_response)
        run_response.created_at = full_model_response.created_at
        if full_model_response.content is not None:
            run_response.content = full_model_response.content
        if full_model_response.thinking is not None:
            run_response.thinking = full_model_response.thinking
        if full_model_response.audio is not None:
            run_response.response_audio = full_model_response.audio
        if full_model_response.citations is not None:
            run_response.citations = full_model_response.citations
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        run_response.messages = messages_for_run_response
        run_response.metrics = self._aggregate_metrics_from_messages(messages_for_run_response)
        if run_messages.system_message is not None:
            self.memory.add_system_message(run_messages.system_message, system_message_role='system')
        messages_for_memory: List[Message] = (
            [run_messages.user_message] if run_messages.user_message is not None else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        team_run = TeamRun(response=run_response)
        team_run.message = run_messages.user_message
        if self.memory is not None and self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message is not None:
            self.memory.update_memory(input=run_messages.user_message.get_content_string())
        self.memory.add_team_run(team_run)
        self.write_to_storage()
        self._log_team_run()
        if stream_intermediate_steps:
            yield self._create_run_response(from_run_response=run_response, event=RunEvent.run_completed)
        print(f'Team Run End: {self.run_id}')

    async def arun(self, message: Union[str, List, Dict, Message], *, stream: bool = False,
                   stream_intermediate_steps: bool = False, retries: Optional[int] = None,
                   audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None,
                   videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None, **kwargs: Any) -> \
    Union[TeamRunResponse, AsyncIterator[TeamRunResponse]]:
        self._initialize_team()
        retries = retries or 3
        if retries < 1:
            raise ValueError('Retries must be at least 1')
        show_tool_calls = self.show_tool_calls
        self.read_from_storage()
        if self.memory is None:
            self.memory = TeamMemory()
        if self.context is not None:
            self._resolve_run_context()
        if self.response_model is not None and self.parse_response:
            stream = False
            print('Disabling stream as response_model is set')
        self._configure_model(show_tool_calls=show_tool_calls)
        last_exception = None
        num_attempts = retries + 1
        for attempt in range(num_attempts):
            run_id = str(uuid4())
            self.run_id = run_id
            print(f'Team Run Start: {self.run_id}')
            print(f'Mode: "{self.mode}"')
            if message is not None:
                if isinstance(message, str):
                    self.run_input = message
                elif isinstance(message, Message):
                    self.run_input = message.to_dict()
                else:
                    self.run_input = message
            _tools: List[Union[Function, Callable, Toolkit, Dict]] = []
            if self.tools is not None:
                for tool in self.tools:
                    _tools.append(tool)
            if self.read_team_history:
                _tools.append(self.get_team_history)
            if self.mode == 'route':
                user_message = self._get_user_message(message, audio=audio, images=images, videos=videos, files=files)
                forward_task_func: Function = self.get_forward_task_function(message=user_message, stream=stream,
                                                                             async_mode=True, images=images,
                                                                             videos=videos, audio=audio, files=files)
                _tools.append(forward_task_func)
                self.model.tool_choice = 'required'
            elif self.mode == 'coordinate':
                _tools.append(
                    self.get_transfer_task_function(stream=stream, async_mode=True, images=images, videos=videos,
                                                    audio=audio, files=files))
                self.model.tool_choice = 'auto'
                if self.enable_agentic_context:
                    _tools.append(self.set_team_context)
            elif self.mode == 'collaborate':
                run_member_agents_func = self.get_run_member_agents_function(stream=stream, async_mode=True,
                                                                             images=images, videos=videos, audio=audio,
                                                                             files=files)
                _tools.append(run_member_agents_func)
                self.model.tool_choice = 'auto'
                if self.enable_agentic_context:
                    _tools.append(self.set_team_context)
            self._add_tools_to_model(self.model, tools=_tools)
            self.run_response = TeamRunResponse(run_id=self.run_id, session_id=self.session_id, team_id=self.team_id)
            self.run_response.model = self.model.id if self.model is not None else None
            if self.mode == 'route':
                run_messages: RunMessages = self.get_run_messages(run_response=self.run_response, message=message,
                                                                  audio=audio, images=images, videos=videos,
                                                                  files=files, **kwargs)
            else:
                run_messages = self.get_run_messages(run_response=self.run_response, message=message, audio=audio,
                                                     images=images, videos=videos, files=files, **kwargs)
            if stream:
                resp = self._arun_stream(run_response=self.run_response, run_messages=run_messages,
                                         stream_intermediate_steps=stream_intermediate_steps)
                return resp
            else:
                await self._arun(run_response=self.run_response, run_messages=run_messages)
                return self.run_response
        if last_exception is not None:
            print(
                f'Failed after {num_attempts} attempts. Last error using {last_exception.model_name}({last_exception.model_id})')
            raise last_exception
        else:
            raise Exception(f'Failed after {num_attempts} attempts.')

    async def _arun(self, run_response: TeamRunResponse, run_messages: RunMessages) -> None:
        """运行团队并返回响应。
        步骤：
        1.如果启用推理，则说明任务的原因
        2.从模型中获取响应
        3.更新run_response
        4.更新团队记忆
        5.计算会话度量
        6.将会话保存到存储
        7.解析任何结构化输出
        8.记录团队运行记录
        """
        self.memory = cast(TeamMemory, self.memory)
        self.model = cast(Model, self.model)
        if self.reasoning or self.reasoning_model is not None:
            reasoning_generator = self._areason(run_response=run_response, run_messages=run_messages)
            async for _ in reasoning_generator:
                pass
        self.run_messages = run_messages
        index_of_last_user_message = len(run_messages.messages)
        model_response = await self.model.aresponse(messages=run_messages.messages)
        if (self.response_model is not None) and not self.use_json_mode and (model_response.parsed is not None):
            run_response.content = model_response.parsed
            run_response.content_type = self.response_model.__name__
        else:
            if not run_response.content:
                run_response.content = model_response.content
            else:
                run_response.content += model_response.content
        if model_response.thinking is not None:
            if not run_response.thinking:
                run_response.thinking = model_response.thinking
            else:
                run_response.thinking += model_response.thinking
        if model_response.citations is not None:
            run_response.citations = model_response.citations
        if model_response.tool_calls is not None:
            if run_response.tools is None:
                run_response.tools = model_response.tool_calls
            else:
                run_response.tools.extend(model_response.tool_calls)
        run_response.formatted_tool_calls = format_tool_calls(run_response.tools)
        if model_response.audio is not None:
            run_response.response_audio = model_response.audio
        run_response.created_at = model_response.created_at
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        run_response.messages = messages_for_run_response
        run_response.metrics = self._aggregate_metrics_from_messages(messages_for_run_response)
        if run_messages.system_message is not None:
            self.memory.add_system_message(run_messages.system_message, system_message_role='system')
        messages_for_memory: List[Message] = (
            [run_messages.user_message] if run_messages.user_message is not None else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        team_run = TeamRun(response=run_response)
        team_run.message = run_messages.user_message
        if self.memory is not None and self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message is not None:
            await self.memory.aupdate_memory(input=run_messages.user_message.get_content_string())
        self.memory.add_team_run(team_run)
        self.write_to_storage()
        if self.response_model is not None:
            if isinstance(run_response.content, str) and self.parse_response:
                try:
                    parsed_response_content = parse_response_model_str(run_response.content, self.response_model)
                    if parsed_response_content is not None:
                        run_response.content = parsed_response_content
                        run_response.content_type = self.response_model.__name__
                    else:
                        print('Failed to convert response to response_model')
                except Exception as e:
                    print(f'Failed to convert response to output model: {e}')
            else:
                print('Something went wrong. Run response content is not a string')
        elif self._member_response_model is not None:
            if isinstance(run_response.content, str):
                try:
                    parsed_response_content = parse_response_model_str(run_response.content,
                                                                       self._member_response_model)
                    if parsed_response_content is not None:
                        run_response.content = parsed_response_content
                        run_response.content_type = self._member_response_model.__name__
                    else:
                        print('Failed to convert response to response_model')
                except Exception as e:
                    print(f'Failed to convert response to output model: {e}')
            else:
                print('Something went wrong. Run response content is not a string')
        await self._alog_team_run()
        print(f'Team Run End: {self.run_id}')

    async def _arun_stream(self, run_response: TeamRunResponse, run_messages: RunMessages,
                           stream_intermediate_steps: bool = False) -> AsyncIterator[TeamRunResponse]:
        """运行团队并返回响应。
        步骤：
        1.如果启用推理，则说明任务的原因
        2.从模型中获取响应
        3.更新run_response
        4.更新团队记忆
        5.计算会话度量
        6.将会话保存到存储
        7.记录团队运行日志
        """
        self.memory = cast(TeamMemory, self.memory)
        self.model = cast(Model, self.model)
        if self.reasoning or self.reasoning_model is not None:
            reasoning_generator = self._areason(run_response=run_response, run_messages=run_messages)
            async for reasoning_response in reasoning_generator:
                yield reasoning_response
        self.run_messages = run_messages
        index_of_last_user_message = len(run_messages.messages)
        if stream_intermediate_steps:
            yield self._create_run_response(content='Run started', event=RunEvent.run_started)
        full_model_response = ModelResponse()
        model_stream = self.model.aresponse_stream(messages=run_messages.messages)
        async for model_response_chunk in model_stream:
            if model_response_chunk.event == ModelResponseEvent.assistant_response.value:
                should_yield = False
                if model_response_chunk.content is not None:
                    if not full_model_response.content:
                        full_model_response.content = model_response_chunk.content
                    else:
                        full_model_response.content += model_response_chunk.content
                    should_yield = True
                if model_response_chunk.thinking is not None:
                    if not full_model_response.thinking:
                        full_model_response.thinking = model_response_chunk.thinking
                    else:
                        full_model_response.thinking += model_response_chunk.thinking
                    should_yield = True
                if model_response_chunk.citations is not None:
                    full_model_response.citations = model_response_chunk.citations
                    should_yield = True
                if model_response_chunk.audio is not None:
                    if full_model_response.audio is None:
                        full_model_response.audio = AudioResponse(id=str(uuid4()), content='', transcript='')
                    if model_response_chunk.audio.id is not None:
                        full_model_response.audio.id = model_response_chunk.audio.id
                    if model_response_chunk.audio.content is not None:
                        full_model_response.audio.content += model_response_chunk.audio.content
                    if model_response_chunk.audio.transcript is not None:
                        full_model_response.audio.transcript += model_response_chunk.audio.transcript
                    if model_response_chunk.audio.expires_at is not None:
                        full_model_response.audio.expires_at = model_response_chunk.audio.expires_at
                    if model_response_chunk.audio.mime_type is not None:
                        full_model_response.audio.mime_type = model_response_chunk.audio.mime_type
                    if model_response_chunk.audio.sample_rate is not None:
                        full_model_response.audio.sample_rate = model_response_chunk.audio.sample_rate
                    if model_response_chunk.audio.channels is not None:
                        full_model_response.audio.channels = model_response_chunk.audio.channels
                    should_yield = True
                if should_yield:
                    yield self._create_run_response(content=model_response_chunk.content,
                                                    thinking=model_response_chunk.thinking,
                                                    response_audio=model_response_chunk.audio,
                                                    citations=model_response_chunk.citations,
                                                    created_at=model_response_chunk.created_at)
            elif model_response_chunk.event == ModelResponseEvent.tool_call_started.value:
                tool_calls_list = model_response_chunk.tool_calls
                if tool_calls_list is not None:
                    if run_response.tools is None:
                        run_response.tools = tool_calls_list
                    else:
                        run_response.tools.extend(tool_calls_list)
                run_response.formatted_tool_calls = format_tool_calls(run_response.tools)
                if stream_intermediate_steps:
                    yield self._create_run_response(content=model_response_chunk.content,
                                                    event=RunEvent.tool_call_started, from_run_response=run_response)
            elif model_response_chunk.event == ModelResponseEvent.tool_call_completed.value:
                tool_calls_list = model_response_chunk.tool_calls
                if tool_calls_list is not None:
                    if run_response.tools:
                        tool_call_index_map = {tc['tool_call_id']: i
                                               for i, tc in enumerate(run_response.tools)
                                               if tc.get('tool_call_id') is not None}
                        for tool_call_dict in tool_calls_list:
                            tool_call_id = tool_call_dict.get('tool_call_id')
                            index = tool_call_index_map.get(tool_call_id)
                            if index is not None:
                                run_response.tools[index] = tool_call_dict
                    else:
                        run_response.tools = tool_calls_list
                    if stream_intermediate_steps:
                        yield self._create_run_response(content=model_response_chunk.content,
                                                        event=RunEvent.tool_call_completed,
                                                        from_run_response=run_response)
        if (self.response_model is not None) and not self.use_json_mode and (full_model_response.parsed is not None):
            run_response.content = full_model_response.parsed
        run_response.created_at = full_model_response.created_at
        if full_model_response.content is not None:
            run_response.content = full_model_response.content
        if full_model_response.thinking is not None:
            run_response.thinking = full_model_response.thinking
        if full_model_response.audio is not None:
            run_response.response_audio = full_model_response.audio
        if full_model_response.citations is not None:
            run_response.citations = full_model_response.citations
        messages_for_run_response = [m for m in run_messages.messages if m.add_to_agent_memory]
        run_response.messages = messages_for_run_response
        run_response.metrics = self._aggregate_metrics_from_messages(messages_for_run_response)
        if run_messages.system_message is not None:
            self.memory.add_system_message(run_messages.system_message, system_message_role='system')
        messages_for_memory: List[Message] = (
            [run_messages.user_message] if run_messages.user_message is not None else [])
        for _rm in run_messages.messages[index_of_last_user_message:]:
            if _rm.add_to_agent_memory:
                messages_for_memory.append(_rm)
        if len(messages_for_memory) > 0:
            self.memory.add_messages(messages=messages_for_memory)
        team_run = TeamRun(response=run_response)
        team_run.message = run_messages.user_message
        if self.memory is not None and self.memory.create_user_memories and self.memory.update_user_memories_after_run and run_messages.user_message is not None:
            await self.memory.aupdate_memory(input=run_messages.user_message.get_content_string())
        self.memory.add_team_run(team_run)
        self.write_to_storage()
        await self._alog_team_run()
        if stream_intermediate_steps:
            yield self._create_run_response(from_run_response=run_response, event=RunEvent.run_completed)
        print(f'Team Run End: {self.run_id}')

    def print_response(self, message: Optional[Union[List, Dict, str, Message]] = None, *, stream: bool = False,
                       stream_intermediate_steps: bool = False, show_message: bool = True, show_reasoning: bool = True,
                       show_reasoning_verbose: bool = False, console: Optional[Any] = None,
                       tags_to_include_in_markdown: Optional[Set[str]] = None, audio: Optional[Sequence[Audio]] = None,
                       images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None,
                       files: Optional[Sequence[File]] = None, markdown: Optional[bool] = None, **kwargs: Any) -> None:
        if not tags_to_include_in_markdown:
            tags_to_include_in_markdown = {'think', 'thinking'}
        if markdown is None:
            markdown = self.markdown
        else:
            self.markdown = markdown
        if not tags_to_include_in_markdown:
            tags_to_include_in_markdown = {'think', 'thinking'}
        with Live(console=console) as live_console:
            status = Status('Thinking...', spinner='aesthetic', speed=0.4, refresh_per_second=10)
            live_console.update(status)
            response_timer = Timer()
            response_timer.start()
            panels = [status]
            if message and show_message:
                message_content = get_text_from_message(message)
                message_panel = create_panel(content=Text(message_content, style='green'), title='Message',
                                             border_style='cyan')
                panels.append(message_panel)
                live_console.update(Group(*panels))
            run_response: TeamRunResponse = self.run(message=message, images=images, audio=audio, videos=videos,
                                                     files=files, stream=False, **kwargs)
            response_timer.stop()
            team_markdown = False
            member_markdown = {}
            if markdown:
                for member in self.members:
                    if isinstance(member, Agent) and member.agent_id is not None:
                        member_markdown[member.agent_id] = True
                    if isinstance(member, Team) and member.team_id is not None:
                        member_markdown[member.team_id] = True
                team_markdown = True
            if self.response_model is not None:
                team_markdown = False
            for member in self.members:
                if member.response_model is not None and isinstance(member, Agent) and member.agent_id is not None:
                    member_markdown[member.agent_id] = False
                if member.response_model is not None and isinstance(member, Team) and member.team_id is not None:
                    member_markdown[member.team_id] = False
            reasoning_steps = []
            if isinstance(run_response,
                          TeamRunResponse) and run_response.extra_data is not None and run_response.extra_data.reasoning_steps is not None:
                reasoning_steps = run_response.extra_data.reasoning_steps
            if len(reasoning_steps) > 0 and show_reasoning:
                for i, step in enumerate(reasoning_steps, 1):
                    reasoning_panel = self._build_reasoning_step_panel(i, step, show_reasoning_verbose)
                    panels.append(reasoning_panel)
                live_console.update(Group(*panels))
            if isinstance(run_response, TeamRunResponse) and run_response.thinking is not None:
                thinking_panel = create_panel(content=Text(run_response.thinking),
                                              title=f'Thinking ({response_timer.elapsed:.1f}s)', border_style='green')
                panels.append(thinking_panel)
                live_console.update(Group(*panels))
            if isinstance(run_response, TeamRunResponse):
                if self.show_members_responses:
                    for member_response in run_response.member_responses:
                        reasoning_steps = []
                        if isinstance(member_response,
                                      RunResponse) and member_response.extra_data is not None and member_response.extra_data.reasoning_steps is not None:
                            reasoning_steps.extend(member_response.extra_data.reasoning_steps)
                        if len(reasoning_steps) > 0 and show_reasoning:
                            for i, step in enumerate(reasoning_steps, 1):
                                member_reasoning_panel = self._build_reasoning_step_panel(i, step,
                                                                                          show_reasoning_verbose,
                                                                                          color='magenta')
                                panels.append(member_reasoning_panel)
                        if self.show_tool_calls and hasattr(member_response, 'tools') and member_response.tools:
                            member_name = None
                            if isinstance(member_response, RunResponse) and member_response.agent_id is not None:
                                member_name = self._get_member_name(member_response.agent_id)
                            elif isinstance(member_response, TeamRunResponse) and member_response.team_id is not None:
                                member_name = self._get_member_name(member_response.team_id)
                            if member_name:
                                formatted_calls = format_tool_calls(member_response.tools)
                                if formatted_calls:
                                    console_width = console.width if console else 80
                                    panel_width = console_width + 30
                                    lines = []
                                    for call in formatted_calls:
                                        wrapped_call = textwrap.fill(f'• {call}', width=panel_width,
                                                                     subsequent_indent='  ')
                                        lines.append(wrapped_call)
                                    tool_calls_text = '\n\n'.join(lines)
                                    member_tool_calls_panel = create_panel(content=tool_calls_text,
                                                                           title=f'{member_name} Tool Calls',
                                                                           border_style='yellow')
                                    panels.append(member_tool_calls_panel)
                                    live_console.update(Group(*panels))
                        show_markdown = False
                        if member_markdown:
                            if isinstance(member_response, RunResponse) and member_response.agent_id is not None:
                                show_markdown = member_markdown.get(member_response.agent_id, False)
                            elif isinstance(member_response, TeamRunResponse) and member_response.team_id is not None:
                                show_markdown = member_markdown.get(member_response.team_id, False)
                        member_response_content: Union[str, JSON, Markdown] = self._parse_response_content(
                            member_response, tags_to_include_in_markdown, show_markdown=show_markdown)
                        if isinstance(member_response, RunResponse) and member_response.agent_id is not None:
                            member_response_panel = create_panel(content=member_response_content,
                                                                 title=f'{self._get_member_name(member_response.agent_id)} Response',
                                                                 border_style='magenta')
                        elif isinstance(member_response, TeamRunResponse) and member_response.team_id is not None:
                            member_response_panel = create_panel(content=member_response_content,
                                                                 title=f'{self._get_member_name(member_response.team_id)} Response',
                                                                 border_style='magenta')
                        panels.append(member_response_panel)
                        if member_response.citations is not None and member_response.citations.urls is not None:
                            md_content = '\n'.join(
                                f'{i + 1}. [{citation.title or citation.url}]({citation.url})' for i, citation in
                                enumerate(member_response.citations.urls) if citation.url)
                            if md_content:
                                citations_panel = create_panel(content=Markdown(md_content), title='Citations',
                                                               border_style='magenta')
                                panels.append(citations_panel)
                    live_console.update(Group(*panels))
                if self.show_tool_calls and run_response.tools:
                    formatted_calls = format_tool_calls(run_response.tools)
                    if formatted_calls:
                        console_width = console.width if console else 80
                        panel_width = console_width + 30
                        lines = []
                        for call in formatted_calls:
                            wrapped_call = textwrap.fill(f'• {call}', width=panel_width, subsequent_indent='  ')
                            lines.append(wrapped_call)
                        tool_calls_text = '\n\n'.join(lines)
                        team_tool_calls_panel = create_panel(content=tool_calls_text, title='Team Tool Calls',
                                                             border_style='yellow')
                        panels.append(team_tool_calls_panel)
                        live_console.update(Group(*panels))
                response_content_batch: Union[str, JSON, Markdown] = self._parse_response_content(run_response,
                                                                                                  tags_to_include_in_markdown,
                                                                                                  show_markdown=team_markdown)
                response_panel = create_panel(content=response_content_batch,
                                              title=f'Response ({response_timer.elapsed:.1f}s)', border_style='blue')
                panels.append(response_panel)
                if run_response.citations is not None and run_response.citations.urls is not None:
                    md_content = '\n'.join(
                        f'{i + 1}. [{citation.title or citation.url}]({citation.url})' for i, citation in
                        enumerate(run_response.citations.urls) if citation.url)
                    if md_content:
                        citations_panel = create_panel(content=Markdown(md_content), title='Citations',
                                                       border_style='green')
                        panels.append(citations_panel)
            panels = [p for p in panels if not isinstance(p, Status)]
            live_console.update(Group(*panels))

    def _build_reasoning_step_panel(self, step_idx: int, step: ReasoningStep, show_reasoning_verbose: bool = False,
                                    color: str = 'green'):
        step_content = Text.assemble()
        if step.title is not None:
            step_content.append(f'{step.title}\n', 'bold')
        if step.action is not None:
            step_content.append(f'{step.action}\n', 'dim')
        if step.result is not None:
            step_content.append(Text.from_markup(step.result, style='dim'))
        if show_reasoning_verbose:
            if step.reasoning is not None:
                step_content.append(Text.from_markup(f'\n[bold]Reasoning:[/bold] {step.reasoning}', style='dim'))
            if step.confidence is not None:
                step_content.append(Text.from_markup(f'\n[bold]Confidence:[/bold] {step.confidence}', style='dim'))
        return create_panel(content=step_content, title=f'Reasoning step {step_idx}', border_style=color)

    def _get_member_name(self, entity_id: str) -> str:
        for member in self.members:
            if isinstance(member, Agent):
                if member.agent_id == entity_id:
                    return member.name or entity_id
            elif isinstance(member, Team):
                if member.team_id == entity_id:
                    return member.name or entity_id
        return entity_id

    def _parse_response_content(self, run_response: Union[TeamRunResponse, RunResponse],
                                tags_to_include_in_markdown: Set[str], show_markdown: bool = True) -> Any:
        if isinstance(run_response.content, str):
            if show_markdown:
                escaped_content = escape_markdown_tags(run_response.content, tags_to_include_in_markdown)
                return Markdown(escaped_content)
            else:
                return run_response.get_content_as_string(indent=4)
        elif isinstance(run_response.content, BaseModel):
            try:
                return JSON(run_response.content.model_dump_json(exclude_none=True), indent=2)
            except Exception as e:
                print(f'Failed to convert response to JSON: {e}')
        else:
            try:
                return JSON(json.dumps(run_response.content), indent=4)
            except Exception as e:
                print(f'Failed to convert response to JSON: {e}')

    def cli_app(self, message: Optional[str] = None, user: str = 'User', emoji: str = ':sunglasses:',
                stream: bool = False, markdown: bool = False, exit_on: Optional[List[str]] = None,
                **kwargs: Any) -> None:
        if message:
            self.print_response(message=message, stream=stream, markdown=markdown, **kwargs)
        _exit_on = exit_on or ['exit', 'quit', 'bye']
        while True:
            message = Prompt.ask(f'[bold] {emoji} {user} [/bold]')
            if message in _exit_on:
                break
            self.print_response(message=message, stream=stream, markdown=markdown, **kwargs)

    def _aggregate_metrics_from_messages(self, messages: List[Message]) -> Dict[str, Any]:
        aggregated_metrics: Dict[str, Any] = defaultdict(list)
        assistant_message_role = self.model.assistant_message_role if self.model is not None else 'assistant'
        for m in messages:
            if m.role == assistant_message_role and m.metrics is not None:
                for k, v in m.metrics.__dict__.items():
                    if k == 'timer':
                        continue
                    if v is not None:
                        aggregated_metrics[k].append(v)
        if aggregated_metrics is not None:
            aggregated_metrics = dict(aggregated_metrics)
        return aggregated_metrics

    def _reason(self, run_response: TeamRunResponse, run_messages: RunMessages,
                stream_intermediate_steps: bool = False) -> Iterator[TeamRunResponse]:
        if stream_intermediate_steps:
            yield self._create_run_response(from_run_response=run_response, content='Reasoning started',
                                            event=RunEvent.reasoning_started)
        reasoning_agent = self.reasoning_agent
        reasoning_message = get_deepseek_reasoning(reasoning_agent=reasoning_agent,
                                                   messages=run_messages.get_input_messages())
        if reasoning_message is None:
            print('Reasoning error. Reasoning response is None, continuing regular session...')
            return
        if reasoning_message:
            run_messages.messages.append(reasoning_message)
            update_run_response_with_reasoning(run_response=run_response,
                                               reasoning_steps=[ReasoningStep(result=reasoning_message.content)],
                                               reasoning_agent_messages=[reasoning_message])
        if stream_intermediate_steps:
            yield self._create_run_response(
                content=ReasoningSteps(reasoning_steps=[ReasoningStep(result=reasoning_message.content)]),
                content_type=ReasoningSteps.__class__.__name__, event=RunEvent.reasoning_completed)

    async def _areason(self, run_response: TeamRunResponse, run_messages: RunMessages,
                       stream_intermediate_steps: bool = False) -> AsyncIterator[TeamRunResponse]:
        if stream_intermediate_steps:
            yield self._create_run_response(from_run_response=run_response, content='Reasoning started',
                                            event=RunEvent.reasoning_started)
        reasoning_agent = self.reasoning_agent
        reasoning_message = await aget_deepseek_reasoning(reasoning_agent=reasoning_agent,
                                                          messages=run_messages.get_input_messages())
        if reasoning_message is None:
            print('Reasoning error. Reasoning response is None, continuing regular session...')
            return
        if reasoning_message:
            run_messages.messages.append(reasoning_message)
            update_run_response_with_reasoning(run_response=run_response,
                                               reasoning_steps=[ReasoningStep(result=reasoning_message.content)],
                                               reasoning_agent_messages=[reasoning_message])
        if stream_intermediate_steps:
            yield self._create_run_response(
                content=ReasoningSteps(reasoning_steps=[ReasoningStep(result=reasoning_message.content)]),
                content_type=ReasoningSteps.__class__.__name__, event=RunEvent.reasoning_completed)

    def _create_run_response(self, content: Optional[Any] = None, content_type: Optional[str] = None,
                             thinking: Optional[str] = None, event: RunEvent = RunEvent.run_response,
                             tools: Optional[List[Dict[str, Any]]] = None, audio: Optional[List[AudioArtifact]] = None,
                             images: Optional[List[ImageArtifact]] = None, videos: Optional[List[VideoArtifact]] = None,
                             response_audio: Optional[AudioResponse] = None, citations: Optional[Citations] = None,
                             model: Optional[str] = None, messages: Optional[List[Message]] = None,
                             created_at: Optional[int] = None,
                             from_run_response: Optional[TeamRunResponse] = None) -> TeamRunResponse:
        extra_data = None
        member_responses = None
        formatted_tool_calls = None
        if from_run_response:
            content = from_run_response.content
            content_type = from_run_response.content_type
            tools = from_run_response.tools
            audio = from_run_response.audio
            images = from_run_response.images
            videos = from_run_response.videos
            response_audio = from_run_response.response_audio
            model = from_run_response.model
            messages = from_run_response.messages
            extra_data = from_run_response.extra_data
            member_responses = from_run_response.member_responses
            citations = from_run_response.citations
            tools = from_run_response.tools
            formatted_tool_calls = from_run_response.formatted_tool_calls
        rr = TeamRunResponse(run_id=self.run_id, session_id=self.session_id, team_id=self.team_id, content=content,
                             thinking=thinking, tools=tools, audio=audio, images=images, videos=videos,
                             response_audio=response_audio, citations=citations, model=model, messages=messages,
                             extra_data=extra_data, event=event.value)
        if formatted_tool_calls:
            rr.formatted_tool_calls = formatted_tool_calls
        if member_responses:
            rr.member_responses = member_responses
        if content_type is not None:
            rr.content_type = content_type
        if created_at is not None:
            rr.created_at = created_at
        return rr

    def _resolve_run_context(self) -> None:
        print('Resolving context')
        if self.context is not None:
            if isinstance(self.context, dict):
                for ctx_key, ctx_value in self.context.items():
                    if callable(ctx_value):
                        try:
                            sig = inspect.signature(ctx_value)
                            if 'agent' in sig.parameters:
                                resolved_ctx_value = ctx_value(agent=self)
                            else:
                                resolved_ctx_value = ctx_value()
                            if resolved_ctx_value is not None:
                                self.context[ctx_key] = resolved_ctx_value
                        except Exception as e:
                            print(f'Failed to resolve context for {ctx_key}: {e}')
                    else:
                        self.context[ctx_key] = ctx_value
            else:
                print('Context is not a dict')

    def _configure_model(self, show_tool_calls: bool = False) -> None:
        if self.model is None:
            print('Setting default model to OpenAI Chat')
            self.model = Ollama()
        if self.response_model is None:
            self.model.response_format = None
        else:
            json_response_format = {'type': 'json_object'}
            if self.model.supports_native_structured_outputs:
                if self.use_json_mode:
                    self.model.response_format = json_response_format
                    self.model.structured_outputs = False
                else:
                    print('Setting Model.response_format to Agent.response_model')
                    self.model.response_format = self.response_model
                    self.model.structured_outputs = True
            elif self.model.supports_json_schema_outputs:
                if self.use_json_mode:
                    print('Setting Model.response_format to JSON response mode')
                    self.model.response_format = {'type': 'json_schema',
                                                  'json_schema': {'name': self.response_model.__name__,
                                                                  'schema': self.response_model.model_json_schema()}}
                else:
                    self.model.response_format = None
                self.model.structured_outputs = False
            else:
                self.model.response_format = json_response_format if self.use_json_mode else None
                self.model.structured_outputs = False
        self.model.show_tool_calls = show_tool_calls
        if self.tool_choice is not None:
            self.model.tool_choice = self.tool_choice
        if self.tool_call_limit is not None:
            self.model.tool_call_limit = self.tool_call_limit

    def _add_tools_to_model(self, model: Model, tools: List[Union[Function, Callable, Toolkit, Dict]]) -> None:
        self._functions_for_model = {}
        self._tools_for_model = []
        if len(tools) > 0:
            print('Processing tools for model')
            strict = False
            if self.response_model is not None and not self.use_json_mode and model.supports_native_structured_outputs:
                strict = True
            for tool in tools:
                if isinstance(tool, Dict):
                    self._tools_for_model.append(tool)
                    print(f'Included builtin tool {tool}')
                elif isinstance(tool, Toolkit):
                    for name, func in tool.functions.items():
                        if name not in self._functions_for_model:
                            func._agent = self
                            func.process_entrypoint(strict=strict)
                            if strict:
                                func.strict = True
                            self._functions_for_model[name] = func
                            self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                            print(f'Included function {name} from {tool.name}')
                elif isinstance(tool, Function):
                    if tool.name not in self._functions_for_model:
                        tool._agent = self
                        tool.process_entrypoint(strict=strict)
                        if strict and tool.strict is None:
                            tool.strict = True
                        self._functions_for_model[tool.name] = tool
                        self._tools_for_model.append({'type': 'function', 'function': tool.to_dict()})
                        print(f'Included function {tool.name}')
                elif callable(tool):
                    try:
                        func = Function.from_callable(tool, strict=strict)
                        func._agent = self
                        if strict:
                            func.strict = True
                        self._functions_for_model[func.name] = func
                        self._tools_for_model.append({'type': 'function', 'function': func.to_dict()})
                        print(f'Included function {func.name}')
                    except Exception as e:
                        print(f'Could not add function {tool}: {e}')
            model.set_tools(tools=self._tools_for_model)
            model.set_functions(functions=self._functions_for_model)

    def get_members_system_message_content(self, indent: int = 0) -> str:
        system_message_content = ''
        for idx, member in enumerate(self.members):
            if isinstance(member, Team):
                system_message_content += f'{indent * " "} - Team: {member.name}\n'
                if member.members is not None:
                    system_message_content += member.get_members_system_message_content(indent=indent + 2)
            else:
                system_message_content += f'{indent * " "} - Agent {idx + 1}:\n'
                if member.name is not None:
                    system_message_content += f'{indent * " "}   - Name: {member.name}\n'
                if member.role is not None:
                    system_message_content += f'{indent * " "}   - Role: {member.role}\n'
                if member.description is not None:
                    system_message_content += f'{indent * " "}   - Description: {member.description}\n'
                if member.tools is not None:
                    system_message_content += f'{indent * " "}   - Available tools:\n'
                    tool_name_and_description = []
                    for _tool in member.tools:
                        if isinstance(_tool, Toolkit):
                            for _func in _tool.functions.values():
                                if _func.entrypoint:
                                    tool_name_and_description.append(
                                        (_func.name, Function.get_entrypoint_docstring(_func.entrypoint)))
                        elif isinstance(_tool, Function) and _tool.entrypoint:
                            tool_name_and_description.append(
                                (_tool.name, Function.get_entrypoint_docstring(_tool.entrypoint)))
                        elif callable(_tool):
                            tool_name_and_description.append((_tool.__name__, Function.get_entrypoint_docstring(_tool)))
                    for _tool_name, _tool_description in tool_name_and_description:
                        system_message_content += f'{indent * " "}    - {_tool_name}: {_tool_description}\n'
        return system_message_content

    def get_system_message(self, audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None,
                           videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None) -> \
    Optional[Message]:
        self.model = cast(Model, self.model)
        instructions: List[str] = []
        if self.instructions is not None:
            _instructions = self.instructions
            if callable(self.instructions):
                _instructions = self.instructions(agent=self)
            if isinstance(_instructions, str):
                instructions.append(_instructions)
            elif isinstance(_instructions, list):
                instructions.extend(_instructions)
        _model_instructions = self.model.get_instructions_for_model()
        if _model_instructions is not None:
            instructions.extend(_model_instructions)
        additional_information: List[str] = []
        if self.markdown and self.response_model is None:
            additional_information.append('Use markdown to format your answers.')
        if self.add_datetime_to_instructions:
            additional_information.append(f'The current time is {datetime.now()}')
        system_message_content: str = ''
        if self.mode == 'coordinate':
            system_message_content += 'You are the leader of a team of AI Agents and possible Sub-Teams:\n'
            system_message_content += self.get_members_system_message_content()
            system_message_content += (
                '''\n-您可以直接响应，也可以将任务转移给团队中的其他代理，具体取决于他们可用的工具及其角色。\n-如果将任务转移给另一个代理，请确保包括：\n-agent_name（str）：要将任务传输到的代理的名称。\n-task_description（str）：任务的清晰描述。\n-expected_output（str）：预期输出。\n-您可以同时将任务传递给多个成员。\n-在响应用户之前，您必须始终验证其他代理的输出。\n-评估其他代理人的反应。如果你觉得任务已经完成，你可以停下来回应用户。\n如果你对结果不满意，可以重新分配任务。\n''')
        elif self.mode == 'route':
            system_message_content += 'You are the leader of a team of AI Agents and possible Sub-Teams:\n'
            system_message_content += self.get_members_system_message_content()
            system_message_content += '-您充当用户请求的路由器。您必须选择正确的代理来转发用户的请求。这应该是完成任务可能性最高的代理。\n-将任务转发给另一个代理时，请确保包括：\n-agent_name（str）：要将任务传输到的代理的名称。\n-expected_output（str）：预期输出。\n你应该尽最大努力把任务交给一个代理人。\n-如果用户请求需要它（例如，如果他们要求多个东西），您可以一次转发给多个代理。\n'
        elif self.mode == 'collaborate':
            system_message_content += 'You are leading a collaborative team of Agents and possible Sub-Teams:\n'
            system_message_content += self.get_members_system_message_content()
            system_message_content += '-对于团队中的所有代理，只调用run_member_agent一次。\n-考虑其他代理的所有回复，并评估任务是否已完成。\n如果你觉得任务已经完成，你可以停下来回应用户。\n'
            system_message_content += '\n'
        if self.enable_agentic_context:
            system_message_content += 'You can and should update the context of the team. Use the `set_team_context` tool to update the shared team context.\n'
        if self.name is not None:
            system_message_content += f'Your name is: {self.name}.\n\n'
        if self.success_criteria:
            system_message_content += f'<success_criteria>\nThe team will be considered successful if the following criteria are met: {self.success_criteria}\nStop the team run when the criteria are met.\n</success_criteria>\n\n'
        if self.description is not None:
            system_message_content += f'<description>\n{self.description}\n</description>\n\n'
        if len(instructions) > 0:
            system_message_content += '<instructions>'
            if len(instructions) > 1:
                for _upi in instructions:
                    system_message_content += f'\n- {_upi}'
            else:
                system_message_content += '\n' + instructions[0]
            system_message_content += '\n</instructions>\n\n'
        if len(additional_information) > 0:
            system_message_content += '<additional_information>'
            for _ai in additional_information:
                system_message_content += f'\n- {_ai}'
            system_message_content += '\n</additional_information>\n\n'
        if audio is not None or images is not None or videos is not None or files is not None:
            system_message_content += '<attached_media>\n'
            system_message_content += 'You have the following media attached to your message:\n'
            if audio is not None and len(audio) > 0:
                system_message_content += ' - Audio\n'
            if images is not None and len(images) > 0:
                system_message_content += ' - Images\n'
            if videos is not None and len(videos) > 0:
                system_message_content += ' - Videos\n'
            if files is not None and len(files) > 0:
                system_message_content += ' - Files\n'
            system_message_content += '</attached_media>\n\n'
        if self.add_state_in_messages:
            system_message_content = self._format_message_with_state_variables(system_message_content)
        system_message_from_model = self.model.get_system_message_for_model()
        if system_message_from_model is not None:
            system_message_content += system_message_from_model
        if self.expected_output is not None:
            system_message_content += f'<expected_output>\n{self.expected_output.strip()}\n</expected_output>\n\n'
        if self.response_model is not None and self.use_json_mode and self.model and self.model.supports_native_structured_outputs:
            system_message_content += f'{self._get_json_output_prompt()}'
        return Message(role='system', content=system_message_content.strip())

    def get_run_messages(self, *, run_response: TeamRunResponse, message: Union[str, List, Dict, Message],
                         audio: Optional[Sequence[Audio]] = None, images: Optional[Sequence[Image]] = None,
                         videos: Optional[Sequence[Video]] = None, files: Optional[Sequence[File]] = None,
                         **kwargs: Any) -> RunMessages:
        """此函数返回具有以下属性的RunMessages对象：
        -system_message：此运行的系统消息
        -user_message：此运行的用户消息
        -messages：要发送到模型的消息列表
        要构建RunMessages对象，请执行以下操作：
        1.将系统消息添加到run_message
        2.向run_message添加历史记录
        3.将用户消息添加到run_message
        """
        self.memory = cast(TeamMemory, self.memory)
        run_messages = RunMessages()
        system_message = self.get_system_message(images=images, audio=audio, videos=videos, files=files)
        if system_message is not None:
            run_messages.system_message = system_message
            run_messages.messages.append(system_message)
        if self.enable_team_history:
            history: List[Message] = self.memory.get_messages_from_last_n_runs(
                last_n=self.num_of_interactions_from_history, skip_role='system')
            if len(history) > 0:
                history_copy = [deepcopy(msg) for msg in history]
                for _msg in history_copy:
                    _msg.from_history = True
                print(f'Adding {len(history_copy)} messages from history')
                run_messages.messages += history_copy
        user_message = self._get_user_message(message, audio=audio, images=images, videos=videos, files=files, **kwargs)
        if user_message is not None:
            run_messages.user_message = user_message
            run_messages.messages.append(user_message)
        return run_messages

    def _get_user_message(self, message: Union[str, List, Dict, Message], audio: Optional[Sequence[Audio]] = None,
                          images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None,
                          files: Optional[Sequence[File]] = None, **kwargs):
        user_message_content: str = ''
        if isinstance(message, str) or isinstance(message, list):
            if self.add_state_in_messages:
                if isinstance(message, str):
                    user_message_content = self._format_message_with_state_variables(message)
                elif isinstance(message, list):
                    user_message_content = '\n'.join(
                        [self._format_message_with_state_variables(msg) for msg in message])
            else:
                if isinstance(message, str):
                    user_message_content = message
                else:
                    user_message_content = '\n'.join(message)
            if self.add_context and self.context is not None:
                user_message_content += '\n\n<context>\n'
                user_message_content += self._convert_context_to_string(self.context) + '\n'
                user_message_content += '</context>'
            return Message(role='user', content=user_message_content, audio=audio, images=images, videos=videos,
                           files=files, **kwargs)
        elif isinstance(message, Message):
            return message
        elif isinstance(message, dict):
            try:
                return Message.model_validate(message)
            except Exception as e:
                print(f'Failed to validate message: {e}')

    def _format_message_with_state_variables(self, message: str) -> Any:
        format_variables = ChainMap(self.session_state or {}, self.context or {}, self.extra_data or {},
                                    {'user_id': self.user_id} if self.user_id is not None else {})
        return SafeFormatter().format(message, **format_variables)

    def _convert_context_to_string(self, context: Dict[str, Any]) -> str:
        if context is None:
            return ''
        try:
            return json.dumps(context, indent=2, default=str)
        except (TypeError, ValueError, OverflowError) as e:
            print(f'Failed to convert context to JSON: {e}')
            sanitized_context = {}
            for key, value in context.items():
                try:
                    json.dumps({key: value}, default=str)
                    sanitized_context[key] = value
                except Exception as e:
                    print(f'Failed to serialize to JSON: {e}')
                    sanitized_context[key] = str(value)
            try:
                return json.dumps(sanitized_context, indent=2)
            except Exception as e:
                print(f'Failed to convert sanitized context to JSON: {e}')
                return str(context)

    def _get_json_output_prompt(self) -> str:
        json_output_prompt = 'Provide your output as a JSON containing the following fields:'
        if self.response_model is not None:
            if isinstance(self.response_model, str):
                json_output_prompt += '\n<json_fields>'
                json_output_prompt += f'\n{self.response_model}'
                json_output_prompt += '\n</json_fields>'
            elif isinstance(self.response_model, list):
                json_output_prompt += '\n<json_fields>'
                json_output_prompt += f'\n{json.dumps(self.response_model)}'
                json_output_prompt += '\n</json_fields>'
            elif issubclass(self.response_model, BaseModel):
                json_schema = self.response_model.model_json_schema()
                if json_schema is not None:
                    response_model_properties = {}
                    json_schema_properties = json_schema.get('properties')
                    if json_schema_properties is not None:
                        for field_name, field_properties in json_schema_properties.items():
                            formatted_field_properties = {prop_name: prop_value
                                                          for prop_name, prop_value in field_properties.items()
                                                          if prop_name != 'title'}
                            response_model_properties[field_name] = formatted_field_properties
                    json_schema_defs = json_schema.get('$defs')
                    if json_schema_defs is not None:
                        response_model_properties['$defs'] = {}
                        for def_name, def_properties in json_schema_defs.items():
                            def_fields = def_properties.get('properties')
                            formatted_def_properties = {}
                            if def_fields is not None:
                                for field_name, field_properties in def_fields.items():
                                    formatted_field_properties = {prop_name: prop_value
                                                                  for prop_name, prop_value in field_properties.items()
                                                                  if prop_name != 'title'}
                                    formatted_def_properties[field_name] = formatted_field_properties
                            if len(formatted_def_properties) > 0:
                                response_model_properties['$defs'][def_name] = formatted_def_properties
                    if len(response_model_properties) > 0:
                        json_output_prompt += '\n<json_fields>'
                        json_output_prompt += (
                            f'\n{json.dumps([key for key in response_model_properties.keys() if key != "$defs"])}')
                        json_output_prompt += '\n</json_fields>'
                        json_output_prompt += '\n\nHere are the properties for each field:'
                        json_output_prompt += '\n<json_field_properties>'
                        json_output_prompt += f'\n{json.dumps(response_model_properties, indent=2)}'
                        json_output_prompt += '\n</json_field_properties>'
            else:
                print(f'Could not build json schema for {self.response_model}')
        else:
            json_output_prompt += 'Provide the output as JSON.'
        json_output_prompt += '\nStart your response with `{` and end it with `}`.'
        json_output_prompt += '\nYour output will be passed to json.loads() to convert it to a Python object.'
        json_output_prompt += '\nMake sure it only contains valid JSON.'
        return json_output_prompt

    def _update_team_state(self, run_response: Union[TeamRunResponse, RunResponse]) -> None:
        if run_response.images is not None:
            if self.images is None:
                self.images = []
            self.images.extend(run_response.images)
        if run_response.videos is not None:
            if self.videos is None:
                self.videos = []
            self.videos.extend(run_response.videos)
        if run_response.audio is not None:
            if self.audio is None:
                self.audio = []
            self.audio.extend(run_response.audio)

    def get_team_history(self, num_chats: Optional[int] = None) -> str:
        history: List[Dict[str, Any]] = []
        all_chats = self.memory.get_all_messages()
        if len(all_chats) == 0:
            return ''
        chats_added = 0
        for chat in all_chats[::-1]:
            history.insert(0, chat[1].to_dict())
            history.insert(0, chat[0].to_dict())
            chats_added += 1
            if num_chats is not None and chats_added >= num_chats:
                break
        return json.dumps(history)

    def set_team_context(self, state: Union[str, dict]) -> str:
        if isinstance(state, str):
            self.memory.set_team_context_text(state)
        elif isinstance(state, dict):
            self.memory.set_team_context_text(json.dumps(state))
        msg = f'Current team context: {self.memory.get_team_context_str()}'
        print(msg)
        return msg

    def get_run_member_agents_function(self, stream: bool = False, async_mode: bool = False,
                                       images: Optional[List[Image]] = None, videos: Optional[List[Video]] = None,
                                       audio: Optional[List[Audio]] = None,
                                       files: Optional[List[File]] = None) -> Function:
        if not images:
            images = []
        if not videos:
            videos = []
        if not audio:
            audio = []
        if not files:
            files = []

        def run_member_agents(task_description: str, expected_output: Optional[str] = None) -> Iterator[str]:
            self.memory = cast(TeamMemory, self.memory)
            team_context_str = None
            if self.enable_agentic_context:
                team_context_str = self.memory.get_team_context_str()
            team_member_interactions_str = None
            if self.share_member_interactions:
                team_member_interactions_str = self.memory.get_team_member_interactions_str()
                if context_images := self.memory.get_team_context_images():
                    images.extend([Image.from_artifact(img) for img in context_images])
                if context_videos := self.memory.get_team_context_videos():
                    videos.extend([Video.from_artifact(vid) for vid in context_videos])
                if context_audio := self.memory.get_team_context_audio():
                    audio.extend([Audio.from_artifact(aud) for aud in context_audio])
            member_agent_task = 'You are a member of a team of agents that collaborate to complete a task.'
            if expected_output is not None:
                member_agent_task += f'\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if team_context_str:
                member_agent_task += f'\n\n{team_context_str}'
            if team_member_interactions_str:
                member_agent_task += f'\n\n{team_member_interactions_str}'
            member_agent_task += f'\n\n<task>\n{task_description}\n</task>'
            for member_agent_index, member_agent in enumerate(self.members):
                if stream:
                    member_agent_run_response_stream = member_agent.run(member_agent_task, images=images, videos=videos,
                                                                        audio=audio, files=files, stream=True)
                    for member_agent_run_response_chunk in member_agent_run_response_stream:
                        check_if_run_cancelled(member_agent_run_response_chunk)
                        yield member_agent_run_response_chunk.content or ''
                else:
                    member_agent_run_response = member_agent.run(member_agent_task, images=images, videos=videos,
                                                                 audio=audio, files=files, stream=False)
                    check_if_run_cancelled(member_agent_run_response)
                    if member_agent_run_response.content is None:
                        yield 'No response from the member agent.'
                    elif isinstance(member_agent_run_response.content, str):
                        yield member_agent_run_response.content
                    elif issubclass(type(member_agent_run_response.content), BaseModel):
                        try:
                            yield member_agent_run_response.content.model_dump_json(indent=2)
                        except Exception as e:
                            yield str(e)
                    else:
                        try:
                            yield json.dumps(member_agent_run_response.content, indent=2)
                        except Exception as e:
                            yield str(e)
                member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
                self.memory = cast(TeamMemory, self.memory)
                self.memory.add_interaction_to_team_context(member_name=member_name, task=task_description,
                                                            run_response=member_agent.run_response)
                self.run_response = cast(TeamRunResponse, self.run_response)
                self.run_response.add_member_run(member_agent.run_response)
                self._update_team_state(member_agent.run_response)

        async def arun_member_agents(task_description: str, expected_output: Optional[str] = None) -> AsyncIterator[
            str]:
            self.memory = cast(TeamMemory, self.memory)
            team_context_str = None
            if self.enable_agentic_context:
                team_context_str = self.memory.get_team_context_str()
            team_member_interactions_str = None
            if self.share_member_interactions:
                team_member_interactions_str = self.memory.get_team_member_interactions_str()
                if context_images := self.memory.get_team_context_images():
                    images.extend([Image.from_artifact(img) for img in context_images])
                if context_videos := self.memory.get_team_context_videos():
                    videos.extend([Video.from_artifact(vid) for vid in context_videos])
                if context_audio := self.memory.get_team_context_audio():
                    audio.extend([Audio.from_artifact(aud) for aud in context_audio])
            member_agent_task = 'You are a member of a team of agents that collaborate to complete a task.'
            if expected_output is not None:
                member_agent_task += f'\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if team_context_str:
                member_agent_task += f'\n\n{team_context_str}'
            if team_member_interactions_str:
                member_agent_task += f'\n\n{team_member_interactions_str}'
            member_agent_task += f'\n\n<task>\n{task_description}\n</task>'
            tasks = []
            for member_agent_index, member_agent in enumerate(self.members):
                current_agent = member_agent
                current_index = member_agent_index

                async def run_member_agent(agent=current_agent, idx=current_index) -> str:
                    response = await agent.arun(member_agent_task, images=images, videos=videos, audio=audio,
                                                files=files, stream=False)
                    check_if_run_cancelled(response)
                    member_name = agent.name if agent.name else f'agent_{idx}'
                    self.memory = cast(TeamMemory, self.memory)
                    self.memory.add_interaction_to_team_context(member_name=member_name, task=task_description,
                                                                run_response=agent.run_response)
                    self.run_response = cast(TeamRunResponse, self.run_response)
                    self.run_response.add_member_run(agent.run_response)
                    self._update_team_state(agent.run_response)
                    if response.content is None:
                        return f'Agent {member_name}: No response from the member agent.'
                    elif isinstance(response.content, str):
                        return f'Agent {member_name}: {response.content}'
                    elif issubclass(type(response.content), BaseModel):
                        try:
                            return f'Agent {member_name}: {response.content.model_dump_json(indent=2)}'
                        except Exception as e:
                            return f'Agent {member_name}: Error - {str(e)}'
                    else:
                        try:
                            return f'Agent {member_name}: {json.dumps(response.content, indent=2)}'
                        except Exception as e:
                            return f'Agent {member_name}: Error - {str(e)}'

                tasks.append(run_member_agent)
            results = await asyncio.gather(*[task() for task in tasks])
            for result in results:
                yield result

        if async_mode:
            run_member_agents_function = arun_member_agents
        else:
            run_member_agents_function = run_member_agents
        run_member_agents_func = Function.from_callable(run_member_agents_function, strict=True)
        return run_member_agents_func

    def get_transfer_task_function(self, stream: bool = False, async_mode: bool = False,
                                   images: Optional[List[Image]] = None, videos: Optional[List[Video]] = None,
                                   audio: Optional[List[Audio]] = None, files: Optional[List[File]] = None) -> Function:
        if not images:
            images = []
        if not videos:
            videos = []
        if not audio:
            audio = []
        if not files:
            files = []

        def transfer_task_to_member(agent_name: str, task_description: str, expected_output: str) -> Iterator[str]:
            self.memory = cast(TeamMemory, self.memory)
            result = self._find_member_by_name(agent_name)
            if result is None:
                yield f'Agent with name {agent_name} not found in the team or any subteams. Please choose the correct agent from the list of agents.'
                return
            member_agent_index, member_agent = result
            self._initialize_member(member_agent)
            team_context_str = None
            if self.enable_agentic_context:
                team_context_str = self.memory.get_team_context_str()
            team_member_interactions_str = None
            if self.share_member_interactions:
                team_member_interactions_str = self.memory.get_team_member_interactions_str()
                if context_images := self.memory.get_team_context_images():
                    images.extend([Image.from_artifact(img) for img in context_images])
                if context_videos := self.memory.get_team_context_videos():
                    videos.extend([Video.from_artifact(vid) for vid in context_videos])
                if context_audio := self.memory.get_team_context_audio():
                    audio.extend([Audio.from_artifact(aud) for aud in context_audio])
            member_agent_task = f'You are a member of a team of agents. Your goal is to complete the following task:\n\n{task_description}\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if team_context_str:
                member_agent_task += f'\n\n{team_context_str}'
            if team_member_interactions_str:
                member_agent_task += f'\n\n{team_member_interactions_str}'
            if stream:
                member_agent_run_response_stream = member_agent.run(member_agent_task, images=images, videos=videos,
                                                                    audio=audio, files=files, stream=True)
                for member_agent_run_response_chunk in member_agent_run_response_stream:
                    check_if_run_cancelled(member_agent_run_response_chunk)
                    yield member_agent_run_response_chunk.content or ''
            else:
                member_agent_run_response = member_agent.run(member_agent_task, images=images, videos=videos,
                                                             audio=audio, files=files, stream=False)
                check_if_run_cancelled(member_agent_run_response)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), BaseModel):
                    try:
                        yield member_agent_run_response.content.model_dump_json(indent=2)
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
            self.memory = cast(TeamMemory, self.memory)
            self.memory.add_interaction_to_team_context(member_name=member_name, task=task_description,
                                                        run_response=member_agent.run_response)
            self.run_response = cast(TeamRunResponse, self.run_response)
            self.run_response.add_member_run(member_agent.run_response)
            self._update_team_state(member_agent.run_response)

        async def atransfer_task_to_member(agent_name: str, task_description: str, expected_output: str) -> \
        AsyncIterator[str]:
            self.memory = cast(TeamMemory, self.memory)
            result = self._find_member_by_name(agent_name)
            if result is None:
                yield f'Agent with name {agent_name} not found in the team or any subteams. Please choose the correct agent from the list of agents.'
                return
            member_agent_index, member_agent = result
            self._initialize_member(member_agent)
            team_context_str = None
            if self.enable_agentic_context:
                team_context_str = self.memory.get_team_context_str()
            team_member_interactions_str = None
            if self.share_member_interactions:
                team_member_interactions_str = self.memory.get_team_member_interactions_str()
                if context_images := self.memory.get_team_context_images():
                    images.extend([Image.from_artifact(img) for img in context_images])
                if context_videos := self.memory.get_team_context_videos():
                    videos.extend([Video.from_artifact(vid) for vid in context_videos])
                if context_audio := self.memory.get_team_context_audio():
                    audio.extend([Audio.from_artifact(aud) for aud in context_audio])
            member_agent_task = f'You are a member of a team of agents. Your goal is to complete the following task:\n\n{task_description}\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if team_context_str:
                member_agent_task += f'\n\n{team_context_str}'
            if team_member_interactions_str:
                member_agent_task += f'\n\n{team_member_interactions_str}'
            if stream:
                member_agent_run_response_stream = await member_agent.arun(member_agent_task, images=images,
                                                                           videos=videos, audio=audio, files=files,
                                                                           stream=True)
                async for member_agent_run_response_chunk in member_agent_run_response_stream:
                    check_if_run_cancelled(member_agent_run_response_chunk)
                    yield member_agent_run_response_chunk.content or ''
            else:
                member_agent_run_response = await member_agent.arun(member_agent_task, images=images, videos=videos,
                                                                    audio=audio, files=files, stream=False)
                check_if_run_cancelled(member_agent_run_response)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), BaseModel):
                    try:
                        yield member_agent_run_response.content.model_dump_json(indent=2)
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
            self.memory.add_interaction_to_team_context(member_name=member_name, task=task_description,
                                                        run_response=member_agent.run_response)
            self.run_response = cast(TeamRunResponse, self.run_response)
            self.run_response.add_member_run(member_agent.run_response)
            self._update_team_state(member_agent.run_response)

        if async_mode:
            transfer_function = atransfer_task_to_member
        else:
            transfer_function = transfer_task_to_member
        transfer_func = Function.from_callable(transfer_function, strict=True)
        return transfer_func

    def _find_member_by_name(self, agent_name: str):
        for i, member in enumerate(self.members):
            if member.name == agent_name:
                return (i, member)
            if isinstance(member, Team):
                result = member._find_member_by_name(agent_name)
                if result is not None:
                    return (i, member)
        return None

    def get_forward_task_function(self, message: Message, stream: bool = False, async_mode: bool = False,
                                  images: Optional[Sequence[Image]] = None, videos: Optional[Sequence[Video]] = None,
                                  audio: Optional[Sequence[Audio]] = None,
                                  files: Optional[Sequence[File]] = None) -> Function:
        if not images:
            images = []
        if not videos:
            videos = []
        if not audio:
            audio = []
        if not files:
            files = []

        def forward_task_to_member(agent_name: str, expected_output: Optional[str] = None) -> Iterator[str]:
            self.memory = cast(TeamMemory, self.memory)
            self._member_response_model = None
            result = self._find_member_by_name(agent_name)
            if result is None:
                yield f'Agent with name {agent_name} not found in the team or any subteams. Please choose the correct agent from the list of agents.'
                return
            member_agent_index, member_agent = result
            self._initialize_member(member_agent)
            if member_agent.response_model is not None:
                self._member_response_model = member_agent.response_model
            member_agent_task = message.get_content_string()
            if expected_output:
                member_agent_task += f'\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if stream:
                member_agent_run_response_stream = member_agent.run(member_agent_task, images=images, videos=videos,
                                                                    audio=audio, files=files, stream=True)
                for member_agent_run_response_chunk in member_agent_run_response_stream:
                    yield member_agent_run_response_chunk.content or ''
            else:
                member_agent_run_response = member_agent.run(member_agent_task, images=images, videos=videos,
                                                             audio=audio, files=files, stream=False)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), BaseModel):
                    try:
                        yield member_agent_run_response.content.model_dump_json(indent=2)
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
            self.memory = cast(TeamMemory, self.memory)
            self.memory.add_interaction_to_team_context(member_name=member_name, task=message.get_content_string(),
                                                        run_response=member_agent.run_response)
            self.run_response = cast(TeamRunResponse, self.run_response)
            self.run_response.add_member_run(member_agent.run_response)
            self._update_team_state(member_agent.run_response)

        async def aforward_task_to_member(agent_name: str, expected_output: Optional[str] = None) -> AsyncIterator[str]:
            self.memory = cast(TeamMemory, self.memory)
            self._member_response_model = None
            result = self._find_member_by_name(agent_name)
            if result is None:
                yield f'Agent with name {agent_name} not found in the team or any subteams. Please choose the correct agent from the list of agents.'
                return
            member_agent_index, member_agent = result
            self._initialize_member(member_agent)
            if member_agent.response_model is not None:
                self._member_response_model = member_agent.response_model
            member_agent_task = message.get_content_string()
            if expected_output:
                member_agent_task += f'\n\n<expected_output>\n{expected_output}\n</expected_output>'
            if stream:
                member_agent_run_response_stream = await member_agent.arun(member_agent_task, images=images,
                                                                           videos=videos, audio=audio, files=files,
                                                                           stream=True)
                async for member_agent_run_response_chunk in member_agent_run_response_stream:
                    check_if_run_cancelled(member_agent_run_response_chunk)
                    yield member_agent_run_response_chunk.content or ''
            else:
                member_agent_run_response = await member_agent.arun(member_agent_task, images=images, videos=videos,
                                                                    audio=audio, files=files, stream=False)
                if member_agent_run_response.content is None:
                    yield 'No response from the member agent.'
                elif isinstance(member_agent_run_response.content, str):
                    yield member_agent_run_response.content
                elif issubclass(type(member_agent_run_response.content), BaseModel):
                    try:
                        yield member_agent_run_response.content.model_dump_json(indent=2)
                    except Exception as e:
                        yield str(e)
                else:
                    try:
                        yield json.dumps(member_agent_run_response.content, indent=2)
                    except Exception as e:
                        yield str(e)
            member_name = member_agent.name if member_agent.name else f'agent_{member_agent_index}'
            self.memory = cast(TeamMemory, self.memory)
            self.memory.add_interaction_to_team_context(member_name=member_name, task=message.get_content_string(),
                                                        run_response=member_agent.run_response)
            self.run_response = cast(TeamRunResponse, self.run_response)
            self.run_response.add_member_run(member_agent.run_response)
            self._update_team_state(member_agent.run_response)

        if async_mode:
            forward_function = aforward_task_to_member
        else:
            forward_function = forward_task_to_member
        forward_func = Function.from_callable(forward_function, strict=True)
        forward_func.stop_after_tool_call = True
        forward_func.show_result = True
        return forward_func

    def load_user_memories(self) -> None:
        self.memory = cast(TeamMemory, self.memory)
        if self.memory and self.memory.create_user_memories:
            if self.user_id is not None:
                self.memory.user_id = self.user_id
            self.memory.load_user_memories()
            if self.user_id is not None:
                print(f'Memories loaded for user: {self.user_id}')
            else:
                print('Memories loaded')

    def read_from_storage(self) -> Optional[TeamSession]:
        if self.storage is not None and self.session_id is not None:
            self.team_session = cast(TeamSession, self.storage.read(session_id=self.session_id))
            if self.team_session is not None:
                self.load_team_session(session=self.team_session)
            self.load_user_memories()
        return self.team_session

    def write_to_storage(self) -> Optional[TeamSession]:
        if self.storage is not None:
            self.team_session = cast(TeamSession, self.storage.upsert(session=self._get_team_session()))
        return self.team_session

    def rename_session(self, session_name: str) -> None:
        self.read_from_storage()
        self.session_name = session_name
        self.write_to_storage()
        self._log_team_session()

    def delete_session(self, session_id: str) -> None:
        if self.storage is not None:
            self.storage.delete_session(session_id=session_id)

    def load_team_session(self, session: TeamSession):
        if self.team_id is None and session.team_id is not None:
            self.team_id = session.team_id
        if self.user_id is None and session.user_id is not None:
            self.user_id = session.user_id
        if self.session_id is None and session.session_id is not None:
            self.session_id = session.session_id
        if session.team_data is not None:
            if self.name is None and 'name' in session.team_data:
                self.name = session.team_data.get('name')
        if session.session_data is not None:
            if self.session_name is None and 'session_name' in session.session_data:
                self.session_name = session.session_data.get('session_name')
            if 'session_state' in session.session_data:
                session_state_from_db = session.session_data.get('session_state')
                if session_state_from_db is not None and isinstance(session_state_from_db, dict) and len(
                        session_state_from_db) > 0:
                    if self.session_state is not None and len(self.session_state) > 0:
                        merge_dictionaries(session_state_from_db, self.session_state)
                    self.session_state = session_state_from_db
            if 'images' in session.session_data:
                images_from_db = session.session_data.get('images')
                if images_from_db is not None and isinstance(images_from_db, list):
                    if self.images is None:
                        self.images = []
                    self.images.extend([ImageArtifact.model_validate(img) for img in images_from_db])
            if 'videos' in session.session_data:
                videos_from_db = session.session_data.get('videos')
                if videos_from_db is not None and isinstance(videos_from_db, list):
                    if self.videos is None:
                        self.videos = []
                    self.videos.extend([VideoArtifact.model_validate(vid) for vid in videos_from_db])
            if 'audio' in session.session_data:
                audio_from_db = session.session_data.get('audio')
                if audio_from_db is not None and isinstance(audio_from_db, list):
                    if self.audio is None:
                        self.audio = []
                    self.audio.extend([AudioArtifact.model_validate(aud) for aud in audio_from_db])
        if session.extra_data is not None:
            if self.extra_data is not None:
                merge_dictionaries(session.extra_data, self.extra_data)
            self.extra_data = session.extra_data
        if self.memory is None:
            self.memory = session.memory
        if not isinstance(self.memory, TeamMemory):
            if isinstance(self.memory, dict):
                self.memory = TeamMemory(**self.memory)
            elif self.memory is not None:
                raise TypeError(f'Expected memory to be a dict or TeamMemory, but got {type(self.memory)}')
        if session.memory is not None and self.memory is not None:
            try:
                if 'runs' in session.memory:
                    try:
                        self.memory.runs = [TeamRun.from_dict(m) for m in session.memory['runs']]
                    except Exception as e:
                        print(f'Failed to load runs from memory: {e}')
                if 'messages' in session.memory:
                    try:
                        self.memory.messages = [Message.model_validate(m) for m in session.memory['messages']]
                    except Exception as e:
                        print(f'Failed to load messages from memory: {e}')
                if 'memories' in session.memory:
                    try:
                        self.memory.memories = [Memory.model_validate(m) for m in session.memory['memories']]
                    except Exception as e:
                        print(f'Failed to load user memories: {e}')
            except Exception as e:
                print(f'Failed to load AgentMemory: {e}')
        print(f'-*- TeamSession loaded: {session.session_id}')

    def _create_run_data(self) -> Dict[str, Any]:
        run_response_format = 'text'
        if self.response_model is not None:
            run_response_format = 'json'
        elif self.markdown:
            run_response_format = 'markdown'
        functions = {}
        if self.model is not None and self.model._functions is not None:
            functions = {f_name: func.to_dict() for f_name, func in self.model._functions.items() if
                         isinstance(func, Function)}
        run_data: Dict[str, Any] = {'functions': functions, 'metrics': self.run_response.metrics}
        if self.monitoring:
            run_data.update({'run_input': self.run_input, 'run_response': self.run_response.to_dict(),
                             'run_response_format': run_response_format})
        return run_data

    def _get_team_data(self) -> Dict[str, Any]:
        team_data: Dict[str, Any] = {}
        if self.name is not None:
            team_data['name'] = self.name
        if self.team_id is not None:
            team_data['team_id'] = self.team_id
        if self.model is not None:
            team_data['model'] = self.model.to_dict()
        return team_data

    def _get_session_data(self) -> Dict[str, Any]:
        session_data: Dict[str, Any] = {}
        if self.session_name is not None:
            session_data['session_name'] = self.session_name
        if self.session_state is not None and len(self.session_state) > 0:
            session_data['session_state'] = self.session_state
        if self.images is not None:
            session_data['images'] = [img.model_dump() for img in self.images]
        if self.videos is not None:
            session_data['videos'] = [vid.model_dump() for vid in self.videos]
        if self.audio is not None:
            session_data['audio'] = [aud.model_dump() for aud in self.audio]
        return session_data

    def _get_team_session(self) -> TeamSession:
        return TeamSession(session_id=self.session_id, team_id=self.team_id, user_id=self.user_id,
                           team_session_id=self.team_session_id,
                           memory=self.memory.to_dict() if self.memory is not None else None,
                           team_data=self._get_team_data(), session_data=self._get_session_data(),
                           extra_data=self.extra_data, created_at=int(time()))

    def _log_team_run(self) -> None:
        if not self.telemetry and not self.monitoring:
            return
        try:
            run_data = self._create_run_data()
            team_session: TeamSession = self.team_session or self._get_team_session()
        except Exception as e:
            print(f'Could not create team event: {e}')

    async def _alog_team_run(self) -> None:
        if not self.telemetry and not self.monitoring:
            return
        try:
            run_data = self._create_run_data()
            team_session: TeamSession = self.team_session or self._get_team_session()
        except Exception as e:
            print(f'Could not create team event: {e}')

    def _log_team_session(self):
        if not (self.telemetry or self.monitoring):
            return
        try:
            team_session: TeamSession = self.team_session or self._get_team_session()
        except Exception as e:
            print(f'Could not create team monitor: {e}')


class Workflow:
    """工作流"""
    def __init__(self, name: Optional[str] = None, workflow_id: Optional[str] = None, description: Optional[str] = None, user_id: Optional[str] = None, session_id: Optional[str] = None, session_name: Optional[str] = None, session_state: Optional[Dict[str, Any]] = None, memory: Optional[WorkflowMemory] = None, extra_data: Optional[Dict[str, Any]] = None, debug_mode: bool = False, monitoring: bool = False, telemetry: bool = True):
        self.name = name or self.__class__.__name__
        self.workflow_id = workflow_id
        self.description = description or self.__class__.__doc__
        self.user_id = user_id
        self.session_id = session_id
        self.session_name = session_name
        self.session_state: Dict[str, Any] = session_state or {}
        self.memory = memory
        self.extra_data = extra_data
        self.debug_mode = debug_mode
        self.monitoring = monitoring
        self.telemetry = telemetry
        self.run_id = None
        self.run_input = None
        self.run_response = None
        self.images = None
        self.videos = None
        self.audio = None
        self.workflow_session: Optional[WorkflowSession] = None
        self._subclass_run: Optional[Callable] = None
        self._run_parameters: Optional[Dict[str, Any]] = None
        self._run_return_type: Optional[str] = None
        self.update_run_method()
        for field_name, value in self.__class__.__dict__.items():
            if isinstance(value, Agent):
                value.session_id = self.session_id

    def run(self, **kwargs: Any):
        print(f'{self.__class__.__name__}.run() method not implemented.')
        return

    def run_workflow(self, **kwargs: Any):
        self.set_debug()
        self.set_workflow_id()
        self.set_session_id()
        self.initialize_memory()
        self.memory = cast(WorkflowMemory, self.memory)
        self.run_id = str(uuid4())
        self.run_input = kwargs
        self.run_response = RunResponse(run_id=self.run_id, session_id=self.session_id, workflow_id=self.workflow_id)
        self.update_agent_session_ids()
        print(f'*********** Workflow Run Start: {self.run_id} ***********')
        try:
            self._subclass_run = cast(Callable, self._subclass_run)
            result = self._subclass_run(**kwargs)
        except Exception as e:
            print(f'Workflow.run() failed: {e}')
            raise e
        if isinstance(result, (GeneratorType, collections.abc.Iterator)):
            self.run_response.content = ''
            def result_generator():
                self.run_response = cast(RunResponse, self.run_response)
                self.memory = cast(WorkflowMemory, self.memory)
                for item in result:
                    if isinstance(item, RunResponse):
                        item.run_id = self.run_id
                        item.session_id = self.session_id
                        item.workflow_id = self.workflow_id
                        if item.content is not None and isinstance(item.content, str):
                            self.run_response.content += item.content
                    else:
                        print(f'Workflow.run() should only yield RunResponse objects, got: {type(item)}')
                    yield item
                self.memory.add_run(WorkflowRun(input=self.run_input, response=self.run_response))
                print(f'*********** Workflow Run End: {self.run_id} ***********')
            return result_generator()
        elif isinstance(result, RunResponse):
            result.run_id = self.run_id
            result.session_id = self.session_id
            result.workflow_id = self.workflow_id
            if result.content is not None and isinstance(result.content, str):
                self.run_response.content = result.content
            self.memory.add_run(WorkflowRun(input=self.run_input, response=self.run_response))
            print(f'*********** Workflow Run End: {self.run_id} ***********')
            return result
        else:
            print(f'Workflow.run() should only return RunResponse objects, got: {type(result)}')
            return None

    def set_workflow_id(self) -> str:
        if self.workflow_id is None:
            self.workflow_id = str(uuid4())
        print(f'*********** Workflow ID: {self.workflow_id} ***********')
        return self.workflow_id

    def set_session_id(self) -> str:
        if self.session_id is None:
            self.session_id = str(uuid4())
        print(f'*********** Session ID: {self.session_id} ***********')
        return self.session_id

    def set_debug(self) -> None:
        if self.debug_mode or os.getenv('AGNO_DEBUG', 'false').lower() == 'true':
            self.debug_mode = True
            print('Debug logs enabled')

    def set_monitoring(self) -> None:
        if self.monitoring or os.getenv('AGNO_MONITOR', 'false').lower() == 'true':
            self.monitoring = True
        else:
            self.monitoring = False
        if self.telemetry or os.getenv('AGNO_TELEMETRY', 'true').lower() == 'true':
            self.telemetry = True
        else:
            self.telemetry = False

    def initialize_memory(self) -> None:
        if self.memory is None:
            self.memory = WorkflowMemory()

    def update_run_method(self):
        if self.__class__.run is not Workflow.run:
            self._subclass_run = self.__class__.run.__get__(self)
            sig = inspect.signature(self.__class__.run)
            self._run_parameters = {param_name: {'name': param_name, 'default': param.default.default
                    if hasattr(param.default, '__class__') and param.default.__class__.__name__ == 'FieldInfo'
                    else (param.default if param.default is not inspect.Parameter.empty else None), 'annotation': (param.annotation.__name__
                        if hasattr(param.annotation, '__name__')
                        else (str(param.annotation).replace('typing.Optional[', '').replace(']', '')
                            if 'typing.Optional' in str(param.annotation)
                            else str(param.annotation)))
                    if param.annotation is not inspect.Parameter.empty
                    else None, 'required': param.default is inspect.Parameter.empty}
                for param_name, param in sig.parameters.items()
                if param_name != 'self'}
            return_annotation = sig.return_annotation
            self._run_return_type = (return_annotation.__name__
                if return_annotation is not inspect.Signature.empty and hasattr(return_annotation, '__name__')
                else str(return_annotation)
                if return_annotation is not inspect.Signature.empty
                else None)
            object.__setattr__(self, 'run', self.run_workflow.__get__(self))
        else:
            self._subclass_run = self.run
            self._run_parameters = {}
            self._run_return_type = None

    def update_agent_session_ids(self):
        for field_name, value in self.__class__.__dict__.items():
            if isinstance(value, Agent):
                field_value = getattr(self, field_name)
                field_value.session_id = self.session_id

    def get_workflow_data(self) -> Dict[str, Any]:
        workflow_data: Dict[str, Any] = {}
        if self.name is not None:
            workflow_data['name'] = self.name
        if self.workflow_id is not None:
            workflow_data['workflow_id'] = self.workflow_id
        if self.description is not None:
            workflow_data['description'] = self.description
        return workflow_data

    def get_session_data(self) -> Dict[str, Any]:
        session_data: Dict[str, Any] = {}
        if self.session_name is not None:
            session_data['session_name'] = self.session_name
        if self.session_state and len(self.session_state) > 0:
            session_data['session_state'] = self.session_state
        if self.images is not None:
            session_data['images'] = [img.model_dump() for img in self.images]
        if self.videos is not None:
            session_data['videos'] = [vid.model_dump() for vid in self.videos]
        if self.audio is not None:
            session_data['audio'] = [aud.model_dump() for aud in self.audio]
        return session_data

    def get_workflow_session(self) -> WorkflowSession:
        self.memory = cast(WorkflowMemory, self.memory)
        self.session_id = cast(str, self.session_id)
        self.workflow_id = cast(str, self.workflow_id)
        return WorkflowSession(session_id=self.session_id, workflow_id=self.workflow_id, user_id=self.user_id, memory=self.memory.to_dict() if self.memory is not None else None, workflow_data=self.get_workflow_data(), session_data=self.get_session_data(), extra_data=self.extra_data)

    def load_workflow_session(self, session: WorkflowSession):
        if self.workflow_id is None and session.workflow_id is not None:
            self.workflow_id = session.workflow_id
        if self.user_id is None and session.user_id is not None:
            self.user_id = session.user_id
        if self.session_id is None and session.session_id is not None:
            self.session_id = session.session_id
        if session.workflow_data is not None:
            if self.name is None and 'name' in session.workflow_data:
                self.name = session.workflow_data.get('name')
        if session.session_data is not None:
            if self.session_name is None and 'session_name' in session.session_data:
                self.session_name = session.session_data.get('session_name')
            if 'session_state' in session.session_data:
                session_state_from_db = session.session_data.get('session_state')
                if session_state_from_db is not None and isinstance(session_state_from_db, dict) and len(session_state_from_db) > 0:
                    if len(self.session_state) > 0:
                        merge_dictionaries(session_state_from_db, self.session_state)
                    self.session_state = session_state_from_db
            if 'images' in session.session_data:
                images_from_db = session.session_data.get('images')
                if images_from_db is not None and isinstance(images_from_db, list):
                    if self.images is None:
                        self.images = []
                    self.images.extend([ImageArtifact.model_validate(img) for img in images_from_db])
            if 'videos' in session.session_data:
                videos_from_db = session.session_data.get('videos')
                if videos_from_db is not None and isinstance(videos_from_db, list):
                    if self.videos is None:
                        self.videos = []
                    self.videos.extend([VideoArtifact.model_validate(vid) for vid in videos_from_db])
            if 'audio' in session.session_data:
                audio_from_db = session.session_data.get('audio')
                if audio_from_db is not None and isinstance(audio_from_db, list):
                    if self.audio is None:
                        self.audio = []
                    self.audio.extend([AudioArtifact.model_validate(aud) for aud in audio_from_db])
        if session.extra_data is not None:
            if self.extra_data is not None:
                merge_dictionaries(session.extra_data, self.extra_data)
            self.extra_data = session.extra_data
        if session.memory is not None:
            if self.memory is None:
                self.memory = WorkflowMemory()
            try:
                if 'runs' in session.memory:
                    try:
                        self.memory.runs = [WorkflowRun(**m) for m in session.memory['runs']]
                    except Exception as e:
                        print(f'Failed to load runs from memory: {e}')
            except Exception as e:
                print(f'Failed to load WorkflowMemory: {e}')
        print(f'-*- WorkflowSession loaded: {session.session_id}')

    def load_session(self, force: bool = False) -> Optional[str]:
        if self.workflow_session is not None and not force:
            if self.session_id is not None and self.workflow_session.session_id == self.session_id:
                return self.workflow_session.session_id
        return self.session_id

    def new_session(self) -> None:
        self.workflow_session = None
        self.session_id = str(uuid4())
        self.load_session(force=True)

    def log_workflow_session(self):
        print(f'*********** Logging WorkflowSession: {self.session_id} ***********')

    def rename(self, name: str) -> None:
        self.name = name
        self.log_workflow_session()

    def rename_session(self, session_name: str):
        self.session_name = session_name
        self.log_workflow_session()


def merge_dictionaries(a: Dict[str, Any], b: Dict[str, Any]):
    for key in b:
        if key in a and isinstance(a[key], dict) and isinstance(b[key], dict):
            merge_dictionaries(a[key], b[key])
        else:
            a[key] = b[key]


def create_panel(content, title, border_style='blue'):
    return Panel(content, title=title, title_align='left', border_style=border_style, box=HEAVY, expand=True, padding=(1, 1))


def escape_markdown_tags(content: str, tags: Set[str]) -> str:
    escaped_content = content
    for tag in tags:
        escaped_content = escaped_content.replace(f'<{tag}>', f'&lt;{tag}&gt;')
        escaped_content = escaped_content.replace(f'</{tag}>', f'&lt;/{tag}&gt;')
    return escaped_content


def update_run_response_with_reasoning(run_response: Union[RunResponse, TeamRunResponse], reasoning_steps: List[ReasoningStep], reasoning_agent_messages: List[Message]) -> None:
    if run_response.extra_data is None:
        run_response.extra_data = RunResponseExtraData()
    if run_response.extra_data.reasoning_steps is None:
        run_response.extra_data.reasoning_steps = reasoning_steps
    else:
        run_response.extra_data.reasoning_steps.extend(reasoning_steps)
    if run_response.extra_data.reasoning_messages is None:
        run_response.extra_data.reasoning_messages = reasoning_agent_messages
    else:
        run_response.extra_data.reasoning_messages.extend(reasoning_agent_messages)


def format_tool_calls(tool_calls: List[Dict[str, Any]]) -> List[str]:
    formatted_tool_calls = []
    for tool_call in tool_calls:
        if 'tool_name' in tool_call and 'tool_args' in tool_call:
            tool_name = tool_call['tool_name']
            args_str = ''
            if 'tool_args' in tool_call and tool_call['tool_args'] is not None:
                args_str = ', '.join(f'{k}={v}' for k, v in tool_call['tool_args'].items())
            formatted_tool_calls.append(f'{tool_name}({args_str})')
    return formatted_tool_calls


def parse_response_model_str(content: str, response_model: Type[BaseModel]) -> Optional[BaseModel]:
    structured_output = None
    try:
        structured_output = response_model.model_validate_json(content)
    except (ValidationError, json.JSONDecodeError):
        content = content
        if '```json' in content:
            content = content.split('```json')[-1].split('```')[0].strip()
        elif '```' in content:
            content = content.split('```')[1].strip()
        content = re.sub(r'[*_`#]', '', content)
        content = content.replace('\n', ' ').replace('\r', '')
        content = re.sub(r'[\x00-\x1F\x7F]', '', content)
        def escape_quotes_in_values(match):
            key = match.group(1)
            value = match.group(2)
            escaped_value = value.replace('"', '\\"')
            return f'"{key}": {escaped_value}'
        content = re.sub(r'"(?P<key>[^"]+)"\s*:\s*"(?P<value>.*?)(?="\s*(?:,|\}))', escape_quotes_in_values, content)
        try:
            structured_output = response_model.model_validate_json(content)
        except (ValidationError, json.JSONDecodeError) as e:
            print(f'Failed to parse cleaned JSON: {e}')
            try:
                data = json.loads(content)
                structured_output = response_model.model_validate(data)
            except (ValidationError, json.JSONDecodeError) as e:
                print(f'Failed to parse as Python dict: {e}')
    return structured_output


def get_text_from_message(message: Union[List, Dict, str, Message]) -> str:
    if isinstance(message, str):
        return message
    if isinstance(message, list):
        text_messages = []
        if len(message) == 0:
            return ''
        if 'type' in message[0]:
            for m in message:
                m_type = m.get('type')
                if m_type is not None and isinstance(m_type, str):
                    m_value = m.get(m_type)
                    if m_value is not None and isinstance(m_value, str):
                        if m_type == 'text':
                            text_messages.append(m_value)
                        if m_type == 'image_url':
                            text_messages.append(f'Image: {m_value}')
                        else:
                            text_messages.append(f'{m_type}: {m_value}')
        elif 'role' in message[0]:
            for m in message:
                m_role = m.get('role')
                if m_role is not None and isinstance(m_role, str):
                    m_content = m.get('content')
                    if m_content is not None and isinstance(m_content, str):
                        if m_role == 'user':
                            text_messages.append(m_content)
        if len(text_messages) > 0:
            return '\n'.join(text_messages)
    if isinstance(message, dict):
        if 'content' in message:
            return get_text_from_message(message['content'])
    if isinstance(message, Message) and message.content is not None:
        return get_text_from_message(message.content)
    return ''
